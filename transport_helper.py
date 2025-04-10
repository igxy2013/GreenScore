from docx import Document
import base64
import io
from docx.shared import Inches
import os
from datetime import datetime
from word_template import replace_placeholders

def add_image_to_document(doc, image_data, bookmark_name=None, width=6.0):
    """
    向Word文档中添加图片
    
    参数:
        doc: Document对象
        image_data: base64编码的图片数据
        bookmark_name: 图片应该插入的书签名
        width: 图片宽度（英寸）
    """
    try:
        # 将base64编码的图片数据转换为二进制
        if image_data.startswith('data:image'):
            # 处理完整的Data URL
            image_data = image_data.split(',')[1]
        
        image_binary = base64.b64decode(image_data)
        image_stream = io.BytesIO(image_binary)
        
        # 遍历文档中的段落，查找包含{地图截图}的段落
        image_added = False
        for paragraph in doc.paragraphs:
            if "{地图截图}" in paragraph.text:
                # 清空段落内容
                paragraph.clear()
                # 添加图片到段落
                run = paragraph.add_run()
                run.add_picture(image_stream, width=Inches(width))
                image_added = True
                return True
        
        # 如果没有找到占位符，则尝试在文档末尾添加图片
        if not image_added:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(image_stream, width=Inches(width))
            return True
            
        return False
    except Exception as e:
        print(f"添加图片时出错: {str(e)}")
        return False

def add_stations_table(doc, stations):
    """
    向Word文档中添加站点表格
    
    参数:
        doc: Document对象
        stations: 站点列表，每个站点为字典，包含name, type, distance, detail字段
    """
    try:
        # 遍历文档中的段落，查找包含{公交站点列表}的段落
        for paragraph in doc.paragraphs:
            if "{公交站点列表}" in paragraph.text:
                # 在段落后添加表格
                table = doc.add_table(rows=1, cols=5)
                table.style = 'Table Grid'
                
                # 设置表头
                header_cells = table.rows[0].cells
                header_cells[0].text = "序号"
                header_cells[1].text = "站点名称"
                header_cells[2].text = "类型"
                header_cells[3].text = "距离(m)"
                header_cells[4].text = "详细信息"
                
                # 添加站点数据
                for i, station in enumerate(stations):
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(i + 1)
                    row_cells[1].text = station.get('name', '')
                    row_cells[2].text = station.get('type', '')
                    row_cells[3].text = str(station.get('distance', ''))
                    row_cells[4].text = station.get('detail', '')
                
                # 删除原段落
                p = paragraph._element
                p.getparent().remove(p)
                paragraph._p = None
                paragraph._element = None
                
                return True
        
        return False
    except Exception as e:
        print(f"添加站点表格时出错: {str(e)}")
        return False

def generate_transport_report(data):
    """
    生成公共交通站点分析报告
    
    参数:
        data: 包含以下字段的字典：
            - address: 项目地址
            - stations: 站点列表，每个站点包含 {name, type, distance, detail}
            - map_image: 地图截图的base64编码
            - conclusion: 分析结论，包含 {result6_1_2, result6_2_1, totalScore}
            - project_info: 项目基本信息
    
    返回:
        生成的文档路径
    """
    try:
        print("开始生成公共交通站点分析报告...")
        
        # 获取模板文件路径
        template_path = os.path.join("static", "templates", "公共交通站点分析报告.docx")
        if not os.path.exists(template_path):
            raise Exception(f"模板文件不存在: {template_path}")
        
        # 准备数据
        address = data.get('address', '')
        stations = data.get('stations', [])
        map_image = data.get('map_image', '')
        conclusion = data.get('conclusion', {})
        project_info = data.get('project_info', {})
        
        # 准备模板数据
        template_data = [
            # 项目基本信息
            {
                "项目名称": project_info.get('项目名称', ''),
                "项目地点": project_info.get('项目地点', address),
                "项目编号": project_info.get('项目编号', ''),
                "建设单位": project_info.get('建设单位', ''),
                "设计单位": project_info.get('设计单位', ''),
                "总用地面积": project_info.get('总用地面积', ''),
                "总建筑面积": project_info.get('总建筑面积', ''),
                "建筑密度": project_info.get('建筑密度', ''),
                "绿地率": project_info.get('绿地率', ''),
                "容积率": project_info.get('容积率', ''),
                "项目地址": address,
                "设计日期": datetime.now().strftime("%Y年%m月%d日"),
                # "地图截图": map_image,  # 这个在后面单独处理
                # "公交站点列表": stations_table,  # 这个在后面单独处理
                "结论": conclusion.get('result6_1_2', '') + '\n\n' + conclusion.get('result6_2_1', '') + '\n\n总得分：' + str(conclusion.get('totalScore', 0)) + '分'
            }
        ]
        
        # 添加条文得分数据
        template_data.append({
            "条文号": "6.1.2",
            "得分": "符合" if conclusion.get('result6_1_2', '').startswith('符合') else "不符合",
            "技术措施": conclusion.get('result6_1_2', '')
        })
        
        template_data.append({
            "条文号": "6.2.1",
            "得分": str(conclusion.get('totalScore', 0)),
            "技术措施": conclusion.get('result6_2_1', '')
        })
        
        # 使用word_template模块处理基本占位符
        print("调用replace_placeholders处理基本占位符...")
        output_path = replace_placeholders(template_path, template_data)
        
        # 处理图片和表格
        print("处理图片和表格...")
        doc = Document(output_path)
        
        # 添加地图截图
        if map_image:
            print("添加地图截图...")
            add_image_to_document(doc, map_image)
        
        # 添加站点表格
        if stations:
            print("添加站点表格...")
            add_stations_table(doc, stations)
        
        # 保存修改后的文档
        doc.save(output_path)
        
        # 返回生成的文档路径
        return output_path
        
    except Exception as e:
        print(f"生成公共交通站点分析报告失败: {str(e)}")
        import traceback
        print(traceback.format_exc())
        raise Exception(f"生成公共交通站点分析报告失败: {str(e)}")

# 如果需要测试，可以添加测试代码
if __name__ == "__main__":
    # 测试数据
    from datetime import datetime
    
    test_data = {
        "address": "成都市武侯区天府大道北段1700号",
        "stations": [
            {"name": "天府五街站", "type": "公交站", "distance": "120", "detail": "1路、2路、3路"},
            {"name": "天府三街站", "type": "公交站", "distance": "350", "detail": "4路、5路"},
            {"name": "天府二街站", "type": "地铁站", "distance": "480", "detail": "1号线、7号线"}
        ],
        "map_image": "", # 实际使用时需要提供base64编码的图片
        "conclusion": {
            "result6_1_2": "符合规范6.1.2要求。最近公交站距离为120m，最近地铁站距离为480m。",
            "result6_2_1": "按照规范6.2.1评分，总得分为8分，其中：\n- 最近公交站距离为120m，最近地铁站距离为480m，得4分；\n- 周边800m内有3个公交站（包括天府五街站、天府三街站等），有1个地铁站（包括天府二街站），总计5条线路，得4分。",
            "totalScore": 8
        },
        "project_info": {
            "项目名称": "测试项目",
            "项目地点": "成都市武侯区天府大道北段1700号",
            "项目编号": "TEST-2024-001",
            "建设单位": "测试建设公司",
            "设计单位": "测试设计院",
            "总用地面积": "5000",
            "总建筑面积": "20000",
            "建筑密度": "35",
            "绿地率": "30",
            "容积率": "4.0"
        }
    }
    
    output_path = generate_transport_report(test_data)
    print(f"报告已生成: {output_path}") 