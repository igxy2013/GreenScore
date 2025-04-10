import os
import traceback
import json
import uuid
import base64
from flask import request, jsonify, Response, url_for, render_template
from datetime import datetime
from flask_login import login_required

# 将所有路由函数放入一个init_routes函数中
def init_routes(app):
    """初始化公共交通分析相关的路由，将app对象作为参数传入"""
    
    # 设置全局日志记录器
    global logger
    logger = app.logger
    
    # 从app导入必要的模型类
    from models import Project
    
    @app.route('/api/fill_transport_report_template', methods=['POST'])
    @login_required
    def fill_transport_report_template():
        """填充公共交通站点分析报告模板"""
        try:
            data = request.json
            if not data:
                return jsonify({'success': False, 'message': '缺少必要的数据'}), 400
            
            # 获取必要的数据
            address = data.get('address', '未指定地址')
            stations = data.get('stations', [])
            map_image_data = data.get('mapImage', '')
            map_info = data.get('mapInfo', {})
            project_id = data.get('project_id')  # 获取项目ID
            
            # 检查必要的数据是否存在
            if not stations:
                return jsonify({'success': False, 'message': '缺少公交站点数据'}), 400
                
            # 创建临时目录（如果不存在）
            exports_dir = os.path.join(app.static_folder, 'exports')
            os.makedirs(exports_dir, exist_ok=True)
            
            # 为当前报告创建一个唯一的文件名
            timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
            unique_id = str(uuid.uuid4())[:8]
            report_filename = f"公交站点分析报告_{timestamp}_{unique_id}.docx"
            output_path = os.path.join(exports_dir, report_filename)
            
            # 模板文件路径
            template_path = os.path.join(app.static_folder, 'templates', '公共交通站点分析报告.docx')
            if not os.path.exists(template_path):
                app.logger.error(f"模板文件不存在: {template_path}")
                return jsonify({'success': False, 'message': f'模板文件不存在，请确认 {template_path} 文件已准备好'}), 500
            
            # 复制模板文件到输出路径
            import shutil
            shutil.copy2(template_path, output_path)
            app.logger.info(f"已将模板复制到: {output_path}")
            
            # 获取项目信息
            project_info = {}
            if project_id:
                try:
                    # 从数据库获取项目信息
                    project = Project.query.get(project_id)
                    if project:
                        project_info = {
                            '项目名称': project.name or '',
                            '项目地点': project.location or '',
                            '项目编号': project.code or '',  # 使用项目编号字段
                            '建设单位': project.construction_unit or '',
                            '设计单位': project.design_unit or '',
                            '总建筑面积': str(project.total_building_area or '') + ' 平方米' if project.total_building_area else '',
                            '总用地面积': str(project.total_land_area or '') + ' 平方米' if project.total_land_area else '',
                            '建筑密度': str(project.building_density or '') + ' %' if project.building_density else '',
                            '容积率': str(project.plot_ratio or '') if project.plot_ratio else '',
                            '绿地率': str(project.green_ratio or '') + ' %' if project.green_ratio else '',
                            '设计日期': datetime.now().strftime('%Y年%m月%d日')
                        }
                        
                        # 表单数据中可能有其他字段，但我们已经直接使用了project.code
                        # 所以不需要再次从form_data获取设计编号
                        
                        app.logger.info(f"成功获取项目信息: {project_info}")
                    else:
                        app.logger.warning(f"未找到ID为{project_id}的项目")
                except Exception as e:
                    app.logger.error(f"获取项目信息时出错: {str(e)}")
                    app.logger.error(traceback.format_exc())
                    # 继续执行，即使没有项目信息
            else:
                app.logger.info("未提供项目ID，将使用默认项目信息")
                
            # 确保项目信息中的占位符值不是None
            for key in list(project_info.keys()):
                if project_info[key] is None:
                    project_info[key] = ''
                    
            app.logger.info(f"最终项目信息: {json.dumps(project_info, ensure_ascii=False)}")
            
            try:
                from docx import Document
                from docx.shared import Inches, Pt
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                import io
                import base64
                import re
                
                # 打开复制的模板文件
                doc = Document(output_path)
                
                # 查找并替换文本占位符
                app.logger.info(f"开始查找文本占位符...")
                
                # 初始化标记
                stations_table_added = False
                placeholder_found = False
                map_image_added = False
                temp_image_path = None
                
                # 在处理段落占位符前，先扫描整个文档以识别所有可能的占位符
                all_placeholders = []
                for i, paragraph in enumerate(doc.paragraphs):
                    original_text = paragraph.text
                    
                    # 使用正则表达式查找所有花括号格式的文本
                    placeholders = re.findall(r'\{([^}]+)\}', original_text)
                    for p in placeholders:
                        if p not in all_placeholders:
                            all_placeholders.append(p)
                            
                # 也检查表格中的占位符
                for table_idx, table in enumerate(doc.tables):
                    for row_idx, row in enumerate(table.rows):
                        for cell_idx, cell in enumerate(row.cells):
                            for para_idx, paragraph in enumerate(cell.paragraphs):
                                original_text = paragraph.text
                                placeholders = re.findall(r'\{([^}]+)\}', original_text)
                                for p in placeholders:
                                    if p not in all_placeholders:
                                        all_placeholders.append(p)
                
                app.logger.info(f"在整个文档中找到的所有占位符: {all_placeholders}")
                
                # 更新项目信息字典，确保包含所有可能的占位符变体
                extended_project_info = project_info.copy()
                
                # 添加潜在的格式变体
                for placeholder in all_placeholders:
                    # 检查是否有相似的键
                    for key in project_info.keys():
                        # 简单的相似度检查 - 移除空格和标点后比较
                        cleaned_placeholder = re.sub(r'[\s\W]+', '', placeholder.lower())
                        cleaned_key = re.sub(r'[\s\W]+', '', key.lower())
                        
                        if cleaned_placeholder == cleaned_key and placeholder not in extended_project_info:
                            extended_project_info[placeholder] = project_info[key]
                            app.logger.info(f"添加占位符变体映射: {placeholder} -> {key} = {project_info[key]}")
                
                app.logger.info(f"扩展后的项目信息字典: {json.dumps(extended_project_info, ensure_ascii=False)}")
                
                # 处理地图图片
                if map_image_data and map_image_data.startswith('data:image/png;base64,'):
                    try:
                        # 从base64数据中提取图片内容
                        map_image_data = map_image_data.replace('data:image/png;base64,', '')
                        image_binary = base64.b64decode(map_image_data)
                        
                        # 保存图片到临时文件
                        temp_image_path = os.path.join(exports_dir, f"map_image_{unique_id}.png")
                        with open(temp_image_path, 'wb') as img_file:
                            img_file.write(image_binary)
                            
                        app.logger.info(f"地图图片已保存到临时文件: {temp_image_path}")
                    except Exception as e:
                        app.logger.error(f"处理地图图片数据时出错: {str(e)}")
                        # 如果Canvas生成的图片处理失败，尝试使用百度地图静态图API
                        if map_info and map_info.get('center'):
                            try:
                                app.logger.info("尝试使用百度地图静态图API生成地图...")
                                # 这里可以通过调用百度地图静态图API来获取地图图片
                                # 但需要额外的处理
                            except Exception as e2:
                                app.logger.error(f"使用百度地图静态图API生成地图时出错: {str(e2)}")
                else:
                    app.logger.warning("未收到有效的地图图片数据")
                    # 可以在这里添加使用百度地图静态图API的备用方案
                
                # 遍历所有段落查找占位符
                for i, paragraph in enumerate(doc.paragraphs):
                    original_text = paragraph.text
                    app.logger.info(f"检查段落 {i+1}: \"{original_text}\"")
                    
                    # 替换地址占位符
                    if '{地址}' in original_text:
                        # 获取整个段落的文本并替换占位符
                        new_text = original_text.replace('{地址}', address)
                        app.logger.info(f"替换地址占位符: '{original_text}' -> '{new_text}'")
                        
                        # 清空段落中所有runs，但保持段落本身
                        p = paragraph._p
                        for run in list(paragraph.runs):  # 使用列表复制，因为我们在修改集合
                            p.remove(run._r)
                            
                        # 添加新的文本
                        paragraph.add_run(new_text)
                    
                    # 替换日期占位符
                    if '{设计日期}' in paragraph.text:
                        # 获取整个段落的文本并替换占位符
                        new_text = paragraph.text.replace('{设计日期}', datetime.now().strftime('%Y年%m月%d日'))
                        app.logger.info(f"替换日期占位符: '{paragraph.text}' -> '{new_text}'")
                        
                        # 清空段落中所有runs，但保持段落本身
                        p = paragraph._p
                        for run in list(paragraph.runs):  # 使用列表复制，因为我们在修改集合
                            p.remove(run._r)
                            
                        # 添加新的文本
                        paragraph.add_run(new_text)
                    
                    # 替换项目信息占位符
                    # 创建一个副本，因为我们将在迭代过程中修改文本
                    original_text = paragraph.text
                    new_text = original_text
                    text_changed = False
                    
                    # 先检查段落文本中是否包含任何花括号
                    if '{' in original_text and '}' in original_text:
                        app.logger.info(f"段落 {i+1} 包含花括号，可能包含占位符")
                        
                        # 使用正则表达式查找所有占位符
                        import re
                        placeholders = re.findall(r'\{([^}]+)\}', original_text)
                        app.logger.info(f"找到的占位符: {placeholders}")
                        
                        # 检查是否包含任何项目信息占位符
                        for placeholder, value in extended_project_info.items():
                            placeholder_text = '{' + placeholder + '}'
                            if placeholder_text in new_text:
                                new_text = new_text.replace(placeholder_text, value)
                                text_changed = True
                                app.logger.info(f"替换了项目信息占位符 {placeholder_text} -> {value}")
                    
                    # 如果文本已更改，应用新文本
                    if text_changed:
                        app.logger.info(f"段落 {i+1} 文本已更改: '{original_text}' -> '{new_text}'")
                        # 清空段落中所有runs，但保持段落本身
                        p = paragraph._p
                        for run in list(paragraph.runs):  # 使用列表复制，因为我们在修改集合
                            p.remove(run._r)
                            
                        # 添加新的文本
                        paragraph.add_run(new_text)
                    
                    # 替换地图占位符，并插入地图图片
                    if '{地图截图}' in paragraph.text and temp_image_path:
                        app.logger.info(f"找到地图截图占位符")
                        
                        # 设置段落居中对齐
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # 清除段落中的所有runs，但保持段落本身
                        p = paragraph._p
                        for run in list(paragraph.runs):  # 使用列表复制，因为我们在修改集合
                            p.remove(run._r)
                        
                        # 在清空后的段落中添加图片
                        try:
                            run = paragraph.add_run()
                            run.add_picture(temp_image_path, width=Inches(6))
                            map_image_added = True
                            app.logger.info(f"成功在占位符位置添加了地图图片")
                        except Exception as e:
                            app.logger.error(f"添加地图图片时出错: {str(e)}")
                            # 如果添加图片失败，添加文本提示
                            paragraph.add_run(f"{address} 周边公交站位置情况（图片加载失败）")
                        
                    elif '{地图截图}' in paragraph.text:
                        # 如果没有图片，只添加标题
                        # 清空段落中所有runs
                        p = paragraph._p
                        for run in list(paragraph.runs):  # 使用列表复制，因为我们在修改集合
                            p.remove(run._r)
                        # 添加新的标题文本
                        run = paragraph.add_run(f"{address} 周边公交站位置情况")
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        app.logger.info(f"替换了地图截图占位符（无图片）")
                    
                    # 处理公交站点列表占位符
                    if '{公交站点列表}' in paragraph.text:
                        app.logger.info(f"找到公交站点列表占位符")
                        placeholder_found = True
                        
                        # 创建表格
                        table = doc.add_table(rows=1, cols=5)
                        table.style = 'Table Grid'
                        
                        # 设置表头
                        header_cells = table.rows[0].cells
                        header_cells[0].text = "序号"
                        header_cells[1].text = "站点名称"
                        header_cells[2].text = "类型"
                        header_cells[3].text = "距离（米）"
                        header_cells[4].text = "详细信息"
                        
                        # 添加数据行
                        for idx, station in enumerate(stations):
                            row = table.add_row()
                            cells = row.cells
                            cells[0].text = str(idx + 1)
                            cells[1].text = station.get('name', '')
                            cells[2].text = station.get('distance', '')
                            cells[3].text = station.get('address', '无详细信息')
                        
                        # 直接用表格替换段落，而不是添加到段落后面
                        try:
                            # 获取段落的父元素
                            parent = paragraph._p.getparent()
                            # 获取段落在父元素中的索引
                            index = parent.index(paragraph._p)
                            # 在段落的位置插入表格
                            parent.insert(index, table._tbl)
                            # 移除原始段落
                            parent.remove(paragraph._p)
                            
                            stations_table_added = True
                            app.logger.info(f"成功用表格替换了占位符段落")
                        except Exception as e:
                            app.logger.error(f"替换段落为表格时出错: {str(e)}")
                            os.remove(output_path)  # 删除临时文件
                            return jsonify({'success': False, 'message': f'替换段落为表格失败: {str(e)}'}), 500
                
                # 如果没有找到占位符，返回错误
                if not placeholder_found:
                    app.logger.warning(f"在模板中未找到'{公交站点列表}'占位符")
                    os.remove(output_path)  # 删除临时文件
                    return jsonify({'success': False, 'message': '在模板中未找到"{公交站点列表}"占位符，请检查模板文件'}), 400
                
                # 如果找到占位符但无法插入表格，返回错误
                if placeholder_found and not stations_table_added:
                    app.logger.warning(f"无法在占位符位置插入表格")
                    os.remove(output_path)  # 删除临时文件
                    return jsonify({'success': False, 'message': '虽然找到了占位符，但无法在此位置插入表格，请检查模板格式'}), 500
                
                # 保存文档
                doc.save(output_path)
                app.logger.info(f"已保存文档到: {output_path}")
                
                # 清理临时文件
                if temp_image_path and os.path.exists(temp_image_path):
                    os.remove(temp_image_path)
                    app.logger.info(f"删除临时图片文件: {temp_image_path}")
                
                # 创建下载链接
                download_url = url_for('static', filename=f'exports/{report_filename}')
                app.logger.info(f"生成的下载链接: {download_url}")
                
                return jsonify({
                    'success': True, 
                    'message': '公交站点分析报告生成成功',
                    'download_url': download_url
                })
                
            except Exception as e:
                app.logger.error(f"处理Word文档时出错: {str(e)}")
                app.logger.error(traceback.format_exc())
                # 删除临时文件
                if os.path.exists(output_path):
                    os.remove(output_path)
                # 清理临时图片文件
                if temp_image_path and os.path.exists(temp_image_path):
                    os.remove(temp_image_path)
                return jsonify({'success': False, 'message': f'处理Word文档时出错: {str(e)}'}), 500
                
        except Exception as e:
            app.logger.error(f"生成公交站点分析报告时出错: {str(e)}")
            app.logger.error(traceback.format_exc())
            return jsonify({'success': False, 'message': f'处理请求失败: {str(e)}'}), 500

    @app.route('/public_transport_analysis')
    @login_required
    def public_transport_analysis():
        """公共交通分析报告生成页面"""
        # 获取可能的项目ID参数
        project_id = request.args.get('project_id')
        project = None
        
        if project_id:
            try:
                project = Project.query.get(project_id)
            except Exception as e:
                app.logger.error(f"获取项目信息失败: {str(e)}")
                
        return render_template('public_transport_analysis.html', project=project)

    @app.route('/api/generate_transport_report', methods=['POST'])
    @login_required
    def generate_transport_report():
        """生成公共交通分析报告"""
        try:
            # 获取请求数据
            data = request.json
            if not data:
                return jsonify({'success': False, 'message': '缺少必要的数据'}), 400
            
            # 提取项目信息
            project_info = data.get('project', {})
            project_name = project_info.get('name', '未命名项目')
            project_location = project_info.get('location', '')
            project_coordinates = project_info.get('coordinates', {})
            
            # 提取分析结果
            analysis_results = data.get('analysis', {})
            stations = data.get('stations', [])
            
            # 记录处理信息
            app.logger.info(f"开始生成公共交通分析报告: 项目 '{project_name}'")
            
            # 准备替换模板中的数据
            template_data = [{
                # 项目基本信息
                '项目名称': project_name,
                '项目地点': project_location,
                '经度': project_coordinates.get('lng', ''),
                '纬度': project_coordinates.get('lat', ''),
                '分析日期': datetime.now().strftime('%Y年%m月%d日'),
                
                # 分析结果数据
                '公交站总数': str(analysis_results.get('busStations', {}).get('total', 0)),
                '地铁站总数': str(analysis_results.get('subwayStations', {}).get('total', 0)),
                '500米内公交站': str(analysis_results.get('busStations', {}).get('qualified', 0)),
                '800米内地铁站': str(analysis_results.get('subwayStations', {}).get('qualified', 0)),
                
                # 最近站点信息
                '最近公交站名称': analysis_results.get('busStations', {}).get('nearest', {}).get('name', '无'),
                '最近公交站距离': str(analysis_results.get('busStations', {}).get('nearest', {}).get('distance', 0)) + '米',
                '最近地铁站名称': analysis_results.get('subwayStations', {}).get('nearest', {}).get('name', '无'),
                '最近地铁站距离': str(analysis_results.get('subwayStations', {}).get('nearest', {}).get('distance', 0)) + '米',
                
                # 评价结果
                '评价结果': analysis_results.get('evaluation', {}).get('result', '不符合'),
                '评分': str(analysis_results.get('evaluation', {}).get('score', 0)),
            }]
            
            # 添加站点列表数据
            for i, station in enumerate(stations, 1):
                if i > 10:  # 最多处理10个站点
                    break
                    
                template_data[0][f'站点{i}名称'] = station.get('name', '')
                template_data[0][f'站点{i}类型'] = station.get('type', '')
                template_data[0][f'站点{i}距离'] = str(station.get('distance', 0)) + '米'
            
            # 生成Word文档
            timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
            template_path = '公共交通分析报告.docx'
            output_filename = f"公共交通分析报告_{project_name}_{timestamp}.docx"
            output_dir = os.path.join(app.config['EXPORT_FOLDER'])
            os.makedirs(output_dir, exist_ok=True)
            output_path = os.path.join(output_dir, output_filename)
            
            # 使用现有的Word模板处理函数
            from word_template import process_template
            result = process_template(template_data, template_path=template_path, output_path=output_path)
            
            if result and os.path.exists(output_path):
                app.logger.info(f"公共交通分析报告生成成功: {output_path}")
                
                # 返回下载链接
                download_url = url_for('static', filename=f'exports/{output_filename}')
                return jsonify({
                    'success': True, 
                    'message': '公共交通分析报告生成成功',
                    'downloadUrl': download_url
                })
            else:
                app.logger.error(f"公共交通分析报告生成失败")
                return jsonify({'success': False, 'message': '报告生成失败，请稍后重试'}), 500
                
        except Exception as e:
            app.logger.error(f"生成公共交通分析报告时出错: {str(e)}")
            app.logger.error(traceback.format_exc())
            return jsonify({'success': False, 'message': f'处理请求失败: {str(e)}'}), 500
    
    # 返回注册的路由函数，以便在需要时可以单独访问
    return {
        'fill_transport_report_template': fill_transport_report_template,
        'public_transport_analysis': public_transport_analysis,
        'generate_transport_report': generate_transport_report
    }
