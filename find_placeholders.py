import os
import sys
import re
import json
from docx import Document
from collections import defaultdict
import pandas as pd
from datetime import datetime

def analyze_template(template_path):
    """
    分析Word模板中的所有占位符，返回找到的所有占位符及其位置信息
    """
    if not os.path.exists(template_path):
        print(f"错误: 模板文件 {template_path} 不存在")
        return None
        
    print(f"\n正在分析模板: {template_path}")
    placeholders = defaultdict(list)
    
    try:
        # 加载Word文档
        doc = Document(template_path)
        
        # 检查段落中的占位符
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text
            if '{' in text and '}' in text:
                # 使用正则表达式查找所有占位符
                matches = re.findall(r'\{([^}]+)\}', text)
                if matches:
                    for placeholder in matches:
                        placeholders[placeholder].append({
                            'type': 'paragraph',
                            'index': i,
                            'text': text
                        })
        
        # 检查表格中的占位符
        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    for p_idx, paragraph in enumerate(cell.paragraphs):
                        text = paragraph.text
                        if '{' in text and '}' in text:
                            # 使用正则表达式查找所有占位符
                            matches = re.findall(r'\{([^}]+)\}', text)
                            if matches:
                                for placeholder in matches:
                                    placeholders[placeholder].append({
                                        'type': 'table_cell',
                                        'table_index': t_idx,
                                        'row_index': r_idx,
                                        'column_index': c_idx,
                                        'paragraph_index': p_idx,
                                        'text': text
                                    })
        
        return dict(placeholders)
    except Exception as e:
        print(f"分析模板时出错: {str(e)}")
        return None

def generate_report(templates_dir, output_format='html'):
    """
    分析指定目录下的所有Word模板，并生成包含所有占位符信息的报告
    """
    if not os.path.exists(templates_dir):
        print(f"错误: 模板目录 {templates_dir} 不存在")
        return
    
    all_results = {}
    template_files = [f for f in os.listdir(templates_dir) if f.endswith('.docx')]
    
    if not template_files:
        print(f"在 {templates_dir} 目录中未找到Word模板文件")
        return
    
    print(f"找到 {len(template_files)} 个Word模板文件")
    
    for template_file in template_files:
        template_path = os.path.join(templates_dir, template_file)
        template_results = analyze_template(template_path)
        if template_results:
            all_results[template_file] = template_results
    
    # 生成报告
    timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
    
    # 创建HTML报告
    if output_format == 'html' or output_format == 'both':
        html_output = f"<html><head><title>Word模板占位符分析报告</title>"
        html_output += "<style>body{font-family:Arial,sans-serif;margin:20px;} "
        html_output += "table{border-collapse:collapse;width:100%;margin-bottom:30px;} "
        html_output += "th,td{border:1px solid #ddd;padding:8px;text-align:left;} "
        html_output += "th{background-color:#f2f2f2;} "
        html_output += "h1,h2{color:#333;} "
        html_output += "h3{color:#555;margin-top:20px;} "
        html_output += ".highlight{background-color:#ffffcc;} "
        html_output += "</style></head><body>"
        html_output += f"<h1>Word模板占位符分析报告</h1>"
        html_output += f"<p>生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>"
        
        # 首先生成占位符汇总表
        html_output += "<h2>所有占位符汇总</h2>"
        html_output += "<table><tr><th>占位符</th><th>出现在模板</th><th>出现次数</th></tr>"
        
        all_placeholders = set()
        for template, placeholders in all_results.items():
            all_placeholders.update(placeholders.keys())
        
        for placeholder in sorted(all_placeholders):
            template_counts = []
            for template, placeholders in all_results.items():
                if placeholder in placeholders:
                    template_counts.append(f"{template} ({len(placeholders[placeholder])}次)")
            
            html_output += f"<tr><td>{{{{ {placeholder} }}}}</td><td>{', '.join(template_counts)}</td><td>{sum(len(all_results[t][placeholder]) for t in all_results if placeholder in all_results[t])}</td></tr>"
        
        html_output += "</table>"
        
        # 为每个模板生成详细报告
        html_output += "<h2>模板详细分析</h2>"
        
        for template, placeholders in all_results.items():
            html_output += f"<h3>模板: {template}</h3>"
            html_output += "<table><tr><th>占位符</th><th>位置</th><th>上下文</th></tr>"
            
            for placeholder, occurrences in placeholders.items():
                for occurrence in occurrences:
                    location = ""
                    if occurrence['type'] == 'paragraph':
                        location = f"段落 #{occurrence['index'] + 1}"
                    else:
                        location = f"表格 #{occurrence['table_index'] + 1}, 行 #{occurrence['row_index'] + 1}, 列 #{occurrence['column_index'] + 1}"
                    
                    # 高亮显示占位符
                    context = occurrence['text'].replace(f"{{{placeholder}}}", f"<span class='highlight'>{{{placeholder}}}</span>")
                    
                    html_output += f"<tr><td>{{{{ {placeholder} }}}}</td><td>{location}</td><td>{context}</td></tr>"
            
            html_output += "</table>"
        
        html_output += "</body></html>"
        
        # 保存HTML报告
        html_report_path = f"placeholder_report_{timestamp}.html"
        with open(html_report_path, 'w', encoding='utf-8') as f:
            f.write(html_output)
        
        print(f"HTML报告已生成: {html_report_path}")
    
    # 创建Excel报告
    if output_format == 'excel' or output_format == 'both':
        # 创建一个ExcelWriter对象
        excel_report_path = f"placeholder_report_{timestamp}.xlsx"
        with pd.ExcelWriter(excel_report_path) as writer:
            # 创建摘要工作表
            summary_data = []
            for placeholder in sorted(set(ph for result in all_results.values() for ph in result.keys())):
                for template, placeholders in all_results.items():
                    if placeholder in placeholders:
                        for occurrence in placeholders[placeholder]:
                            if occurrence['type'] == 'paragraph':
                                location = f"段落 #{occurrence['index'] + 1}"
                            else:
                                location = f"表格 #{occurrence['table_index'] + 1}, 行 #{occurrence['row_index'] + 1}, 列 #{occurrence['column_index'] + 1}"
                            
                            summary_data.append({
                                '模板': template,
                                '占位符': f"{{{placeholder}}}",
                                '位置': location,
                                '上下文': occurrence['text']
                            })
            
            # 创建摘要DataFrame并保存到Excel
            if summary_data:
                summary_df = pd.DataFrame(summary_data)
                summary_df.to_excel(writer, sheet_name='占位符摘要', index=False)
            
            # 为每个模板创建单独的工作表
            for template, placeholders in all_results.items():
                template_data = []
                for placeholder, occurrences in placeholders.items():
                    for occurrence in occurrences:
                        if occurrence['type'] == 'paragraph':
                            location = f"段落 #{occurrence['index'] + 1}"
                        else:
                            location = f"表格 #{occurrence['table_index'] + 1}, 行 #{occurrence['row_index'] + 1}, 列 #{occurrence['column_index'] + 1}"
                        
                        template_data.append({
                            '占位符': f"{{{placeholder}}}",
                            '位置': location,
                            '上下文': occurrence['text']
                        })
                
                # 创建模板的DataFrame并保存到Excel
                if template_data:
                    # 限制工作表名长度为31个字符（Excel限制）
                    sheet_name = template[:31].replace('[', '').replace(']', '').replace('*', '')
                    template_df = pd.DataFrame(template_data)
                    template_df.to_excel(writer, sheet_name=sheet_name, index=False)
        
        print(f"Excel报告已生成: {excel_report_path}")
    
    # 创建JSON报告
    if output_format == 'json' or output_format == 'both':
        json_report_path = f"placeholder_report_{timestamp}.json"
        with open(json_report_path, 'w', encoding='utf-8') as f:
            json.dump(all_results, f, ensure_ascii=False, indent=2)
        
        print(f"JSON报告已生成: {json_report_path}")
    
    return all_results

if __name__ == "__main__":
    print("Word模板占位符分析工具")
    print("----------------------")
    
    # 默认目录为static/templates
    templates_dir = "static/templates"
    
    # 如果提供了命令行参数，使用指定目录
    if len(sys.argv) > 1:
        templates_dir = sys.argv[1]
    
    # 允许通过命令行参数指定输出格式
    output_format = 'html'
    if len(sys.argv) > 2:
        if sys.argv[2].lower() in ['html', 'excel', 'json', 'both']:
            output_format = sys.argv[2].lower()
    
    print(f"分析目录: {templates_dir}")
    print(f"输出格式: {output_format}")
    
    # 运行分析并生成报告
    generate_report(templates_dir, output_format) 