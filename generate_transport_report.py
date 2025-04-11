from word_template import replace_placeholders
import os
import json
import base64
import io
from datetime import datetime
from flask import current_app
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
# 导入我们的自定义占位符处理函数
from process_custom_placeholders import process_custom_placeholders

def generate_transport_report(data):
    """
    生成公共交通站点分析报告
    
    参数:
        data: 包含以下字段的字典：
            - address: 详细地址
            - stations: 站点列表，每个站点包含 {name, type, distance, detail}
            - mapImage: 地图截图的base64编码
            - conclusion: 分析结论，包含 {result6_1_2, result6_2_1, totalScore}
            - project_info: 项目基本信息
    
    返回:
        生成的文档路径
    """
    try:
        print("\n=== 开始生成公共交通站点分析报告 ===")
        
        # 获取模板文件路径
        template_path = os.path.join("static", "templates", "公共交通站点分析报告.docx")
        if not os.path.exists(template_path):
            raise Exception(f"模板文件不存在: {template_path}")
        
        print(f"使用模板: {template_path}")
        
        # 提取数据
        address = data.get('address', '')
        stations = data.get('stations', [])
        map_image = data.get('mapImage', '')
        conclusion = data.get('conclusion', {})
        project_info = data.get('project_info', {})
        project_id = data.get('project_id', '')
        
        print(f"项目ID: {project_id}")
        print(f"地址: {address}")
        print(f"站点数量: {len(stations)}")
        print(f"结论: {conclusion}")
        
        # 输出项目信息
        print("\n项目信息内容:")
        print(json.dumps(project_info, ensure_ascii=False, indent=2))
        
        # 准备数据 - 符合word_template.py期望的格式
        # 创建项目基础信息字典作为列表第一个元素
        project_data = {}
        
        # 添加项目基本信息
        if isinstance(project_info, dict):
            for key, value in project_info.items():
                project_data[key] = value
        
        # 用户提到的实际存在于模板中的占位符
        # 1. 项目名称等基础信息 - 已从project_info中添加
        # 2. 地址相关
        project_data['详细地址'] = address
        project_data['地址'] = address
        
        # 3. 结论
        if isinstance(conclusion, dict):
            conclusion_text = f"{conclusion.get('result6_1_2', '')}\n\n{conclusion.get('result6_2_1', '')}总得分：{conclusion.get('totalScore', 0)}分"
            project_data['结论'] = conclusion_text
        
        # 4. 设计日期
        current_date = datetime.now().strftime("%Y年%m月%d日")
        project_data['设计日期'] = current_date
        
        # 5. 地图截图 - 为了兼容现有的replace_placeholders函数，提供文本占位符
        # 但实际图片会在后续步骤中处理
        project_data['地图截图'] = "{地图截图}"
        
        # 6. 公交站点列表 - 同样提供文本占位符
        project_data['公交站点列表'] = "{公交站点列表}"
        
        # 创建最终的数据列表 - 第一个元素包含所有信息
        final_data = [project_data]
        
        # 将站点列表添加为单独的数据项（为表格处理准备）
        for station in stations:
            final_data.append(station)
        
        print("\n将传递给模板系统的数据结构:")
        print(f"数据列表长度: {len(final_data)}")
        print(f"第一个元素包含的字段: {', '.join(k for k in final_data[0].keys() if k not in ['地图截图', '公交站点列表'])}")
        
        # 第一步：调用基本的Word模板处理函数处理文本占位符
        temp_output_path = replace_placeholders(template_path, final_data)
        
        # 检查第一步处理的结果文件是否存在
        if not os.path.exists(temp_output_path):
            raise Exception(f"文本占位符处理失败，输出文件不存在: {temp_output_path}")
        
        print(f"\n基本文本替换完成: {temp_output_path}")
        
        # 设置替换后的文本为宋体小四
        set_replaced_text_font(temp_output_path, project_data)
        
        # 准备自定义占位符处理的数据
        custom_data = {
            '地图截图': map_image,
            'stations': stations
        }
        
        # 第二步：处理特殊占位符（图片和表格）
        final_output_path = process_custom_placeholders(temp_output_path, custom_data)
        
        print(f"\n报告生成完成: {final_output_path}")
        
        return final_output_path
    except Exception as e:
        print(f"生成报告失败: {str(e)}")
        import traceback
        print(traceback.format_exc())
        raise Exception(f"生成报告失败: {str(e)}")

def set_replaced_text_font(docx_path, project_data):
    """
    设置被替换的文本为宋体小四
    
    参数:
        docx_path: Word文档路径
        project_data: 包含替换值的字典
    """
    try:
        # 加载文档
        doc = Document(docx_path)
        
        # 获取替换值列表，用于识别哪些文本是我们替换的
        replacement_values = []
        for key, value in project_data.items():
            if key not in ['地图截图', '公交站点列表'] and value:
                replacement_values.append(str(value))
        
        # 遍历所有段落
        for paragraph in doc.paragraphs:
            # 检查段落文本是否包含我们的替换值
            for value in replacement_values:
                if value in paragraph.text:
                    # 找到了可能包含替换值的段落，设置所有run的字体
                    for run in paragraph.runs:
                        if any(val in run.text for val in replacement_values):
                            # 这个run包含我们替换的值
                            run.font.name = '宋体'
                            run.font.size = Pt(12)  # 小四字号为12磅
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 遍历所有表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        # 检查段落文本是否包含我们的替换值
                        for value in replacement_values:
                            if value in paragraph.text:
                                # 找到了可能包含替换值的段落，设置所有run的字体
                                for run in paragraph.runs:
                                    if any(val in run.text for val in replacement_values):
                                        # 这个run包含我们替换的值
                                        run.font.name = '宋体'
                                        run.font.size = Pt(12)  # 小四字号为12磅
                                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 保存修改后的文档
        doc.save(docx_path)
        print("设置替换文本字体为宋体小四完成")
        
    except Exception as e:
        print(f"设置替换文本字体时出错: {str(e)}")
        import traceback
        print(traceback.format_exc())