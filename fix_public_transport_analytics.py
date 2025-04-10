import os
import re
import json
import logging
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import base64
from datetime import datetime

# 配置日志记录
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('transport_report_fix.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

def fix_transport_report_template(
    template_path="static/templates/公共交通站点分析报告.docx",
    output_path="fixed_公共交通站点分析报告.docx",
    test_mode=True,
    map_image_path=None
):
    """
    修复公共交通站点分析报告模板，确保占位符正确替换和表格正确创建
    
    Parameters:
    - template_path: 模板文件路径
    - output_path: 输出文件路径
    - test_mode: 是否使用测试数据
    - map_image_path: 可选的地图图片路径
    """
    if not os.path.exists(template_path):
        logger.error(f"错误: 模板文件 {template_path} 不存在")
        return False
    
    logger.info(f"开始修复公共交通站点分析报告模板")
    logger.info(f"模板文件: {template_path}")
    logger.info(f"输出文件: {output_path}")
    
    try:
        # 准备测试数据
        if test_mode:
            # 模拟项目信息
            project_info = {
                '项目名称': '测试项目',
                '项目地点': '四川省成都市高新区',
                '项目编号': 'TEST-2023-001',
                '建设单位': '成都测试建设有限公司',
                '设计单位': '成都测试设计有限公司',
                '总用地面积': '5000',
                '总建筑面积': '20000',
                '建筑密度': '40',
                '绿地率': '30',
                '容积率': '4.0',
                '项目地址': '四川省成都市高新区天府大道1234号'
            }
            
            # 模拟站点数据
            stations = [
                {
                    "index": 1,
                    "name": "高新西站",
                    "type": "公交站",
                    "distance": "150",
                    "detail": "1路; 2路; 快速公交1号线",
                    "location": {"lng": 104.12345, "lat": 30.12345}
                },
                {
                    "index": 2,
                    "name": "天府三街",
                    "type": "地铁站",
                    "distance": "300",
                    "detail": "地铁1号线; 地铁7号线",
                    "location": {"lng": 104.54321, "lat": 30.54321}
                },
                {
                    "index": 3,
                    "name": "南湖北路北",
                    "type": "公交站",
                    "distance": "267",
                    "detail": "512路; T202路; TG07路",
                    "location": {"lng": 104.04978622590484, "lat": 30.510216076931695}
                }
            ]
            
            # 模拟结论数据
            conclusion = {
                "result6_1_2": "符合规范6.1.2要求。最近公交站距离为150m，最近地铁站距离为300m。",
                "result6_2_1": "按照规范6.2.1评分，总得分为8分，其中：\n- 最近公交站距离为150m，最近地铁站距离为300m，得4分；\n- 周边800m内有2个公交站（包括高新西站、南湖北路北），有1个地铁站（包括天府三街），总计5条线路，得4分。",
                "totalScore": 8
            }
            
            logger.info("使用测试数据进行修复")
        else:
            # 在实际使用中，这里应该从参数中接收真实数据
            logger.error("非测试模式下需要提供真实数据，但目前未提供")
            return False
        
        # 加载Word文档
        doc = Document(template_path)
        logger.info(f"成功加载Word模板")
        
        # 查找文档中的所有占位符
        all_placeholders = set()
        
        # 检查段落中的占位符
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text
            if '{' in text and '}' in text:
                matches = re.findall(r'\{([^}]+)\}', text)
                all_placeholders.update(matches)
        
        # 检查表格中的占位符
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        text = paragraph.text
                        if '{' in text and '}' in text:
                            matches = re.findall(r'\{([^}]+)\}', text)
                            all_placeholders.update(matches)
        
        logger.info(f"在文档中找到的所有占位符: {all_placeholders}")
        
        # 检查项目信息中是否缺少占位符对应的值
        missing_placeholders = []
        for placeholder in all_placeholders:
            if placeholder not in project_info and placeholder not in ['地图截图', '公交站点列表', '结论']:
                missing_placeholders.append(placeholder)
        
        if missing_placeholders:
            logger.warning(f"以下占位符在项目信息中没有对应的值: {missing_placeholders}")
        
        # 创建扩展的项目信息字典，确保包含所有占位符
        extended_project_info = project_info.copy()
        
        # 添加结论信息
        if '结论' not in extended_project_info and conclusion:
            conclusion_text = conclusion.get("result6_1_2", "") + "\n\n" + conclusion.get("result6_2_1", "") + "\n\n总得分：" + str(conclusion.get("totalScore", 0)) + "分"
            extended_project_info['结论'] = conclusion_text
            logger.info(f"添加结论信息: {conclusion_text}")
        
        # 明确的占位符映射 - 对应项目数据库字段名
        placeholder_mappings = {
            '项目名称': 'name',
            '项目地点': 'location',
            '项目编号': 'code',
            '项目地址': 'location',
            '建设单位': 'construction_unit',
            '设计单位': 'design_unit',
            '总用地面积': 'total_land_area',
            '总建筑面积': 'total_building_area',
            '建筑密度': 'building_density',
            '绿地率': 'green_ratio',
            '容积率': 'plot_ratio',
            '设计日期': None  # 这个会特殊处理为当前日期
        }
        
        # STEP 1: 替换段落中的占位符 (除了特殊占位符)
        logger.info("开始替换段落中的占位符")
        stations_table_added = False
        map_image_added = False
        temp_image_path = None
        
        for i, paragraph in enumerate(doc.paragraphs):
            original_text = paragraph.text
            
            # 检查是否包含特殊占位符
            special_placeholders = ['地图截图', '公交站点列表', '结论']
            contains_special = any(f"{{{ph}}}" in original_text for ph in special_placeholders)
            
            # 检查该段落是否包含任何占位符
            if '{' in original_text and '}' in original_text:
                logger.info(f"处理段落 #{i+1}: '{original_text}'")
                
                # 替换地址占位符
                if '{地址}' in original_text:
                    address_value = extended_project_info.get('项目地址', '') or extended_project_info.get('项目地点', '')
                    new_text = original_text.replace('{地址}', address_value)
                    logger.info(f"替换地址占位符: '{original_text}' -> '{new_text}'")
                    
                    # 清空段落中所有runs，但保持段落本身
                    p = paragraph._p
                    for run in list(paragraph.runs):
                        p.remove(run._r)
                    
                    # 添加新的文本
                    paragraph.add_run(new_text)
                    continue  # 处理完特殊占位符后，跳过常规替换
                
                # 替换日期占位符 (在段落中)
                if '{设计日期}' in original_text:
                    date_value = datetime.now().strftime('%Y年%m月%d日')
                    new_text = original_text.replace('{设计日期}', date_value)
                    logger.info(f"替换日期占位符: '{original_text}' -> '{new_text}'")
                    
                    # 清空段落中所有runs，但保持段落本身
                    p = paragraph._p
                    for run in list(paragraph.runs):
                        p.remove(run._r)
                    
                    # 添加新的文本
                    paragraph.add_run(new_text)
                    continue  # 处理完特殊占位符后，跳过常规替换
                
                # 替换结论占位符
                if '{结论}' in original_text:
                    logger.info(f"找到结论占位符")
                    conclusion_value = extended_project_info.get('结论', '未提供结论')
                    
                    # 清空段落中所有runs，但保持段落本身
                    p = paragraph._p
                    for run in list(paragraph.runs):
                        p.remove(run._r)
                    
                    # 分行添加结论文本
                    lines = conclusion_value.split('\n')
                    for line_idx, line in enumerate(lines):
                        if line_idx > 0:  # 非首行添加换行
                            paragraph.add_run().add_break()
                        if line.strip():  # 只添加非空行
                            paragraph.add_run(line)
                    
                    logger.info(f"替换结论占位符完成")
                    continue  # 处理完特殊占位符后，跳过常规替换
                
                # 替换地图截图占位符
                if '{地图截图}' in original_text:
                    logger.info(f"找到地图截图占位符")
                    
                    # 设置段落居中对齐
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # 清除段落中的所有runs，但保持段落本身
                    p = paragraph._p
                    for run in list(paragraph.runs):
                        p.remove(run._r)
                    
                    # 添加地图图片
                    if map_image_path and os.path.exists(map_image_path):
                        try:
                            run = paragraph.add_run()
                            run.add_picture(map_image_path, width=Inches(6))
                            map_image_added = True
                            logger.info(f"成功添加地图图片")
                            
                            # 添加黑色编号说明 - 在图片下方
                            explanation_paragraph = doc.add_paragraph()
                            explanation_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            explanation_run = explanation_paragraph.add_run("（图中编号对应下表站点序号）")
                            explanation_run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文字为黑色
                            explanation_run.font.size = Pt(10)  # 设置字体大小
                        except Exception as e:
                            logger.error(f"添加地图图片时出错: {str(e)}")
                            # 如果添加图片失败，添加文本提示
                            paragraph.add_run(f"项目周边公交站位置情况（图片加载失败）")
                    else:
                        # 如果没有图片，只添加标题
                        run = paragraph.add_run(f"项目周边公交站位置情况")
                        logger.info(f"未提供地图图片，添加标题文本")
                    
                    continue  # 处理完特殊占位符后，跳过常规替换
                
                # 处理公交站点列表占位符
                if '{公交站点列表}' in original_text:
                    logger.info(f"找到公交站点列表占位符")
                    
                    # 创建表格
                    table = doc.add_table(rows=1, cols=5)
                    table.style = 'Table Grid'
                    logger.info(f"创建表格: 行数={len(table.rows)}, 列数={len(table.rows[0].cells) if table.rows else 0}")
                    
                    # 设置表头
                    header_cells = table.rows[0].cells
                    header_cells[0].text = "序号"
                    header_cells[1].text = "站点名称"
                    header_cells[2].text = "类型"
                    header_cells[3].text = "距离（米）"
                    header_cells[4].text = "详细信息"
                    logger.info(f"设置表头: {[cell.text for cell in header_cells]}")
                    
                    # 添加数据行
                    for idx, station in enumerate(stations):
                        row = table.add_row()
                        cells = row.cells
                        
                        logger.info(f"\n处理表格行 {idx+1}:")
                        logger.info(f"站点数据: {json.dumps(station, ensure_ascii=False)}")
                        
                        # 逐个单元格填充并记录
                        logger.info("开始填充单元格:")
                        
                        # 序号列
                        cells[0].text = str(idx + 1)
                        logger.info(f"单元格[0](序号): 已填充值 '{cells[0].text}'")
                        
                        # 站点名称列
                        name_value = station.get('name', '')
                        cells[1].text = name_value
                        logger.info(f"单元格[1](站点名称): 从字段'name'获取值 '{name_value}'")
                        
                        # 类型列
                        type_value = station.get('type', '')
                        logger.info(f"字段'type'的值: '{type_value}'")
                        cells[2].text = type_value
                        logger.info(f"单元格[2](类型): 已填充值 '{cells[2].text}'")
                        
                        # 距离列
                        distance_value = station.get('distance', '')
                        logger.info(f"字段'distance'的值: '{distance_value}'")
                        cells[3].text = str(distance_value)
                        logger.info(f"单元格[3](距离): 已填充值 '{cells[3].text}'")
                        
                        # 详细信息列
                        detail_value = station.get('detail', station.get('address', '无详细信息'))
                        logger.info(f"详细信息使用值: '{detail_value}'")
                        cells[4].text = detail_value
                        logger.info(f"单元格[4](详细信息): 已填充值 '{cells[4].text}'")
                        
                        # 验证填充结果
                        logger.info(f"行 {idx+1} 填充结果:")
                        for i, cell in enumerate(cells):
                            logger.info(f"  单元格[{i}]: '{cell.text}'")
                    
                    # 直接用表格替换段落，而不是添加到段落后面
                    try:
                        # 获取段落的父元素
                        parent = paragraph._p.getparent()
                        # 获取段落在父元素中的索引
                        index = parent.index(paragraph._p)
                        # 在段落的位置插入表格
                        parent.insert(index, table._tbl)
                        # 移除原始段落
                        parent.remove(paragraph._p)
                        
                        stations_table_added = True
                        logger.info(f"成功用表格替换了占位符段落")
                    except Exception as e:
                        logger.error(f"替换段落为表格时出错: {str(e)}")
                        raise
                    
                    continue  # 处理完特殊占位符后，跳过常规替换
                
                # 对于普通占位符，检查每个可能的占位符并替换
                new_text = original_text
                text_changed = False
                
                for placeholder, value in extended_project_info.items():
                    if placeholder in ['地图截图', '公交站点列表', '结论']:
                        continue  # 跳过特殊占位符
                    
                    placeholder_text = '{' + placeholder + '}'
                    if placeholder_text in new_text:
                        if value is not None:
                            new_text = new_text.replace(placeholder_text, str(value))
                            text_changed = True
                            logger.info(f"替换了占位符 {placeholder_text} -> {value}")
                
                if text_changed:
                    logger.info(f"段落 #{i+1} 文本已更改: '{original_text}' -> '{new_text}'")
                    # 清空段落中所有runs，但保持段落本身
                    p = paragraph._p
                    for run in list(paragraph.runs):
                        p.remove(run._r)
                    
                    # 添加新的文本
                    paragraph.add_run(new_text)
        
        # STEP 2: 替换表格中的占位符
        logger.info("\n开始替换表格中的占位符")
        
        for t_idx, table in enumerate(doc.tables):
            logger.info(f"处理表格 #{t_idx+1}")
            
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    original_text = cell.text
                    
                    if '{' in original_text and '}' in original_text:
                        logger.info(f"表格 #{t_idx+1}, 行 #{r_idx+1}, 列 #{c_idx+1} 包含占位符: '{original_text}'")
                        
                        # 特殊处理日期占位符
                        if '{设计日期}' in original_text:
                            date_value = datetime.now().strftime('%Y年%m月%d日')
                            new_text = original_text.replace('{设计日期}', date_value)
                            
                            # 替换单元格文本
                            cell.text = new_text
                            logger.info(f"替换表格中的日期占位符: '{original_text}' -> '{new_text}'")
                            continue
                        
                        # 替换其他占位符
                        new_text = original_text
                        text_changed = False
                        
                        for placeholder, value in extended_project_info.items():
                            if placeholder in ['地图截图', '公交站点列表', '结论']:
                                continue  # 跳过特殊占位符
                            
                            placeholder_text = '{' + placeholder + '}'
                            if placeholder_text in new_text:
                                if value is not None:
                                    new_text = new_text.replace(placeholder_text, str(value))
                                    text_changed = True
                                    logger.info(f"替换了表格中的占位符 {placeholder_text} -> {value}")
                        
                        if text_changed:
                            cell.text = new_text
                            logger.info(f"表格单元格文本已更改: '{original_text}' -> '{new_text}'")
        
        # 检查是否成功创建公交站点表格
        if not stations_table_added:
            logger.warning("未成功创建公交站点表格!")
        
        # 保存文档
        doc.save(output_path)
        logger.info(f"已保存修复后的文档: {output_path}")
        
        return True
        
    except Exception as e:
        logger.error(f"修复模板时出错: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return False

if __name__ == "__main__":
    print("公共交通站点分析报告模板修复工具")
    print("----------------------------")
    
    # 默认值
    template_path = "static/templates/公共交通站点分析报告.docx"
    output_path = "fixed_公共交通站点分析报告.docx"
    
    # 提示修复开始
    print(f"开始修复: {template_path}")
    print(f"输出路径: {output_path}")
    print("修复日志将记录到 transport_report_fix.log 文件")
    
    # 执行修复
    success = fix_transport_report_template(
        template_path=template_path,
        output_path=output_path,
        test_mode=True
    )
    
    if success:
        print(f"\n✅ 修复成功! 已生成文件: {output_path}")
        print("请检查生成的文件，确认所有占位符已正确替换，表格结构正确。")
    else:
        print(f"\n❌ 修复失败! 请查看日志文件了解详细信息。") 