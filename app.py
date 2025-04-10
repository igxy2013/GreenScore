# 在文件顶部添加以下代码
import platform
import os

# 检查是否在WSL环境中
IS_WSL = 'WSL' in platform.uname().release or \
         (os.path.exists('/proc/version') and 'microsoft' in open('/proc/version').read().lower())

# 根据环境导入相应的模块
import pymysql
pymysql.install_as_MySQLdb()

if not IS_WSL:
    try:
        import pyodbc
    except ImportError:
        print("警告: pyodbc模块未安装，某些功能可能不可用")
        print("请运行 'pip install pyodbc' 安装所需模块")

from flask import Flask, render_template, request, redirect, url_for, jsonify, send_file, send_from_directory, session, flash, Response, Blueprint
from flask_caching import Cache
from flask_migrate import Migrate
from functools import wraps
import os
from dotenv import load_dotenv
import urllib.parse
import traceback
import time
import sqlite3
from sqlalchemy import text
import json
from datetime import datetime, timedelta
import logging
from logging.handlers import RotatingFileHandler
from flask_cors import CORS
from word_template import process_template
from export import generate_word, generate_dwg, generate_self_assessment_report
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from models import db, User, InvitationCode, LogRecord, Project  # 导入Project模型
import random
import string
import werkzeug.exceptions
from flask_sqlalchemy import SQLAlchemy
from routes import auth_bp
from admin import admin_app  # 使用admin.py中的admin_app
from utils.extract_word_info import extract_project_info  # 导入Word信息提取函数
import uuid
import base64

# 定义等级到属性的映射
LEVEL_TO_ATTRIBUTE = {
    '基本级': '控制项',
    '提高级': '评分项'
}

# 加载环境变量
load_dotenv()

# 创建应用
app = Flask(__name__, static_folder='static', static_url_path='/static')

# 设置配置
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['EXPORT_FOLDER'] = 'static/exports'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 最大上传文件限制为16MB

# 配置日志
logging.basicConfig(level=logging.WARNING)  # 将日志级别从INFO改为WARNING
logger = logging.getLogger('greenscore')
handler = RotatingFileHandler('logs/app.log', maxBytes=10000000, backupCount=5)
handler.setFormatter(logging.Formatter(
    '%(asctime)s %(levelname)s: %(message)s [in %(pathname)s:%(lineno)d]'
))
handler.setLevel(logging.WARNING)  # 将处理器的级别从INFO改为WARNING
logger.addHandler(handler)
app.logger = logger

# 禁用其他库的冗余日志
logging.getLogger('sqlalchemy.engine').setLevel(logging.WARNING)
logging.getLogger('sqlalchemy.pool').setLevel(logging.WARNING)
logging.getLogger('sqlalchemy.dialects').setLevel(logging.WARNING)
logging.getLogger('sqlalchemy.orm').setLevel(logging.WARNING)
logging.getLogger('werkzeug').setLevel(logging.WARNING)
logging.getLogger('flask_cors').setLevel(logging.WARNING)

# 配置 session
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'dev_key_123')  # 添加一个默认的密钥
app.config['SESSION_TYPE'] = 'filesystem'
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(days=7)

# 配置CORS
CORS(app, resources={r"/api/*": {"origins": "*"}})

# 判断是否为生产环境
is_production = os.environ.get('FLASK_ENV') == 'production'

# 配置缓存
cache_config = {
    "DEBUG": not is_production,
    "CACHE_TYPE": "FileSystemCache" if is_production else "SimpleCache",
    "CACHE_DIR": "cache" if is_production else None,
    "CACHE_DEFAULT_TIMEOUT": 3600  # 缓存过期时间，单位秒（1小时）
}
cache = Cache(app, config=cache_config)

# 配置数据库连接
db_uri = os.environ.get('DATABASE_URL')
if not db_uri:
    # MySQL数据库连接配置
    mysql_host = os.environ.get('MYSQL_HOST', 'localhost')
    mysql_port = os.environ.get('MYSQL_PORT', '3306')
    mysql_database = os.environ.get('MYSQL_DATABASE', '绿色建筑')
    mysql_username = os.environ.get('MYSQL_USERNAME', 'root')
    mysql_password = os.environ.get('MYSQL_PASSWORD', 'password')
    
    # 构建MySQL连接字符串
    db_uri = f'mysql+pymysql://{mysql_username}:{mysql_password}@{mysql_host}:{mysql_port}/{mysql_database}?charset=utf8mb4'
    app.logger.warning("警告: DATABASE_URL 环境变量未设置，使用默认连接字符串")

app.config['SQLALCHEMY_DATABASE_URI'] = db_uri
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['SQLALCHEMY_ECHO'] = False  # 禁用SQLAlchemy的ECHO功能，不管是开发环境还是生产环境

# 初始化数据库
db.init_app(app)

# 初始化数据库迁移
migrate = Migrate(app, db)

# 初始化login manager
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'admin.admin_login'  # 修改这里，指向管理员登录页面
login_manager.login_message = '请先登录'
login_manager.login_message_category = 'info'

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

# 添加根路由重定向到index.html
@app.route('/')
def home():
    return render_template('index.html')

@app.route('/privacy_policy')
def privacy_policy():
    return render_template('privacy_policy.html')

@app.route('/terms_of_service')
def terms_of_service():
    return render_template('terms_of_service.html')

# 添加请求处理器来更新用户的last_seen时间
@app.before_request
def update_last_seen():
    if current_user.is_authenticated:
        current_user.last_seen = datetime.utcnow()
        db.session.commit()

# 添加请求日志中间件
@app.before_request
def log_request_info():
    # 只记录GET和POST请求
    if request.method in ['GET', 'POST', 'PUT', 'DELETE']:
        try:
            # 获取用户ID（如果已登录）
            user_id = current_user.id if current_user.is_authenticated else None
            
            # 记录请求信息，但排除静态文件和某些特定路径
            path = request.path
            if not path.startswith('/static/') and not path.startswith('/favicon.ico'):
                LogRecord.add_log(
                    level="INFO",
                    message=f"请求: {request.method} {path}",
                    source="HTTP_REQUEST",
                    user_id=user_id,
                    ip_address=request.remote_addr,
                    path=path,
                    method=request.method,
                    user_agent=request.user_agent.string
                )
        except Exception as e:
            app.logger.error(f"记录请求日志时出错: {str(e)}")

# 添加响应日志中间件
@app.after_request
def log_response_info(response):
    try:
        # 只记录错误响应
        if response.status_code >= 400:
            path = request.path
            if not path.startswith('/static/') and not path.startswith('/favicon.ico'):
                user_id = current_user.id if current_user.is_authenticated else None
                
                LogRecord.add_log(
                    level="ERROR" if response.status_code >= 500 else "WARNING",
                    message=f"响应错误: {request.method} {path} 状态码: {response.status_code}",
                    source="HTTP_RESPONSE",
                    user_id=user_id,
                    ip_address=request.remote_addr,
                    path=path,
                    method=request.method,
                    user_agent=request.user_agent.string
                )
    except Exception as e:
        app.logger.error(f"记录响应日志时出错: {str(e)}")
    
    return response

# 添加异常日志中间件
@app.errorhandler(Exception)
def log_exception(e):
    try:
        # 获取请求相关信息
        path = request.path if request else "未知路径"
        method = request.method if request else "未知方法"
        user_id = current_user.id if current_user and current_user.is_authenticated else None
        
        # 对静态文件的404错误进行特殊处理，减少日志记录
        if isinstance(e, werkzeug.exceptions.NotFound) and (path.startswith('/static/') or path == '/favicon.ico'):
            # 静态文件404不记录详细日志，只在DEBUG模式下记录
            if app.debug:
                app.logger.debug(f"静态文件未找到: {path}")
            return render_template('error.html', error="文件未找到"), 404
        
        # 记录异常信息
        LogRecord.add_log(
            level="ERROR",
            message=f"系统异常: {str(e)}",
            source="EXCEPTION",
            user_id=user_id,
            ip_address=request.remote_addr if request else None,
            path=path,
            method=method,
            user_agent=request.user_agent.string if request else None
        )
        
        # 记录详细的异常堆栈到应用日志
        app.logger.error(f"系统异常: {str(e)}", exc_info=True)
        
    except Exception as log_err:
        app.logger.error(f"记录异常日志时出错: {str(log_err)}")
    
    # 返回通用的错误页面
    return render_template('error.html', error=str(e)), 500

# 添加专门的404错误处理器
@app.errorhandler(404)
def page_not_found(e):
    # 获取请求路径
    path = request.path
    
    # 对静态文件的请求进行特殊处理
    if path.startswith('/static/') or path == '/favicon.ico':
        # 静态文件不存在的情况，返回简单的404响应而不记录详细日志
        return "File not found", 404
    
    # 非静态文件的404，记录日志
    app.logger.warning(f"页面未找到: {request.method} {path}")
    
    # 返回自定义404页面
    return render_template('error.html', error="页面未找到"), 404

# 添加405错误处理器
@app.errorhandler(405)
def method_not_allowed(e):
    # 检查是否为非应用路径的请求
    non_app_paths = ['/dns-query', '/robots.txt', '/favicon.ico']
    if request.path in non_app_paths:
        app.logger.debug(f"忽略对非应用路径的请求: {request.method} {request.path}")
        return '', 204  # 返回无内容状态码
    
    # 记录其他405错误
    app.logger.warning(f"方法不允许: {request.method} {request.path}")
    return render_template('error.html', 
                          error_code=405, 
                          error_message="请求方法不被允许"), 405

# 邀请码相关API路由
@app.route('/api/invite-codes', methods=['POST'])
@login_required
def generate_invite_code():
    if not current_user.is_admin():
        return jsonify({'success': False, 'message': '只有管理员可以生成邀请码'}), 403
    
    try:
        # 生成随机邀请码
        code = ''.join(random.choices(string.ascii_uppercase + string.digits, k=8))
        
        # 创建新的邀请码记录
        new_code = InvitationCode(code=code)
        db.session.add(new_code)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': '邀请码生成成功',
            'code': code
            
        })
    except Exception as e:
        db.session.rollback()
        app.logger.error(f'生成邀请码失败: {str(e)}')
        return jsonify({'success': False, 'message': '生成邀请码失败'}), 500

@app.route('/api/invite-codes/<int:code_id>', methods=['DELETE'])
@login_required
def delete_invite_code(code_id):
    if not current_user.is_admin():
        return jsonify({'success': False, 'message': '只有管理员可以删除邀请码'}), 403
    
    try:
        code = InvitationCode.query.get_or_404(code_id)
        db.session.delete(code)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': '邀请码删除成功'
        })
    except Exception as e:
        db.session.rollback()
        app.logger.error(f'删除邀请码失败: {str(e)}')
        return jsonify({'success': False, 'message': '删除邀请码失败'}), 500

# 添加登录要求装饰器
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            return redirect(url_for('login_page'))
        return f(*args, **kwargs)
    return decorated_function

def sync_score_tables(project_id):
    """同步得分表和project_scores表的数据"""
    if not project_id:
        logger.warning("项目ID不能为空，无法同步数据")
        return False
    
    try:
        # 记录开始时间
        start_time = datetime.now()
        logger.info(f"开始同步项目ID={project_id}的得分表和project_scores表数据")
        
        # 获取项目信息
        project = get_project(project_id)
        if not project:
            logger.error(f"项目不存在: ID={project_id}")
            return False
            
        project_standard = project.standard
        
        # 从得分表导入数据到project_scores表
        count = 0
        try:
            result = db.session.execute(
                text("SELECT COUNT(*) FROM `得分表` WHERE `项目ID` = :project_id"),
                {"project_id": project_id}
            )
            count = result.scalar()
            logger.info(f"得分表中项目ID={project_id}的记录数: {count}")
        except Exception as e:
            logger.error(f"查询得分表记录数时出错: {str(e)}")
            
        # 如果不存在记录，跳过同步
        if count == 0:
            logger.warning(f"得分表中不存在项目ID={project_id}的记录，跳过同步")
            return False
            
        # 从得分表中获取数据，考虑缺失的列
        try:
            # 使用安全的查询方式
            result = db.session.execute(
                text("""
                    SELECT 
                        `项目ID`, `项目名称`, `专业`, `评价等级`, `条文号`, 
                        `分类`, `是否达标`, `得分`, `技术措施`, 
                        COALESCE(`评价标准`, :default_standard) as `评价标准`
                    FROM `得分表` 
                    WHERE `项目ID` = :project_id
                """),
                {"project_id": project_id, "default_standard": project_standard}
            )
            
            scores = result.fetchall()
            logger.info(f"从得分表中获取到 {len(scores)} 条记录")
            
            if not scores:
                logger.warning(f"得分表中不存在项目ID={project_id}的记录，跳过同步")
                return False
                
            # 先删除project_scores表中的相关记录
            try:
                delete_result = db.session.execute(
                    text("DELETE FROM project_scores WHERE project_id = :project_id"),
                    {"project_id": project_id}
                )
                logger.info(f"已从project_scores表中删除 {delete_result.rowcount} 条记录")
                db.session.flush()
            except Exception as e:
                logger.error(f"删除project_scores表记录时出错: {str(e)}")
                db.session.rollback()
                raise
                
            # 批量插入数据到project_scores表
            inserted_count = 0
            for score in scores:
                try:
                    # 解包得分数据，避免数组越界错误
                    项目ID = score[0] if len(score) > 0 else project_id
                    项目名称 = score[1] if len(score) > 1 else ""
                    专业 = score[2] if len(score) > 2 else ""
                    评价等级 = score[3] if len(score) > 3 else "基本级"
                    条文号 = score[4] if len(score) > 4 else ""
                    分类 = score[5] if len(score) > 5 else ""
                    是否达标 = score[6] if len(score) > 6 else "是"
                    得分 = score[7] if len(score) > 7 else 0
                    技术措施 = score[8] if len(score) > 8 else ""
                    评价标准 = score[9] if len(score) > 9 else project_standard
                    
                    # 转换得分为浮点数
                    try:
                        if 得分 and str(得分).strip():
                            score_float = float(得分)
                        else:
                            score_float = 0
                    except (ValueError, TypeError):
                        logger.warning(f"无法将得分转换为浮点数: {得分}，使用默认值0")
                        score_float = 0
                    
                    # 插入数据到project_scores表
                    db.session.execute(
                        text("""
                            INSERT INTO project_scores (
                                project_id, project_name, specialty, level, clause_number, 
                                category, is_achieved, score, technical_measures, standard
                            ) VALUES (
                                :project_id, :project_name, :specialty, :level, :clause_number,
                                :category, :is_achieved, :score, :technical_measures, :standard
                            )
                        """),
                        {
                            "project_id": 项目ID,
                            "project_name": 项目名称,
                            "specialty": 专业,
                            "level": 评价等级,
                            "clause_number": 条文号,
                            "category": 分类,
                            "is_achieved": 是否达标,
                            "score": score_float,
                            "technical_measures": 技术措施,
                            "standard": 评价标准
                        }
                    )
                    inserted_count += 1
                except Exception as e:
                    logger.error(f"插入project_scores表时出错: {str(e)}, 条文号: {条文号}")
                    # 继续处理下一条记录
                    
            # 提交事务
            db.session.commit()
            end_time = datetime.now()
            logger.info(f"同步完成，共插入 {inserted_count} 条记录，耗时: {(end_time - start_time).total_seconds()} 秒")
            return True
            
        except Exception as e:
            db.session.rollback()
            logger.error(f"同步数据时出错: {str(e)}")
            logger.error(traceback.format_exc())
            return False
            
    except Exception as e:
        logger.error(f"同步得分表和project_scores表数据时出错: {str(e)}")
        logger.error(traceback.format_exc())
        return False

# 添加评价标准的模型
class review_standard(db.Model):
    __tablename__ = '评价标准'
    
    # 使用中文字段名
    序号 = db.Column(db.Integer, primary_key=True)
    条文号 = db.Column(db.String(20))
    分类 = db.Column(db.String(50))
    专业 = db.Column(db.String(50))
    条文内容 = db.Column(db.Text)
    分值 = db.Column(db.String(10))
    审查材料 = db.Column(db.Text)
    属性 = db.Column(db.String(20))  # 属性字段，包含控制项、评分项
    标准名称 = db.Column(db.String(20))  # 标准名称字段

class FormData(db.Model):
    __tablename__ = 'form_data'
    
    id = db.Column(db.Integer, primary_key=True)
    project_id = db.Column(db.Integer, db.ForeignKey('projects.id'), nullable=True)
    project_name = db.Column(db.String(100))
    building_no = db.Column(db.String(50))
    project_location = db.Column(db.String(200))
    design_no = db.Column(db.String(50))
    construction_unit = db.Column(db.String(100))
    design_unit = db.Column(db.String(100))
    standard_selection = db.Column(db.String(20))
    form_data = db.Column(db.Text)  # 存储JSON格式的表单数据
    created_at = db.Column(db.DateTime, default=db.func.current_timestamp())
    updated_at = db.Column(db.DateTime, default=db.func.current_timestamp(), onupdate=db.func.current_timestamp())

# 缓存所有数据的函数
@cache.cached(key_prefix='all_standards')
def get_all_standards():
    print("从数据库获取所有标准数据...")
    start_time = time.time()
    standards = review_standard.query.all()
    end_time = time.time()
    print(f"查询耗时: {end_time - start_time:.2f}秒，获取到 {len(standards)} 条记录")
    return standards

# 按专业和级别缓存数据的函数
@cache.memoize()
def get_filtered_standards(level, specialty):
    print(f"从数据库获取筛选数据: 级别={level}, 专业={specialty}")
    start_time = time.time()
    # 将级别转换为对应的属性值
    attribute = LEVEL_TO_ATTRIBUTE.get(level, '')
    if attribute:
        print(f"级别 '{level}' 对应属性值 '{attribute}'")
    
    # 获取当前项目
    project = get_project()
    standard_name = project.standard if project and project.standard else '成都市标'
    print(f"当前评价标准: {standard_name}")
    
    # 使用新的评价标准模型
    model_class = review_standard
    
    # 尝试多种筛选方法
    standards = []
    
    # 方法1: 使用属性字段精确匹配并按标准名称筛选
    if attribute and specialty:
        try:
            query1 = model_class.query.filter(getattr(model_class, '属性') == attribute)
            query1 = query1.filter(getattr(model_class, '专业').like(f'%{specialty}%'))
            query1 = query1.filter(getattr(model_class, '标准名称') == standard_name)
            # 添加按序号排序
            query1 = query1.order_by(getattr(model_class, '序号'))
            standards1 = query1.all()
            print(f"方法1 (属性精确匹配): 找到 {len(standards1)} 条记录")
            if standards1:
                standards = standards1
        except Exception as e:
            print(f"方法1查询错误: {str(e)}")
    
    # 方法2: 使用属性字段模糊匹配并按标准名称筛选
    if not standards and attribute:
        try:
            query2 = model_class.query
            query2 = query2.filter(getattr(model_class, '属性').like(f'%{attribute}%'))
            if specialty:
                query2 = query2.filter(getattr(model_class, '专业').like(f'%{specialty}%'))
            query2 = query2.filter(getattr(model_class, '标准名称') == standard_name)
            # 添加按序号排序
            query2 = query2.order_by(getattr(model_class, '序号'))
            standards2 = query2.all()
            print(f"方法2 (属性模糊匹配): 找到 {len(standards2)} 条记录")
            if standards2:
                standards = standards2
        except Exception as e:
            print(f"方法2查询错误: {str(e)}")
    
    # 方法3: 只按专业和标准名称筛选
    if not standards and specialty:
        try:
            query3 = model_class.query.filter(getattr(model_class, '专业').like(f'%{specialty}%'))
            query3 = query3.filter(getattr(model_class, '标准名称') == standard_name)
            # 添加按序号排序
            query3 = query3.order_by(getattr(model_class, '序号'))
            standards3 = query3.all()
            print(f"方法3 (仅按专业): 找到 {len(standards3)} 条记录")
            if standards3:
                standards = standards3
        except Exception as e:
            print(f"方法3查询错误: {str(e)}")
    
    # 方法4: 使用SQL文本查询
    if not standards:
        try:
            sql = f"SELECT * FROM `{model_class.__tablename__}`"
            conditions = []
            if attribute:
                conditions.append(f"属性 = '{attribute}'")
            if specialty:
                conditions.append(f"专业 LIKE '%{specialty}%'")
            conditions.append(f"标准名称 = '{standard_name}'")
            
            if conditions:
                sql += " WHERE " + " AND ".join(conditions)
            
            # 添加按序号排序
            sql += " ORDER BY 序号"
            
            print(f"执行SQL: {sql}")
            result = db.session.execute(text(sql))
            standards = [dict(zip(result.keys(), row)) for row in result]
            print(f"方法4 (SQL文本查询): 找到 {len(standards)} 条记录")
        except Exception as e:
            print(f"方法4查询错误: {str(e)}")
    
    # 如果所有方法都失败，返回当前标准下的所有记录（同样按序号排序）
    if not standards:
        try:
            standards = model_class.query.filter(getattr(model_class, '标准名称') == standard_name).order_by(getattr(model_class, '序号')).all()
            print(f"所有筛选方法都失败，返回所有 {len(standards)} 条记录")
        except Exception as e:
            print(f"获取所有记录错误: {str(e)}")
            standards = []
    
    end_time = time.time()
    print(f"筛选查询耗时: {end_time - start_time:.2f}秒")
    
    return standards

# 根据标准名称获取对应的标准数据
def get_standards_by_name(standard_name):
    print(f"获取标准数据: {standard_name}")
    return review_standard.query.filter(review_standard.标准名称 == standard_name).all()

# 获取项目信息
def get_project(project_id=None):
    try:
        if project_id:
            # 获取指定ID的项目
            project = Project.query.get(project_id)
        else:
            # 获取第一个项目
            project = Project.query.first()
        return project
    except Exception as e:
        print(f"获取项目信息时出错: {str(e)}")
        return None

# 保存项目信息
def save_project_info(form_data):
    try:
        # 获取项目ID
        project_id = form_data.get('project_id', '')
        if project_id:
            try:
                project_id = int(project_id)
            except ValueError:
                project_id = None
        
        # 如果有项目ID，获取现有项目；否则创建新项目
        if project_id:
            project = Project.query.get(project_id)
            if not project:
                print(f"未找到ID为{project_id}的项目，创建新项目")
                project = Project()
        else:
            print("未提供项目ID，创建新项目")
            project = Project()
        
        # 检查是否是详细信息表单
        is_detail_form = form_data.get('detail_form') == '1'
        print(f"表单类型: {'详细信息表单' if is_detail_form else '基本信息表单'}")
        
        # 获取新的项目名称
        new_project_name = form_data.get('project_name', '')
        if not new_project_name:
            print("错误：未提供项目名称")
            return None
        
        # 获取当前登录用户的ID
        user_id = session.get('user_id')
        if not user_id:
            print("错误：用户未登录")
            return None
        
        # 检查项目标准
        standard = form_data.get('standard_selection', '')
        if not standard:
            print("错误：未提供评价标准")
            return None
            
        # 检查建筑类型
        building_type = form_data.get('building_type', '')
        
        # 检查星级目标
        star_rating_target = form_data.get('star_rating_target', '')
        
        # 基本信息表单数据
        project.name = new_project_name
        project.user_id = user_id
        project.standard = standard
        project.building_type = building_type
        project.star_rating_target = star_rating_target
        project.code = form_data.get('project_code', '')
        project.construction_unit = form_data.get('construction_unit', '')
        project.design_unit = form_data.get('design_unit', '')
        project.location = form_data.get('project_location', '')
        project.climate_zone = form_data.get('climate_zone', '')
        # 设置项目状态为"进行中"
        project.status = '进行中'
        
        # 始终处理详细字段，不再依赖is_detail_form标志
        # 尝试将表单值转换为浮点数
        try_parse_float(form_data, 'total_land_area', project, 'total_land_area')
        try_parse_float(form_data, 'total_building_area', project, 'total_building_area')
        try_parse_float(form_data, 'above_ground_area', project, 'above_ground_area')
        try_parse_float(form_data, 'underground_area', project, 'underground_area')
        try_parse_float(form_data, 'first_floor_underground_area', project, 'underground_floor_area')
        try_parse_float(form_data, 'plot_ratio', project, 'plot_ratio')
        try_parse_float(form_data, 'building_base_area', project, 'building_base_area')
        try_parse_float(form_data, 'building_density', project, 'building_density')
        
        # 特别记录绿地面积的处理
        green_area_value = form_data.get('green_area', '')
        print(f"处理绿地面积字段: 原始值={green_area_value}")
        try_parse_float(form_data, 'green_area', project, 'green_area')
        print(f"绿地面积处理后的值: {project.green_area}")
        
        try_parse_float(form_data, 'green_ratio', project, 'green_ratio')
        try_parse_float(form_data, 'building_height', project, 'building_height')
        
        # 尝试将表单值转换为整数
        try_parse_int(form_data, 'ground_parking_spaces', project, 'ground_parking_spaces')
        try_parse_int(form_data, 'residential_units', project, 'residential_units')
        
        # 设置其他字段
        project.building_floors = form_data.get('building_floors', '')
        project.air_conditioning_type = form_data.get('ac_type', '')
        project.average_floors = form_data.get('avg_floors', '')
        project.has_garbage_room = form_data.get('has_garbage_room', '')
        project.has_elevator = form_data.get('has_elevator', '')
        project.has_underground_garage = form_data.get('has_underground_garage', '')
        project.construction_type = form_data.get('construction_type', '')
        project.has_water_landscape = form_data.get('has_water_landscape', '')
        project.is_fully_decorated = form_data.get('is_fully_decorated', '')
        project.public_building_type = form_data.get('public_building_type', '')
        project.public_green_space = form_data.get('public_green_space', '')
        
        # 记录将要保存的数据内容
        print(f"正在保存项目信息，详细字段包括：总用地面积={project.total_land_area}, 总建筑面积={project.total_building_area}, 地上建筑面积={project.above_ground_area}, 地下建筑面积={project.underground_area}")
        
        # 保存项目信息到数据库
        db.session.add(project)
        db.session.commit()
        db.session.refresh(project)  # 刷新对象以获取最新值
        
        print(f"项目信息保存成功: ID={project.id}, 名称={project.name}")
        # 打印详细字段是否成功保存
        print(f"详细字段保存结果：总用地面积={project.total_land_area}, 总建筑面积={project.total_building_area}, 地上建筑面积={project.above_ground_area}, 地下建筑面积={project.underground_area}")
        
        return project
    except Exception as e:
        db.session.rollback()
        print(f"保存项目信息时出错: {str(e)}")
        print(traceback.format_exc())
        return None

# 辅助函数：尝试解析浮点数
def try_parse_float(form_data, form_field, model_obj, model_field):
    """尝试将表单字段解析为浮点数并赋值给模型对象"""
    try:
        value = form_data.get(form_field, '')
        if value is not None and value != '':
            # 清理数值字符串，移除可能的非数字字符
            clean_value = str(value).replace(',', '').strip()
            # 尝试转换为浮点数
            float_value = float(clean_value)
            # 确保值为0也被正确保存，而不是被当作空值处理
            model_obj.__setattr__(model_field, float_value)
            
            if form_field == 'green_area':
                print(f"成功将绿地面积 '{value}' 转换为浮点数: {float_value}")
        else:
            if form_field == 'green_area':
                print(f"绿地面积字段为空或无效: '{value}'")
            model_obj.__setattr__(model_field, None)
    except (ValueError, TypeError) as e:
        print(f"无法将字段 {form_field} 值 '{value}' 转换为浮点数: {str(e)}")
        # 如果转换失败，设置为None
        model_obj.__setattr__(model_field, None)

# 辅助函数：尝试解析整数
def try_parse_int(form_data, form_field, model_obj, model_field):
    """尝试将表单字段解析为整数并赋值给模型对象"""
    try:
        value = form_data.get(form_field, '')
        if value is not None and value != '':
            # 清理数值字符串，移除可能的非数字字符
            clean_value = str(value).replace(',', '').strip()
            # 尝试转换为整数
            int_value = int(float(clean_value))  # 先转为float再转为int，处理可能的小数点
            # 确保值为0也被正确保存，而不是被当作空值处理
            model_obj.__setattr__(model_field, int_value)
            print(f"成功将字段 {form_field} 值 '{value}' 转换为整数: {int_value}")
        else:
            model_obj.__setattr__(model_field, None)
            print(f"字段 {form_field} 为空，设置为None")
    except (ValueError, TypeError) as e:
        print(f"无法将字段 {form_field} 值 '{value}' 转换为整数: {str(e)}")
        # 如果转换失败，设置为None
        model_obj.__setattr__(model_field, None)

# 项目管理页面路由
@app.route('/projects')
@app.route('/project_management')
@login_required
def project_management():
    try:
        username = session.get('username', '未登录')
        return render_template('project_management.html', username=username)
    except Exception as e:
        print(f"项目管理页面出错: {str(e)}")
        print(traceback.format_exc())
        return render_template('error.html', error=str(e))

# 创建项目路由
@app.route('/create_project', methods=['POST'])
@login_required
def create_project():
    try:
        # 获取当前用户ID
        user_id = session.get('user_id')
        if not user_id:
            return jsonify({'error': '用户未登录'}), 401
        
        data = request.get_json()
        project_name = data.get('name')
        standard = data.get('standard')
        star_rating_target = data.get('star_rating_target')
        building_type = data.get('building_type')
        
        if not project_name:
            return jsonify({'error': '项目名称不能为空'}), 400
        
        if not standard:
            return jsonify({'error': '请选择评价标准'}), 400
        
        try:
            # 创建新项目
            project = Project(
                name=project_name,
                user_id=user_id,
                standard=standard,
                building_type=building_type,
                star_rating_target=star_rating_target,
                code=data.get('code'),
                construction_unit=data.get('construction_unit'),
                design_unit=data.get('design_unit'),
                location=data.get('location'),
                climate_zone=data.get('climate_zone'),
                total_land_area=data.get('total_land_area'),
                total_building_area=data.get('total_building_area'),
                above_ground_area=data.get('above_ground_area'),
                underground_area=data.get('underground_area'),
                plot_ratio=data.get('plot_ratio'),
                building_base_area=data.get('building_base_area'),
                building_density=data.get('building_density'),
                green_area=data.get('green_area'),
                green_ratio=data.get('green_ratio'),
                building_height=data.get('building_height'),
                building_floors=data.get('building_floors'),
                air_conditioning_type=data.get('air_conditioning_type'),
                has_garbage_room=data.get('has_garbage_room'),
                has_elevator=data.get('has_elevator'),
                has_underground_garage=data.get('has_underground_garage'),
                construction_type=data.get('construction_type'),
                has_water_landscape=data.get('has_water_landscape'),
                is_fully_decorated=data.get('is_fully_decorated'),
                public_building_type=data.get('public_building_type'),
                public_green_space=data.get('public_green_space')
            )
            
            # 尝试添加项目并提交事务
            db.session.add(project)
            db.session.flush()  # 获取项目ID但不提交事务
            
            project_id = project.id
            app.logger.info(f"项目创建成功: ID={project_id}, 名称={project_name}")
            
            # 创建成功后尝试创建默认评分数据
            try:
                # 直接创建默认评分数据
                app.logger.info(f"创建项目默认评分数据: 项目ID={project_id}, 名称={project_name}, 标准={standard}")
                create_default_scores(project_id, project_name, standard)
                
                # 提交事务
                db.session.commit()
                return jsonify({
                    'id': project_id,
                    'name': project.name,
                    'message': '项目创建成功'
                })
            except Exception as e:
                db.session.rollback()
                app.logger.error(f"创建默认评分数据失败: {str(e)}")
                app.logger.error(traceback.format_exc())
                
                # 虽然评分数据创建失败，但项目创建已成功，返回项目ID
                return jsonify({
                    'id': project_id,
                    'name': project_name,
                    'error': '创建默认评分数据失败，但项目已创建'
                })
                
        except Exception as db_error:
            db.session.rollback()
            app.logger.error(f"数据库操作失败: {str(db_error)}")
            app.logger.error(traceback.format_exc())
            
            # 检查是否是ID字段无默认值错误
            if "Field 'id' doesn't have a default value" in str(db_error):
                return jsonify({
                    'error': '数据库配置错误：ID字段未配置为自动递增。请检查数据库表结构或联系管理员。'
                }), 500
            else:
                return jsonify({
                    'error': f'数据库写入失败: {str(db_error)}'
                }), 500
        
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"创建项目失败: {str(e)}")
        app.logger.error(traceback.format_exc())
        return jsonify({'error': f'创建项目失败: {str(e)}'}), 500

# 修改项目访问权限检查
def check_project_access(project_id):
    """检查当前用户是否有权限访问指定项目"""
    try:
        user_id = session.get('user_id')
        if not user_id:
            return False
        
        project = Project.query.filter_by(id=project_id, user_id=user_id).first()
        return project is not None
    except Exception as e:
        print(f"检查项目访问权限失败: {str(e)}")
        return False

# 修改项目详情页面访问
@app.route('/project/<int:project_id>')
@login_required
def project_detail(project_id):
    # 检查访问权限
    if not check_project_access(project_id):
        flash('您没有权限访问该项目', 'error')
        return redirect(url_for('project_management'))
    
    try:
        project = Project.query.get_or_404(project_id)
        # 获取page参数，默认为project_info
        page = request.args.get('page', 'project_info')
        
        # 如果页面类型为公共交通分析，使用iframe加载页面
        if page == 'public_transport_analysis':
            app.logger.info(f"访问项目 ID: {project_id}, 名称: {project.name}, 页面: {page}")
            return render_template('dashboard.html', project=project, current_page=page)
        
        app.logger.info(f"访问项目 ID: {project_id}, 名称: {project.name}, 页面: {page}")
        return render_template('dashboard.html', project=project, current_page=page)
    except Exception as e:
        app.logger.error(f"获取项目详情失败: {str(e)}")
        app.logger.error(traceback.format_exc())
        return render_template('error.html', error=f"获取项目详情失败: {str(e)}")

# 修改项目删除函数
@app.route('/delete_project/<int:project_id>', methods=['DELETE'])
@login_required
def delete_project(project_id):
    try:
        # 检查访问权限
        if not check_project_access(project_id):
            return jsonify({'error': '您没有权限删除该项目'}), 403
        
        project = Project.query.get_or_404(project_id)
        db.session.delete(project)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': '项目删除成功'
        })
    except Exception as e:
        db.session.rollback()
        print(f"删除项目失败: {str(e)}")
        return jsonify({'error': '删除项目失败'}), 500

@app.route('/')
def index():
    return render_template('login.html')

@app.route('/reset_password')
def reset_password_page():
    """处理密码重置页面路由"""
    try:
        return render_template('reset_password.html')
    except Exception as e:
        app.logger.error(f"渲染密码重置页面错误: {str(e)}")
        return "页面加载失败", 500
@app.route('/login')
def login_page():
    return render_template('login.html')

@app.route('/register')
def register_page():
    return render_template('register.html')

@app.route('/', methods=['POST'])
def handle_form():
    try:
        form_type = request.form.get('form_type', '')
        print(f"接收到表单提交: form_type={form_type}")
        
        # 检查是否为AJAX请求
        is_ajax = request.headers.get('X-Requested-With') == 'XMLHttpRequest' or request.is_json or 'application/json' in request.headers.get('Accept', '')
        print(f"是否是AJAX请求: {is_ajax}")
        
        # 打印所有表单字段
        print("表单数据:")
        for key, value in request.form.items():
            print(f"  {key}: {value}")
        
        if form_type == 'project_info':
            # 处理项目信息表单提交
            print("处理项目信息表单...")
            # 检查是否是详细信息表单
            is_detail_form = request.form.get('detail_form') == '1'
            print(f"表单类型: {'详细信息表单' if is_detail_form else '基本信息表单'}")
            
            # 检查public_green_space字段
            public_green_space = request.form.get('public_green_space', '')
            print(f"绿地向公众开放: {public_green_space}")
            
            project = save_project_info(request.form)
            if project:
                print(f"项目信息保存成功: ID={project.id}, 名称={project.name}, 标准={project.standard}, 绿地向公众开放={project.public_green_space}")
                
                if is_ajax:
                    # 如果是AJAX请求，返回JSON响应
                    return jsonify({
                        'success': True,
                        'message': '项目信息保存成功',
                        'project_id': project.id,
                        'project_name': project.name,
                        'is_detail_form': is_detail_form  # 添加这个字段以便前端知道是详细表单
                    })
                else:
                    # 如果是传统表单提交，重定向回项目详情页面
                    return redirect(url_for('project_detail', project_id=project.id))
            else:
                print("项目信息保存失败")
                
                if is_ajax:
                    # 如果是AJAX请求，返回错误JSON响应
                    return jsonify({
                        'success': False,
                        'message': '项目信息保存失败'
                    }), 400
                else:
                    # 如果是传统表单提交，重定向到项目管理页面
                    return redirect(url_for('project_management'))
        else:
            print(f"未知的表单类型: {form_type}")
            
            if is_ajax:
                return jsonify({
                    'success': False,
                    'message': f'未知的表单类型: {form_type}'
                }), 400
        
        # 如果不是已知的表单类型且不是AJAX请求，返回首页
        return redirect(url_for('index'))
    except Exception as e:
        print(f"处理表单时发生错误: {str(e)}")
        print(traceback.format_exc())
        
        # 检查是否为AJAX请求
        is_ajax = request.headers.get('X-Requested-With') == 'XMLHttpRequest' or request.is_json or 'application/json' in request.headers.get('Accept', '')
        
        if is_ajax:
            # 如果是AJAX请求，返回错误JSON响应
            return jsonify({
                'success': False,
                'message': f'处理表单时发生错误: {str(e)}'
            }), 500
        else:
            # 如果是传统表单提交，渲染错误页面
            return render_template('error.html', error=str(e))

@app.route('/filter')
def filter_standards():
    try:
        level = request.args.get('level', '基本级')
        specialty = request.args.get('specialty', '建筑')
        project_id = request.args.get('project_id')
        
        # 获取项目信息
        project = None
        if project_id:
            project = Project.query.get(project_id)
        
        # 获取项目对应的标准名称
        standard_name = None
        if project and project.standard:
            standard_name = project.standard
            print(f"使用项目设置的评价标准: {standard_name}")
        else:
            standard_name = request.args.get('standard', '成都市标')
            print(f"项目未设置评价标准，使用默认标准: {standard_name}")
        
        # 获取过滤后的标准
        standards = []
        
        # 使用新的评价标准模型
        model_class = review_standard
        
        # 获取属性值
        attribute = LEVEL_TO_ATTRIBUTE.get(level, '')
        
        # 查询数据
        try:
            query = model_class.query
            if attribute:
                query = query.filter(getattr(model_class, '属性') == attribute)
            if specialty:
                query = query.filter(getattr(model_class, '专业').like(f'%{specialty}%'))
            # 按标准名称筛选
            query = query.filter(getattr(model_class, '标准名称') == standard_name)
            
            # 添加按序号排序
            query = query.order_by(getattr(model_class, '序号').asc())
            standards = query.all()
            print(f"从 {standard_name} 获取标准数据: level={level}, specialty={specialty}, 找到{len(standards)}条记录")
        except Exception as e:
            print(f"查询标准数据出错: {str(e)}")
            print(traceback.format_exc())
        
        # 如果没有找到数据，尝试使用更宽松的条件
        if not standards:
            try:
                query = model_class.query
                if specialty:
                    query = query.filter(getattr(model_class, '专业').like(f'%{specialty}%'))
                # 按标准名称筛选
                query = query.filter(getattr(model_class, '标准名称') == standard_name)
                
                # 添加按序号排序
                query = query.order_by(getattr(model_class, '序号').asc())
                standards = query.all()
                print(f"使用宽松条件从 {standard_name} 获取标准数据: specialty={specialty}, 找到{len(standards)}条记录")
            except Exception as e:
                print(f"宽松查询标准数据出错: {str(e)}")
                print(traceback.format_exc())
        
        # 打印调试信息
        print(f"过滤标准: level={level}, specialty={specialty}, standard={standard_name}, 找到{len(standards)}条记录")
        
        # 渲染模板
        return render_template(
            'dashboard.html', 
            standards=standards, 
            current_level=level, 
            current_specialty=specialty,
            current_page='specialty',  # 设置为specialty页面，确保保存按钮正常工作
            project=project
        )
    except Exception as e:
        app.logger.error(f"过滤标准时出错: {str(e)}")
        traceback.print_exc()
        return render_template('error.html', error=str(e))

@app.route('/calculator')
@login_required
def calculator():
    return render_template('calculator.html')

@app.route('/solar_calculator')
@login_required
def solar_calculator():
    # 获取项目ID参数
    project_id = request.args.get('project_id')
    project = None
    
    # 如果提供了项目ID，获取项目信息
    if project_id:
        project = Project.query.get(project_id)
    
    return render_template('solar_calculator.html', project=project)

# 添加清除缓存的路由（可选，用于管理员手动刷新缓存）
@app.route('/clear_cache')
def clear_cache():
    try:
        cache.clear()
        return "缓存已清除"
    except Exception as e:
        return f"清除缓存时出错: {str(e)}"

@app.route('/api/save_form', methods=['POST'])
def save_form():
    try:
        import json
        data = request.json
        
        if not data:
            return jsonify({'error': '没有收到数据'}), 400
            
        # 检查当前项目
        project = Project.query.first()
        project_id = project.id if project else None
        
        # 查找现有表单数据或创建新的
        form_data = FormData.query.first()
        if not form_data:
            form_data = FormData()
            form_data.project_id = project_id
        
        # 更新表单数据
        form_data.project_name = data.get('projectName', '')
        form_data.building_no = data.get('buildingNo', '')
        form_data.project_location = data.get('projectLocation', '')
        form_data.design_no = data.get('designNo', '')
        form_data.construction_unit = data.get('constructionUnit', '')
        form_data.design_unit = data.get('designUnit', '')
        form_data.standard_selection = data.get('standardSelection', 'municipal')
        
        # 将form_data对象转换为JSON字符串存储
        form_data.form_data = json.dumps(data.get('formData', {}), ensure_ascii=False)
        
        # 保存到数据库
        db.session.add(form_data)
        db.session.commit()
        
        return jsonify({'success': True, 'message': '数据保存成功'})
    except Exception as e:
        db.session.rollback()
        print(f"保存表单数据时出错: {str(e)}")
        print(traceback.format_exc())
        return jsonify({'error': str(e)}), 500

@app.route('/api/load_form', methods=['GET'])
def load_form():
    try:
        import json
        
        # 获取最新的表单数据
        form_data = FormData.query.order_by(FormData.updated_at.desc()).first()
        
        if not form_data:
            return jsonify({'error': '没有找到保存的数据'}), 404
            
        result = {
            'projectName': form_data.project_name,
            'buildingNo': form_data.building_no,
            'projectLocation': form_data.project_location,
            'designNo': form_data.design_no,
            'constructionUnit': form_data.construction_unit,
            'designUnit': form_data.design_unit,
            'standardSelection': form_data.standard_selection,
            'formData': json.loads(form_data.form_data) if form_data.form_data else {}
        }
        
        return jsonify(result)
    except Exception as e:
        print(f"加载表单数据时出错: {str(e)}")
        print(traceback.format_exc())
        return jsonify({'error': str(e)}), 500

@app.route('/api/project_info', methods=['GET'])
def get_project_info():
    from flask import request
    try:
        # 获取项目ID参数
        project_id = request.args.get('project_id')
        
        # 根据项目ID获取项目信息
        if project_id:
            project = Project.query.filter_by(id=project_id).first()
        else:
            # 如果没有提供项目ID，则获取第一个项目
            project = Project.query.first()
        
        # 获取最新的表单数据
        form_data = FormData.query.order_by(FormData.updated_at.desc()).first()
        
        # 辅助函数：格式化浮点数，保留2位小数
        def format_float(value):
            if value is not None:
                return round(value, 2)
            return ''
        
        result = {
            'projectName': project.name if project and project.name is not None else '',
            'projectCode': project.code if project and project.code is not None else '',
            'constructionUnit': project.construction_unit if project and project.construction_unit is not None else '',
            'designUnit': project.design_unit if project and project.design_unit is not None else '',
            'projectLocation': project.location if project and project.location is not None else '',
            'buildingArea': format_float(project.building_area) if project else '',
            'buildingType': project.building_type if project and project.building_type is not None else '',  # 添加建筑类型
            'buildingNo': form_data.building_no if form_data and form_data.building_no is not None else '',
            'designNo': form_data.design_no if form_data and form_data.design_no is not None else '',
            'standardSelection': form_data.standard_selection if form_data and form_data.standard_selection is not None else 'municipal',
            'buildingHeight': format_float(project.building_height) if project else '',
            'buildingFloors': project.building_floors if project and project.building_floors is not None else ''
        }
        
        return jsonify(result)
    except Exception as e:
        print(f"获取项目信息时出错: {str(e)}")
        print(traceback.format_exc())
        return jsonify({'error': str(e)}), 500

@app.route('/api/get_template', methods=['GET'])
def get_template():
    try:
        template_path = os.path.join(app.static_folder, '绿色建材应用比例计算书.docx')
        if not os.path.exists(template_path):
            return jsonify({'error': '模板文件不存在'}), 404
            
        return send_file(template_path, as_attachment=True)
    except Exception as e:
        print(f"获取模板文件时出错: {str(e)}")
        print(traceback.format_exc())
        return jsonify({'error': str(e)}), 500

# 添加一个函数来生成得分表缓存的键
def get_scores_cache_key(level, specialty, project_id=None, standard=None):
    """
    生成评分数据的缓存键
    
    参数:
        level: 评价等级
        specialty: 专业
        project_id: 项目ID
        standard: 评价标准
        
    返回值:
        缓存键字符串
    """
    # 构建基本缓存键
    cache_key = f"scores_{level}_{specialty}"
    
    # 如果有项目ID，添加到缓存键
    if project_id:
        cache_key += f"_{project_id}"
    
    # 如果有评价标准，添加到缓存键
    if standard:
        cache_key += f"_{standard}"
    
    return cache_key

@app.route('/api/save_score', methods=['POST'])
def save_score():
    """保存评分数据的API端点"""
    try:
        # 获取请求数据
        data = request.get_json()
        
        # 验证数据
        if not data:
            return jsonify({'error': '请求数据为空'}), 400
        
        # 获取评价等级和专业
        level = data.get('level')
        specialty = data.get('specialty')
        project_id = data.get('project_id')
        scores = data.get('scores', [])
        standard = data.get('standard', '成都市标')  # 获取评价标准，默认为成都市标
        
        # 验证必要字段
        if not level or not specialty:
            return jsonify({'error': '缺少必要字段: level, specialty'}), 400
        
        # 如果没有提供项目ID，尝试从第一个评分记录中获取项目名称
        project_name = None
        if not project_id and scores and len(scores) > 0:
            project_name = scores[0].get('project_name')
        
        # 记录请求信息
        app.logger.info(f"保存评分数据: 级别={level}, 专业={specialty}, 项目ID={project_id}, 项目名称={project_name}, 评价标准={standard}, 评分数量={len(scores)}")
        
        # 如果提供了项目ID，获取项目信息
        if project_id:
            try:
                project = get_project(project_id)
                if project:
                    project_name = project.name
                    # 如果项目有评价标准，使用项目的评价标准
                    if project.standard:
                        standard = project.standard
                        app.logger.info(f"使用项目的评价标准: {standard}")
            except Exception as e:
                app.logger.error(f"获取项目信息失败: {str(e)}")
        
        try:
            # 开始数据库事务
            # 如果提供了项目ID，先删除该项目该专业该级别的所有评分记录
            if project_id:
                delete_query = """
                DELETE FROM `得分表`
                WHERE `项目ID` = :project_id AND `专业` = :specialty AND `评价等级` = :level
                """
                result = db.session.execute(
                    text(delete_query), 
                    {"project_id": project_id, "specialty": specialty, "level": level}
                )
                app.logger.info(f"删除项目 {project_id} 的 {specialty} 专业 {level} 级别的评分记录: {result.rowcount} 条")
                
                # 同时删除project_scores表中的记录
                delete_ps_query = """
                DELETE FROM project_scores
                WHERE project_id = :project_id AND specialty = :specialty AND level = :level
                """
                result = db.session.execute(
                    text(delete_ps_query), 
                    {"project_id": project_id, "specialty": specialty, "level": level}
                )
                app.logger.info(f"删除project_scores表中项目 {project_id} 的 {specialty} 专业 {level} 级别的评分记录: {result.rowcount} 条")
            
            # 如果提供了项目名称但没有项目ID，先删除该项目名称该专业该级别的所有评分记录
            elif project_name:
                delete_query = """
                DELETE FROM `得分表`
                WHERE `项目名称` = :project_name AND `专业` = :specialty AND `评价等级` = :level
                """
                result = db.session.execute(
                    text(delete_query), 
                    {"project_name": project_name, "specialty": specialty, "level": level}
                )
                app.logger.info(f"删除项目 '{project_name}' 的 {specialty} 专业 {level} 级别的评分记录: {result.rowcount} 条")
            
            # 插入新的评分记录
            insert_count = 0
            for score_data in scores:
                # 获取评分数据
                clause_number = score_data.get('clause_number')
                category = score_data.get('category')
                is_achieved = score_data.get('is_achieved')
                score = score_data.get('score', '0')
                technical_measures = score_data.get('technical_measures', '')
                
                # 如果没有条文号，跳过
                if not clause_number:
                    continue
                
                # 插入评分记录到得分表
                insert_query = """
                INSERT INTO `得分表` (
                    `项目ID`, `项目名称`, `专业`, `评价等级`, `条文号`, `分类`, `是否达标`, `得分`, `技术措施`, `评价标准`
                ) VALUES (:project_id, :project_name, :specialty, :level, :clause_number, :category, 
                         :is_achieved, :score, :technical_measures, :standard)
                """
                
                try:
                    db.session.execute(
                        text(insert_query),
                        {
                            "project_id": project_id,
                            "project_name": project_name,
                            "specialty": specialty,
                            "level": level,
                            "clause_number": clause_number,
                            "category": category,
                            "is_achieved": is_achieved,
                            "score": score,
                            "technical_measures": technical_measures,
                            "standard": standard
                        }
                    )
                    
                    # 同时插入到project_scores表
                    if project_id:
                        # 尝试将得分转换为浮点数
                        try:
                            if score and score.strip():
                                score_float = float(score)
                            else:
                                score_float = 0
                        except (ValueError, TypeError):
                            score_float = 0
                        
                        insert_ps_query = """
                        INSERT INTO project_scores (
                            project_id, level, specialty, clause_number, category, is_achieved, score, technical_measures
                        ) VALUES (:project_id, :level, :specialty, :clause_number, :category, :is_achieved, :score_float, :technical_measures)
                        """
                        
                        db.session.execute(
                            text(insert_ps_query),
                            {
                                "project_id": project_id,
                                "level": level,
                                "specialty": specialty,
                                "clause_number": clause_number,
                                "category": category,
                                "is_achieved": is_achieved,
                                "score_float": score_float,
                                "technical_measures": technical_measures
                            }
                        )
                    
                    insert_count += 1
                except Exception as insert_error:
                    print(f"插入评分记录失败: {str(insert_error)}, 条文号: {clause_number}")
                    continue
            
            # 提交事务
            db.session.commit()
            app.logger.info(f"成功插入 {insert_count} 条评分记录")
            
            # 清除缓存
            cache_key = get_scores_cache_key(level, specialty, project_id, standard)
            if cache.has(cache_key):
                cache.delete(cache_key)
                app.logger.info(f"清除缓存: {cache_key}")
            
            return jsonify({
                'success': True,
                'message': f'成功保存 {insert_count} 条评分记录'
            }), 200
        
        except Exception as e:
            # 回滚事务
            db.session.rollback()
            app.logger.error(f"保存评分数据失败: {str(e)}")
            app.logger.error(traceback.format_exc())
            
            return jsonify({'error': f'保存评分数据失败: {str(e)}'}), 500
    
    except Exception as e:
        app.logger.error(f"处理保存评分请求失败: {str(e)}")
        app.logger.error(traceback.format_exc())
        return jsonify({'error': f'处理请求失败: {str(e)}'}), 500

@app.route('/api/load_scores', methods=['GET'])
def load_scores():
    """加载评分数据的API端点"""
    try:
        # 获取请求参数
        level = request.args.get('level')
        specialty = request.args.get('specialty')
        project_id = request.args.get('project_id')
        standard = request.args.get('standard', '成都市标')  # 获取评价标准，默认为成都市标
        
        # 验证必要参数
        if not level or not specialty:
            return jsonify({'error': '缺少必要参数: level, specialty'}), 400
        
        # 记录请求信息
        app.logger.info(f"加载评分数据: 级别={level}, 专业={specialty}, 项目ID={project_id}, 评价标准={standard}")
        
        # 如果提供了项目ID，获取项目信息
        project_name = None
        if project_id:
            try:
                project = get_project(project_id)
                if project:
                    project_name = project.name
                    # 如果项目有评价标准，使用项目的评价标准
                    if project.standard:
                        standard = project.standard
                        app.logger.info(f"使用项目的评价标准: {standard}")
            except Exception as e:
                app.logger.error(f"获取项目信息失败: {str(e)}")
        
        # 构建缓存键
        cache_key = get_scores_cache_key(level, specialty, project_id, standard)
        
        # 尝试从缓存获取数据
        cached_data = cache.get(cache_key)
        if cached_data:
            app.logger.info(f"从缓存获取评分数据: {cache_key}")
            return jsonify(cached_data)
        
        # 构建查询条件
        query_conditions = []
        query_params = {}
        
        # 添加必要条件
        query_conditions.append("`评价等级` = :level")
        query_params["level"] = level
        
        query_conditions.append("`专业` = :specialty")
        query_params["specialty"] = specialty
        
        # 如果指定了项目ID，添加到查询条件
        if project_id:
            query_conditions.append("`项目ID` = :project_id")
            query_params["project_id"] = project_id
            
            # 如果项目有评价标准，添加到查询条件
            if standard:
                query_conditions.append("`评价标准` = :standard")
                query_params["standard"] = standard
        # 如果指定了项目名称，添加到查询条件
        elif project_name:
            query_conditions.append("`项目名称` = :project_name")
            query_params["project_name"] = project_name
            
            # 如果项目有评价标准，添加到查询条件
            if standard:
                query_conditions.append("`评价标准` = :standard")
                query_params["standard"] = standard
        # 如果没有指定项目ID和项目名称，但指定了评价标准，添加到查询条件
        elif standard:
            query_conditions.append("`评价标准` = :standard")
            query_params["standard"] = standard
        
        # 构建查询语句
        query = f"""
        SELECT `条文号`, `分类`, `是否达标`, `得分`, `技术措施`
        FROM `得分表`
        WHERE {' AND '.join(query_conditions)}
        ORDER BY `条文号`
        """
        
        # 打印查询语句和参数，用于调试
        app.logger.info(f"执行查询1 (所有条件): {query}")
        app.logger.info(f"查询参数: {query_params}")
        
        try:
            # 执行查询
            result = db.session.execute(text(query), query_params)
            rows = result.fetchall()
            
            # 打印查询结果的行数，用于调试
            app.logger.info(f"查询结果1: {len(rows)} 行")
            
            # 如果没有结果，尝试只按项目ID查询
            if len(rows) == 0 and project_id:
                query_conditions = ["`项目ID` = :project_id"]
                query_params = {"project_id": project_id}
                
                query = f"""
                SELECT `条文号`, `分类`, `是否达标`, `得分`, `技术措施`
                FROM `得分表`
                WHERE {' AND '.join(query_conditions)}
                ORDER BY `条文号`
                """
                
                app.logger.info(f"执行查询2 (只按项目ID): {query}")
                app.logger.info(f"查询参数: {query_params}")
                
                result = db.session.execute(text(query), query_params)
                rows = result.fetchall()
                
                app.logger.info(f"查询结果2: {len(rows)} 行")
            
            # 如果仍然没有结果，尝试不加任何条件查询
            if len(rows) == 0:
                query = """
                SELECT `条文号`, `分类`, `是否达标`, `得分`, `技术措施`
                FROM `得分表`
                LIMIT 10
                """
                
                app.logger.info(f"执行查询3 (不加条件): {query}")
                
                result = db.session.execute(text(query))
                rows = result.fetchall()
                
                app.logger.info(f"查询结果3: {len(rows)} 行")
            
            # 处理查询结果
            scores = []
            for row in rows:
                scores.append({
                    'clause_number': row[0],
                    'category': row[1],
                    'is_achieved': row[2],
                    'score': row[3],
                    'technical_measures': row[4]
                })
            
            # 构建响应数据
            response_data = {
                'success': True,
                'scores': scores,
                'count': len(scores),
                'level': level,
                'specialty': specialty,
                'project_id': project_id,
                'standard': standard
            }
            
            # 缓存响应数据
            cache.set(cache_key, response_data)
            app.logger.info(f"缓存评分数据: {cache_key}, 记录数: {len(scores)}")
            
            return jsonify(response_data)
        
        except Exception as db_error:
            app.logger.error(f"数据库查询失败: {str(db_error)}")
            app.logger.error(traceback.format_exc())
            return jsonify({'error': f'数据库查询失败: {str(db_error)}'}), 500
        
    except Exception as e:
        app.logger.error(f"加载评分数据失败: {str(e)}")
        traceback.print_exc()
        return jsonify({'error': f'加载评分数据失败: {str(e)}'}), 500

@app.route('/api/project_scores', methods=['GET'])
def get_project_scores():
    try:
        project_id = request.args.get('project_id')
        by_category = request.args.get('by_category', 'false').lower() == 'true'
        
        # 如果提供了项目ID，转换为整数
        if project_id:
            project_id = int(project_id)
        
        # 计算各专业得分
        specialty_scores = calculate_specialty_scores(project_id, by_category)
        
        # 计算总得分
        if by_category:
            total_score = sum(specialty_scores.values())
            
            # 计算各分类的总分
            categories = ['安全耐久', '健康舒适', '生活便利', '资源节约', '环境宜居', '提高与创新']
            category_totals = {category: 0 for category in categories}
            
            for specialty, scores in specialty_scores.items():
                for category in categories:
                    category_totals[category] += scores
            
            # 确保所有分类得分都四舍五入到两位小数
            for category in categories:
                category_totals[category] = round(category_totals[category], 2)
        else:
            total_score = sum(specialty_scores.values())
            category_totals = None
        
        return jsonify({
            'success': True,
            'specialty_scores': specialty_scores,
            'total_score': round(total_score, 2),
            'category_totals': category_totals
        }), 200
    
    except Exception as e:
        app.logger.error(f"获取专业得分失败: {str(e)}")
        traceback.print_exc()
        return jsonify({'error': f'获取专业得分失败: {str(e)}'}), 500

@app.route('/api/get_score_summary', methods=['GET'])
def api_get_score_summary():
    """获取评分汇总数据的API端点"""
    try:
        # 获取项目ID参数
        project_id = request.args.get('project_id')
        
        # 获取force_refresh参数
        force_refresh = request.args.get('force_refresh', 'false').lower() == 'true'
        
        # 如果没有提供项目ID，返回错误
        if not project_id:
            return jsonify({'error': '缺少项目ID参数'}), 400
        
        # 转换项目ID为整数
        project_id = int(project_id)
        
        # 获取项目信息，确定评价标准
        project = get_project(project_id)
        if not project:
            return jsonify({'error': f'找不到ID为{project_id}的项目'}), 404
            
        project_standard = project.standard if project and project.standard else '成都市标'
        
        # 只在强制刷新时记录详细日志
        if force_refresh:
            app.logger.info(f"获取评分汇总数据: 项目ID={project_id}, 评价标准={project_standard}, 强制刷新={force_refresh}")
        
        # 缓存键
        cache_key = f"score_summary_{project_id}_{project_standard}"
        
        # 如果强制刷新，清除缓存
        if force_refresh and cache.has(cache_key):
            cache.delete(cache_key)
            if app.debug:
                app.logger.info(f"清除评分汇总缓存: {cache_key}")
        
        # 获取评分汇总数据
        score_summary = get_score_summary(project_id, force_refresh=force_refresh)
        
        # 检查返回的评分汇总数据评价标准是否与项目标准一致（只在强制刷新时检查）
        if force_refresh and score_summary.get('project_standard') != project_standard:
            if app.debug:
                app.logger.warning(f"评分汇总数据的评价标准({score_summary.get('project_standard')})与项目评价标准({project_standard})不一致，重新获取")
            cache.delete(cache_key)
            score_summary = get_score_summary(project_id, force_refresh=True)
        
        # 返回评分汇总数据
        return jsonify(score_summary)
    
    except ValueError as e:
        app.logger.error(f"参数错误: {str(e)}")
        return jsonify({'error': f'参数错误: {str(e)}'}), 400
    except Exception as e:
        app.logger.error(f"获取评分汇总数据失败: {str(e)}")
        if app.debug:
            app.logger.error(traceback.format_exc())
        # 返回错误信息
        return jsonify({
            'error': '获取评分汇总数据失败',
            'message': str(e) if app.debug else '服务器内部错误'
        }), 500

# 移除缓存装饰器，确保每次都从数据库获取最新数据
def get_min_scores():
    """
    获取各专业的最低得分要求，直接返回硬编码的值
    
    返回值:
        包含各专业最低得分要求的字典
    """
    # 直接返回硬编码的最低得分要求
    return {
        '建筑专业': 16,
        '结构专业': 8,
        '给排水专业': 8,
        '电气专业': 8,
        '暖通专业': 8,
        '景观专业': 8
    }

# 在应用启动时检查并添加缺失的列
def check_and_add_missing_columns():
    """检查并添加缺失的字段"""
    try:
        logger.info("开始检查缺失的字段...")

        # 检查并添加项目地点字段到星级案例表
        try:
            # 使用MySQL语法检查和添加列
            db.session.execute(text("""
            ALTER TABLE `星级案例` 
            ADD COLUMN IF NOT EXISTS `项目地点` VARCHAR(200)
            """))
            db.session.commit()
            logger.info("检查或添加星级案例表的项目地点字段成功")
        except Exception as e:
            logger.error(f"检查或添加星级案例表的项目地点字段时出错: {str(e)}")
            db.session.rollback()

        # 检查并添加唯一索引，防止数据重复
        try:
            # 先检查索引是否存在
            result = db.session.execute(text("""
            SHOW INDEX FROM `星级案例` WHERE Key_name = 'IX_星级案例_Unique'
            """))
            if not result.fetchone():
                # 创建唯一索引
                db.session.execute(text("""
                CREATE UNIQUE INDEX `IX_星级案例_Unique` ON `星级案例`(`评价标准`, `星级目标`, `条文号`, `建筑类型`)
                """))
                db.session.commit()
                logger.info("添加星级案例表的唯一索引成功")
            else:
                logger.info("星级案例表的唯一索引已存在")
        except Exception as e:
            logger.error(f"检查或添加星级案例表的唯一索引时出错: {str(e)}")
            db.session.rollback()

        # 检查并为得分表添加索引，提高查询速度
        try:
            # 先检查索引是否存在
            result = db.session.execute(text("""
            SHOW INDEX FROM `得分表` WHERE Key_name = 'IX_得分表_项目ID_评价标准'
            """))
            if not result.fetchone():
                # 创建索引
                db.session.execute(text("""
                CREATE INDEX `IX_得分表_项目ID_评价标准` ON `得分表`(`项目ID`, `评价标准`)
                """))
                db.session.commit()
                logger.info("为得分表添加项目ID和评价标准的复合索引成功")
            else:
                logger.info("得分表的项目ID和评价标准复合索引已存在")
        except Exception as e:
            logger.error(f"为得分表添加索引时出错: {str(e)}")
            db.session.rollback()
            
        # 检查并为得分表添加专业索引，进一步提高查询速度
        try:
            # 先检查索引是否存在
            result = db.session.execute(text("""
            SHOW INDEX FROM `得分表` WHERE Key_name = 'IX_得分表_专业'
            """))
            if not result.fetchone():
                # 创建索引
                db.session.execute(text("""
                CREATE INDEX `IX_得分表_专业` ON `得分表`(`专业`)
                """))
                db.session.commit()
                logger.info("为得分表添加专业索引成功")
            else:
                logger.info("得分表的专业索引已存在")
        except Exception as e:
            logger.error(f"为得分表添加专业索引时出错: {str(e)}")
            db.session.rollback()

        # 其他字段检查...
        
    except Exception as e:
        logger.error(f"检查和添加缺失字段时出错: {str(e)}")
        logger.error(traceback.format_exc())

# 在init_db函数中调用此函数
def init_db():
    """初始化数据库"""
    try:
        db.create_all()
        check_and_add_missing_columns()  # 添加这一行
        return True
    except Exception as e:
        logger.error(f"初始化数据库失败: {str(e)}")
        return False

@app.route('/update_score', methods=['POST'])
def update_score():
    """
    接收并处理更新项目得分的请求
    """
    try:
        # 获取JSON数据
        data = request.get_json()
        
        if not data:
            return jsonify({'success': False, 'message': '未接收到数据'}), 400
        
        # 提取必要的字段
        project_id = data.get('project_id')
        standard = data.get('standard')
        clause_number = data.get('clause_number')
        score = data.get('score')
        page = data.get('page')
        level = data.get('level')
        
        # 验证必要字段
        if not all([project_id, standard, clause_number, score is not None]):
            return jsonify({'success': False, 'message': '缺少必要的字段'}), 400
        
        try:
            # 检查项目是否存在
            result = db.session.execute(
                text("SELECT id FROM projects WHERE id = :project_id"),
                {"project_id": project_id}
            )
            project = result.fetchone()
            
            if not project:
                return jsonify({'success': False, 'message': '项目不存在'}), 404
            
            # 检查得分记录是否已存在
            result = db.session.execute(
                text("SELECT id FROM project_scores WHERE project_id = :project_id AND standard = :standard AND clause_number = :clause_number"),
                {"project_id": project_id, "standard": standard, "clause_number": clause_number}
            )
            existing_score = result.fetchone()
            
            if existing_score:
                # 更新现有记录
                db.session.execute(
                    text("""
                    UPDATE project_scores 
                    SET score = :score, page = :page, level = :level, updated_at = CURRENT_TIMESTAMP
                    WHERE project_id = :project_id AND standard = :standard AND clause_number = :clause_number
                    """),
                    {
                        "score": score, 
                        "page": page, 
                        "level": level, 
                        "project_id": project_id, 
                        "standard": standard, 
                        "clause_number": clause_number
                    }
                )
                app.logger.info(f"更新项目 {project_id} 的得分记录，条文号: {clause_number}, 分值: {score}")
            else:
                # 创建新记录
                db.session.execute(
                    text("""
                    INSERT INTO project_scores (project_id, standard, clause_number, score, page, level, created_at, updated_at)
                    VALUES (:project_id, :standard, :clause_number, :score, :page, :level, CURRENT_TIMESTAMP, CURRENT_TIMESTAMP)
                    """),
                    {
                        "project_id": project_id, 
                        "standard": standard, 
                        "clause_number": clause_number, 
                        "score": score, 
                        "page": page, 
                        "level": level
                    }
                )
                app.logger.info(f"创建项目 {project_id} 的新得分记录，条文号: {clause_number}, 分值: {score}")
            
            # 提交事务
            db.session.commit()
            
            return jsonify({'success': True, 'message': '得分更新成功'})
        
        except Exception as e:
            db.session.rollback()
            app.logger.error(f"数据库操作失败: {str(e)}")
            return jsonify({'success': False, 'message': f'数据库操作失败: {str(e)}'}), 500
        
    except Exception as e:
        app.logger.error(f"更新得分时出错: {str(e)}")
        return jsonify({'success': False, 'message': f'服务器错误: {str(e)}'}), 500

# 添加新的路由用于直接更新得分
@app.route('/api/update_score_direct', methods=['POST'])
def update_score_direct():

    try:
        # 获取请求数据
        data = request.get_json()
        if not data:
            return jsonify({
                'success': False,
                'message': '未接收到JSON数据'
            }), 400
        
        # 验证必要参数
        project_id = data.get('project_id')
        clause_number = data.get('clause_number')
        score = data.get('score')
        standard = data.get('standard', '成都市标')
        
        # 可选参数
        specialty = data.get('specialty', '建筑专业')
        level = data.get('level', '提高级')
        category = data.get('category', '资源节约')
        is_achieved = data.get('is_achieved', 'true')
        technical_measures = data.get('technical_measures', '')
        
        # 记录请求信息
        app.logger.info(f"接收到更新请求: 项目ID={project_id}, 条文号={clause_number}, 得分={score}, 标准={standard}")
        
        # 验证必要参数
        if not all([project_id, clause_number, score is not None]):
            return jsonify({
                'success': False,
                'message': '缺少必要参数: project_id, clause_number, score'
            }), 400
        
        max_retries = 3
        retry_count = 0
        
        while retry_count < max_retries:
            try:
                # 使用SQLAlchemy事务管理
                app.logger.info("开始数据库操作")
                
                with db.session.begin():
                    # 先检查记录是否存在
                    check_query = """
                    SELECT COUNT(*) FROM `得分表`
                    WHERE `项目ID` = :project_id AND `条文号` = :clause_number AND `评价标准` = :standard
                    """
                    result = db.session.execute(
                        text(check_query), 
                        {"project_id": project_id, "clause_number": clause_number, "standard": standard}
                    )
                    count = result.fetchone()[0]
                    
                    if count > 0:
                        # 更新现有记录，只修改得分字段
                        update_query = """
                        UPDATE `得分表`
                        SET `得分` = :score
                        WHERE `项目ID` = :project_id AND `条文号` = :clause_number AND `评价标准` = :standard
                        """
                        result = db.session.execute(
                            text(update_query),
                            {"score": score, "project_id": project_id, "clause_number": clause_number, "standard": standard}
                        )
                        app.logger.info(f"更新记录: 影响行数={result.rowcount}")
                    else:
                        # 插入新记录
                        insert_query = """
                        INSERT INTO `得分表` (
                            `项目ID`, `项目名称`, `专业`, `评价等级`, `条文号`, 
                            `分类`, `是否达标`, `得分`, `技术措施`, `评价标准`
                        )
                        VALUES (:project_id, :project_name, :specialty, :level, :clause_number, 
                               :category, :is_achieved, :score, :technical_measures, :standard)
                        """
                        result = db.session.execute(
                            text(insert_query),
                            {
                                "project_id": project_id, 
                                "project_name": f'项目{project_id}', 
                                "specialty": specialty, 
                                "level": level, 
                                "clause_number": clause_number,
                                "category": category, 
                                "is_achieved": is_achieved, 
                                "score": score, 
                                "technical_measures": technical_measures, 
                                "standard": standard
                            }
                        )
                        app.logger.info(f"插入记录: 影响行数={result.rowcount}")
                    
                    # 验证更新是否成功
                    verify_query = """
                    SELECT `得分` FROM `得分表`
                    WHERE `项目ID` = :project_id AND `条文号` = :clause_number AND `评价标准` = :standard
                    """
                    result = db.session.execute(
                        text(verify_query), 
                        {"project_id": project_id, "clause_number": clause_number, "standard": standard}
                    )
                    row = result.fetchone()
                    
                    if not row:
                        raise Exception("验证失败：未找到更新后的记录")
                    
                    actual_score = row[0]
                    app.logger.info(f"验证成功: 条文 {clause_number} 的得分为 {actual_score}")
                
                # 清除评分汇总缓存
                cache_key = f"score_summary_{project_id}_{standard}"
                if cache.has(cache_key):
                    cache.delete(cache_key)
                    app.logger.info(f"清除评分汇总缓存: {cache_key}")
                
                # 清除专业得分缓存
                specialty_cache_key = get_scores_cache_key('提高级', specialty.split('专业')[0], project_id, standard)
                if cache.has(specialty_cache_key):
                    cache.delete(specialty_cache_key)
                    app.logger.info(f"清除专业得分缓存: {specialty_cache_key}")
                
                # 返回成功响应
                return jsonify({
                    'success': True,
                    'message': '更新成功',
                    'data': {
                        'project_id': project_id,
                        'clause_number': clause_number,
                        'score': actual_score,
                        'standard': standard
                    }
                })
                
            except Exception as e:
                # 回滚事务
                db.session.rollback()
                app.logger.error(f"数据库操作失败: {str(e)}")
                retry_count += 1
                if retry_count < max_retries:
                    app.logger.info(f"正在进行第{retry_count}次重试...")
                    time.sleep(1)  # 等待1秒后重试
                else:
                    app.logger.error("已达到最大重试次数，操作失败")
                    app.logger.error(traceback.format_exc())
                    return jsonify({
                        'success': False,
                        'message': f'数据库操作失败: {str(e)}'
                    }), 500
    
    except Exception as e:
        app.logger.error(f"处理请求失败: {str(e)}")
        app.logger.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': f'处理请求失败: {str(e)}'
        }), 500

# 添加新的路由用于根据条文号查询得分
@app.route('/api/get_score_by_clause', methods=['POST'])
def get_score_by_clause():
    """
    根据条文号查询得分
    
    请求参数:
    {
        "project_id": 1,                 // 项目ID
        "clause_number": "3.1.2.14",     // 条文号
        "standard": "成都市标"            // 评价标准
    }
    
    响应:
    {
        "success": true,                 // 是否成功
        "message": "查询成功",            // 消息
        "clause_number": "3.1.2.14",     // 条文号
        "score": 12,                     // 得分值
        "category": "资源节约",           // 分类
        "specialty": "建筑专业",          // 专业
        "is_achieved": "true"            // 是否达标
    }
    """
    try:
        # 获取请求数据
        data = request.get_json()
        if not data:
            return jsonify({
                'success': False,
                'message': '未接收到JSON数据'
            }), 400
        
        # 验证必要参数
        project_id = data.get('project_id')
        clause_number = data.get('clause_number')
        standard = data.get('standard', '成都市标')
        
        # 记录请求信息
        app.logger.info(f"接收到查询请求: 项目ID={project_id}, 条文号={clause_number}, 标准={standard}")
        
        # 验证必要参数
        if not all([project_id, clause_number]):
            return jsonify({
                'success': False,
                'message': '缺少必要参数: project_id, clause_number'
            }), 400
        
        try:
            # 查询记录
            query = """
            SELECT `条文号`, `分类`, `专业`, `是否达标`, `得分`
            FROM `得分表`
            WHERE `项目ID` = :project_id AND `条文号` = :clause_number AND `评价标准` = :standard
            """
            result = db.session.execute(
                text(query), 
                {"project_id": project_id, "clause_number": clause_number, "standard": standard}
            )
            row = result.fetchone()
            
            if row:
                # 返回成功响应
                return jsonify({
                    'success': True,
                    'message': '查询成功',
                    'clause_number': row[0],
                    'category': row[1],
                    'specialty': row[2],
                    'is_achieved': row[3],
                    'score': row[4]
                })
            else:
                return jsonify({
                    'success': False,
                    'message': f'未找到条文 {clause_number} 的记录'
                }), 404
        
        except Exception as e:
            app.logger.error(f"数据库查询失败: {str(e)}")
            app.logger.error(traceback.format_exc())
            return jsonify({
                'success': False,
                'message': f'数据库查询失败: {str(e)}'
            }), 500
    
    except Exception as e:
        app.logger.error(f"处理请求失败: {str(e)}")
        app.logger.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': f'处理请求失败: {str(e)}'
        }), 500

# 添加新的路由用于执行SQL语句
@app.route('/api/execute_sql', methods=['POST'])
def execute_sql():
    """
    执行SQL语句
    
    请求参数:
    {
        "sql": "SQL语句",
        "params": [参数1, 参数2, ...] (可选)
    }
    
    响应:
    {
        "success": true,
        "message": "执行成功",
        "results": [
            {
                "column1": "value1",
                "column2": "value2"
            }
        ]
    }
    """
    try:
        # 获取请求数据
        data = request.get_json()
        if not data:
            return jsonify({
                'success': False,
                'message': '未接收到JSON数据'
            }), 400
        
        # 获取SQL语句和参数
        sql = data.get('sql')
        params = data.get('params', {})
        
        if not sql:
            return jsonify({
                'success': False,
                'message': '缺少必要参数: sql'
            }), 400
        
        # 记录请求信息
        app.logger.info(f"接收到SQL执行请求: {sql}, 参数: {params}")
        
        # 安全检查：禁止执行危险操作
        sql_lower = sql.lower().strip()
        
        # 禁止执行的操作列表
        forbidden_operations = [
            'drop table', 'drop database', 'truncate table', 
            'alter table', 'create table', 'create database',
            'exec ', 'execute ', 'sp_', 'xp_'
        ]
        
        # 检查是否包含禁止的操作
        for op in forbidden_operations:
            if op in sql_lower:
                app.logger.warning(f"尝试执行危险SQL操作: {sql}")
                return jsonify({
                    'success': False,
                    'message': f'禁止执行危险操作: {op}'
                }), 403
        
        try:
            # 使用SQLAlchemy执行SQL语句
            result = db.session.execute(text(sql), params)
            
            # 如果是SELECT语句，获取结果
            results = []
            if sql_lower.startswith('select'):
                # 获取结果集的列名
                column_names = result.keys()
                # 处理返回的结果集
                for row in result.fetchall():
                    results.append(dict(zip(column_names, row)))
            
            # 如果不是SELECT语句，则提交事务
            if not sql_lower.startswith('select'):
                db.session.commit()
            
            app.logger.info("SQL执行成功")
            
            # 返回成功响应
            return jsonify({
                'success': True,
                'message': '执行成功',
                'results': results
            })
        
        except Exception as e:
            # 回滚事务
            db.session.rollback()
            app.logger.error(f"SQL执行失败: {str(e)}")
            app.logger.error(traceback.format_exc())
            
            return jsonify({
                'success': False,
                'message': f'SQL执行失败: {str(e)}'
            }), 500
    
    except Exception as e:
        app.logger.error(f"处理请求失败: {str(e)}")
        app.logger.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': f'处理请求失败: {str(e)}'
        }), 500

@app.route('/api/project_scores', methods=['GET'])
def check_project_scores():
    """检查项目是否有得分数据的API端点"""
    try:
        # 获取请求参数
        project_id = request.args.get('project_id')
        standard = request.args.get('standard', '成都市标')
        
        if not project_id:
            return jsonify({'error': '缺少必要参数: project_id'}), 400
        
        # 使用SQLAlchemy查询项目得分数据
        # 查询项目得分数据
        query = """
        SELECT COUNT(*) AS count
        FROM `得分表`
        WHERE `项目ID` = :project_id
        """
        
        result = db.session.execute(text(query), {"project_id": project_id})
        row = result.fetchone()
        
        # 获取得分记录数
        score_count = row[0] if row else 0
        
        # 如果有得分记录，返回一些示例得分数据
        scores = []
        if score_count > 0:
            sample_query = """
            SELECT `条文号`, `分类`, `是否达标`, `得分`, `技术措施`, `专业`, `评价等级`
            FROM `得分表`
            WHERE `项目ID` = :project_id
            LIMIT 5
            """
            
            result = db.session.execute(text(sample_query), {"project_id": project_id})
            for row in result.fetchall():
                scores.append({
                    'clause_number': row[0],
                    'category': row[1],
                    'is_achieved': row[2],
                    'score': row[3],
                    'technical_measures': row[4],
                    'specialty': row[5],
                    'level': row[6]
                })
        
        # 返回结果
        return jsonify({
            'success': True,
            'has_scores': score_count > 0,
            'score_count': score_count,
            'sample_scores': scores
        })
    
    except Exception as e:
        app.logger.error(f"检查项目得分数据时出错: {str(e)}")
        app.logger.error(traceback.format_exc())
        return jsonify({'error': f'服务器错误: {str(e)}'}), 500

def save_score_for_new_project(data):
    """为新创建的项目保存初始评分数据"""
    try:
        # 获取数据
        level = data.get('level')
        specialty = data.get('specialty')
        project_id = data.get('project_id')
        scores = data.get('scores', [])
        standard = data.get('standard', '成都市标')  # 获取评价标准，默认为成都市标
        
        # 验证必要字段
        if not level or not specialty:
            print(f"缺少必要字段: level={level}, specialty={specialty}")
            return False
        
        # 如果没有提供项目ID，尝试从第一个评分记录中获取项目名称
        project_name = None
        if not project_id and scores and len(scores) > 0:
            project_name = scores[0].get('project_name')
            print(f"从评分记录中获取项目名称: {project_name}")
        
        # 如果没有提供评分数据，则生成默认评分数据
        if not scores or len(scores) == 0:
            print("没有提供评分数据，生成默认评分数据")
            try:
                # 获取标准数据
                standards_data = get_standards_by_name(standard)
                
                # 过滤出当前专业和等级的标准
                filtered_standards = []
                for std in standards_data:
                    std_specialty = std.get('专业', '')
                    std_level = '基本级'  # 默认为基本级
                    
                    # 根据条文号判断等级
                    clause_number = std.get('条文号', '')
                    if clause_number and clause_number.startswith('9'):
                        std_level = '提高级'
                    
                    # 检查专业和等级是否匹配
                    if (std_specialty and specialty in std_specialty) and std_level == level:
                        filtered_standards.append(std)
                
                # 生成默认评分数据
                scores = []
                for std in filtered_standards:
                    score_data = {
                        'clause_number': std.get('条文号', ''),
                        'category': std.get('分类', ''),
                        'is_achieved': '否',  # 默认为否
                        'score': '0',  # 默认为0分
                        'technical_measures': ''  # 默认为空
                    }
                    scores.append(score_data)
                
                print(f"生成了 {len(scores)} 条默认评分数据")
            except Exception as e:
                print(f"生成默认评分数据失败: {str(e)}")
                traceback.print_exc()
                return False
        
        # 使用SQLAlchemy进行数据库操作
        try:
            print("开始数据库操作")
            
            # 检查得分表是否存在
            result = db.session.execute(text("SHOW TABLES LIKE '得分表'"))
            if not result.fetchone():
                print("得分表不存在，尝试创建")
                # 尝试创建得分表
                create_table_query = """
                CREATE TABLE `得分表` (
                    `ID` INT AUTO_INCREMENT PRIMARY KEY,
                    `项目ID` INT,
                    `项目名称` VARCHAR(100),
                    `专业` VARCHAR(50),
                    `评价等级` VARCHAR(20),
                    `条文号` VARCHAR(50),
                    `分类` VARCHAR(50),
                    `是否达标` VARCHAR(10),
                    `得分` VARCHAR(10),
                    `技术措施` TEXT,
                    `评价标准` VARCHAR(50),
                    created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                    updated_at DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci
                """
                db.session.execute(text(create_table_query))
                db.session.commit()
                print("得分表创建成功")
            
            # 检查project_scores表是否存在
            result = db.session.execute(text("SHOW TABLES LIKE 'project_scores'"))
            if not result.fetchone():
                print("project_scores表不存在，尝试创建")
                # 尝试创建project_scores表
                create_table_query = """
                CREATE TABLE project_scores (
                    id INT AUTO_INCREMENT PRIMARY KEY,
                    project_id INT NOT NULL,
                    level VARCHAR(10) NOT NULL,
                    specialty VARCHAR(20) NOT NULL,
                    clause_number VARCHAR(20) NOT NULL,
                    category VARCHAR(20) NULL,
                    is_achieved VARCHAR(10) NOT NULL,
                    score FLOAT NULL,
                    technical_measures TEXT NULL,
                    created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                    updated_at DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci
                """
                db.session.execute(text(create_table_query))
                db.session.commit()
                print("project_scores表创建成功")
                
                # 创建索引
                db.session.execute(text("""
                    CREATE INDEX idx_project_scores_project_id ON project_scores (project_id)
                """))
                db.session.commit()
                print("成功创建project_scores表索引")
        except Exception as e:
            print(f"检查或创建project_scores表失败: {str(e)}")
            traceback.print_exc()
            return False
        
        # 开始事务
        try:
            # 如果提供了项目ID，先删除该项目该专业该级别的所有评分记录
            if project_id:
                delete_query = """
                DELETE FROM `得分表`
                WHERE `项目ID` = :project_id AND `专业` = :specialty AND `评价等级` = :level
                """
                result = db.session.execute(
                    text(delete_query),
                    {"project_id": project_id, "specialty": specialty, "level": level}
                )
                print(f"删除项目 {project_id} 的 {specialty} 专业 {level} 级别的评分记录: {result.rowcount} 条")
                
                # 同时删除project_scores表中的记录
                delete_ps_query = """
                DELETE FROM project_scores
                WHERE project_id = :project_id AND specialty = :specialty AND level = :level
                """
                result = db.session.execute(
                    text(delete_ps_query),
                    {"project_id": project_id, "specialty": specialty, "level": level}
                )
                print(f"删除project_scores表中项目 {project_id} 的 {specialty} 专业 {level} 级别的评分记录: {result.rowcount} 条")
            
            # 插入评分数据
            insert_count = 0
            for score_data in scores:
                # 获取评分数据
                clause_number = score_data.get('clause_number')
                category = score_data.get('category')
                is_achieved = score_data.get('is_achieved')
                score = score_data.get('score', '0')
                technical_measures = score_data.get('technical_measures', '')
                
                # 插入评分记录到得分表
                insert_query = """
                INSERT INTO `得分表` (
                    `项目ID`, `项目名称`, `专业`, `评价等级`, `条文号`, `分类`, `是否达标`, `得分`, `技术措施`, `评价标准`
                ) VALUES (:project_id, :project_name, :specialty, :level, :clause_number, :category, 
                         :is_achieved, :score, :technical_measures, :standard)
                """
                
                try:
                    db.session.execute(
                        text(insert_query),
                        {
                            "project_id": project_id,
                            "project_name": project_name,
                            "specialty": specialty,
                            "level": level,
                            "clause_number": clause_number,
                            "category": category,
                            "is_achieved": is_achieved,
                            "score": score,
                            "technical_measures": technical_measures,
                            "standard": standard
                        }
                    )
                    
                    # 同时插入到project_scores表
                    if project_id:
                        # 尝试将得分转换为浮点数
                        try:
                            if score and score.strip():
                                score_float = float(score)
                            else:
                                score_float = 0
                        except (ValueError, TypeError):
                            score_float = 0
                        
                        insert_ps_query = """
                        INSERT INTO project_scores (
                            project_id, level, specialty, clause_number, category, is_achieved, score, technical_measures
                        ) VALUES (:project_id, :level, :specialty, :clause_number, :category, :is_achieved, :score_float, :technical_measures)
                        """
                        
                        db.session.execute(
                            text(insert_ps_query),
                            {
                                "project_id": project_id,
                                "level": level,
                                "specialty": specialty,
                                "clause_number": clause_number,
                                "category": category,
                                "is_achieved": is_achieved,
                                "score_float": score_float,
                                "technical_measures": technical_measures
                            }
                        )
                    
                    insert_count += 1
                except Exception as insert_error:
                    print(f"插入评分记录失败: {str(insert_error)}, 条文号: {clause_number}")
                    continue
            
            # 提交事务
            db.session.commit()
            
            # 同时也保存到缓存
            cache_key = get_scores_cache_key(level, specialty, project_id, standard)
            cache.set(cache_key, scores)
            
            print(f"已为项目 {project_id} 创建 {level}-{specialty} 的初始评分记录，共 {insert_count} 条")
            return True
        except Exception as e:
            # 回滚事务
            try:
                db.session.rollback()
            except:
                pass
            
            try:
                db.session.close()
            except:
                pass
            
            print(f"保存评分记录失败: {str(e)}")
            traceback.print_exc()
            return False
    except Exception as e:
        print(f"保存初始评分记录失败: {str(e)}")
        traceback.print_exc()
        return False

@app.route('/api/star_case_scores', methods=['GET'])
def get_star_case_scores():
    try:
        # 获取目标项目ID
        target_project_id = request.args.get('target_project_id')
        if not target_project_id:
            return jsonify({
                'success': False,
                'message': '请提供目标项目ID'
            }), 400

        # 获取项目信息
        target_project = Project.query.get(target_project_id)
        if not target_project:
            return jsonify({
                'success': False,
                'message': f'目标项目(ID={target_project_id})不存在'
            }), 404

        # 获取项目的评价标准和星级目标
        standard = target_project.standard or '成都市标'
        star_rating_target = target_project.star_rating_target or '一星级'
        building_type = target_project.building_type
        project_location = target_project.location  # 获取项目地点

        try:
            # 首先获取当前项目适用的条文号列表
            app.logger.info(f"获取项目 {target_project_id} 适用的条文列表")
            result = db.session.execute(text(f"""
                SELECT 条文号 
                FROM 评价标准 
                WHERE 标准名称 = :standard_name
            """), {"standard_name": standard})
            valid_clauses = set([row[0] for row in result.fetchall()])
            app.logger.info(f"找到 {len(valid_clauses)} 条适用的条文")

            # 构建基本查询条件
            query_conditions = ["评价标准 = :standard"]
            query_params = {"standard": standard}

            # 添加星级目标条件
            if star_rating_target:
                query_conditions.append("星级目标 = :star_rating_target")
                query_params["star_rating_target"] = star_rating_target

            # 添加建筑类型条件
            if building_type:
                query_conditions.append("建筑类型 = :building_type")
                query_params["building_type"] = building_type
            
            # 首先尝试使用项目地点进行匹配
            location_matched_case_data = None
            if project_location:
                app.logger.info(f"尝试使用项目地点进行模糊匹配: {project_location}")
                
                # 构建模糊匹配查询
                location_query = f"""
                    SELECT DISTINCT 条文号, 分类, 是否达标, 得分, 技术措施, 专业, 评价等级
                    FROM 星级案例
                    WHERE {' AND '.join(query_conditions)} 
                    AND (项目地点 LIKE :location_pattern OR :project_location LIKE CONCAT(项目地点, '%'))
                """
                location_params = {
                    **query_params,
                    "location_pattern": f'%{project_location}%',
                    "project_location": project_location
                }
                
                app.logger.info(f"执行项目地点匹配查询: {location_query}")
                result = db.session.execute(text(location_query), location_params)
                location_matched_case_data = result.fetchall()
                
                if location_matched_case_data:
                    app.logger.info(f"基于项目地点找到 {len(location_matched_case_data)} 条匹配的星级案例数据")
            
            # 如果没有通过地点匹配到数据，则使用基本条件查询
            if not location_matched_case_data:
                # 标准查询
                query = f"""
                    SELECT DISTINCT 条文号, 分类, 是否达标, 得分, 技术措施, 专业, 评价等级
                    FROM 星级案例
                    WHERE {' AND '.join(query_conditions)}
                """
                
                app.logger.info(f"执行标准查询: {query}")
                result = db.session.execute(text(query), query_params)
                case_data = result.fetchall()
            else:
                # 使用基于地点匹配的数据
                case_data = location_matched_case_data
                app.logger.info("使用基于项目地点匹配的星级案例数据")

            if case_data:
                app.logger.info(f"找到 {len(case_data)} 条匹配的星级案例数据")
                
                # 过滤出当前标准中有效的条文
                filtered_case_data = []
                for record in case_data:
                    if record[0] in valid_clauses:
                        filtered_case_data.append(record)
                    else:
                        app.logger.info(f"过滤掉不适用的条文: {record[0]}")
                
                app.logger.info(f"过滤后剩余 {len(filtered_case_data)} 条有效的星级案例数据")
                
                # 先删除目标项目的所有得分数据
                delete_query = """
                DELETE FROM 得分表
                WHERE 项目ID = :project_id
                """
                result = db.session.execute(text(delete_query), {"project_id": target_project_id})
                deleted_count = result.rowcount
                app.logger.info(f"删除目标项目的 {deleted_count} 条得分数据")

                # 插入新的得分数据
                inserted_count = 0
                # 用于跟踪已插入的条文号，防止重复
                inserted_clauses = set()
                
                for record in filtered_case_data:
                    clause_number, category, is_achieved, score, technical_measures, specialty, level = record
                    
                    # 跳过已经插入过的条文号
                    if clause_number in inserted_clauses:
                        app.logger.info(f"跳过重复条文: {clause_number}")
                        continue
                    
                    inserted_clauses.add(clause_number)

                    insert_query = """
                    INSERT INTO 得分表 (
                        项目ID, 项目名称, 专业, 评价等级, 条文号, 
                        分类, 是否达标, 得分, 技术措施, 评价标准
                    ) VALUES (:project_id, :project_name, :specialty, :level, :clause_number,
                            :category, :is_achieved, :score, :technical_measures, :standard)
                    """
                    db.session.execute(
                        text(insert_query),
                        {
                            "project_id": target_project_id,
                            "project_name": target_project.name,
                            "specialty": specialty,
                            "level": level,
                            "clause_number": clause_number,
                            "category": category,
                            "is_achieved": is_achieved,
                            "score": score,
                            "technical_measures": technical_measures,
                            "standard": standard
                        }
                    )
                    inserted_count += 1

                    # 每100条提交一次，避免事务过大
                    if inserted_count % 100 == 0:
                        db.session.commit()
                        app.logger.info(f"已提交 {inserted_count} 条记录")

                # 提交事务
                db.session.commit()
                app.logger.info(f"事务提交成功，共导入 {inserted_count} 条记录")

                # 清除目标项目的缓存
                cache_key = f"score_summary_{target_project_id}_{standard}"
                if cache.has(cache_key):
                    cache.delete(cache_key)
                    app.logger.info(f"清除目标项目的评分汇总缓存: {cache_key}")

                return jsonify({
                    'success': True,
                    'message': f'成功导入 {inserted_count} 条星级案例数据',
                    'data': {
                        'standard': standard,
                        'star_rating_target': star_rating_target,
                        'imported_count': inserted_count,
                        'location_matched': bool(location_matched_case_data)
                    }
                })
            else:
                app.logger.warning(f"未找到匹配的星级案例数据")
                return jsonify({
                    'success': False,
                    'message': f'未找到匹配的星级案例数据'
                }), 404

        except Exception as e:
            # 回滚事务
            db.session.rollback()
            app.logger.error(f"数据库操作失败: {str(e)}")
            app.logger.error(traceback.format_exc())
            return jsonify({
                'success': False,
                'message': f'数据库操作失败: {str(e)}'
            }), 500

    except Exception as e:
        app.logger.error(f"处理请求失败: {str(e)}")
        app.logger.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': f'处理请求失败: {str(e)}'
        }), 500

def create_default_scores(project_id, project_name, standard_selection):
    """
    为项目创建默认评分记录
    
    参数:
    - project_id: 项目ID
    - project_name: 项目名称
    - standard_selection: 评价标准
    """
    try:
        print("开始创建默认评分记录...")
        
        # 获取标准数据
        print(f"获取标准数据: {standard_selection}")
        standards_data = get_standards_by_name(standard_selection)
        if standards_data:
            print(f"获取到 {len(standards_data)} 条 {standard_selection} 标准数据")
            # 打印前5条标准数据的示例
            if len(standards_data) > 0:
                print("标准数据示例:")
                for i, std in enumerate(standards_data[:5]):
                    try:
                        print(f"  {i+1}. 条文号: {std.条文号}, 专业: {std.专业}, 属性: {std.属性}, 分类: {std.分类}")
                    except Exception as attr_error:
                        print(f"  {i+1}. 无法显示标准数据: {str(attr_error)}")
        else:
            print(f"未获取到 {standard_selection} 标准数据")
            # 尝试获取默认标准数据
            print("尝试获取默认标准数据")
            standards_data = get_standards_by_name('成都市标')
            if standards_data:
                print(f"获取到 {len(standards_data)} 条默认标准数据")
            else:
                print("未获取到任何标准数据")
                raise Exception("未获取到任何标准数据")
        
        # 定义要处理的专业和级别
        specialties = ['建筑专业', '结构专业', '给排水专业', '暖通专业', '电气专业']
        levels = ['基本级', '提高级']
        
        # 为每个专业和级别生成评分数据
        total_inserted = 0
        for specialty in specialties:
            for level in levels:
                # 过滤出该专业该级别的条文
                filtered_standards = []
                for std in standards_data:
                    try:
                        if std.专业 == specialty:
                            if (level == '基本级' and std.属性 == '控制项') or \
                               (level == '提高级' and std.属性 == '评分项'):
                                filtered_standards.append(std)
                    except Exception as attr_error:
                        print(f"处理标准数据时出错: {str(attr_error)}, 标准数据: {std}")
                        continue
                
                print(f"找到 {len(filtered_standards)} 条 {standard_selection} 的 {specialty} 专业 {level} 级别条文")
                
                # 如果找到了条文，则插入评分数据
                if filtered_standards:
                    # 先删除该项目该专业该级别的所有评分记录
                    try:
                        result = db.session.execute(
                            text("DELETE FROM `得分表` WHERE `项目ID` = :project_id AND `专业` = :specialty AND `评价等级` = :level"),
                            {"project_id": project_id, "specialty": specialty, "level": level}
                        )
                        print(f"删除项目 {project_id} 的 {specialty} 专业 {level} 级别的评分记录: {result.rowcount} 条")
                    except Exception as delete_error:
                        print(f"删除评分记录失败: {str(delete_error)}")
                        traceback.print_exc()
                    
                    # 插入新的评分记录
                    insert_count = 0
                    insert_errors = 0
                    for std in filtered_standards:
                        # 基本级默认达标，提高级默认不达标
                        is_achieved = '是' if level == '基本级' else '否'
                        
                        # 插入评分记录
                        try:
                            db.session.execute(
                                text("""
                                INSERT INTO `得分表` (
                                    `项目ID`, `项目名称`, `专业`, `评价等级`, `条文号`, `分类`, 
                                    `是否达标`, `得分`, `技术措施`, `评价标准`
                                ) VALUES (
                                    :project_id, :project_name, :specialty, :level, :clause_number, 
                                    :category, :is_achieved, :score, :technical_measures, :standard
                                )
                                """),
                                {
                                    "project_id": project_id,
                                    "project_name": project_name,
                                    "specialty": specialty,
                                    "level": level,
                                    "clause_number": std.条文号,
                                    "category": std.分类,
                                    "is_achieved": is_achieved,
                                    "score": '0',
                                    "technical_measures": '',
                                    "standard": standard_selection
                                }
                            )
                            insert_count += 1
                            
                            # 每插入10条记录提交一次事务，避免事务过大
                            if insert_count % 10 == 0:
                                db.session.commit()
                                print(f"已提交 {insert_count} 条记录")
                        except Exception as insert_error:
                            insert_errors += 1
                            print(f"插入评分记录失败: {str(insert_error)}, 条文号: {std.条文号}")
                            if insert_errors <= 3:  # 只打印前3个错误的详细信息
                                traceback.print_exc()
                            continue
                    
                    print(f"为项目 {project_id} 的 {specialty} 专业 {level} 级别插入了 {insert_count} 条评分记录，失败 {insert_errors} 条")
                    total_inserted += insert_count
        
        # 提交最后的事务
        try:
            db.session.commit()
            print("最终提交事务成功")
        except Exception as commit_error:
            print(f"提交事务失败: {str(commit_error)}")
            db.session.rollback()
            traceback.print_exc()
        
        print(f"为项目 {project_id} 总共插入了 {total_inserted} 条评分记录")
        return True
    
    except Exception as e:
        print(f"生成项目评分数据失败: {str(e)}")
        traceback.print_exc()
        return False

# 注册export.py中的generate_word函数为app的路由
@app.route('/api/generate_word', methods=['POST'])
def handle_generate_word():
    """
    处理生成Word文档的请求
    """
    try:
        # 获取请求数据
        data = request.get_json()
        if not data:
            return jsonify({"error": "请求数据为空"}), 400

        # 提取必要参数
        project_id = data.get('project_id')
        if not project_id:
            return jsonify({"error": "缺少项目ID参数"}), 400
            
        # 获取标准参数，默认为成都市标
        standard = data.get('standard', '成都市标')
        
        # 添加use_cache参数，默认为False，强制从数据库获取最新数据
        request_data = {
            'project_id': project_id,
            'standard': standard,
            'use_cache': False
        }
        
        # 调用generate_word函数
        return generate_word(request_data)
    except Exception as e:
        app.logger.error(f"处理生成Word请求失败: {str(e)}")
        return jsonify({"error": f"处理请求失败: {str(e)}"}), 500

@app.route('/api/generate_dwg', methods=['POST'])
def handle_generate_dwg():
    """
    处理生成DWG文档的请求
    """
    try:
        # 获取请求数据
        data = request.get_json()
        if not data:
            return jsonify({"error": "请求数据为空"}), 400
        
        # 获取项目ID
        project_id = data.get('project_id')
        if not project_id:
            return jsonify({"error": "请提供项目ID"}), 400
        
        # 先计算最新的评分数据
        app.logger.info(f"生成DWG文件前，先计算项目 {project_id} 的最新评分数据")
        try:
            # 获取项目信息
            project = Project.query.get(project_id)
            if not project:
                return jsonify({'success': False, 'message': '项目不存在'}), 404
            
            # 同步得分表和project_scores表的数据
            sync_result = sync_score_tables(project_id)
            if not sync_result:
                app.logger.warning(f"同步得分表和project_scores表数据失败或无需同步，继续使用现有数据")
            
            # 获取项目的所有评分记录
            result = db.session.execute(text("""
                SELECT ps.clause_number, ps.score, s.分类, s.专业, s.属性
                FROM project_scores ps
                LEFT JOIN 
                (
                    SELECT 条文号, 分类, 专业, 属性 FROM 成都市标
                    UNION ALL
                    SELECT 条文号, 分类, 专业, 属性 FROM 四川省标
                    UNION ALL
                    SELECT 条文号, 分类, 专业, 属性 FROM 国标
                ) s ON ps.clause_number = s.条文号
                WHERE ps.project_id = :project_id
            """), {"project_id": project_id})
            scores = result.fetchall()
            
            # 将更新后的项目添加到session并保存
            db.session.add(project)
            db.session.commit()
            
            app.logger.info(f"项目 {project_id} 的评分数据已更新")
        except Exception as e:
            db.session.rollback()
            app.logger.error(f"计算项目评分时出错: {str(e)}")
            app.logger.error(traceback.format_exc())
            # 继续执行，不影响后续操作

        # 添加use_cache参数，默认为False，强制从数据库获取最新数据
        data['use_cache'] = False
        
        # 调用generate_dwg函数
        return generate_dwg(data)
    except Exception as e:
        app.logger.error(f"处理生成DWG请求失败: {str(e)}")
        return jsonify({"error": f"处理请求失败: {str(e)}"}), 500

@app.route('/api/self-assessment-report', methods=['POST'])
def handle_self_assessment_report():
    """
    处理生成绿建自评估报告的请求
    """
    try:
        # 获取请求数据
        data = request.get_json()
        if not data:
            return jsonify({"error": "请求数据为空"}), 400

        # 提取必要参数
        project_id = data.get('project_id')
        if not project_id:
            return jsonify({"error": "缺少项目ID参数"}), 400
        
        # 添加use_cache参数，默认为False，强制从数据库获取最新数据
        request_data = {
            'project_id': project_id,
            'use_cache': False
        }
        
        # 调用generate_self_assessment_report函数
        return generate_self_assessment_report(request_data)
    except Exception as e:
        app.logger.error(f"处理生成绿建自评估报告请求失败: {str(e)}")
        return jsonify({"error": f"处理请求失败: {str(e)}"}), 500

@app.route('/api/calculate_scores/<int:project_id>', methods=['POST'])
def calculate_project_scores(project_id):
    """计算并更新项目的各项评分"""
    try:
        # 添加更详细的日志
        logger.info(f"开始计算项目评分: 项目ID={project_id}")
        
        # 获取项目信息
        project = Project.query.get(project_id)
        if not project:
            logger.error(f"项目不存在: ID={project_id}")
            return jsonify({'success': False, 'message': '项目不存在'}), 404
        
        # 获取项目评价标准
        project_standard = project.standard if project and project.standard else '成都市标'
        logger.info(f"计算项目评分: 项目ID={project_id}, 评价标准={project_standard}")
        
        # 同步得分表和project_scores表的数据
        try:
            sync_result = sync_score_tables(project_id)
            if not sync_result:
                logger.warning(f"同步得分表和project_scores表数据失败或无需同步，继续使用现有数据")
        except Exception as e:
            logger.error(f"同步得分表和project_scores表数据出错: {str(e)}")
            # 继续执行，不中断流程
        
        # 获取数据库会话
        session = db.session
        
        try:
            # 查询项目的所有评分记录，使用项目评价标准
            sql_query = """
                SELECT `专业`, `分类`, `是否达标`, `得分`, `评价等级`
                FROM `得分表`
                WHERE `项目ID` = :project_id 
                AND (`评价标准` = :standard 
                     OR (:standard = '国标' AND `评价标准` IN ('国标', '通用国标'))
                     OR (:standard = '通用国标' AND `评价标准` IN ('国标', '通用国标')))
            """
            
            logger.info(f"执行SQL查询: 项目ID={project_id}, 标准={project_standard}")
            
            result = session.execute(
                text(sql_query),
                {"project_id": project_id, "standard": project_standard}
            )
            
            scores = result.fetchall()
            logger.info(f"查询到 {len(scores)} 条评分记录")
            
            if len(scores) == 0:
                logger.warning(f"未查询到评分记录，尝试使用默认标准")
                # 尝试使用默认标准再次查询
                result = session.execute(
                    text(sql_query),
                    {"project_id": project_id, "standard": "成都市标"}
                )
                scores = result.fetchall()
                logger.info(f"使用默认标准查询到 {len(scores)} 条评分记录")
                
        except Exception as e:
            logger.error(f"查询评分记录时出错: {str(e)}")
            logger.error(traceback.format_exc())
            return jsonify({'success': False, 'message': f'查询评分记录失败: {str(e)}'}), 500
        
        # 初始化各专业分数
        专业分数 = {
            '建筑': 0,
            '结构': 0,
            '给排水': 0,
            '电气': 0,
            '暖通': 0,
            '景观': 0,
            '环境健康与节能': 0,
            '环境健康与节能创新': 0,
            '建筑创新': 0,
            '结构创新': 0,
            '暖通创新': 0,
            '景观创新': 0,
        }
        
        # 初始化各章节分数
        章节分数 = {
            '安全耐久': 0,
            '健康舒适': 0,
            '生活便利': 0,
            '资源节约': 0,
            '环境宜居': 0,
            '提高与创新': 0,
        }
        
        # 初始化专业分类得分
        specialty_scores_by_category = {}
        specialties = ['建筑专业', '结构专业', '给排水专业', '电气专业', '暖通专业', '景观专业']
        if project_standard == '四川省标':
            specialties.append('环境健康与节能专业')
        
        for specialty in specialties:
            specialty_scores_by_category[specialty] = {
                '安全耐久': 0,
                '健康舒适': 0,
                '生活便利': 0,
                '资源节约': 0,
                '环境宜居': 0,
                '提高与创新': 0,
                '总分': 0
            }
        
        # 计算各专业和章节的分数
        for idx, score in enumerate(scores):
            try:
                专业 = score[0] if score[0] else ""  # 专业
                分类 = score[1] if score[1] else ""  # 分类
                是否达标 = score[2] if score[2] else ""  # 是否达标
                score_value = score[3] if score[3] is not None else 0  # 得分
                评价等级 = score[4] if score[4] else ""  # 评价等级
                
                # 确保score_value是浮点数
                if isinstance(score_value, str):
                    try:
                        score_value = float(score_value)
                    except (ValueError, TypeError):
                        score_value = 0
                        logger.warning(f"无法将得分转换为浮点数: {score[3]}, 使用默认值0")
                
                print(f"处理评分记录 {idx+1}/{len(scores)}: 专业={专业}, 分类={分类}, 是否达标={是否达标}, 得分={score_value}, 评价等级={评价等级}")
                
                # 基本级条文必须达标才计分，提高级条文有得分就计分
                if (评价等级 == '基本级' and 是否达标.lower() in ['是', 'yes', 'true', '1', 'y']) or \
                   (评价等级 == '提高级' and score_value > 0):
                    # 判断专业
                    if 专业 and '建' in 专业:
                        if '创新' in 专业:
                            专业分数['建筑创新'] += score_value
                            print(f"建筑创新得分 +{score_value}, 累计={专业分数['建筑创新']}")
                        else:
                            专业分数['建筑'] += score_value
                            print(f"建筑得分 +{score_value}, 累计={专业分数['建筑']}")
                            
                            # 更新专业分类得分
                            if '安全' in 分类 or '耐久' in 分类:
                                specialty_scores_by_category['建筑专业']['安全耐久'] += score_value
                                章节分数['安全耐久'] += score_value
                                print(f"安全耐久得分 +{score_value}, 累计={章节分数['安全耐久']}")
                            elif '健康' in 分类 or '舒适' in 分类:
                                specialty_scores_by_category['建筑专业']['健康舒适'] += score_value
                                章节分数['健康舒适'] += score_value
                                print(f"健康舒适得分 +{score_value}, 累计={章节分数['健康舒适']}")
                            elif '生活' in 分类 or '便利' in 分类:
                                specialty_scores_by_category['建筑专业']['生活便利'] += score_value
                                章节分数['生活便利'] += score_value
                                print(f"生活便利得分 +{score_value}, 累计={章节分数['生活便利']}")
                            elif '资源' in 分类 or '节约' in 分类:
                                specialty_scores_by_category['建筑专业']['资源节约'] += score_value
                                章节分数['资源节约'] += score_value
                                print(f"资源节约得分 +{score_value}, 累计={章节分数['资源节约']}")
                            elif '环境' in 分类 or '宜居' in 分类:
                                specialty_scores_by_category['建筑专业']['环境宜居'] += score_value
                                章节分数['环境宜居'] += score_value
                                print(f"环境宜居得分 +{score_value}, 累计={章节分数['环境宜居']}")
                            elif '提高' in 分类 or '创新' in 分类:
                                specialty_scores_by_category['建筑专业']['提高与创新'] += score_value
                                章节分数['提高与创新'] += score_value
                                print(f"提高与创新得分 +{score_value}, 累计={章节分数['提高与创新']}")
                            # 更新总分
                            specialty_scores_by_category['建筑专业']['总分'] += score_value
                            
                    elif 专业 and '结' in 专业:
                        if '创新' in 专业:
                            专业分数['结构创新'] += score_value
                            print(f"结构创新得分 +{score_value}, 累计={专业分数['结构创新']}")
                        else:
                            专业分数['结构'] += score_value
                            print(f"结构得分 +{score_value}, 累计={专业分数['结构']}")
                            
                            # 更新专业分类得分
                            if '安全' in 分类 or '耐久' in 分类:
                                specialty_scores_by_category['结构专业']['安全耐久'] += score_value
                                章节分数['安全耐久'] += score_value
                                print(f"安全耐久得分 +{score_value}, 累计={章节分数['安全耐久']}")
                            elif '健康' in 分类 or '舒适' in 分类:
                                specialty_scores_by_category['结构专业']['健康舒适'] += score_value
                                章节分数['健康舒适'] += score_value
                                print(f"健康舒适得分 +{score_value}, 累计={章节分数['健康舒适']}")
                            elif '生活' in 分类 or '便利' in 分类:
                                specialty_scores_by_category['结构专业']['生活便利'] += score_value
                                章节分数['生活便利'] += score_value
                                print(f"生活便利得分 +{score_value}, 累计={章节分数['生活便利']}")
                            elif '资源' in 分类 or '节约' in 分类:
                                specialty_scores_by_category['结构专业']['资源节约'] += score_value
                                章节分数['资源节约'] += score_value
                                print(f"资源节约得分 +{score_value}, 累计={章节分数['资源节约']}")
                            elif '环境' in 分类 or '宜居' in 分类:
                                specialty_scores_by_category['结构专业']['环境宜居'] += score_value
                                章节分数['环境宜居'] += score_value
                                print(f"环境宜居得分 +{score_value}, 累计={章节分数['环境宜居']}")
                            elif '提高' in 分类 or '创新' in 分类:
                                specialty_scores_by_category['结构专业']['提高与创新'] += score_value
                                章节分数['提高与创新'] += score_value
                                print(f"提高与创新得分 +{score_value}, 累计={章节分数['提高与创新']}")
                            # 更新总分
                            specialty_scores_by_category['结构专业']['总分'] += score_value
                            
                    elif 专业 and ('给' in 专业 or '排' in 专业 or '水' in 专业):
                        专业分数['给排水'] += score_value
                        print(f"给排水得分 +{score_value}, 累计={专业分数['给排水']}")
                        
                        # 更新专业分类得分
                        if '安全' in 分类 or '耐久' in 分类:
                            specialty_scores_by_category['给排水专业']['安全耐久'] += score_value
                            章节分数['安全耐久'] += score_value
                            print(f"安全耐久得分 +{score_value}, 累计={章节分数['安全耐久']}")
                        elif '健康' in 分类 or '舒适' in 分类:
                            specialty_scores_by_category['给排水专业']['健康舒适'] += score_value
                            章节分数['健康舒适'] += score_value
                            print(f"健康舒适得分 +{score_value}, 累计={章节分数['健康舒适']}")
                        elif '生活' in 分类 or '便利' in 分类:
                            specialty_scores_by_category['给排水专业']['生活便利'] += score_value
                            章节分数['生活便利'] += score_value
                            print(f"生活便利得分 +{score_value}, 累计={章节分数['生活便利']}")
                        elif '资源' in 分类 or '节约' in 分类:
                            specialty_scores_by_category['给排水专业']['资源节约'] += score_value
                            章节分数['资源节约'] += score_value
                            print(f"资源节约得分 +{score_value}, 累计={章节分数['资源节约']}")
                        elif '环境' in 分类 or '宜居' in 分类:
                            specialty_scores_by_category['给排水专业']['环境宜居'] += score_value
                            章节分数['环境宜居'] += score_value
                            print(f"环境宜居得分 +{score_value}, 累计={章节分数['环境宜居']}")
                        elif '提高' in 分类 or '创新' in 分类:
                            specialty_scores_by_category['给排水专业']['提高与创新'] += score_value
                            章节分数['提高与创新'] += score_value
                            print(f"提高与创新得分 +{score_value}, 累计={章节分数['提高与创新']}")
                        # 更新总分
                        specialty_scores_by_category['给排水专业']['总分'] += score_value
                        
                    elif 专业 and '电' in 专业:
                        专业分数['电气'] += score_value
                        print(f"电气得分 +{score_value}, 累计={专业分数['电气']}")
                        
                        # 更新专业分类得分
                        if '安全' in 分类 or '耐久' in 分类:
                            specialty_scores_by_category['电气专业']['安全耐久'] += score_value
                            章节分数['安全耐久'] += score_value
                            print(f"安全耐久得分 +{score_value}, 累计={章节分数['安全耐久']}")
                        elif '健康' in 分类 or '舒适' in 分类:
                            specialty_scores_by_category['电气专业']['健康舒适'] += score_value
                            章节分数['健康舒适'] += score_value
                            print(f"健康舒适得分 +{score_value}, 累计={章节分数['健康舒适']}")
                        elif '生活' in 分类 or '便利' in 分类:
                            specialty_scores_by_category['电气专业']['生活便利'] += score_value
                            章节分数['生活便利'] += score_value
                            print(f"生活便利得分 +{score_value}, 累计={章节分数['生活便利']}")
                        elif '资源' in 分类 or '节约' in 分类:
                            specialty_scores_by_category['电气专业']['资源节约'] += score_value
                            章节分数['资源节约'] += score_value
                            print(f"资源节约得分 +{score_value}, 累计={章节分数['资源节约']}")
                        elif '环境' in 分类 or '宜居' in 分类:
                            specialty_scores_by_category['电气专业']['环境宜居'] += score_value
                            章节分数['环境宜居'] += score_value
                            print(f"环境宜居得分 +{score_value}, 累计={章节分数['环境宜居']}")
                        elif '提高' in 分类 or '创新' in 分类:
                            specialty_scores_by_category['电气专业']['提高与创新'] += score_value
                            章节分数['提高与创新'] += score_value
                            print(f"提高与创新得分 +{score_value}, 累计={章节分数['提高与创新']}")
                        # 更新总分
                        specialty_scores_by_category['电气专业']['总分'] += score_value
                        
                    elif 专业 and ('暖' in 专业 or '通' in 专业 or '空调' in 专业):
                        if '创新' in 专业:
                            专业分数['暖通创新'] += score_value
                            print(f"暖通创新得分 +{score_value}, 累计={专业分数['暖通创新']}")
                        else:
                            专业分数['暖通'] += score_value
                            print(f"暖通得分 +{score_value}, 累计={专业分数['暖通']}")
                            
                            # 更新专业分类得分
                            if '安全' in 分类 or '耐久' in 分类:
                                specialty_scores_by_category['暖通专业']['安全耐久'] += score_value
                                章节分数['安全耐久'] += score_value
                                print(f"安全耐久得分 +{score_value}, 累计={章节分数['安全耐久']}")
                            elif '健康' in 分类 or '舒适' in 分类:
                                specialty_scores_by_category['暖通专业']['健康舒适'] += score_value
                                章节分数['健康舒适'] += score_value
                                print(f"健康舒适得分 +{score_value}, 累计={章节分数['健康舒适']}")
                            elif '生活' in 分类 or '便利' in 分类:
                                specialty_scores_by_category['暖通专业']['生活便利'] += score_value
                                章节分数['生活便利'] += score_value
                                print(f"生活便利得分 +{score_value}, 累计={章节分数['生活便利']}")
                            elif '资源' in 分类 or '节约' in 分类:
                                specialty_scores_by_category['暖通专业']['资源节约'] += score_value
                                章节分数['资源节约'] += score_value
                                print(f"资源节约得分 +{score_value}, 累计={章节分数['资源节约']}")
                            elif '环境' in 分类 or '宜居' in 分类:
                                specialty_scores_by_category['暖通专业']['环境宜居'] += score_value
                                章节分数['环境宜居'] += score_value
                                print(f"环境宜居得分 +{score_value}, 累计={章节分数['环境宜居']}")
                            elif '提高' in 分类 or '创新' in 分类:
                                specialty_scores_by_category['暖通专业']['提高与创新'] += score_value
                                章节分数['提高与创新'] += score_value
                                print(f"提高与创新得分 +{score_value}, 累计={章节分数['提高与创新']}")
                            # 更新总分
                            specialty_scores_by_category['暖通专业']['总分'] += score_value
                            
                    elif 专业 and ('景' in 专业 or '园' in 专业):
                        if '创新' in 专业:
                            专业分数['景观创新'] += score_value
                            print(f"景观创新得分 +{score_value}, 累计={专业分数['景观创新']}")
                        else:
                            专业分数['景观'] += score_value
                            print(f"景观得分 +{score_value}, 累计={专业分数['景观']}")
                            
                            # 更新专业分类得分
                            if '安全' in 分类 or '耐久' in 分类:
                                specialty_scores_by_category['景观专业']['安全耐久'] += score_value
                                章节分数['安全耐久'] += score_value
                                print(f"安全耐久得分 +{score_value}, 累计={章节分数['安全耐久']}")
                            elif '健康' in 分类 or '舒适' in 分类:
                                specialty_scores_by_category['景观专业']['健康舒适'] += score_value
                                章节分数['健康舒适'] += score_value
                                print(f"健康舒适得分 +{score_value}, 累计={章节分数['健康舒适']}")
                            elif '生活' in 分类 or '便利' in 分类:
                                specialty_scores_by_category['景观专业']['生活便利'] += score_value
                                章节分数['生活便利'] += score_value
                                print(f"生活便利得分 +{score_value}, 累计={章节分数['生活便利']}")
                            elif '资源' in 分类 or '节约' in 分类:
                                specialty_scores_by_category['景观专业']['资源节约'] += score_value
                                章节分数['资源节约'] += score_value
                                print(f"资源节约得分 +{score_value}, 累计={章节分数['资源节约']}")
                            elif '环境' in 分类 or '宜居' in 分类:
                                specialty_scores_by_category['景观专业']['环境宜居'] += score_value
                                章节分数['环境宜居'] += score_value
                                print(f"环境宜居得分 +{score_value}, 累计={章节分数['环境宜居']}")
                            elif '提高' in 分类 or '创新' in 分类:
                                specialty_scores_by_category['景观专业']['提高与创新'] += score_value
                                章节分数['提高与创新'] += score_value
                                print(f"提高与创新得分 +{score_value}, 累计={章节分数['提高与创新']}")
                            # 更新总分
                            specialty_scores_by_category['景观专业']['总分'] += score_value
                            
                    elif 专业 and ('环境' in 专业 or '健康' in 专业 or '节能' in 专业):
                        if '创新' in 专业:
                            专业分数['环境健康与节能创新'] += score_value
                            print(f"环境健康与节能创新得分 +{score_value}, 累计={专业分数['环境健康与节能创新']}")
                        else:
                            专业分数['环境健康与节能'] += score_value
                            print(f"环境健康与节能得分 +{score_value}, 累计={专业分数['环境健康与节能']}")
                            
                            # 四川省标才有环境健康与节能专业
                            if project_standard == '四川省标' and '环境健康与节能专业' in specialty_scores_by_category:
                                # 更新专业分类得分
                                if '安全' in 分类 or '耐久' in 分类:
                                    specialty_scores_by_category['环境健康与节能专业']['安全耐久'] += score_value
                                    章节分数['安全耐久'] += score_value
                                    print(f"安全耐久得分 +{score_value}, 累计={章节分数['安全耐久']}")
                                elif '健康' in 分类 or '舒适' in 分类:
                                    specialty_scores_by_category['环境健康与节能专业']['健康舒适'] += score_value
                                    章节分数['健康舒适'] += score_value
                                    print(f"健康舒适得分 +{score_value}, 累计={章节分数['健康舒适']}")
                                elif '生活' in 分类 or '便利' in 分类:
                                    specialty_scores_by_category['环境健康与节能专业']['生活便利'] += score_value
                                    章节分数['生活便利'] += score_value
                                    print(f"生活便利得分 +{score_value}, 累计={章节分数['生活便利']}")
                                elif '资源' in 分类 or '节约' in 分类:
                                    specialty_scores_by_category['环境健康与节能专业']['资源节约'] += score_value
                                    章节分数['资源节约'] += score_value
                                    print(f"资源节约得分 +{score_value}, 累计={章节分数['资源节约']}")
                                elif '环境' in 分类 or '宜居' in 分类:
                                    specialty_scores_by_category['环境健康与节能专业']['环境宜居'] += score_value
                                    章节分数['环境宜居'] += score_value
                                    print(f"环境宜居得分 +{score_value}, 累计={章节分数['环境宜居']}")
                                elif '提高' in 分类 or '创新' in 分类:
                                    specialty_scores_by_category['环境健康与节能专业']['提高与创新'] += score_value
                                    章节分数['提高与创新'] += score_value
                                    print(f"提高与创新得分 +{score_value}, 累计={章节分数['提高与创新']}")
                                # 更新总分
                                specialty_scores_by_category['环境健康与节能专业']['总分'] += score_value
            except Exception as e:
                logger.error(f"处理评分记录时出错 (索引 {idx}): {str(e)}")
                # 继续处理下一条记录
        
        print("\n=== 评分计算结果 ===")
        print("\n各专业得分：")
        for 专业, 得分 in 专业分数.items():
            if 得分 > 0:  # 只显示有得分的专业
                print(f"{专业}: {得分:.2f}")
        
        print("\n各章节得分：")
        for 章节, 得分 in 章节分数.items():
            if 得分 > 0:  # 只显示有得分的章节
                print(f"{章节}: {得分:.2f}")
        
        # 计算各分类的总得分
        分类总分 = {
            '安全耐久': 0,
            '健康舒适': 0,
            '生活便利': 0,
            '资源节约': 0,
            '环境宜居': 0,
            '提高与创新': 0
        }
        
        # 遍历所有专业的分类得分，累加到分类总分
        for 专业, 分类得分 in specialty_scores_by_category.items():
            for 分类, 得分 in 分类得分.items():
                if 分类 != '总分':  # 跳过总分字段
                    分类总分[分类] += 得分
        
        # 计算项目总分
        total_score = round(sum(专业分数.values()), 2)
        print(f"\n项目总分：{total_score:.2f}")
        
        # 对专业分数和章节分数进行四舍五入处理
        专业分数_rounded = {k: round(v, 2) for k, v in 专业分数.items()}
        章节分数_rounded = {k: round(v, 2) for k, v in 章节分数.items()}
        
        # 直接使用SQL更新项目得分，确保数据库更新成功
        try:
            # 构造更新SQL
            update_sql = """
            UPDATE projects
            SET 
                architecture_score = :architecture_score,
                structure_score = :structure_score,
                water_supply_score = :water_supply_score,
                electrical_score = :electrical_score,
                hvac_score = :hvac_score,
                landscape_score = :landscape_score,
                env_health_energy_score = :env_health_energy_score,
                env_health_energy_innovation_score = :env_health_energy_innovation_score,
                architecture_innovation_score = :architecture_innovation_score,
                structure_innovation_score = :structure_innovation_score,
                hvac_innovation_score = :hvac_innovation_score,
                landscape_innovation_score = :landscape_innovation_score,
                safety_durability_score = :safety_durability_score,
                health_comfort_score = :health_comfort_score,
                life_convenience_score = :life_convenience_score,
                resource_saving_score = :resource_saving_score,
                environment_livability_score = :environment_livability_score,
                improvement_innovation_score = :improvement_innovation_score,
                total_score = :total_score,
                evaluation_result = :evaluation_result
            WHERE id = :project_id
            """
            
            # 确定评定结果
            evaluation_result = "未知"
            if project.standard == '成都市标':
                if total_score >= 80:
                    evaluation_result = '绿色建筑'
                else:
                    evaluation_result = '未达标'
            elif project.standard == '四川省标':
                if total_score >= 60 and total_score < 70:
                    evaluation_result = '基本级'
                elif total_score >= 70 and total_score < 80:
                    evaluation_result = '一星级'
                elif total_score >= 80 and total_score < 90:
                    evaluation_result = '二星级'
                elif total_score >= 90:
                    evaluation_result = '三星级'
                else:
                    evaluation_result = '未达标'
            elif project.standard == '国标':
                if total_score >= 60 and total_score < 70:
                    evaluation_result = '基本级'
                elif total_score >= 70 and total_score < 80:
                    evaluation_result = '一星级'
                elif total_score >= 80 and total_score < 90:
                    evaluation_result = '二星级'
                elif total_score >= 90:
                    evaluation_result = '三星级'
                else:
                    evaluation_result = '未达标'
            
            # 执行SQL更新
            params = {
                "architecture_score": 专业分数_rounded['建筑'],
                "structure_score": 专业分数_rounded['结构'],
                "water_supply_score": 专业分数_rounded['给排水'],
                "electrical_score": 专业分数_rounded['电气'],
                "hvac_score": 专业分数_rounded['暖通'],
                "landscape_score": 专业分数_rounded['景观'],
                "env_health_energy_score": 专业分数_rounded['环境健康与节能'],
                "env_health_energy_innovation_score": 专业分数_rounded['环境健康与节能创新'],
                "architecture_innovation_score": 专业分数_rounded['建筑创新'],
                "structure_innovation_score": 专业分数_rounded['结构创新'],
                "hvac_innovation_score": 专业分数_rounded['暖通创新'],
                "landscape_innovation_score": 专业分数_rounded['景观创新'],
                "safety_durability_score": 章节分数_rounded['安全耐久'],
                "health_comfort_score": 章节分数_rounded['健康舒适'],
                "life_convenience_score": 章节分数_rounded['生活便利'],
                "resource_saving_score": 章节分数_rounded['资源节约'],
                "environment_livability_score": 章节分数_rounded['环境宜居'],
                "improvement_innovation_score": 章节分数_rounded['提高与创新'],
                "total_score": total_score,
                "evaluation_result": evaluation_result,
                "project_id": project_id
            }
            
            # 显示参数
            logger.info(f"更新项目得分，参数: {params}")
            
            # 执行更新
            update_result = db.session.execute(text(update_sql), params)
            db.session.commit()
            
            # 检查更新结果
            logger.info(f"更新项目得分结果: 影响 {update_result.rowcount} 行")
            
            # 验证更新是否成功
            # 重新查询项目以确认更新成功
            project = db.session.query(Project).get(project_id)
            logger.info(f"更新后的项目得分: architecture_score={project.architecture_score}, total_score={project.total_score}")
            
            # 更新本地项目对象的值，便于后续构建响应
            project.architecture_score = 专业分数_rounded['建筑']
            project.structure_score = 专业分数_rounded['结构']
            project.water_supply_score = 专业分数_rounded['给排水']
            project.electrical_score = 专业分数_rounded['电气']
            project.hvac_score = 专业分数_rounded['暖通']
            project.landscape_score = 专业分数_rounded['景观']
            project.env_health_energy_score = 专业分数_rounded['环境健康与节能']
            project.env_health_energy_innovation_score = 专业分数_rounded['环境健康与节能创新']
            project.architecture_innovation_score = 专业分数_rounded['建筑创新']
            project.structure_innovation_score = 专业分数_rounded['结构创新']
            project.hvac_innovation_score = 专业分数_rounded['暖通创新']
            project.landscape_innovation_score = 专业分数_rounded['景观创新']
            project.safety_durability_score = 章节分数_rounded['安全耐久']
            project.health_comfort_score = 章节分数_rounded['健康舒适']
            project.life_convenience_score = 章节分数_rounded['生活便利']
            project.resource_saving_score = 章节分数_rounded['资源节约']
            project.environment_livability_score = 章节分数_rounded['环境宜居']
            project.improvement_innovation_score = 章节分数_rounded['提高与创新']
            project.total_score = total_score
            project.evaluation_result = evaluation_result
            
            # 清除评分汇总缓存
            cache_key = f"score_summary_{project_id}_{project_standard}"
            if cache.has(cache_key):
                cache.delete(cache_key)
                logger.info(f"清除评分汇总缓存: {cache_key}")
            
            # 清除专业得分缓存
            for specialty in ['建筑', '结构', '给排水', '电气', '暖通', '景观', '环境健康与节能']:
                specialty_cache_key = get_scores_cache_key('提高级', specialty, project_id, project_standard)
                if cache.has(specialty_cache_key):
                    cache.delete(specialty_cache_key)
                    logger.info(f"清除专业得分缓存: {specialty_cache_key}")
                    
            # 准备专业得分数据结构用于前端展示
            specialties = ['建筑专业', '结构专业', '给排水专业', '电气专业', '暖通专业', '景观专业']
            if project_standard == '四川省标':
                specialties.append('环境健康与节能专业')
                
            # 专业到字段的映射
            specialty_to_field = {
                '建筑专业': 专业分数_rounded['建筑'],
                '结构专业': 专业分数_rounded['结构'],
                '给排水专业': 专业分数_rounded['给排水'],
                '电气专业': 专业分数_rounded['电气'],
                '暖通专业': 专业分数_rounded['暖通'],
                '景观专业': 专业分数_rounded['景观'],
                '环境健康与节能专业': 专业分数_rounded['环境健康与节能']
            }
            
            specialty_scores = {specialty: specialty_to_field.get(specialty, 0) for specialty in specialties}
            
            # 对specialty_scores_by_category中的值进行四舍五入处理
            rounded_specialty_scores_by_category = {}
            for specialty, categories in specialty_scores_by_category.items():
                rounded_specialty_scores_by_category[specialty] = {
                    category: round(score, 2) for category, score in categories.items()
                }
            
            # 返回计算结果和项目详细评分数据
            return jsonify({
                'success': True,
                'message': '项目评分计算并更新成功',
                'total_score': total_score,
                'evaluation_result': evaluation_result,
                'specialty_scores': specialty_scores,
                'specialty_scores_by_category': rounded_specialty_scores_by_category,
                'project_standard': project_standard,
                'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                'scores_detail': {
                    '专业分数': 专业分数_rounded,
                    '章节分数': 章节分数_rounded
                }
            })
            
        except Exception as e:
            db.session.rollback()
            logger.error(f"更新项目评分时出错: {str(e)}")
            logger.error(traceback.format_exc())
            return jsonify({'success': False, 'message': f'更新项目评分时出错: {str(e)}'}), 500
        
    except Exception as e:
        logger.error(f"计算项目评分时出错: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'success': False, 'message': f'计算项目评分时出错: {str(e)}'}), 500

@app.route('/api/login', methods=['POST'])
def login():
    try:
        data = request.get_json()
        email = data.get('email')
        password = data.get('password')
        
        if not email or not password:
            return jsonify({'error': '请输入邮箱和密码'}), 400
            
        user = User.query.filter_by(email=email).first()
        
        if user and user.check_password(password):
            # 更新最后在线时间
            user.last_seen = datetime.utcnow()
            db.session.commit()
            
            session['user_id'] = user.id
            session['username'] = user.email  # 添加用户邮箱到 session
            session['role'] = user.role  # 添加用户角色到 session
            return jsonify({'success': True, 'redirect': '/project_management'})
        else:
            return jsonify({'error': '邮箱或密码错误'}), 401
            
    except Exception as e:
        app.logger.error(f"登录错误: {str(e)}")
        return jsonify({'error': '登录失败，请稍后重试'}), 500

@app.route('/api/register', methods=['POST'])
def register():
    """用户注册处理"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'message': '请求数据无效'}), 400
            
        email = data.get('email')
        password = data.get('password')
        invitation_code = data.get('invitation_code')
        
        app.logger.info(f"收到注册请求: email={email}")
        
        # 基本验证
        if not email or not password or not invitation_code:
            return jsonify({'message': '请输入邮箱、密码和邀请码'}), 400
            
        # 邮箱格式验证
        if not '@' in email or not '.' in email:
            return jsonify({'message': '请输入有效的邮箱地址'}), 400
            
        # 密码长度验证
        if len(password) < 8:
            return jsonify({'message': '密码长度不能少于8位'}), 400
            
        # 检查邮箱是否已存在
        if User.query.filter_by(email=email).first():
            return jsonify({'message': '该邮箱已被注册'}), 400

        # 验证邀请码
        invite = InvitationCode.query.filter_by(code=invitation_code).first()
        if not invite:
            return jsonify({'message': '无效的邀请码'}), 400
            
        if invite.usage_count >= invite.max_usage:
            return jsonify({'message': '邀请码已达到使用上限'}), 400

        try:
            # 创建新用户
            new_user = User(
                email=email,
                role='user'
            )
            new_user.set_password(password)
            db.session.add(new_user)
            db.session.flush()  # 获取新用户的ID
            
            # 更新邀请码使用情况
            invite.usage_count += 1
            invite.used_at = datetime.utcnow()
            invite.used_by = new_user.id
            
            db.session.commit()
            app.logger.info(f"用户注册成功: {email}")
            return jsonify({'success': True, 'message': '注册成功'}), 201
            
        except Exception as e:
            db.session.rollback()
            app.logger.error(f"数据库操作失败: {str(e)}")
            return jsonify({'message': '注册失败，请稍后重试'}), 500
            
    except Exception as e:
        app.logger.error(f"注册过程出错: {str(e)}")
        return jsonify({'message': '注册失败，请稍后重试'}), 500

@app.route('/api/check_user', methods=['POST'])
def check_user():
    try:
        data = request.get_json()
        username = data.get('username')
        
        app.logger.info(f"检查用户是否存在: username={username}")
        
        if not username:
            return jsonify({'error': '请输入用户名'}), 400
            
        user = User.query.filter_by(email=username).first()
        return jsonify({'exists': bool(user)})
        
    except Exception as e:
        app.logger.error(f"检查用户错误: {str(e)}")
        return jsonify({'error': '服务器错误'}), 500

@app.route('/api/reset_password', methods=['POST'])
def reset_password():
    try:
        data = request.get_json()
        username = data.get('username')
        password = data.get('password')
        
        app.logger.info(f"收到重置密码请求: username={username}")
        
        if not username or not password:
            return jsonify({'error': '请输入用户名和新密码'}), 400
            
        user = User.query.filter_by(email=username).first()
        if not user:
            return jsonify({'error': '用户不存在'}), 404
            
        # 更新密码
        try:
            user.set_password(password)
            db.session.commit()
            app.logger.info(f"密码重置成功: username={username}")
            return jsonify({'success': True})
        except Exception as db_error:
            db.session.rollback()
            app.logger.error(f"数据库操作失败: {str(db_error)}")
            return jsonify({'error': '重置密码失败，请稍后重试'}), 500
            
    except Exception as e:
        app.logger.error(f"重置密码错误: {str(e)}")
        return jsonify({'error': '重置密码失败，请稍后重试'}), 500

@app.route('/logout')
def logout():
    session.pop('user_id', None)
    return redirect(url_for('login_page'))

# 添加登录要求装饰器
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            return redirect(url_for('login_page'))
        return f(*args, **kwargs)
    return decorated_function

# 添加更新数据库表结构的函数
def update_database_structure():
    """更新数据库表结构，添加 user_id 字段和 role 字段"""
    try:
        # 检查 user_id 列是否存在
        result = db.session.execute(text("""
            SELECT COUNT(*)
            FROM INFORMATION_SCHEMA.COLUMNS
            WHERE TABLE_NAME = 'projects'
            AND COLUMN_NAME = 'user_id'
        """))
        
        if result.scalar() == 0:
            # 添加 user_id 列
            db.session.execute(text("""
                ALTER TABLE projects
                ADD user_id INT NOT NULL DEFAULT 1
            """))
            
            # 添加外键约束
            db.session.execute(text("""
                ALTER TABLE projects
                ADD CONSTRAINT FK_Projects_Users
                FOREIGN KEY (user_id) REFERENCES users (id)
            """))
            
            db.session.commit()
            print("成功添加 user_id 列和外键约束")
        
        # 检查 role 列是否存在
        result = db.session.execute(text("""
            SELECT COUNT(*)
            FROM INFORMATION_SCHEMA.COLUMNS
            WHERE TABLE_NAME = 'users'
            AND COLUMN_NAME = 'role'
        """))
        
        if result.scalar() == 0:
            # 添加 role 列
            db.session.execute(text("""
                ALTER TABLE users
                ADD COLUMN role VARCHAR(20) NOT NULL DEFAULT 'user'
            """))
            
            db.session.commit()
            print("成功添加 role 列")
            
        print("数据库表结构更新完成")
        
    except Exception as e:
        db.session.rollback()
        print(f"更新数据库表结构失败: {str(e)}")
        raise

# 在应用启动时更新数据库结构
with app.app_context():
    update_database_structure()

@app.route('/api/projects', methods=['GET'])
@login_required
def get_projects():
    try:
        # 获取当前用户ID
        user_id = session.get('user_id')
        
        if not user_id:
            return jsonify({'error': '用户未登录'}), 401
        
        # 获取该用户的所有项目
        projects = Project.query.filter_by(user_id=user_id).order_by(Project.created_at.desc()).all()
        
        # 将项目数据转换为JSON格式
        projects_data = []
        for project in projects:
            # 确保项目对象有效
            if project is None:
                continue
                
            # 使用to_dict方法获取项目数据
            project_data = project.to_dict()
            projects_data.append(project_data)
        
        return jsonify({
            'success': True,
            'projects': projects_data
        })
    except Exception as e:
        app.logger.error(f"获取项目列表时出错: {str(e)}")
        app.logger.error(traceback.format_exc())
        return jsonify({'error': '获取项目列表失败'}), 500

@app.route('/api/projects/<int:project_id>', methods=['DELETE'])
@login_required
def delete_project_api(project_id):
    try:
        app.logger.info(f"开始删除项目 {project_id}")
        
        # 获取当前用户ID和角色
        user_id = session.get('user_id')
        user_role = session.get('role')
        
        app.logger.info(f"当前用户ID: {user_id}, 角色: {user_role}")
        
        if not user_id:
            app.logger.warning("用户未登录")
            return jsonify({'error': '用户未登录'}), 401
        
        # 获取项目
        project = Project.query.get(project_id)
        if not project:
            app.logger.warning(f"项目 {project_id} 不存在")
            return jsonify({'error': '项目不存在'}), 404
            
        app.logger.info(f"找到项目: {project.name}, 所属用户ID: {project.user_id}")
            
        # 检查权限：管理员可以删除任何项目，普通用户只能删除自己的项目
        if user_role != 'admin' and project.user_id != user_id:
            app.logger.warning(f"用户 {user_id} 无权限删除项目 {project_id}")
            return jsonify({'error': '无权限删除此项目'}), 403
        
        try:
            # 删除项目相关的得分记录
            result = db.session.execute(
                text("DELETE FROM `得分表` WHERE `项目ID` = :project_id"),
                {"project_id": project_id}
            )
            app.logger.info(f"删除得分记录: {result.rowcount} 条")
            
            # 删除项目
            db.session.delete(project)
            db.session.commit()
            app.logger.info(f"项目 {project_id} 删除成功")
            
            return jsonify({
                'success': True,
                'message': '项目删除成功'
            })
            
        except Exception as db_error:
            db.session.rollback()
            app.logger.error(f"数据库操作失败: {str(db_error)}")
            app.logger.error(traceback.format_exc())
            return jsonify({'error': f'数据库操作失败: {str(db_error)}'}), 500
        
    except Exception as e:
        app.logger.error(f"删除项目失败: {str(e)}")
        app.logger.error(traceback.format_exc())
        return jsonify({'error': f'删除项目失败: {str(e)}'}), 500

@app.route('/api/save_star_case', methods=['POST'])
def save_star_case():
    """保存项目数据到星级案例表"""
    try:
        # 从URL参数或请求体中获取项目ID
        data = request.get_json() or {}
        project_id = request.args.get('project_id') or data.get('project_id')
        
        if not project_id:
            return jsonify({
                'success': False,
                'message': '请提供项目ID'
            }), 400

        # 获取项目信息
        project = Project.query.get(project_id)
        if not project:
            return jsonify({
                'success': False,
                'message': f'项目(ID={project_id})不存在'
            }), 404

        try:
            # 获取项目的得分数据
            query = """
                SELECT 条文号, 分类, 是否达标, 得分, 技术措施, 专业, 评价等级
                FROM 得分表
                WHERE 项目ID = :project_id
            """
            result = db.session.execute(text(query), {"project_id": project_id})
            score_data = result.fetchall()

            if not score_data:
                return jsonify({
                    'success': False,
                    'message': '未找到项目得分数据'
                }), 404

            # 获取当前最大序号
            result = db.session.execute(text("SELECT MAX(序号) FROM 星级案例"))
            max_seq = result.scalar() or 0

            # 插入数据到星级案例表
            inserted_count = 0
            for record in score_data:
                clause_number, category, is_achieved, score, technical_measures, specialty, level = record

                # 检查记录是否已存在
                check_query = """
                SELECT COUNT(*) FROM 星级案例
                WHERE 条文号 = :clause_number AND 评价标准 = :standard AND 星级目标 = :star_rating_target AND 建筑类型 = :building_type
                """
                result = db.session.execute(
                    text(check_query), 
                    {
                        "clause_number": clause_number,
                        "standard": project.standard,
                        "star_rating_target": project.star_rating_target,
                        "building_type": project.building_type
                    }
                )
                count = result.scalar()

                if count > 0:
                    # 更新现有记录
                    update_query = """
                    UPDATE 星级案例
                    SET 分类 = :category, 是否达标 = :is_achieved, 得分 = :score, 
                        技术措施 = :technical_measures, 专业 = :specialty, 评价等级 = :level, 
                        项目地点 = :location
                    WHERE 条文号 = :clause_number AND 评价标准 = :standard 
                        AND 星级目标 = :star_rating_target AND 建筑类型 = :building_type
                    """
                    db.session.execute(
                        text(update_query),
                        {
                            "category": category,
                            "is_achieved": is_achieved,
                            "score": score,
                            "technical_measures": technical_measures,
                            "specialty": specialty,
                            "level": level,
                            "location": project.location,
                            "clause_number": clause_number,
                            "standard": project.standard,
                            "star_rating_target": project.star_rating_target,
                            "building_type": project.building_type
                        }
                    )
                else:
                    # 插入新记录，使用自增序号
                    max_seq += 1
                    insert_query = """
                    INSERT INTO 星级案例 (
                        序号, 条文号, 分类, 是否达标, 得分, 技术措施, 专业, 评价等级,
                        评价标准, 星级目标, 建筑类型, 项目地点
                    ) VALUES (:seq, :clause_number, :category, :is_achieved, :score, 
                            :technical_measures, :specialty, :level, :standard, 
                            :star_rating_target, :building_type, :location)
                    """
                    db.session.execute(
                        text(insert_query),
                        {
                            "seq": max_seq,
                            "clause_number": clause_number,
                            "category": category,
                            "is_achieved": is_achieved,
                            "score": score,
                            "technical_measures": technical_measures,
                            "specialty": specialty,
                            "level": level,
                            "standard": project.standard,
                            "star_rating_target": project.star_rating_target,
                            "building_type": project.building_type,
                            "location": project.location
                        }
                    )

                inserted_count += 1

                # 每100条提交一次，避免事务过大
                if inserted_count % 100 == 0:
                    db.session.commit()
                    app.logger.info(f"已提交 {inserted_count} 条记录")

            # 提交事务
            db.session.commit()
            app.logger.info(f"事务提交成功，共处理 {inserted_count} 条记录")

            return jsonify({
                'success': True,
                'message': f'成功保存 {inserted_count} 条星级案例数据',
                'data': {
                    'standard': project.standard,
                    'star_rating_target': project.star_rating_target,
                    'building_type': project.building_type,
                    'location': project.location,
                    'processed_count': inserted_count
                }
            })

        except Exception as e:
            # 回滚事务
            db.session.rollback()
            app.logger.error(f"数据库操作失败: {str(e)}")
            app.logger.error(traceback.format_exc())
            return jsonify({
                'success': False,
                'message': f'数据库操作失败: {str(e)}'
            }), 500

    except Exception as e:
        app.logger.error(f"处理请求失败: {str(e)}")
        app.logger.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': f'处理请求失败: {str(e)}'
        }), 500

@app.route('/user_management')
@login_required
def user_management():
    # 检查当前用户是否为管理员
    if not session.get('role') == 'admin':
        flash('只有管理员可以访问用户管理页面')
        return redirect(url_for('index'))
    
    # 获取所有用户
    users = User.query.all()
    return render_template('user_management.html', users=users)

@app.route('/api/users', methods=['POST'])
@login_required
def create_user():
    # 检查权限
    if not session.get('role') == 'admin':
        return jsonify({'success': False, 'message': '只有管理员可以创建用户'}), 403
    
    data = request.get_json()
    
    # 验证必填字段
    if not all(key in data for key in ['email', 'password', 'role']):
        return jsonify({'success': False, 'message': '缺少必要字段'}), 400
    
    # 检查邮箱是否已存在
    if User.query.filter_by(email=data['email']).first():
        return jsonify({'success': False, 'message': '该邮箱已被注册'}), 400
    
    # 创建新用户
    try:
        new_user = User()
        new_user.email = data['email']
        new_user.set_password(data['password'])
        new_user.role = data['role']
        
        db.session.add(new_user)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': '用户创建成功',
            'user': {
                'id': new_user.id,
                'email': new_user.email,
                'role': new_user.role
            }
        })
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/users/<int:user_id>', methods=['GET'])
@login_required
def get_user(user_id):
    # 检查权限
    if not session.get('role') == 'admin':
        return jsonify({'success': False, 'message': '只有管理员可以查看用户信息'}), 403
    
    user = User.query.get_or_404(user_id)
    return jsonify({
        'id': user.id,
        'email': user.email,
        'role': user.role,
        'created_at': user.created_at.strftime('%Y-%m-%d %H:%M:%S')
    })

@app.route('/api/users/<int:user_id>', methods=['PUT'])
@login_required
def update_user(user_id):
    # 检查权限
    if not session.get('role') == 'admin':
        return jsonify({'success': False, 'message': '只有管理员可以更新用户信息'}), 403
    
    user = User.query.get_or_404(user_id)
    data = request.get_json()
    
    try:
        # 更新邮箱
        if 'email' in data and data['email'] != user.email:
            # 检查新邮箱是否已存在
            if User.query.filter_by(email=data['email']).first():
                return jsonify({'success': False, 'message': '该邮箱已被使用'}), 400
            user.email = data['email']
        
        # 更新密码（如果提供了新密码）
        if 'password' in data and data['password']:
            user.set_password(data['password'])
        
        # 更新角色
        if 'role' in data:
            user.role = data['role']
        
        db.session.commit()
        return jsonify({'success': True, 'message': '用户信息更新成功'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/users/<int:user_id>', methods=['DELETE'])
@login_required
def delete_user(user_id):
    # 检查权限
    if not session.get('role') == 'admin':
        return jsonify({'success': False, 'message': '只有管理员可以删除用户'}), 403
    
    user = User.query.get_or_404(user_id)
    
    # 防止删除自己
    if user.id == session.get('user_id'):
        return jsonify({'success': False, 'message': '不能删除当前登录的用户'}), 400
    
    try:
        db.session.delete(user)
        db.session.commit()
        return jsonify({'success': True, 'message': '用户删除成功'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/validate-invite-code', methods=['POST'])
def validate_invite_code():
    """验证邀请码"""
    try:
        data = request.get_json()
        code = data.get('code')
        
        if not code:
            return jsonify({
                'valid': False,
                'message': '请输入邀请码'
            }), 400
            
        # 检查邀请码是否有效
        invite = InvitationCode.query.filter_by(code=code).first()
        if not invite:
            return jsonify({
                'valid': False,
                'message': '无效的邀请码'
            }), 400
            
        if invite.usage_count >= invite.max_usage:
            return jsonify({
                'valid': False,
                'message': '邀请码已达到使用上限'
            }), 400
            
        return jsonify({
            'valid': True,
            'message': '邀请码验证成功'
        }), 200
            
    except Exception as e:
        app.logger.error(f"邀请码验证失败: {str(e)}")
        return jsonify({
            'valid': False,
            'message': '验证失败，请稍后重试'
        }), 500

# 获取评分汇总数据的函数
def get_score_summary(project_id, force_refresh=False):
    """获取评分汇总数据的函数"""
    try:
        # 获取项目信息，确定评价标准
        project = get_project(project_id)
        project_standard = project.standard if project and project.standard else '成都市标'
        
        # 只在强制刷新或开发环境下记录详细日志
        if force_refresh or app.debug:
            app.logger.info(f"获取评分汇总数据: 项目ID={project_id}, 评价标准={project_standard}, 强制刷新={force_refresh}")
        
        # 构建缓存键
        cache_key = f"score_summary_{project_id}_{project_standard}"
        
        # 标记是否需要更新项目表
        need_update_project = force_refresh
        summary_data = None
        
        # 如果强制刷新或缓存中没有数据，则重新计算
        if force_refresh or not cache.get(cache_key):
            # 专业名称映射，用于处理数据库中的专业名称与预定义专业名称的不完全匹配
            specialty_mapping = {
                '建筑': '建筑专业',
                '结构': '结构专业',
                '给排水': '给排水专业',
                '电气': '电气专业',
                '暖通': '暖通专业',
                '景观': '景观专业',
                '环境健康与节能': '环境健康与节能专业',
                '建筑专业': '建筑专业',
                '结构专业': '结构专业',
                '给排水专业': '给排水专业',
                '电气专业': '电气专业',
                '暖通空调专业': '暖通专业',
                '暖通专业': '暖通专业',
                '景观专业': '景观专业'
            }
            
            # 获取所有专业的得分数据
            specialties = ['建筑专业', '结构专业', '给排水专业', '电气专业', '暖通专业', '景观专业']
            
            # 如果是四川省标，添加环境健康与节能专业
            if project_standard == '四川省标':
                specialties.append('环境健康与节能专业')
            
            # 存储各专业得分
            specialty_scores = {specialty: 0 for specialty in specialties}
            # 存储各专业按分类的得分
            specialty_scores_by_category = {}
            
            # 初始化各专业的分类得分
            for specialty in specialties:
                specialty_scores_by_category[specialty] = {
                    '安全耐久': 0,
                    '健康舒适': 0,
                    '生活便利': 0,
                    '资源节约': 0,
                    '环境宜居': 0,
                    '提高与创新': 0,
                    '总分': 0
                }
            
            try:
                # 优化SQL查询：一次获取所有需要的数据
                sql_query = """
                SELECT `专业`, `分类`, `是否达标`, `得分`, `评价等级`
                FROM `得分表`
                WHERE `项目ID` = :project_id AND `评价标准` = :standard
                """
                
                # 执行查询并获取所有结果
                result = db.session.execute(
                    text(sql_query), 
                    {"project_id": project_id, "standard": project_standard}
                )
                rows = result.fetchall()
                
                # 减少日志输出，只记录结果行数
                if app.debug:
                    app.logger.info(f"查询结果: {len(rows)} 行")
                
                # 处理查询结果
                for row in rows:
                    specialty = row[0]
                    category = row[1]
                    is_achieved = row[2]
                    score = row[3]
                    level = row[4]
                    
                    # 减少日志记录，提高性能
                    # 映射专业名称
                    mapped_specialty = specialty_mapping.get(specialty)
                    if not mapped_specialty:
                        # 如果没有精确匹配，尝试部分匹配
                        for key, value in specialty_mapping.items():
                            if key in specialty or specialty in key:
                                mapped_specialty = value
                                break
                    
                    # 如果仍然没有匹配，跳过
                    if not mapped_specialty:
                        if app.debug:
                            app.logger.warning(f"未能映射专业名称: {specialty}")
                        continue
                    
                    # 处理是否达标字段
                    is_achieved_value = is_achieved.lower() if isinstance(is_achieved, str) else str(is_achieved).lower()
                    is_achieved_flag = is_achieved_value in ['是', 'yes', 'true', '1', 'y']
                    
                    # 处理得分字段
                    score_value = 0
                    if score is not None:
                        try:
                            if isinstance(score, (int, float)):
                                score_value = float(score)
                            elif isinstance(score, str) and score.strip():
                                score_value = float(score)
                        except (ValueError, TypeError):
                            # 简化错误处理，不记录详细日志
                            pass
                    
                    # 基本级条文必须达标才计分，提高级条文有得分就计分
                    if (level == '基本级' and is_achieved_flag) or (level == '提高级' and score_value > 0):
                        specialty_scores[mapped_specialty] += score_value
                        
                        # 按分类累加得分
                        if category in specialty_scores_by_category[mapped_specialty]:
                            specialty_scores_by_category[mapped_specialty][category] += score_value
                
                # 计算各专业的总分
                for specialty in specialties:
                    category_scores = specialty_scores_by_category[specialty]
                    total_score = sum(score for category, score in category_scores.items() if category != '总分')
                    category_scores['总分'] = total_score
                    specialty_scores[specialty] = total_score  # 更新专业总分
                
                # 计算总分（所有专业分数之和）
                total_score = sum(specialty_scores.values())
                
                # 根据总分确定评定结果
                if total_score >= 85:
                    evaluation_result = '三星级绿色建筑'
                elif total_score >= 70:
                    evaluation_result = '二星级绿色建筑'
                elif total_score >= 55:
                    evaluation_result = '一星级绿色建筑'
                else:
                    evaluation_result = '未达标'
                
                # 构建汇总数据
                summary_data = {
                    'specialty_scores': specialty_scores,
                    'specialty_scores_by_category': specialty_scores_by_category,
                    'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                    'total_score': total_score,
                    'project_standard': project_standard,
                    'evaluation_result': evaluation_result
                }
                
                # 缓存结果 - 设置较长的过期时间（8小时）
                cache.set(cache_key, summary_data, timeout=28800)
                
                # 一定需要更新项目表
                need_update_project = True
                
            except Exception as e:
                app.logger.error(f"查询得分表失败: {str(e)}")
                if app.debug:
                    app.logger.error(traceback.format_exc())
                need_update_project = False
        else:
            # 从缓存中获取数据
            if app.debug:
                app.logger.info(f"从缓存中获取评分汇总数据: {cache_key}")
            summary_data = cache.get(cache_key)
            
        # 如果需要更新项目表，并且有有效的汇总数据
        if need_update_project and summary_data and project_id:
            try:
                # 高效更新项目表，不阻塞API响应
                update_project_scores_efficient(project_id, summary_data)
            except Exception as e:
                # 仅记录错误，不影响API响应
                app.logger.error(f"更新项目表评分数据时出错: {str(e)}")
                
        return summary_data
    
    except Exception as e:
        app.logger.error(f"获取评分汇总数据失败: {str(e)}")
        if app.debug:
            app.logger.error(traceback.format_exc())
        # 返回一个基本的数据结构，而不是隐式返回None
        return {
            'specialty_scores': {},
            'specialty_scores_by_category': {},
            'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'total_score': 0,
            'project_standard': '未知',
            'evaluation_result': '未评定'
        }

def update_project_scores_efficient(project_id, scores):
    """高效更新项目表中的评分数据，不使用with_for_update以避免锁定延迟"""
    try:
        # 使用单独的会话避免锁定主会话
        project = Project.query.get(project_id)
        if not project:
            if app.debug:
                app.logger.warning(f"找不到要更新评分的项目: ID={project_id}")
            return False
            
        # 更新专业评分
        if 'specialty_scores' in scores:
            specialty_scores = scores['specialty_scores']
            project.architecture_score = specialty_scores.get('建筑专业', 0)
            project.structure_score = specialty_scores.get('结构专业', 0)
            project.water_supply_score = specialty_scores.get('给排水专业', 0)
            project.electrical_score = specialty_scores.get('电气专业', 0)
            project.hvac_score = specialty_scores.get('暖通专业', 0)
            project.landscape_score = specialty_scores.get('景观专业', 0)
            project.env_health_energy_score = specialty_scores.get('环境健康与节能专业', 0)
            
        # 更新章节评分
        if 'specialty_scores_by_category' in scores:
            categories = ['安全耐久', '健康舒适', '生活便利', '资源节约', '环境宜居', '提高与创新']
            category_scores = {cat: 0 for cat in categories}
            specialty_count = 0
            
            for specialty_data in scores['specialty_scores_by_category'].values():
                specialty_count += 1
                for category in categories:
                    if category in specialty_data:
                        category_scores[category] += float(specialty_data[category])
            
            if specialty_count > 0:
                project.safety_durability_score = category_scores['安全耐久'] 
                project.health_comfort_score = category_scores['健康舒适'] 
                project.life_convenience_score = category_scores['生活便利']
                project.resource_saving_score = category_scores['资源节约'] 
                project.environment_livability_score = category_scores['环境宜居']
                project.improvement_innovation_score = category_scores['提高与创新'] 
            
        # 更新总分
        project.total_score = scores.get('total_score', 0)
        
        # 设置评定结果
        project.evaluation_result = scores.get('evaluation_result', '未评定')
        
        # 保存更改
        db.session.commit()
        
        if app.debug:
            app.logger.debug(f"成功更新项目评分: 项目ID={project_id}, 总分={project.total_score}")
        
        return True
        
    except Exception as e:
        db.session.rollback()
        if app.debug:
            app.logger.error(f"更新项目评分失败: {str(e)}")
        return False

# 数据库连接配置
def get_db_connection():
    try:
        # 解析数据库连接字符串
        db_url = os.getenv('DATABASE_URL')
        if not db_url:
            app.logger.error("数据库连接字符串未配置")
            raise ValueError("数据库连接字符串未配置")

        if IS_WSL:
            # WSL环境下使用MySQL连接
            # 使用正则表达式解析连接字符串
            import re
            pattern = r'mysql\+pymysql://([^:]+):([^@]+)@([^/]+)/([^?]+)'
            match = re.match(pattern, db_url)
            if not match:
                app.logger.error("数据库连接字符串格式错误")
                raise ValueError("数据库连接字符串格式错误")

            username, password, server, database = match.groups()
            
            try:
                conn = pymysql.connect(
                    host=server,
                    user=username,
                    password=password,
                    database=database,
                    charset='utf8mb4'
                )
                app.logger.info("MySQL数据库连接成功")
                return conn
            except Exception as e:
                app.logger.error(f"MySQL数据库连接失败: {str(e)}")
                raise
        else:
            # Windows环境下使用SQL Server连接
            # 使用正则表达式解析连接字符串
            import re
            pattern = r'mssql\+pyodbc://([^:]+):([^@]+)@([^/]+)/([^?]+)'
            match = re.match(pattern, db_url)
            if not match:
                app.logger.error("数据库连接字符串格式错误")
                raise ValueError("数据库连接字符串格式错误")

            username, password, server, database = match.groups()
            database = database.split('?')[0]  # 移除查询参数

            # 构建连接字符串，添加超时和TLS配置
            conn_str = f"DRIVER={{ODBC Driver 17 for SQL Server}};SERVER={server};DATABASE={database};UID={username};PWD={password};Connection Timeout=30;Encrypt=yes;TrustServerCertificate=yes"
            # 尝试建立连接
            try:
                conn = pyodbc.connect(conn_str)
                app.logger.info("SQL Server数据库连接成功")
                return conn
            except Exception as e:
                app.logger.error(f"SQL Server数据库连接失败: {str(e)}")
                raise
    except Exception as e:
        app.logger.error(f"创建数据库连接时出错: {str(e)}")
        raise

@app.route('/user_guide')
@login_required
def user_guide():
    """用户使用指南页面"""
    try:
        return render_template('user_guide.html')
    except Exception as e:
        app.logger.error(f"访问用户指南页面出错: {str(e)}")
        return render_template('error.html', error=str(e))

@app.route('/api/projects/<int:project_id>/update_status', methods=['PUT'])
def update_project_status(project_id):
    """更新项目状态
    
    请求参数:
    {
        "status": "进行中"  // 新的项目状态
    }
    
    响应:
    {
        "success": true,
        "message": "项目状态更新成功",
        "data": {
            "id": 1,
            "name": "项目名称",
            "status": "进行中"
        }
    }
    """
    try:
        # 检查用户是否登录
        if 'user_id' not in session:
            app.logger.warning("未登录用户尝试更新项目状态")
            return jsonify({'success': False, 'message': '请先登录'}), 401
            
        # 获取用户ID和角色
        user_id = session.get('user_id')
        user_role = session.get('role', 'user')
        
        # 获取请求数据
        data = request.get_json()
        if not data or 'status' not in data:
            return jsonify({'success': False, 'message': '请提供新的项目状态'}), 400
            
        new_status = data['status']
        
        # 获取项目
        project = Project.query.get(project_id)
        if not project:
            app.logger.warning(f"项目 {project_id} 不存在")
            return jsonify({'success': False, 'message': '项目不存在'}), 404
            
        app.logger.info(f"找到项目: {project.name}, 所属用户ID: {project.user_id}")
            
        # 检查权限：管理员可以更新任何项目，普通用户只能更新自己的项目
        if user_role != 'admin' and project.user_id != user_id:
            app.logger.warning(f"用户 {user_id} 无权限更新项目 {project_id}")
            return jsonify({'success': False, 'message': '无权限更新此项目'}), 403
        
        # 更新项目状态
        project.status = new_status
        db.session.commit()
        app.logger.info(f"项目 {project_id} 状态更新为 {new_status}")
        
        return jsonify({
            'success': True,
            'message': '项目状态更新成功',
            'data': {
                'id': project.id,
                'name': project.name,
                'status': project.status
            }
        })
        
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"更新项目状态失败: {str(e)}")
        app.logger.error(traceback.format_exc())
        return jsonify({'success': False, 'message': f'更新项目状态失败: {str(e)}'}), 500

@app.route('/api/update_project_scores', methods=['POST'])
def update_project_scores():
    """将评分数据更新到projects表中"""
    try:
        # 获取请求数据
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'message': '未接收到数据'}), 400
            
        # 提取必要字段
        project_id = data.get('project_id')
        scores = data.get('scores')
        
        if not project_id or not scores:
            return jsonify({'success': False, 'message': '缺少必要的字段'}), 400
            
        # 获取项目记录 - 使用带锁的查询避免并发问题
        project = db.session.query(Project).with_for_update().get(project_id)
        if not project:
            return jsonify({'success': False, 'message': f'找不到ID为{project_id}的项目'}), 404
            
        # 更新项目的评分字段
        try:
            # 更新专业评分 - 只在debug模式下记录日志
            if 'specialty_scores' in scores:
                specialty_scores = scores['specialty_scores']
                if app.debug:
                    app.logger.debug(f"更新项目专业评分: 项目ID={project_id}")
                # 批量更新所有字段
                project.architecture_score = specialty_scores.get('建筑专业', 0)
                project.structure_score = specialty_scores.get('结构专业', 0)
                project.water_supply_score = specialty_scores.get('给排水专业', 0)
                project.electrical_score = specialty_scores.get('电气专业', 0)
                project.hvac_score = specialty_scores.get('暖通专业', 0)
                project.landscape_score = specialty_scores.get('景观专业', 0)
                project.env_health_energy_score = specialty_scores.get('环境健康与节能专业', 0)
                
            # 更新章节评分
            if 'specialty_scores_by_category' in scores:
                # 使用所有专业的平均分作为章节分数
                categories = ['安全耐久', '健康舒适', '生活便利', '资源节约', '环境宜居', '提高与创新']
                category_scores = {cat: 0 for cat in categories}
                specialty_count = 0
                
                # 计算所有专业的章节平均分
                for specialty_data in scores['specialty_scores_by_category'].values():
                    specialty_count += 1
                    for category in categories:
                        if category in specialty_data:
                            category_scores[category] += float(specialty_data[category])
                
                # 计算平均分并更新
                if specialty_count > 0:
                    if app.debug:
                        app.logger.debug(f"更新项目章节评分: 项目ID={project_id}, 专业数={specialty_count}")
                    
                    project.safety_durability_score = category_scores['安全耐久'] / specialty_count
                    project.health_comfort_score = category_scores['健康舒适'] / specialty_count
                    project.life_convenience_score = category_scores['生活便利'] / specialty_count
                    project.resource_saving_score = category_scores['资源节约'] / specialty_count
                    project.environment_livability_score = category_scores['环境宜居'] / specialty_count
                    project.improvement_innovation_score = category_scores['提高与创新'] / specialty_count
                
            # 更新总分
            project.total_score = scores.get('total_score', 0)
            
            # 根据总分确定评定结果
            if project.total_score >= 85:
                project.evaluation_result = '三星级绿色建筑'
            elif project.total_score >= 70:
                project.evaluation_result = '二星级绿色建筑'
            elif project.total_score >= 55:
                project.evaluation_result = '一星级绿色建筑'
            else:
                project.evaluation_result = '未达标'
            
            # 保存更改
            db.session.commit()
            
            # 只在成功时记录日志，减少日志量
            if app.debug:
                app.logger.info(f"成功更新项目评分: 项目ID={project_id}, 总分={project.total_score}")
            
            return jsonify({
                'success': True, 
                'message': '成功更新项目评分',
                'project_id': project_id,
                'total_score': project.total_score,
                'evaluation_result': project.evaluation_result
            })
            
        except Exception as e:
            db.session.rollback()
            app.logger.error(f"更新项目评分时出错: {str(e)}")
            return jsonify({'success': False, 'message': f'更新评分失败: {str(e)}'}), 500
            
    except Exception as e:
        app.logger.error(f"处理更新项目评分请求时出错: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500

# 注册蓝图
app.register_blueprint(auth_bp)
app.register_blueprint(admin_app, url_prefix='/admin')  # 使用admin_app，添加/admin前缀

@app.route('/api/extract_word_info', methods=['POST'])
def extract_word_info():
    """处理Word文档并提取项目信息"""
    try:
        # 检查是否有文件上传
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': '未找到上传的文件'}), 400
            
        file = request.files['file']
        if file.filename == '':
            return jsonify({'success': False, 'error': '未选择文件'}), 400
            
        # 检查文件扩展名
        if not file.filename.endswith(('.doc', '.docx')):
            return jsonify({'success': False, 'error': '仅支持.doc和.docx格式的Word文档'}), 400
            
        # 创建临时目录（如果不存在）
        temp_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'temp')
        os.makedirs(temp_dir, exist_ok=True)
        
        # 保存上传的文件
        file_path = os.path.join(temp_dir, werkzeug.utils.secure_filename(file.filename))
        file.save(file_path)
        
        # 提取项目信息
        try:
            project_info = extract_project_info(file_path)
            
            # 删除临时文件
            try:
                os.remove(file_path)
            except Exception as e:
                app.logger.warning(f"删除临时文件失败: {str(e)}")
            
            if project_info:
                # 确保地下层数默认为0
                if not project_info.get("地下层数"):
                    project_info["地下层数"] = "0"
                    
                return jsonify({
                    'success': True,
                    'info': project_info
                })
            else:
                return jsonify({'success': False, 'error': '无法从文档中提取项目信息'}), 500
                
        except Exception as e:
            app.logger.error(f"提取项目信息时出错: {str(e)}")
            # 尝试删除临时文件
            try:
                os.remove(file_path)
            except:
                pass
            return jsonify({'success': False, 'error': f'提取项目信息失败: {str(e)}'}), 500
            
    except Exception as e:
        app.logger.error(f"处理Word文档提取请求时出错: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/extract_project_info', methods=['POST'])
def extract_project_info_api():
    """提取Word文档中的项目信息并返回"""
    try:
        app.logger.info("收到提取项目信息请求")
        app.logger.info(f"请求表单数据: {request.form.keys()}")
        app.logger.info(f"请求文件: {request.files.keys()}")
        
        # 检查是否有文件上传
        if 'word_file' not in request.files and 'image_file' not in request.files and 'file' not in request.files:
            app.logger.error("未找到上传的文件")
            return jsonify({'success': False, 'message': '未找到上传的文件'}), 400
        
        if 'word_file' in request.files:
            # Word文件处理逻辑保持不变...
            file = request.files['word_file']
            if file.filename == '':
                app.logger.error("未选择Word文件")
                return jsonify({'success': False, 'message': '未选择文件'}), 400
                
            # 检查文件扩展名
            if not file.filename.endswith(('.doc', '.docx')):
                app.logger.error(f"不支持的文件格式: {file.filename}")
                return jsonify({'success': False, 'message': '仅支持.doc和.docx格式的Word文档'}), 400
                
            # 创建临时目录（如果不存在）
            temp_dir = os.path.join('static', 'temp')
            os.makedirs(temp_dir, exist_ok=True)
            
            # 保存上传的文件
            timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
            file_path = os.path.join(temp_dir, f"{timestamp}_{werkzeug.utils.secure_filename(file.filename)}")
            file.save(file_path)
            
            app.logger.info(f"已保存Word文件: {file_path}")
            
            # 调用extract_doc_info函数提取信息
            from utils.word_extractor import extract_doc_info
            project_info = extract_doc_info(file_path)
            
            # 文件使用完毕后删除
            try:
                os.remove(file_path)
                app.logger.info(f"已删除临时文件: {file_path}")
            except Exception as e:
                app.logger.warning(f"删除临时文件失败: {str(e)}")
        
        elif 'image_file' in request.files or 'file' in request.files:
            # 支持 'image_file' 或 'file' 参数名
            file = request.files.get('image_file') or request.files.get('file')
            app.logger.info(f"处理图片文件: {file.filename if file else 'None'}")
            
            if file.filename == '':
                app.logger.error("未选择图片文件")
                return jsonify({'success': False, 'message': '未选择文件'}), 400
                
            # 检查文件扩展名
            if not file.filename.lower().endswith(('.jpg', '.jpeg', '.png', '.bmp', '.tiff')):
                app.logger.error(f"不支持的图片格式: {file.filename}")
                return jsonify({'success': False, 'message': '仅支持JPG、PNG、BMP和TIFF格式的图片'}), 400
            
            # 使用图像提取函数处理图片
            try:
                from utils.image_extractor import extract_image_info_with_raw_text
                app.logger.info("调用图像提取函数...")
                result = extract_image_info_with_raw_text(file)
                app.logger.info(f"图像提取函数返回结果: {result.keys() if result else 'None'}")
                
                # 解析结果
                if result and 'raw_text' in result:
                    raw_text = result['raw_text']
                    project_info = result['project_info']
                    app.logger.info(f"成功提取到文本，长度: {len(raw_text)}")
                    app.logger.info(f"提取到项目信息: {len(project_info)} 项")
                else:
                    app.logger.warning("图像提取函数返回无效结果，没有raw_text字段")
                    raw_text = ""
                    project_info = {}
                
                # 无论是否成功提取项目信息，都记录原始文本
                if raw_text:
                    app.logger.info("OCR提取的原始文本长度: " + str(len(raw_text)))
                    app.logger.info("OCR文本前200字符: " + raw_text[:200].replace('\n', ' '))
                    # 将完整文本记录到文件中以便分析
                    try:
                        log_dir = 'logs'
                        os.makedirs(log_dir, exist_ok=True)
                        timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
                        log_file = os.path.join(log_dir, f"ocr_text_{timestamp}.txt")
                        with open(log_file, 'w', encoding='utf-8') as f:
                            f.write(raw_text)
                        app.logger.info(f"OCR完整文本已保存到: {log_file}")
                    except Exception as e:
                        app.logger.warning(f"保存OCR文本到文件失败: {str(e)}")
                else:
                    app.logger.warning("未能提取到任何文本")
                    
                if not project_info:
                    app.logger.warning(f"未能从图片中提取到有效信息")
                    return jsonify({
                        'success': False, 
                        'message': '未能从图片中识别到项目信息，请尝试使用更清晰的图片或确保图片中包含项目相关文字'
                    }), 400
                    
                # 记录成功提取的信息
                app.logger.info(f"成功从图片提取到信息: {len(project_info)} 项")
                app.logger.debug(f"提取的信息内容: {project_info}")
                
            except Exception as e:
                error_msg = str(e)
                app.logger.error(f"图片处理出现错误: {error_msg}")
                app.logger.error(traceback.format_exc())
                return jsonify({'success': False, 'message': f'处理图片时出错: {error_msg}'}), 500
        
        if project_info:
            app.logger.info(f"成功提取项目信息: {project_info}")
            return jsonify({'success': True, 'info': project_info})
        else:
            app.logger.error("提取项目信息失败")
            return jsonify({'success': False, 'message': '无法从文件中提取项目信息'}), 400
            
    except Exception as e:
        app.logger.error(f"处理文件时出错: {str(e)}")
        return jsonify({'success': False, 'message': f'处理请求失败: {str(e)}'}), 500

# 添加获取真实海拔数据的API端点
@app.route('/api/elevation', methods=['GET'])
def get_elevation():
    try:
        lat = request.args.get('lat')
        lng = request.args.get('lng')
        
        if not lat or not lng:
            return jsonify({'error': '缺少经纬度参数'}), 400
        
        # 调用Open-Elevation API获取真实海拔数据
        import requests
        
        # 使用Open-Elevation公共API
        url = f"https://api.open-elevation.com/api/v1/lookup?locations={lat},{lng}"
        
        app.logger.info(f"请求Open-Elevation API: {url}")
        
        response = requests.get(url, timeout=5)
        if response.status_code == 200:
            data = response.json()
            if data and 'results' in data and len(data['results']) > 0:
                elevation = data['results'][0]['elevation']
                app.logger.info(f"获取到海拔数据: {elevation}米")
                return jsonify({'elevation': elevation})
        
        # 如果Open-Elevation失败，尝试其他免费海拔API
        backup_url = f"https://elevation-api.io/api/elevation?points=({lat},{lng})"
        app.logger.info(f"请求备用海拔API: {backup_url}")
        
        backup_response = requests.get(backup_url, timeout=5)
        if backup_response.status_code == 200:
            backup_data = backup_response.json()
            if backup_data and 'elevations' in backup_data and len(backup_data['elevations']) > 0:
                elevation = backup_data['elevations'][0]['elevation']
                app.logger.info(f"从备用API获取到海拔数据: {elevation}米")
                return jsonify({'elevation': elevation})
        
        # 如果所有API都失败，使用SRTM数据集（如果可用）
        try:
            from srtm import get_data
            srtm_data = get_data()
            elevation = srtm_data.get_elevation(float(lat), float(lng))
            if elevation is not None:
                app.logger.info(f"从SRTM数据集获取到海拔数据: {elevation}米")
                return jsonify({'elevation': elevation})
        except Exception as srtm_error:
            app.logger.error(f"SRTM数据获取失败: {str(srtm_error)}")
        
        # 所有方法都失败
        return jsonify({'error': '无法获取海拔数据'}), 500
    except Exception as e:
        app.logger.error(f"获取海拔数据时出错: {str(e)}")
        app.logger.error(traceback.format_exc())
        return jsonify({'error': f'服务器错误: {str(e)}'}), 500

@app.route('/api/google_elevation', methods=['GET'])
def get_google_elevation():
    try:
        lat = request.args.get('lat')
        lng = request.args.get('lng')
        
        if not lat or not lng:
            return jsonify({'error': '缺少经纬度参数'}), 400
        
        # 从环境变量获取Google Maps API密钥
        api_key = os.environ.get('GOOGLE_MAPS_API_KEY')
        
        if not api_key:
            return jsonify({'error': '未配置Google Maps API密钥'}), 500
        
        # 调用Google Maps Elevation API
        import requests
        
        url = f"https://maps.googleapis.com/maps/api/elevation/json?locations={lat},{lng}&key={api_key}"
        
        app.logger.info(f"请求Google Elevation API")
        
        response = requests.get(url, timeout=5)
        if response.status_code == 200:
            data = response.json()
            if data['status'] == 'OK' and len(data['results']) > 0:
                elevation = data['results'][0]['elevation']
                app.logger.info(f"从Google API获取到海拔数据: {elevation}米")
                return jsonify({'elevation': elevation})
        
        return jsonify({'error': '无法从Google获取海拔数据'}), 500
    except Exception as e:
        app.logger.error(f"获取Google海拔数据时出错: {str(e)}")
        app.logger.error(traceback.format_exc())
        return jsonify({'error': f'服务器错误: {str(e)}'}), 500

# 添加获取百度地图API密钥的API端点
@app.route('/api/map_api_key', methods=['GET'])
def get_map_api_key():
    try:
        # 从环境变量获取百度地图API密钥
        api_key = os.environ.get('BAIDU_MAP_API_KEY', 'J6UW18n9sxCMtrxTkjpLE3JkU8pfw3bL')  # 使用环境变量中的密钥，提供一个默认值
        
        # 返回API密钥
        return jsonify({'api_key': api_key})
    except Exception as e:
        app.logger.error(f"获取地图API密钥时出错: {str(e)}")
        return jsonify({'error': f'服务器错误: {str(e)}'}), 500

# 添加百度地图API代理
@app.route('/api/map_proxy', methods=['GET', 'POST'])
def map_api_proxy():
    try:
        app.logger.info(f"百度地图API代理请求: {request.url}")
        
        # 获取百度地图API密钥
        api_key = os.environ.get('BAIDU_MAP_API_KEY', 'J6UW18n9sxCMtrxTkjpLE3JkU8pfw3bL')
        
        # 获取所有请求参数
        params = request.args.to_dict()
        
        # 获取服务路径
        service_path = params.pop('service_path', 'api')
        
        # 判断是否是静态资源请求（图片等）
        is_static_resource = service_path.endswith(('.png', '.jpg', '.jpeg', '.gif', '.css', '.js'))
        
        # 如果不是静态资源请求，则添加API密钥
        if not is_static_resource:
            params['ak'] = api_key
        
        # 构建百度地图API请求URL
        base_url = f"https://api.map.baidu.com/{service_path}"
        app.logger.info(f"代理到URL: {base_url}, 参数: {params}")
        
        # 发送请求
        import requests
        
        if request.method == 'GET':
            response = requests.get(base_url, params=params, timeout=10)
        else:  # POST请求
            post_data = request.get_data()
            headers = {'Content-Type': request.headers.get('Content-Type', 'application/x-www-form-urlencoded')}
            response = requests.post(base_url, params=params, data=post_data, headers=headers, timeout=10)
        
        # 根据不同的资源类型设置不同的Content-Type
        content_type = response.headers.get('Content-Type')
        if is_static_resource:
            if service_path.endswith('.png'):
                content_type = 'image/png'
            elif service_path.endswith('.jpg') or service_path.endswith('.jpeg'):
                content_type = 'image/jpeg'
            elif service_path.endswith('.gif'):
                content_type = 'image/gif'
            elif service_path.endswith('.css'):
                content_type = 'text/css'
            elif service_path.endswith('.js'):
                content_type = 'application/javascript'
        
        app.logger.info(f"代理响应状态码: {response.status_code}, Content-Type: {content_type}")
        
        # 返回百度地图API的响应
        return Response(
            response.content, 
            status=response.status_code,
            content_type=content_type or 'application/json'
        )
    except Exception as e:
        app.logger.error(f"百度地图API代理出错: {str(e)}")
        app.logger.error(traceback.format_exc())
        return jsonify({'error': f'代理请求失败: {str(e)}'}), 500

# 添加百度地图JavaScript API代理
@app.route('/api/map_js_api', methods=['GET'])
def map_js_api_proxy():
    try:
        # 获取百度地图API密钥
        api_key = os.environ.get('BAIDU_MAP_API_KEY', 'J6UW18n9sxCMtrxTkjpLE3JkU8pfw3bL')
        
        # 获取请求参数
        params = request.args.to_dict()
        
        # 添加API密钥
        params['ak'] = api_key
        
        # 构建百度地图JavaScript API的URL
        url = "https://api.map.baidu.com/api"
        
        # 发送请求
        import requests
        response = requests.get(url, params=params, timeout=10)
        
        # 返回JavaScript内容
        return Response(
            response.content, 
            status=response.status_code,
            content_type='application/javascript'
        )
    except Exception as e:
        app.logger.error(f"百度地图JavaScript API代理出错: {str(e)}")
        return Response(
            f"console.error('加载地图API失败: {str(e)}');", 
            status=500,
            content_type='application/javascript'
        )

# 添加项目信息页面路由
@app.route('/project_info_page')
def project_info_page():
    try:
        # 获取项目ID参数
        project_id = request.args.get('project_id')
        
        if project_id:
            # 根据项目ID获取项目信息
            project = Project.query.filter_by(id=project_id).first()
            if not project:
                app.logger.warning(f"未找到项目ID: {project_id}")
                # 如果项目不存在，返回空项目信息
                return render_template('project_info.html', project=None)
            
            app.logger.info(f"访问项目信息页面: 项目ID={project_id}, 名称={project.name}")
            return render_template('project_info.html', project=project)
        else:
            app.logger.warning("访问项目信息页面，但未提供项目ID")
            return render_template('project_info.html', project=None)
    except Exception as e:
        app.logger.error(f"访问项目信息页面出错: {str(e)}")
        app.logger.error(traceback.format_exc())
        return render_template('error.html', error=f"获取项目信息失败: {str(e)}")

# 添加一个新的路由处理评分汇总页面的请求
@app.route('/score_summary_page/<project_id>')
def score_summary_page(project_id):
    # 查询项目信息
    project = Project.query.filter_by(id=project_id).first()
    
    # 将项目信息传递给模板
    return render_template('score_summary.html', project=project, session=session)

@app.route('/public_transport_analysis')
@login_required
def public_transport_analysis():
    """公共交通分析报告生成页面"""
    # 获取可能的项目ID参数
    project_id = request.args.get('project_id')
    project = None
    
    if project_id:
        try:
            project = Project.query.get(project_id)
        except Exception as e:
            app.logger.error(f"获取项目信息失败: {str(e)}")
            
    return render_template('public_transport_analysis.html', project=project)

@app.route('/api/generate_transport_report', methods=['POST'])
@login_required
def generate_transport_report():
    """生成公共交通分析报告"""
    try:
        # 获取请求数据
        data = request.json
        if not data:
            return jsonify({'success': False, 'message': '缺少必要的数据'}), 400
        
        # 提取项目信息
        project_info = data.get('project', {})
        project_name = project_info.get('name', '未命名项目')
        project_location = project_info.get('location', '')
        project_coordinates = project_info.get('coordinates', {})
        
        # 提取分析结果
        analysis_results = data.get('analysis', {})
        stations = data.get('stations', [])
        
        # 记录处理信息
        app.logger.info(f"开始生成公共交通分析报告: 项目 '{project_name}'")
        
        # 准备替换模板中的数据
        template_data = [{
            # 项目基本信息
            '项目名称': project_name,
            '项目地点': project_location,
            '经度': project_coordinates.get('lng', ''),
            '纬度': project_coordinates.get('lat', ''),
            '分析日期': datetime.now().strftime('%Y年%m月%d日'),
            
            # 分析结果数据
            '公交站总数': str(analysis_results.get('busStations', {}).get('total', 0)),
            '地铁站总数': str(analysis_results.get('subwayStations', {}).get('total', 0)),
            '500米内公交站': str(analysis_results.get('busStations', {}).get('qualified', 0)),
            '800米内地铁站': str(analysis_results.get('subwayStations', {}).get('qualified', 0)),
            
            # 最近站点信息
            '最近公交站名称': analysis_results.get('busStations', {}).get('nearest', {}).get('name', '无'),
            '最近公交站距离': str(analysis_results.get('busStations', {}).get('nearest', {}).get('distance', 0)) + '米',
            '最近地铁站名称': analysis_results.get('subwayStations', {}).get('nearest', {}).get('name', '无'),
            '最近地铁站距离': str(analysis_results.get('subwayStations', {}).get('nearest', {}).get('distance', 0)) + '米',
            
            # 评价结果
            '评价结果': analysis_results.get('evaluation', {}).get('result', '不符合'),
            '评分': str(analysis_results.get('evaluation', {}).get('score', 0)),
        }]
        
        # 添加站点列表数据
        for i, station in enumerate(stations, 1):
            if i > 10:  # 最多处理10个站点
                break
                
            template_data[0][f'站点{i}名称'] = station.get('name', '')
            template_data[0][f'站点{i}类型'] = station.get('type', '')
            template_data[0][f'站点{i}距离'] = str(station.get('distance', 0)) + '米'
        
        # 生成Word文档
        timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
        template_path = '公共交通分析报告.docx'
        output_filename = f"公共交通分析报告_{project_name}_{timestamp}.docx"
        output_dir = os.path.join(app.config['EXPORT_FOLDER'])
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, output_filename)
        
        # 使用现有的Word模板处理函数
        from word_template import process_template
        result = process_template(template_data, template_path=template_path, output_path=output_path)
        
        if result and os.path.exists(output_path):
            app.logger.info(f"公共交通分析报告生成成功: {output_path}")
            
            # 返回下载链接
            download_url = url_for('static', filename=f'exports/{output_filename}')
            return jsonify({
                'success': True, 
                'message': '公共交通分析报告生成成功',
                'downloadUrl': download_url
            })
        else:
            app.logger.error(f"公共交通分析报告生成失败")
            return jsonify({'success': False, 'message': '报告生成失败，请稍后重试'}), 500
            
    except Exception as e:
        app.logger.error(f"生成公共交通分析报告时出错: {str(e)}")
        app.logger.error(traceback.format_exc())
        return jsonify({'success': False, 'message': f'处理请求失败: {str(e)}'}), 500

@app.route('/api/fill_transport_report_template', methods=['POST'])
@login_required
def fill_transport_report_template():
    """填充公共交通站点分析报告模板"""
    try:
        data = request.json
        if not data:
            return jsonify({'success': False, 'message': '缺少必要的数据'}), 400
        
        # 获取必要的数据
        address = data.get('address', '未指定地址')
        stations = data.get('stations', [])
        map_image_data = data.get('mapImage', '')
        map_info = data.get('mapInfo', {})
        project_id = data.get('project_id')  # 获取项目ID
        
        # 检查必要的数据是否存在
        if not stations:
            return jsonify({'success': False, 'message': '缺少公交站点数据'}), 400
            
        # 创建临时目录（如果不存在）
        exports_dir = os.path.join(app.static_folder, 'exports')
        os.makedirs(exports_dir, exist_ok=True)
        
        # 为当前报告创建一个唯一的文件名
        timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
        unique_id = str(uuid.uuid4())[:8]
        report_filename = f"公交站点分析报告_{timestamp}_{unique_id}.docx"
        output_path = os.path.join(exports_dir, report_filename)
        
        # 模板文件路径
        template_path = os.path.join(app.static_folder, 'templates', '公共交通站点分析报告.docx')
        if not os.path.exists(template_path):
            app.logger.error(f"模板文件不存在: {template_path}")
            return jsonify({'success': False, 'message': f'模板文件不存在，请确认 {template_path} 文件已准备好'}), 500
        
        # 复制模板文件到输出路径
        import shutil
        shutil.copy2(template_path, output_path)
        app.logger.info(f"已将模板复制到: {output_path}")
        
        # 获取项目信息
        project_info = {}
        if project_id:
            try:
                # 从数据库获取项目信息
                project = Project.query.get(project_id)
                if project:
                    project_info = {
                        '项目名称': project.name or '',
                        '项目地点': project.location or '',
                        '项目编号': project.code or '',  # 使用项目编号字段
                        '建设单位': project.construction_unit or '',
                        '设计单位': project.design_unit or '',
                        '总建筑面积': str(project.total_building_area or '') + ' 平方米' if project.total_building_area else '',
                        '总用地面积': str(project.total_land_area or '') + ' 平方米' if project.total_land_area else '',
                        '建筑密度': str(project.building_density or '') + ' %' if project.building_density else '',
                        '容积率': str(project.plot_ratio or '') if project.plot_ratio else '',
                        '绿地率': str(project.green_ratio or '') + ' %' if project.green_ratio else '',
                        '设计日期': datetime.now().strftime('%Y年%m月%d日')
                    }
                    
                    # 表单数据中可能有其他字段，但我们已经直接使用了project.code
                    # 所以不需要再次从form_data获取设计编号
                    
                    app.logger.info(f"成功获取项目信息: {project_info}")
                else:
                    app.logger.warning(f"未找到ID为{project_id}的项目")
            except Exception as e:
                app.logger.error(f"获取项目信息时出错: {str(e)}")
                app.logger.error(traceback.format_exc())
                # 继续执行，即使没有项目信息
        else:
            app.logger.info("未提供项目ID，将使用默认项目信息")
            
        # 确保项目信息中的占位符值不是None
        for key in list(project_info.keys()):
            if project_info[key] is None:
                project_info[key] = ''
                
        app.logger.info(f"最终项目信息: {json.dumps(project_info, ensure_ascii=False)}")
        
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import io
            import base64
            import re
            
            # 打开复制的模板文件
            doc = Document(output_path)
            
            # 查找并替换文本占位符
            app.logger.info(f"开始查找文本占位符...")
            
            # 初始化标记
            stations_table_added = False
            placeholder_found = False
            map_image_added = False
            temp_image_path = None
            
            # 在处理段落占位符前，先扫描整个文档以识别所有可能的占位符
            all_placeholders = []
            for i, paragraph in enumerate(doc.paragraphs):
                original_text = paragraph.text
                
                # 使用正则表达式查找所有花括号格式的文本
                placeholders = re.findall(r'\{([^}]+)\}', original_text)
                for p in placeholders:
                    if p not in all_placeholders:
                        all_placeholders.append(p)
                        
            # 也检查表格中的占位符
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        for para_idx, paragraph in enumerate(cell.paragraphs):
                            original_text = paragraph.text
                            placeholders = re.findall(r'\{([^}]+)\}', original_text)
                            for p in placeholders:
                                if p not in all_placeholders:
                                    all_placeholders.append(p)
            
            app.logger.info(f"在整个文档中找到的所有占位符: {all_placeholders}")
            
            # 更新项目信息字典，确保包含所有可能的占位符变体
            extended_project_info = project_info.copy()
            
            # 添加潜在的格式变体
            for placeholder in all_placeholders:
                # 检查是否有相似的键
                for key in project_info.keys():
                    # 简单的相似度检查 - 移除空格和标点后比较
                    cleaned_placeholder = re.sub(r'[\s\W]+', '', placeholder.lower())
                    cleaned_key = re.sub(r'[\s\W]+', '', key.lower())
                    
                    if cleaned_placeholder == cleaned_key and placeholder not in extended_project_info:
                        extended_project_info[placeholder] = project_info[key]
                        app.logger.info(f"添加占位符变体映射: {placeholder} -> {key} = {project_info[key]}")
            
            app.logger.info(f"扩展后的项目信息字典: {json.dumps(extended_project_info, ensure_ascii=False)}")
            
            # 处理地图图片
            if map_image_data and map_image_data.startswith('data:image/png;base64,'):
                try:
                    # 从base64数据中提取图片内容
                    map_image_data = map_image_data.replace('data:image/png;base64,', '')
                    image_binary = base64.b64decode(map_image_data)
                    
                    # 保存图片到临时文件
                    temp_image_path = os.path.join(exports_dir, f"map_image_{unique_id}.png")
                    with open(temp_image_path, 'wb') as img_file:
                        img_file.write(image_binary)
                        
                    app.logger.info(f"地图图片已保存到临时文件: {temp_image_path}")
                except Exception as e:
                    app.logger.error(f"处理地图图片数据时出错: {str(e)}")
                    # 如果Canvas生成的图片处理失败，尝试使用百度地图静态图API
                    if map_info and map_info.get('center'):
                        try:
                            app.logger.info("尝试使用百度地图静态图API生成地图...")
                            # 这里可以通过调用百度地图静态图API来获取地图图片
                            # 但需要额外的处理
                        except Exception as e2:
                            app.logger.error(f"使用百度地图静态图API生成地图时出错: {str(e2)}")
            else:
                app.logger.warning("未收到有效的地图图片数据")
                # 可以在这里添加使用百度地图静态图API的备用方案
            
            # 遍历所有段落查找占位符
            for i, paragraph in enumerate(doc.paragraphs):
                original_text = paragraph.text
                app.logger.info(f"检查段落 {i+1}: \"{original_text}\"")
                
                # 替换地址占位符
                if '{地址}' in original_text:
                    # 获取整个段落的文本并替换占位符
                    new_text = original_text.replace('{地址}', address)
                    app.logger.info(f"替换地址占位符: '{original_text}' -> '{new_text}'")
                    
                    # 清空段落中所有runs，但保持段落本身
                    p = paragraph._p
                    for run in list(paragraph.runs):  # 使用列表复制，因为我们在修改集合
                        p.remove(run._r)
                        
                    # 添加新的文本
                    paragraph.add_run(new_text)
                
                # 替换日期占位符
                if '{日期}' in paragraph.text:
                    # 获取整个段落的文本并替换占位符
                    new_text = paragraph.text.replace('{日期}', datetime.now().strftime('%Y年%m月%d日'))
                    app.logger.info(f"替换日期占位符: '{paragraph.text}' -> '{new_text}'")
                    
                    # 清空段落中所有runs，但保持段落本身
                    p = paragraph._p
                    for run in list(paragraph.runs):  # 使用列表复制，因为我们在修改集合
                        p.remove(run._r)
                        
                    # 添加新的文本
                    paragraph.add_run(new_text)
                
                # 替换项目信息占位符
                # 创建一个副本，因为我们将在迭代过程中修改文本
                original_text = paragraph.text
                new_text = original_text
                text_changed = False
                
                # 先检查段落文本中是否包含任何花括号
                if '{' in original_text and '}' in original_text:
                    app.logger.info(f"段落 {i+1} 包含花括号，可能包含占位符")
                    
                    # 使用正则表达式查找所有占位符
                    import re
                    placeholders = re.findall(r'\{([^}]+)\}', original_text)
                    app.logger.info(f"找到的占位符: {placeholders}")
                    
                    # 检查是否包含任何项目信息占位符
                    for placeholder, value in project_info.items():
                        placeholder_text = '{' + placeholder + '}'
                        if placeholder_text in new_text:
                            new_text = new_text.replace(placeholder_text, value)
                            text_changed = True
                            app.logger.info(f"替换了项目信息占位符 {placeholder_text} -> {value}")
                
                # 如果文本已更改，应用新文本
                if text_changed:
                    app.logger.info(f"段落 {i+1} 文本已更改: '{original_text}' -> '{new_text}'")
                    # 清空段落中所有runs，但保持段落本身
                    p = paragraph._p
                    for run in list(paragraph.runs):  # 使用列表复制，因为我们在修改集合
                        p.remove(run._r)
                        
                    # 添加新的文本
                    paragraph.add_run(new_text)
                
                # 替换地图占位符，并插入地图图片
                if '{地图截图}' in paragraph.text and temp_image_path:
                    app.logger.info(f"找到地图截图占位符")
                    
                    # 设置段落居中对齐
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # 清除段落中的所有runs，但保持段落本身
                    p = paragraph._p
                    for run in list(paragraph.runs):  # 使用列表复制，因为我们在修改集合
                        p.remove(run._r)
                    
                    # 在清空后的段落中添加图片
                    try:
                        run = paragraph.add_run()
                        run.add_picture(temp_image_path, width=Inches(6))
                        map_image_added = True
                        app.logger.info(f"成功在占位符位置添加了地图图片")
                    except Exception as e:
                        app.logger.error(f"添加地图图片时出错: {str(e)}")
                        # 如果添加图片失败，添加文本提示
                        paragraph.add_run(f"{address} 周边公交站位置情况（图片加载失败）")
                    
                elif '{地图截图}' in paragraph.text:
                    # 如果没有图片，只添加标题
                    # 清空段落中所有runs
                    p = paragraph._p
                    for run in list(paragraph.runs):  # 使用列表复制，因为我们在修改集合
                        p.remove(run._r)
                    # 添加新的标题文本
                    run = paragraph.add_run(f"{address} 周边公交站位置情况")
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    app.logger.info(f"替换了地图截图占位符（无图片）")
                
                # 处理公交站点列表占位符
                if '{公交站点列表}' in paragraph.text:
                    app.logger.info(f"找到公交站点列表占位符")
                    placeholder_found = True
                    
                    # 创建表格
                    table = doc.add_table(rows=1, cols=4)
                    table.style = 'Table Grid'
                    
                    # 设置表头
                    header_cells = table.rows[0].cells
                    header_cells[0].text = "序号"
                    header_cells[1].text = "站点名称"
                    header_cells[2].text = "距离（米）"
                    header_cells[3].text = "详细信息"
                    
                    # 添加数据行
                    for idx, station in enumerate(stations):
                        row = table.add_row()
                        cells = row.cells
                        cells[0].text = str(idx + 1)
                        cells[1].text = station.get('name', '')
                        cells[2].text = station.get('distance', '')
                        cells[3].text = station.get('address', '无详细信息')
                    
                    # 直接用表格替换段落，而不是添加到段落后面
                    try:
                        # 获取段落的父元素
                        parent = paragraph._p.getparent()
                        # 获取段落在父元素中的索引
                        index = parent.index(paragraph._p)
                        # 在段落的位置插入表格
                        parent.insert(index, table._tbl)
                        # 移除原始段落
                        parent.remove(paragraph._p)
                        
                        stations_table_added = True
                        app.logger.info(f"成功用表格替换了占位符段落")
                    except Exception as e:
                        app.logger.error(f"替换段落为表格时出错: {str(e)}")
                        os.remove(output_path)  # 删除临时文件
                        return jsonify({'success': False, 'message': f'替换段落为表格失败: {str(e)}'}), 500
            
            # 如果没有找到占位符，返回错误
            if not placeholder_found:
                app.logger.warning(f"在模板中未找到'{公交站点列表}'占位符")
                os.remove(output_path)  # 删除临时文件
                return jsonify({'success': False, 'message': '在模板中未找到"{公交站点列表}"占位符，请检查模板文件'}), 400
            
            # 如果找到占位符但无法插入表格，返回错误
            if placeholder_found and not stations_table_added:
                app.logger.warning(f"无法在占位符位置插入表格")
                os.remove(output_path)  # 删除临时文件
                return jsonify({'success': False, 'message': '虽然找到了占位符，但无法在此位置插入表格，请检查模板格式'}), 500
            
            # 保存文档
            doc.save(output_path)
            app.logger.info(f"已保存文档到: {output_path}")
            
            # 清理临时文件
            if temp_image_path and os.path.exists(temp_image_path):
                os.remove(temp_image_path)
                app.logger.info(f"删除临时图片文件: {temp_image_path}")
            
            # 创建下载链接
            download_url = url_for('static', filename=f'exports/{report_filename}')
            app.logger.info(f"生成的下载链接: {download_url}")
            
            return jsonify({
                'success': True, 
                'message': '公交站点分析报告生成成功',
                'download_url': download_url
            })
            
        except Exception as e:
            app.logger.error(f"处理Word文档时出错: {str(e)}")
            app.logger.error(traceback.format_exc())
            # 删除临时文件
            if os.path.exists(output_path):
                os.remove(output_path)
            # 清理临时图片文件
            if temp_image_path and os.path.exists(temp_image_path):
                os.remove(temp_image_path)
            return jsonify({'success': False, 'message': f'处理Word文档时出错: {str(e)}'}), 500
            
    except Exception as e:
        app.logger.error(f"生成公交站点分析报告时出错: {str(e)}")
        app.logger.error(traceback.format_exc())
        return jsonify({'success': False, 'message': f'处理请求失败: {str(e)}'}), 500

if __name__ == '__main__':
    # 初始化数据库
    init_db()
    
    # 根据环境变量决定是否开启调试模式
    debug_mode = not is_production
    app.logger.info(f"应用启动: 调试模式={debug_mode}")
    app.run(debug=debug_mode, host='0.0.0.0', port=int(os.environ.get('PORT', 5000)))