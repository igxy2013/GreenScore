import os
import re
from loguru import logger
import platform

def extract_project_info(file_path):
    """
    从Word文档中提取项目信息
    
    Args:
        file_path (str): Word文档的路径
        
    Returns:
        dict: 包含提取出的项目信息的字典，如果提取失败则返回None
    """
    import os
    from loguru import logger
    
    try:
        # 检查文件是否存在
        if not os.path.exists(file_path):
            logger.error(f"文件不存在: {file_path}")
            return None
            
        # 根据文件扩展名选择处理方式
        file_ext = os.path.splitext(file_path)[1].lower()
        
        # 处理没有扩展名的情况
        if not file_ext:
            logger.warning(f"文件没有扩展名: {file_path}，尝试作为.docx处理")
            file_ext = '.docx'
        
        if file_ext == '.docx':
            # 使用python-docx处理.docx文件
            try:
                from docx import Document
                doc = Document(file_path)
                
                # 初始化结果字典
                project_info = {
                    "项目名称": "",
                    "项目地点": "",
                    "气候分区": "",
                    "建筑分类": "",
                    "结构形式": "",
                    "建筑朝向": "",
                    "建筑面积": "",
                    "建筑层数": "",
                    "地下层数": "",
                    "建筑高度": "",
                    "设计单位": "",
                    "建设单位": ""
                }
                
                # 提取信息
                extract_info_from_docx(doc, project_info)
                
                # 处理地下层数信息
                if "地下层数" not in project_info or not project_info["地下层数"]:
                    # 尝试从建筑层数中提取地下层数信息
                    if project_info["建筑层数"]:
                        # 正确识别"地下X层"的模式
                        underground_match = re.search(r"地下[：:\s]*(\d+)[层\s]", project_info["建筑层数"])
                        if underground_match:
                            project_info["地下层数"] = underground_match.group(1)
                
                # 对提取出的项目名称进行清理和验证
                if project_info["项目名称"]:
                    # 1. 移除可能混入的多余符号和空格
                    project_info["项目名称"] = project_info["项目名称"].strip()
                    project_info["项目名称"] = re.sub(r'^[：:、\s]+', '', project_info["项目名称"])
                    project_info["项目名称"] = re.sub(r'[：:、\s]+$', '', project_info["项目名称"])
                    
                    # 2. 再次清理常见词汇
                    project_info["项目名称"] = re.sub(r'(节能设计|规定性指标|计算报告书|审图答复|专篇|设计说明|建筑节能|\(.*?\))', '', project_info["项目名称"]).strip()
                    
                    # 3. 限制长度，过长可能是错误提取
                    if len(project_info["项目名称"]) > 50:
                        # 尝试截取合理部分
                        name_parts = re.split(r'[,，。;；]', project_info["项目名称"])
                        if name_parts and len(name_parts[0]) > 3 and len(name_parts[0]) < 50:
                            project_info["项目名称"] = name_parts[0].strip()
                    
                    # 4. 验证提取结果是否合理
                    if project_info["项目名称"] in ["设计报告书", "报告书", "审图答复", "专篇", "设计说明", "建筑节能"]:
                        # 这些情况下提取到的是错误结果，置空
                        project_info["项目名称"] = ""
                
                # 保留项目名称字段是否存在的标记
                project_name_field_exists = project_info.get("_项目名称字段存在", False)
                
                # 移除处理过程中的内部标记
                if "_项目名称字段存在" in project_info:
                    del project_info["_项目名称字段存在"]
                
                # 如果项目名称字段确实存在但是为空，应该保持为空
                if project_name_field_exists and not project_info["项目名称"]:
                    logger.info("文档中明确存在项目名称字段但值为空，保持为空值")
                    project_info["项目名称"] = ""
                    
                # 记录提取结果
                logger.info(f"从文件 {os.path.basename(file_path)} 提取到的项目信息: {project_info}")
                return project_info
                
            except Exception as e:
                logger.error(f"处理.docx文件时出错: {str(e)}")
                import traceback
                logger.error(traceback.format_exc())
                return None
                
        elif file_ext == '.doc' and os.name == 'nt':  # 仅在Windows下支持旧版.doc文件
            return extract_from_doc(file_path)
        else:
            logger.error(f"不支持的文件格式: {file_ext}，仅支持.doc和.docx")
            return None
            
    except Exception as e:
        logger.error(f"提取项目信息时出错: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return None

def extract_info_from_docx(doc, project_info):
    """从docx文档对象中提取信息"""
    # 标记是否在文档中明确存在但为空的项目名称字段
    project_name_field_exists = False
    
    # 先尝试从表格中提取工程名称，这是最可靠的方法
    try:
        if doc.tables and len(doc.tables) > 0:
            # 尝试从前三个表格中查找工程名称
            for table_idx in range(min(3, len(doc.tables))):
                table = doc.tables[table_idx]
                if len(table.rows) > 0:
                    # 对每个表格的前5行进行检查
                    for row_idx in range(min(5, len(table.rows))):
                        row = table.rows[row_idx]
                        if len(row.cells) >= 2:
                            # 检查第一个单元格是否为"工程名称"或包含"工程名称"
                            cell0_text = row.cells[0].text.strip().rstrip(':')
                            if cell0_text == "工程名称" or cell0_text == "项目名称":
                                # 此时找到了项目名称字段，无论其值是否为空，都标记字段存在
                                project_name_field_exists = True
                                project_name = row.cells[1].text.strip()
                                project_info["项目名称"] = project_name
                                # 如果字段为空，记录此信息
                                if not project_name:
                                    logger.info("表格中发现项目名称字段但值为空")
                            
                            # 处理建设单位
                            if cell0_text == "建设单位" and not project_info["建设单位"]:
                                project_info["建设单位"] = row.cells[1].text.strip()
                                
                            # 处理设计单位
                            if cell0_text == "设计单位" and not project_info["设计单位"]:
                                project_info["设计单位"] = row.cells[1].text.strip()
                                
                            # 处理工程地点
                            if cell0_text == "工程地点" and not project_info["项目地点"]:
                                project_info["项目地点"] = row.cells[1].text.strip()
                        
                        # 检查所有单元格，寻找包含"工程名称:"的文本
                        for cell_idx, cell in enumerate(row.cells):
                            cell_text = cell.text.strip()
                            if "工程名称:" in cell_text or "项目名称:" in cell_text:
                                project_name_field_exists = True
                                if cell_idx + 1 < len(row.cells):
                                    next_cell = row.cells[cell_idx + 1]
                                    project_name = next_cell.text.strip()
                                    project_info["项目名称"] = project_name
                                    if not project_name:
                                        logger.info("单元格中发现项目名称标记但值为空")
                    
                    # 如果在这个表格中找到了项目名称字段，跳出表格循环
                    if project_name_field_exists:
                        break
    except Exception as e:
        logger.warning(f"尝试从表格提取工程名称时出错: {str(e)}")
    
    # 在project_info中添加一个标记，表示是否在文档中明确找到项目名称字段
    project_info["_项目名称字段存在"] = project_name_field_exists
    
    # 提取段落信息
    paragraphs = [para.text.strip() for para in doc.paragraphs]
    extract_from_paragraphs(paragraphs, project_info)
    
    # 提取表格信息
    tables_data = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        tables_data.append(table_data)
    
    extract_from_tables(tables_data, project_info)

def extract_from_paragraphs(paragraphs, project_info):
    """从段落文本中提取项目信息"""
    # 检查是否有项目名称字段在文档中明确存在
    project_name_field_exists = project_info.get("_项目名称字段存在", False)
    
    # 检查是否有明显的文档标题作为项目名称
    title_candidates = []
    
    # 尝试从文档标题或前几个段落提取项目名称
    for i, text in enumerate(paragraphs[:15]):  # 检查前15个段落
        text = text.strip()
        if not text:
            continue
        
        # 收集可能的标题段落
        if len(text) > 5 and len(text) < 70:  # 合理长度的标题
            # 特征1: 居中的文本通常是标题
            if text == text.strip():
                title_candidates.append(text)
            
            # 特征2: 全大写或首字母大写的文本可能是标题
            if text.isupper() or text.istitle():
                title_candidates.append(text)
        
        # 检查是否包含"项目名称"或"工程名称"字段
        if "项目名称" in text or "工程名称" in text:
            # 标记项目名称字段在段落中存在
            project_name_field_exists = True
            project_info["_项目名称字段存在"] = True
            
            # 尝试提取字段值
            match = re.search(r"(?:项目名称|工程名称)[：:]\s*(.+?)(?:\s*$|，|,|；|;|\(|（)", text)
            if match:
                name_value = match.group(1).strip()
                project_info["项目名称"] = name_value
                # 如果字段存在但值为空，记录此信息
                if not name_value:
                    logger.info("段落中发现项目名称字段但值为空")
            else:
                # 尝试提取冒号后的所有内容
                match = re.search(r"(?:项目名称|工程名称)[：:]\s*(.+)", text)
                if match:
                    name_value = match.group(1).strip()
                    project_info["项目名称"] = name_value
                    if not name_value:
                        logger.info("段落中发现项目名称字段但值为空")
                else:
                    # 如果找到字段但没有提取到值，可能是值为空
                    logger.info("段落中检测到项目名称字段但无法提取值，可能为空")
                    project_info["项目名称"] = ""
        
        # 如果未找到项目名称字段且当前段落文本有合适内容，尝试作为标题提取
        if not project_name_field_exists and not project_info["项目名称"] and text and not project_info["项目名称"]:
            # 如果段落包含"节能"、"设计"、"报告"等关键词，可能是标题，尝试提取项目名称
            if any(keyword in text for keyword in ["节能", "设计", "报告", "规划", "建筑"]):
                # 移除常见的文档标题词，保留项目名称部分
                cleaned_text = re.sub(r'(节能设计|规定性指标|计算报告书|审图答复|专篇|设计说明|建筑节能)', '', text).strip()
                if cleaned_text and len(cleaned_text) > 3 and len(cleaned_text) < 50:  # 合理长度的项目名称
                    project_info["项目名称"] = cleaned_text
                    break
    
    # 如果仍未找到项目名称且没有明确的项目名称字段，尝试处理收集到的标题候选
    if not project_info["项目名称"] and not project_name_field_exists and title_candidates:
        # 选择最可能的标题 (不包含"表"字，且最短的那个)
        valid_titles = [t for t in title_candidates if "表" not in t]
        if valid_titles:
            # 按长度排序，选择最短的有效标题
            shortest_title = min(valid_titles, key=len)
            # 清理常见词汇
            cleaned_title = re.sub(r'(节能设计|规定性指标|计算报告书|审图答复|专篇|设计说明|建筑节能)', '', shortest_title).strip()
            if cleaned_title and len(cleaned_title) > 3:
                project_info["项目名称"] = cleaned_title
    
    # 更新项目名称字段存在标记
    project_info["_项目名称字段存在"] = project_name_field_exists
    
    # 处理其他字段
    for text in paragraphs:
        text = text.strip()
        if not text:
            continue
        
        # 项目地点（城市）- 增加对"项目城市"和"工程地点"的识别
        if "项目地点" in text or "建设地点" in text or "项目城市" in text or "工程地点" in text:
            match = re.search(r"(?:项目地点|建设地点|项目城市|工程地点)[：:]\s*(.+?)(?:\s|$|，|,)", text)
            if match:
                location = match.group(1).strip()
                # 提取城市名称
                city_match = re.search(r"([^\s省市区县]+?[市州])", location)
                if city_match:
                    project_info["项目地点"] = city_match.group(1)
                else:
                    project_info["项目地点"] = location
        
        # 气候分区
        if "气候分区" in text or "气候区划" in text:
            match = re.search(r"气候[分区|区划][：:]\s*(.+?)(?:\s|$|，|,)", text)
            if match:
                project_info["气候分区"] = match.group(1).strip()
        
        # 建筑分类
        if "建筑类型" in text or "建筑分类" in text or "结构类型" in text:
            match = re.search(r"(?:建筑[类型|分类]|结构类型)[：:]\s*(.+?)(?:\s|$|，|,)", text)
            if match:
                project_info["建筑分类"] = match.group(1).strip()
        
        # 结构形式
        if "结构形式" in text or "结构类型" in text:
            match = re.search(r"(?:结构形式|结构类型)[：:]\s*(.+?)(?:\s|$|，|,)", text)
            if match:
                project_info["结构形式"] = match.group(1).strip()
        
        # 建筑朝向
        if "建筑朝向" in text or "北向角度" in text:
            match = re.search(r"(?:建筑朝向|北向角度)[：:]\s*(.+?)(?:\s|$|，|,)", text)
            if match:
                project_info["建筑朝向"] = match.group(1).strip()
        
        # 建筑面积
        if "建筑面积" in text and "总建筑面积" not in text:
            match = re.search(r"建筑面积[：:]\s*([0-9,.]+)\s*(?:m²|平方米)?", text)
            if match:
                project_info["建筑面积"] = match.group(1).strip()
        elif "总建筑面积" in text:
            match = re.search(r"总建筑面积[：:]\s*([0-9,.]+)\s*(?:m²|平方米)?", text)
            if match:
                project_info["建筑面积"] = match.group(1).strip()
        
        # 建筑层数
        if "层数" in text and "建筑层数" not in text and "地下层数" not in text:
            match = re.search(r"(?:建筑)?层数[：:]\s*(.+?)(?:\s|$|，|,)", text)
            if match:
                project_info["建筑层数"] = match.group(1).strip()
                # 尝试提取地下层数 - 正确识别"地下X层"的模式
                underground_match = re.search(r"地下[：:\s]*(\d+)[层\s]", match.group(1))
                if underground_match:
                    project_info["地下层数"] = underground_match.group(1)
        elif "建筑层数" in text:
            match = re.search(r"建筑层数[：:]\s*(.+?)(?:\s|$|，|,)", text)
            if match:
                project_info["建筑层数"] = match.group(1).strip()
                # 尝试提取地下层数 - 正确识别"地下X层"的模式
                underground_match = re.search(r"地下[：:\s]*(\d+)[层\s]", match.group(1))
                if underground_match:
                    project_info["地下层数"] = underground_match.group(1)
        
        # 单独的地下层数
        if "地下层数" in text and not project_info["地下层数"]:
            match = re.search(r"地下层数[：:]\s*(\d+)", text)
            if match:
                project_info["地下层数"] = match.group(1).strip()
        
        # 建筑高度
        if "建筑高度" in text:
            match = re.search(r"建筑高度[：:]\s*([0-9,.]+)\s*(?:m|米)?", text)
            if match:
                project_info["建筑高度"] = match.group(1).strip()
        
        # 设计单位
        if "设计单位" in text and not project_info["设计单位"]:
            match = re.search(r"设计单位[：:]\s*(.+?)(?:\s|$|，|,|。)", text)
            if match:
                project_info["设计单位"] = match.group(1).strip()
        
        # 建设单位
        if "建设单位" in text and not project_info["建设单位"]:
            match = re.search(r"建设单位[：:]\s*(.+?)(?:\s|$|，|,|。)", text)
            if match:
                project_info["建设单位"] = match.group(1).strip()

def extract_from_tables(tables_data, project_info):
    """从表格数据中提取项目信息"""
    # 检查是否已经标记了项目名称字段的存在
    project_name_field_exists = project_info.get("_项目名称字段存在", False)
    
    # 首先特别处理前三个表格，专注提取工程名称
    if tables_data:
        # 遍历前三个表格
        for table_idx in range(min(3, len(tables_data))):
            if table_idx < len(tables_data) and tables_data[table_idx]:
                # 检查表格的前5行
                for row_idx in range(min(5, len(tables_data[table_idx]))):
                    row_data = tables_data[table_idx][row_idx]
                    if len(row_data) >= 2:  # 确保至少有两列
                        # 提取第一列单元格内容，去除冒号
                        cell0_text = row_data[0].strip().rstrip(':')
                        
                        # 处理工程名称 (或项目名称)
                        if cell0_text == "工程名称" or cell0_text == "项目名称":
                            # 标记项目名称字段在表格中存在
                            project_name_field_exists = True
                            project_info["_项目名称字段存在"] = True
                            
                            # 提取值，即使为空
                            project_name = row_data[1].strip()
                            project_info["项目名称"] = project_name
                            
                            # 如果字段为空，记录此信息
                            if not project_name:
                                logger.info("表格提取: 找到项目名称字段但值为空")
                        
                        # 处理建设单位
                        if cell0_text == "建设单位" and not project_info["建设单位"]:
                            project_info["建设单位"] = row_data[1].strip()
                            
                        # 处理设计单位
                        if cell0_text == "设计单位" and not project_info["设计单位"]:
                            project_info["设计单位"] = row_data[1].strip()
                        
                        # 处理工程地点
                        if cell0_text == "工程地点" and not project_info["项目地点"]:
                            project_info["项目地点"] = row_data[1].strip()
    
    # 如果仍未找到项目名称字段，继续检查所有表格的所有行
    if not project_name_field_exists:
        for table_data in tables_data:
            for row_data in table_data:
                if len(row_data) >= 2:  # 确保至少有两列
                    # 第一列是否包含"工程名称"或"项目名称"
                    cell0_text = row_data[0].strip().rstrip(':')
                    if "工程名称" in cell0_text or "项目名称" in cell0_text:
                        # 标记存在项目名称字段
                        project_name_field_exists = True
                        project_info["_项目名称字段存在"] = True
                        
                        # 提取值，即使为空
                        project_name = row_data[1].strip()
                        project_info["项目名称"] = project_name
                        
                        # 如果值为空，记录此信息
                        if not project_name:
                            logger.info("表格深度提取: 找到项目名称字段但值为空")
    
    # 更新项目名称字段存在标记
    project_info["_项目名称字段存在"] = project_name_field_exists
    
    # 继续提取其他信息
    for table_data in tables_data:
        for row_data in table_data:
            if len(row_data) >= 2:  # 确保至少有两列
                cell0_text = row_data[0].strip().rstrip(':')
                
                # 精确匹配第一列是否为"工程地点"
                if cell0_text == "工程地点" and not project_info["项目地点"]:
                    project_info["项目地点"] = row_data[1].strip()
                
                # 精确匹配第一列是否为"气候分区"
                if cell0_text == "气候分区" and not project_info["气候分区"]:
                    project_info["气候分区"] = row_data[1].strip()
                    
                # 精确匹配第一列是否为"结构类型"
                if cell0_text == "结构类型" and not project_info["结构形式"]:
                    project_info["结构形式"] = row_data[1].strip()
                    
                # 精确匹配第一列是否为"北向角度"
                if cell0_text == "北向角度" and not project_info["建筑朝向"]:
                    project_info["建筑朝向"] = row_data[1].strip()
                    
                # 精确匹配第一列是否为"建筑面积"
                if cell0_text == "建筑面积" and not project_info["建筑面积"]:
                    area_text = row_data[1].strip()
                    match = re.search(r"([0-9,.]+)", area_text)
                    if match:
                        project_info["建筑面积"] = match.group(1)
                        
                # 精确匹配第一列是否为"建筑层数"
                if cell0_text == "建筑层数" and not project_info["建筑层数"]:
                    project_info["建筑层数"] = row_data[1].strip()
                    
                # 精确匹配第一列是否为"建筑高度"
                if cell0_text == "建筑高度" and not project_info["建筑高度"]:
                    height_text = row_data[1].strip()
                    match = re.search(r"([0-9,.]+)", height_text)
                    if match:
                        project_info["建筑高度"] = match.group(1)
                        
                # 精确匹配第一列是否为"设计单位"
                if cell0_text == "设计单位" and not project_info["设计单位"]:
                    project_info["设计单位"] = row_data[1].strip()
                    
                # 精确匹配第一列是否为"建设单位"
                if cell0_text == "建设单位" and not project_info["建设单位"]:
                    project_info["建设单位"] = row_data[1].strip()

def extract_from_doc(file_path):
    """
    从.doc文件中提取项目信息
    
    Args:
        file_path (str): .doc文件路径
    
    Returns:
        dict: 包含提取出的项目信息
    """
    from loguru import logger
    
    # 初始化结果字典
    project_info = {
        "项目名称": "",
        "项目地点": "",
        "气候分区": "",
        "建筑分类": "",
        "结构形式": "",
        "建筑朝向": "",
        "建筑面积": "",
        "建筑层数": "",
        "地下层数": "",
        "建筑高度": "",
        "设计单位": "",
        "建设单位": "",
        "_项目名称字段存在": False  # 初始化标记字段
    }
    
    try:
        # 检查操作系统
        if platform.system() != 'Windows':
            logger.error("只有Windows系统可以处理.doc文件")
            return None
            
        try:
            import win32com.client
            import pythoncom
            
            # 初始化COM组件
            pythoncom.CoInitialize()
            
            try:
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                doc = word.Documents.Open(os.path.abspath(file_path))
                
                # 提取段落信息
                all_paragraphs = []
                for i in range(1, doc.Paragraphs.Count + 1):
                    all_paragraphs.append(doc.Paragraphs(i).Range.Text.strip())
                
                # 提取表格信息
                tables_data = []
                for i in range(1, doc.Tables.Count + 1):
                    table = doc.Tables(i)
                    table_data = []
                    for row_idx in range(1, table.Rows.Count + 1):
                        row_data = []
                        for col_idx in range(1, table.Columns.Count + 1):
                            try:
                                cell_text = table.Cell(row_idx, col_idx).Range.Text.strip()
                                cell_text = cell_text.replace('\r', '').replace('\n', ' ').replace('\x07', '')
                                row_data.append(cell_text)
                            except:
                                row_data.append("")
                        table_data.append(row_data)
                    tables_data.append(table_data)
                
                # 关闭文档和Word应用
                doc.Close(False)
                word.Quit()
                
                # 处理提取的文本内容
                extract_from_paragraphs(all_paragraphs, project_info)
                extract_from_tables(tables_data, project_info)
                
                # 处理地下层数信息
                if "地下层数" not in project_info or not project_info["地下层数"]:
                    # 尝试从建筑层数中提取地下层数信息
                    if project_info["建筑层数"]:
                        # 正确识别"地下X层"的模式
                        underground_match = re.search(r"地下[：:\s]*(\d+)[层\s]", project_info["建筑层数"])
                        if underground_match:
                            project_info["地下层数"] = underground_match.group(1)
                
                # 对提取出的项目名称进行清理和验证
                if project_info["项目名称"]:
                    # 1. 移除可能混入的多余符号和空格
                    project_info["项目名称"] = project_info["项目名称"].strip()
                    project_info["项目名称"] = re.sub(r'^[：:、\s]+', '', project_info["项目名称"])
                    project_info["项目名称"] = re.sub(r'[：:、\s]+$', '', project_info["项目名称"])
                    
                    # 2. 再次清理常见词汇
                    project_info["项目名称"] = re.sub(r'(节能设计|规定性指标|计算报告书|审图答复|专篇|设计说明|建筑节能|\(.*?\))', '', project_info["项目名称"]).strip()
                    
                    # 3. 限制长度，过长可能是错误提取
                    if len(project_info["项目名称"]) > 50:
                        # 尝试截取合理部分
                        name_parts = re.split(r'[,，。;；]', project_info["项目名称"])
                        if name_parts and len(name_parts[0]) > 3 and len(name_parts[0]) < 50:
                            project_info["项目名称"] = name_parts[0].strip()
                    
                    # 4. 验证提取结果是否合理
                    if project_info["项目名称"] in ["设计报告书", "报告书", "审图答复", "专篇", "设计说明", "建筑节能"]:
                        # 这些情况下提取到的是错误结果，置空
                        project_info["项目名称"] = ""
                
                # 保留项目名称字段是否存在的标记
                project_name_field_exists = project_info.get("_项目名称字段存在", False)
                
                # 移除处理过程中的内部标记
                if "_项目名称字段存在" in project_info:
                    del project_info["_项目名称字段存在"]
                
                # 如果项目名称字段确实存在但是为空，应该保持为空
                if project_name_field_exists and not project_info["项目名称"]:
                    logger.info("文档中明确存在项目名称字段但值为空，保持为空值")
                    project_info["项目名称"] = ""
                
                logger.info(f"从文件 {os.path.basename(file_path)} 提取到的项目信息: {project_info}")
                return project_info
            finally:
                # 确保在函数结束时释放COM资源
                pythoncom.CoUninitialize()
            
        except ImportError:
            logger.error("处理.doc文件需要pywin32和pythoncom库，请安装: pip install pywin32")
            return None
        except Exception as e:
            logger.error(f"处理.doc文件时出错: {str(e)}")
            import traceback
            logger.error(traceback.format_exc())
            return None
    
    except Exception as e:
        logger.error(f"提取项目信息时出错: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return None

def main():
    """
    测试函数
    """
    import sys
    if len(sys.argv) < 2:
        print("用法: python extract_word_info.py <word文件路径>")
        return
    
    file_path = sys.argv[1]
    info = extract_project_info(file_path)
    
    if info:
        print("\n提取到的项目信息:")
        for key, value in info.items():
            print(f"{key}: {value}")
    else:
        print("提取信息失败")

if __name__ == "__main__":
    main() 