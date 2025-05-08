import re
import os
try:
    import win32com.client
    import pythoncom
    _pywin32_installed = True
except ImportError:
    _pywin32_installed = False
    print("警告：pywin32 模块未安装或导入失败。将无法使用 .doc 到 .docx 的转换功能。请运行 'pip install pywin32' 安装。")

from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table, _Row
from docx.text.paragraph import Paragraph

def convert_doc_to_docx(doc_path: str, docx_path: str = None) -> str | None:
    """
    使用 Microsoft Word 将 .doc 文件转换为 .docx 格式。
    需要安装 Microsoft Word 并且系统为 Windows。
    需要安装 pywin32 库。

    Args:
        doc_path: 输入的 .doc 文件路径。
        docx_path: 可选的输出 .docx 文件路径。如果为 None，则保存在原文件同目录下。

    Returns:
        成功则返回转换后的 .docx 文件路径，否则返回 None。
    """
    if not _pywin32_installed:
        print("错误：pywin32 未安装或导入失败，无法执行转换。")
        return None

    if not doc_path.lower().endswith(".doc"):
        print(f"错误：输入文件 '{doc_path}' 不是 .doc 文件。")
        return None

    if not os.path.exists(doc_path):
        print(f"错误：输入文件 '{doc_path}' 未找到。")
        return None

    # 确保路径对于 COM 是绝对路径
    absolute_doc_path = os.path.abspath(doc_path)

    if docx_path is None:
        # 如果未指定输出路径，则默认在原路径旁边生成 .docx 文件
        absolute_docx_path = os.path.splitext(absolute_doc_path)[0] + ".docx"
    else:
        absolute_docx_path = os.path.abspath(docx_path)

    # 检查输出文件是否已存在
    if os.path.exists(absolute_docx_path):
         try:
             # 尝试删除已存在的输出文件，避免 SaveAs 失败或弹出确认框
             os.remove(absolute_docx_path)
             print(f"信息：已删除已存在的输出文件 '{absolute_docx_path}'")
         except OSError as e:
             print(f"警告：无法删除已存在的输出文件 '{absolute_docx_path}'。转换可能会失败或被阻止: {e}")
             # 可以选择在这里返回失败，或者继续尝试
             # return None

    # 确保输出目录存在
    output_dir = os.path.dirname(absolute_docx_path)
    if not os.path.exists(output_dir):
        try:
            os.makedirs(output_dir, exist_ok=True)
        except OSError as e:
            print(f"错误：创建输出目录 '{output_dir}' 失败: {e}")
            return None

    word = None
    doc = None
    # 初始化当前线程的 COM 库
    pythoncom.CoInitialize()
    try:
        # 创建 Word 应用对象
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False # 不显示 Word 窗口
        # 禁用可能出现的宏警告或文件转换对话框
        word.DisplayAlerts = 0 # wdAlertsNone

        # 打开 .doc 文件
        print(f"正在打开 .doc 文件: {absolute_doc_path}")
        doc = word.Documents.Open(absolute_doc_path, ReadOnly=True, AddToRecentFiles=False)

        # .docx 格式的文件代码为 12
        wdFormatXMLDocument = 12

        # 另存为 .docx
        print(f"正在另存为 .docx 文件: {absolute_docx_path}")
        # 使用 SaveAs2 以便更好地控制兼容性等，但 SaveAs 也常用
        doc.SaveAs(absolute_docx_path, FileFormat=wdFormatXMLDocument)

        print("转换成功。")
        return absolute_docx_path

    except pythoncom.com_error as e:
        print(f"COM 错误导致转换失败: {e}")
        # 常见错误：未安装 Word 或 Word 安全设置限制
        if e.hresult == -2147221005: # CO_E_CLASSSTRING
            print("请确保已正确安装 Microsoft Word。")
        elif e.hresult == -2146822884: # Macro security settings
             print('请在 Word 的信任中心设置中启用 "信任对 VBA 工程对象模型的访问"。')
        elif e.hresult == -2147352567: # Exception occurred.
             print("Word 操作时发生内部错误，可能与文件损坏或权限有关。")
        return None
    except Exception as e:
        print(f"发生意外错误: {e}")
        return None
    finally:
        # 确保关闭 Word 进程和文档
        if doc:
            doc.Close(False) # 关闭时不保存更改
        if word:
            word.Quit()
        # 取消初始化 COM 库
        pythoncom.CoUninitialize()

def iter_block_items(parent):
    """
    Generate a reference to each paragraph and table child within *parent*,
    in document order. Each returned value is an instance of either Table or
    Paragraph. *parent* would most commonly be a reference to a main Document
    object, but also works for a _Cell object, which itself can contain
    paragraphs and tables.
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def extract_air_quality_scores(docx_path: str) -> dict:
    """
    从空气质量评价报告书.docx格式的文档中提取条文号和得分。

    Args:
        docx_path: Word 文档的路径。

    Returns:
        一个字典，键是条文号 (str)，值是得分 (int)。
        例如: {'5.1.1': 0, '5.2.1': 12}
    """
    scores = {}
    try:
        document = Document(docx_path)
        target_table = None
        table_index = -1 # Add index for debugging

        # 查找目标表格 (标题包含 "室内空气质量评价")
        # print("DEBUG [Air Extract]: Starting table search...") # Removed Debug print
        for block in iter_block_items(document):
            if isinstance(block, Table):
                table_index += 1 # Increment index
                # print(f"DEBUG [Air Extract]: Checking Table {table_index}") # Removed Debug print
                # 检查表格前的段落是否是表格标题
                # 或者检查表格第一行作为标题行 (更复杂，先用简单方法)
                # 这里假设表格紧跟在一个包含标题的段落后，或表格自身某单元格包含标题
                # 更可靠的是检查表格标题行，但python-docx直接获取表格标题不直观
                # 我们先尝试匹配表格内容
                if len(block.rows) > 0:
                    first_row_texts = [cell.text.strip() for cell in block.rows[0].cells]
                    # print(f"DEBUG [Air Extract]: Table {table_index} Header Texts: {first_row_texts}") # Removed Debug print
                    # -- Corrected Header Check Logic --
                    criteria_col_found = any("评价依据" in text for text in first_row_texts)
                    score_col_found = any("结论/得分" in text for text in first_row_texts)
                    header_keywords_present = criteria_col_found and score_col_found
                    # ---------------------------------
                    # print(f"DEBUG [Air Extract]: Table {table_index} Criteria Col Found: {criteria_col_found}, Score Col Found: {score_col_found}, Header Keywords Present: {header_keywords_present}") # Removed Debug print

                    if header_keywords_present:
                         # 主要依赖表头关键词来确定目标表格
                         # print(f"DEBUG [Air Extract]: Found potential target table based on header: Table {table_index}") # Removed Debug print
                         target_table = block
                         # 可以考虑在这里重新加上 break，因为很可能第一个匹配的就是目标
                         # break
                    # 移除对 specific_text_present 的依赖
                    # full_table_text = "\n".join(cell.text for row in block.rows for cell in row.cells)
                    # specific_text_present = "室内空气质量评价" in full_table_text
                    # print(f"DEBUG [Air Extract]: Table {table_index} Specific Text ('室内空气质量评价') Present: {specific_text_present}")
                    # if specific_text_present:
                    #      print(f"DEBUG [Air Extract]: Found potential target table: Table {table_index}")
                    #      target_table = block
                    # 暂时移除 break，检查所有表格
                else:
                    # print(f"DEBUG [Air Extract]: Table {table_index} has no rows.") # Removed Debug print
                    pass # Keep pass or remove print

                # 尝试查找紧邻表格上方的段落作为标题（另一种可能）
                # 这需要更复杂的逻辑来关联段落和表格，暂时跳过

        if target_table:
            # print(f"DEBUG [Air Extract]: Processing target table (last potential match found).") # Removed Debug print
            # 从第二行开始遍历 (跳过标题行)
            for r_idx, row in enumerate(target_table.rows[1:], 1): # Add row index for debug
                # print(f"--- Processing Row {r_idx} ---") # Removed Debug print
                try:
                    cells = row.cells
                    if len(cells) >= 3:
                        criteria_text = cells[1].text.strip() # 第 2 列: 评价依据
                        score_text = cells[2].text.strip()   # 第 3 列: 结论/得分
                        # print(f"  Cell 1 Text: '{criteria_text}'") # Removed Debug print
                        # print(f"  Cell 2 Text: '{score_text}'")   # Removed Debug print

                        # 提取条文号 (例如: 5.1.1, 3.2.8)
                        section_match = re.search(r'^(\d+(\.\d+)+)', criteria_text)
                        # print(f"  Section Match: {section_match}") # Removed Debug print
                        if section_match:
                            section_number = section_match.group(1)
                            # print(f"    Extracted Section: {section_number}") # Removed Debug print

                            # 提取得分 (例如: 12分, 3 分)
                            score_match = re.search(r'(\d+)\s*分', score_text)
                            # print(f"  Score Match: {score_match}") # Removed Debug print
                            if score_match:
                                score = int(score_match.group(1))
                                # print(f"    Extracted Score: {score}") # Removed Debug print
                                scores[section_number] = score
                            else:
                                # print(f"    Score not extracted from '{score_text}'") # Removed Debug print
                                pass # Keep pass or remove print
                        else:
                             # print(f"    Section not extracted from '{criteria_text}'") # Removed Debug print
                             pass # Keep pass or remove print
                    else:
                        # print(f"  Skipping row: Found only {len(cells)} cells.") # Removed Debug print
                        pass # Keep pass or remove print

                except IndexError:
                    # 处理可能的行单元格不足问题
                    # print(f"  Skipping row due to IndexError: cells length likely less than 3.") # Removed Debug print
                    continue
                except Exception as e:
                    # print(f"  Error processing row: {e}") # Removed Debug print
                    continue
        else:
             # print("DEBUG [Air Extract]: No target table was identified.") # Removed Debug print
             pass # Keep pass or remove print

    except Exception as e:
        print(f"Error opening or processing document {docx_path}: {e}")
        # 可以考虑抛出异常或返回空字典
        return {}

    return scores

def extract_wind_environment_scores(docx_path: str) -> dict:
    """
    从室外风环境模拟分析报告.docx格式的文档中提取总得分，关联到固定条文号 8.2.8。

    Args:
        docx_path: Word 文档的路径。

    Returns:
        一个字典，键是条文号 '8.2.8' (str)，值是总得分 (int)。
        例如: {'8.2.8': 7}
    """
    scores = {}
    fixed_section_number = "8.2.8"
    try:
        document = Document(docx_path)
        conclusion_found = False
        summary_score = None
        conclusion_text_blocks = [] # 存储结论部分的文本块

        # 1. 查找结论部分，允许更灵活的标题匹配
        for para in reversed(document.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            # 更灵活地匹配结论标题，例如包含"结论"或特定编号模式
            if re.search(r'^(?:\d+(?:\.\d+)*\s*)?结论', text):
                conclusion_found = True
                break # 找到结论标题的大致位置
            conclusion_text_blocks.insert(0, text) # 记录潜在的结论文本

        # 如果没找到明确的结论标题，假定最后几段是结论
        if not conclusion_found:
            conclusion_text_blocks = [p.text.strip() for p in document.paragraphs[-10:] if p.text.strip()]

        # 2. 在结论文本块中查找总结性得分语句 (更宽松的正则)
        full_conclusion_text = "\n".join(conclusion_text_blocks)
        # 匹配包含 "项目"、"得分" 和数字及"分"的模式
        summary_match = re.search(r'项目.*?得分.*?(\d+)\s*分', full_conclusion_text, re.DOTALL)
        if summary_match:
            summary_score = int(summary_match.group(1))
            # print(f"Found summary score in text: {summary_score}") # Removed Debug print

        # 3. 处理得分
        if summary_score is not None:
            scores[fixed_section_number] = summary_score
        else:
            # 备选逻辑：累加特定表格得分
            # print(f"Warning: Could not find summary sentence in conclusion text of {docx_path}.") # Removed Debug print
            total_score_from_tables = 0
            tables_processed = 0
            for table in document.tables:
                # 只处理标题行包含"达标判断表"或简单包含"得分"的表格
                if len(table.rows) > 0:
                    header_text = "\n".join(cell.text.strip() for cell in table.rows[0].cells)
                    if "达标判断表" in header_text or "得分" in header_text:
                        tables_processed += 1
                        # print(f"Processing table with header: {header_text}") # Removed Debug print
                        try:
                            # 获取"得分"列的索引
                            score_col_index = -1
                            for idx, cell_text in enumerate(table.rows[0].cells):
                                if "得分" in cell_text.strip():
                                    score_col_index = idx
                                    break

                            if score_col_index != -1:
                                # print(f"  '得分' column index: {score_col_index}") # Removed Debug print
                                for i, row in enumerate(table.rows[1:], 1): # 从第二行开始
                                    if len(row.cells) > score_col_index:
                                        score_text = row.cells[score_col_index].text.strip()
                                        score_match = re.search(r'^(\d+)\b', score_text) # 匹配开头的数字
                                        if score_match:
                                            row_score = int(score_match.group(1))
                                            # print(f"  Row {i+1}: Found score {row_score} in '{score_text}'") # Removed Debug print
                                            total_score_from_tables += row_score
                                        else:
                                             # print(f"  Row {i+1}: No score found in '{score_text}'") # Removed Debug print
                                             pass # Keep pass or remove print
                            else:
                                # print(f"  Could not find '得分' column in header.") # Removed Debug print
                                pass # Keep pass or remove print

                        except Exception as e:
                             print(f"Error processing score table: {e}")
                             continue

            if tables_processed > 0 and total_score_from_tables >= 0: #允许0分
                 # print(f"Summing scores from {tables_processed} processed table(s): {total_score_from_tables}") # Removed Debug print
                 scores[fixed_section_number] = total_score_from_tables
            else:
                 # print(f"Warning: Could not find summary sentence or valid score tables in {docx_path} for section {fixed_section_number}.") # Removed Debug print
                 pass # Keep pass or remove print


    except Exception as e:
        print(f"Error opening or processing document {docx_path}: {e}")
        return {}

    return scores

def extract_noise_report_scores(docx_path: str) -> dict:
    """
    从室外噪声分析报告.docx格式的文档中提取条文号和得分。
    尝试查找文档中最后一个出现的条文号和最后一个出现的得分，并将它们关联。

    Args:
        docx_path: Word 文档的路径。

    Returns:
        一个字典，键是条文号 (str)，值是得分 (int)，如果找不到则为空字典。
        例如: {'8.2.6': 10}
    """
    scores = {}
    try:
        document = Document(docx_path)
        # 使用 finditer 获取匹配对象，方便提取分组
        section_pattern = r'(?:第|依据|参照)\s*(\d+(?:\.\d+)*)\s*条' # Group 1: the number
        score_pattern = r'得\s*(\d+)\s*分' # Group 1: the score

        full_text = "\n".join([para.text for para in document.paragraphs])

        # 查找最后一个匹配的条文号
        last_section_match_obj = None
        for match in re.finditer(section_pattern, full_text):
            last_section_match_obj = match
        last_section_number = last_section_match_obj.group(1) if last_section_match_obj else None

        # 查找最后一个匹配的得分
        last_score_match_obj = None
        for match in re.finditer(score_pattern, full_text):
             last_score_match_obj = match
        last_score = int(last_score_match_obj.group(1)) if last_score_match_obj else None


        if last_section_number and last_score is not None:
            print(f"[Noise Extract] Found Last Section: {last_section_number}, Found Last Score: {last_score}") # Info print
            scores[last_section_number] = last_score
        else:
            print(f"[Noise Extract] Could not find both section pattern and score pattern (using last match) in document: {docx_path}")
            if not last_section_number:
                 print("[Noise Extract] Section pattern not found.")
            if last_score is None:
                 print("[Noise Extract] Score pattern not found.")

    except Exception as e:
        print(f"Error opening or processing document {docx_path}: {e}")
        return {}

    return scores

def extract_natural_ventilation_scores(docx_path: str) -> dict:
    """
    从室内自然通风模拟分析报告.docx格式的文档中提取条文号和得分。
    主要解析包含"条文"和"得分"列的表格。

    Args:
        docx_path: Word 文档的路径。

    Returns:
        一个字典，键是条文号 (str)，值是得分 (int)。
        例如: {'5.2.10': 8}
    """
    scores = {}
    try:
        document = Document(docx_path)
        target_table = None
        table_index = -1

        # 查找包含 "条文" 和 "得分" 列的表格
        for block in iter_block_items(document):
            if isinstance(block, Table):
                table_index += 1
                if len(block.rows) > 0:
                    first_row_texts = [cell.text.strip() for cell in block.rows[0].cells]
                    # 检查表头是否同时包含"条文"和"得分"
                    if any("条文" in text for text in first_row_texts) and \
                       any("得分" in text for text in first_row_texts):
                        target_table = block
                        print(f"[Ventilation Extract] Found target table {table_index} based on header.")
                        break # 假设第一个匹配的就是目标表格

        if target_table:
            # 确定"条文"和"得分"列的索引
            header_cells = target_table.rows[0].cells
            section_col_idx = -1
            score_col_idx = -1
            for idx, cell in enumerate(header_cells):
                text = cell.text.strip()
                if "条文" in text:
                    section_col_idx = idx
                if "得分" in text:
                    score_col_idx = idx

            if section_col_idx != -1 and score_col_idx != -1:
                 print(f"[Ventilation Extract] Section Col Index: {section_col_idx}, Score Col Index: {score_col_idx}")
                 # 从第二行开始遍历
                 for r_idx, row in enumerate(target_table.rows[1:], 1):
                     try:
                         cells = row.cells
                         if len(cells) > max(section_col_idx, score_col_idx):
                             section_text = cells[section_col_idx].text.strip()
                             score_text = cells[score_col_idx].text.strip()

                             # 提取条文号
                             section_match = re.search(r'^(\d+(\.\d+)*)', section_text)
                             if section_match:
                                 section_number = section_match.group(1)

                                 # 提取得分 (只接受数字)
                                 score_match = re.search(r'(\d+)\s*分?', score_text) # 允许数字后无"分"
                                 if score_match:
                                     score = int(score_match.group(1))
                                     # --- Updated Logic: Sum scores for the same section --- 
                                     if section_number in scores:
                                         scores[section_number] += score
                                     else:
                                         scores[section_number] = score
                                     # ----------------------------------------------------
                                     # print(f"  Row {r_idx}: Extracted {section_number}: {score}") # Optional debug
                                 # else: # 处理 "满足" 等文本或空单元格
                                     # print(f"  Row {r_idx}: Score text '{score_text}' is not numeric.") # Optional debug
                         # else: # 跳过条文号提取失败的行
                             # print(f"  Row {r_idx}: Section number not found in '{section_text}'") # Optional debug

                     except Exception as e:
                         print(f"[Ventilation Extract] Error processing row {r_idx}: {e}")
                         continue
            else:
                print("[Ventilation Extract] Could not find both '条文' and '得分' columns in the identified table header.")
        else:
            print("[Ventilation Extract] No target table found.")

    except Exception as e:
        print(f"Error opening or processing document {docx_path}: {e}")
        return {}

    return scores

def extract_sound_insulation_scores(docx_path: str) -> dict:
    """
    从构件隔声性能分析报告.docx格式的文档中提取条文号和总得分。
    主要查找结论部分提及"总得分"和对应条文号的文本。

    Args:
        docx_path: Word 文档的路径。

    Returns:
        一个字典，键是条文号 (str)，值是总得分 (int)。
        例如: {'5.2.7': 4}
    """
    scores = {}
    try:
        document = Document(docx_path)

        # 查找包含条文号和总得分的总结性文本，通常在末尾
        # 正则: 匹配 "依据 X.Y.Z 条 ... 总得分为 Y 分" 类似的模式
        # 允许中间有其他文字，使用 re.DOTALL 匹配换行符
        pattern = r'依据\s*(\d+(?:\.\d+)*)\s*条.*?总得分为?\s*(\d+)\s*分'

        full_text = "\n".join([para.text for para in document.paragraphs])

        # 从后向前查找，提高找到最终结论的可能性
        match = None
        # 搜索整个文本，因为总结句可能跨越多段
        potential_matches = list(re.finditer(pattern, full_text, re.DOTALL))

        if potential_matches:
            match = potential_matches[-1] # 取最后一个匹配
            section_number = match.group(1)
            total_score = int(match.group(2))
            print(f"[Insulation Extract] Found Section: {section_number}, Total Score: {total_score}")
            scores[section_number] = total_score
        else:
            # 备选：如果找不到明确的总分句，尝试查找最后一个"得X分"并关联最后一个条文号
            print("[Insulation Extract] Could not find summary sentence with section and total score. Trying fallback...")
            section_pattern_fallback = r'(\d+(?:\.\d+)*)\s*条'
            score_pattern_fallback = r'得\s*(\d+)\s*分'

            last_section_match_obj = None
            for m in re.finditer(section_pattern_fallback, full_text):
                last_section_match_obj = m
            last_section_number = last_section_match_obj.group(1) if last_section_match_obj else None

            last_score_match_obj = None
            for m in re.finditer(score_pattern_fallback, full_text):
                 last_score_match_obj = m
            last_score = int(last_score_match_obj.group(1)) if last_score_match_obj else None

            if last_section_number and last_score is not None:
                print(f"[Insulation Extract - Fallback] Found Last Section: {last_section_number}, Found Last Score: {last_score}")
                scores[last_section_number] = last_score
            else:
                 print("[Insulation Extract - Fallback] Could not find section or score pattern.")

    except Exception as e:
        print(f"Error opening or processing document {docx_path}: {e}")
        return {}

    return scores

def extract_green_material_score(docx_path: str) -> dict:
    """
    从绿色建材应用比例分析报告.docx格式的文档中提取得分。
    查找结论部分提及"绿色建材...得X分"的文本。
    将提取的分数关联到固定的条文号 '7.2.18'。

    Args:
        docx_path: Word 文档的路径。

    Returns:
        一个字典，键是 "7.2.18" (str)，值是得分 (int)。
        例如: {'7.2.18': 8}
    """
    scores = {}
    fixed_section_key = "7.2.18" # 固定条文号
    try:
        document = Document(docx_path)
        # 查找包含 "绿色建材...得 X 分" 的文本，允许得分是小数
        pattern = r'绿色建材.*?比例.*?得\s*([\d.]+)\s*分'

        full_text = "\n".join([para.text for para in document.paragraphs])

        # 查找最后一个匹配（通常结论在末尾）
        last_score_match_obj = None
        for match in re.finditer(pattern, full_text, re.DOTALL):
            last_score_match_obj = match

        if last_score_match_obj:
            score_str = last_score_match_obj.group(1)
            try:
                # 尝试将得分转为浮点数再转整数，处理 "8.0" 的情况
                score = int(float(score_str))
                print(f"[Green Material Extract] Found Score: {score} for section {fixed_section_key}")
                scores[fixed_section_key] = score
            except ValueError:
                print(f"[Green Material Extract] Found score text '{score_str}', but could not convert to int.")
        else:
            print(f"[Green Material Extract] Could not find score pattern in document: {docx_path}")

    except Exception as e:
        print(f"Error opening or processing document {docx_path}: {e}")
        return {}

    return scores

def extract_indoor_sound_evaluation_scores(docx_path: str) -> dict:
    """
    从室内声环境评价报告.docx格式的文档中提取条文号和得分。
    主要解析包含"检查项", "评价依据", "结论/得分"列的表格。
    能够处理需要累加子项得分的条文。

    Args:
        docx_path: Word 文档的路径。

    Returns:
        一个字典，键是条文号 (str)，值是得分 (int)。
        例如: {'5.2.6': 4, '5.2.7': 10}
    """
    scores = {}
    try:
        document = Document(docx_path)
        target_table = None
        target_table_conclusion_col_idx = -1 # 新增：存储目标表格的"结论/得分"列索引

        # 查找目标表格
        for table_idx, table in enumerate(document.tables):
            # print(f"[Indoor Sound Eval Debug] 正在检查表格 {table_idx}") # 可以暂时注释掉，除非需要最详细的日志
            if len(table.rows) > 0 and len(table.columns) >= 2:
                header_cells = table.rows[0].cells
                actual_header_texts_all = [cell.text.strip() for cell in header_cells]
                
                expected_h0 = "检查项"
                expected_h1 = "评价依据"
                expected_h2_keyword = "结论/得分"

                if len(actual_header_texts_all) >= 2 and \
                   expected_h0 in actual_header_texts_all[0] and \
                   expected_h1 in actual_header_texts_all[1]:
                    
                    # print(f"[Indoor Sound Eval Debug] 表格 {table_idx}: 前两列表头 ('检查项', '评价依据') 匹配。正在查找 '结论/得分'...")
                    for h_idx in range(2, len(actual_header_texts_all)):
                        if expected_h2_keyword in actual_header_texts_all[h_idx]:
                            target_table = table
                            target_table_conclusion_col_idx = h_idx # 存储找到的列索引
                            # print(f"[Indoor Sound Eval] 找到了目标表格 {table_idx}。"结论/得分"在第 {h_idx} 列 (0-indexed)。")
                            break 
                    if target_table:
                        break
        
        if not target_table:
            print(f"[Indoor Sound Eval] 未找到完整表头匹配的目标表格。")
            return scores

        current_main_section = None
        accumulated_score_for_current_section = 0
        is_accumulating = False

        # 从第二行开始遍历 (跳过标题行)
        print(f"[Indoor Sound Eval] 开始处理目标表格的行。'结论/得分'列索引为: {target_table_conclusion_col_idx}")
        for r_idx, row in enumerate(target_table.rows[1:], 1): # r_idx 是从1开始的行号（相对于数据行）
            cells = row.cells
            if len(cells) <= target_table_conclusion_col_idx: # 确保行中有足够的单元格到达"结论/得分"列
                print(f"[Indoor Sound Eval] 跳过行 {r_idx}: 单元格数量 ({len(cells)})不足以到达结论列 ({target_table_conclusion_col_idx}).")
                continue

            evaluation_basis_text = cells[1].text.strip() # "评价依据"固定在第1列 (0-indexed)
            # 使用之前确定的索引来获取"结论/得分"列的文本
            conclusion_score_text_cell = cells[target_table_conclusion_col_idx].text.strip() 
            
            print(f"[Indoor Sound Eval Row {r_idx}] 评价依据: '{evaluation_basis_text[:50]}...' | 结论/得分单元格文本: '{conclusion_score_text_cell}'")

            main_section_match = re.search(r"^\s*(\d+\.\d+\.\d+)", evaluation_basis_text)
            
            score_value = None
            score_value_match = re.search(r"(\d+)\s*分", conclusion_score_text_cell)
            if score_value_match:
                try:
                    score_value = int(score_value_match.group(1))
                    print(f"[Indoor Sound Eval Row {r_idx}] 从 '{conclusion_score_text_cell}' 提取到分数: {score_value}")
                except ValueError:
                    print(f"[Indoor Sound Eval Row {r_idx}] 提取分数时转换错误: '{score_value_match.group(1)}'")
            # else:
                # print(f"[Indoor Sound Eval Row {r_idx}] 未从 '{conclusion_score_text_cell}' 提取到分数模式。")


            if main_section_match:
                new_section = main_section_match.group(1)
                print(f"[Indoor Sound Eval Row {r_idx}] 识别到新的主条文: {new_section}")
                if current_main_section and is_accumulating:
                    scores[current_main_section] = accumulated_score_for_current_section
                    print(f"[Indoor Sound Eval] 保存条文 {current_main_section} (累加型) 得分: {accumulated_score_for_current_section}")

                current_main_section = new_section
                is_accumulating = "分别评分并累计" in evaluation_basis_text
                accumulated_score_for_current_section = 0 # 关键：新主条文开始，累加器严格清零
                print(f"[Indoor Sound Eval] 条文 {current_main_section} - 累加模式: {is_accumulating}。累加器已清零。")
                
                if score_value is not None:
                    if is_accumulating:
                        accumulated_score_for_current_section = score_value
                        print(f"[Indoor Sound Eval] 条文 {current_main_section} (累加型起点) 初始分数: {score_value} (来自其自身行)。当前累加: {accumulated_score_for_current_section}")
                    else:
                        scores[current_main_section] = score_value
                        print(f"[Indoor Sound Eval] 条文 {current_main_section} (直接型) 得分: {score_value}")
                # else: # main_section_match 为真，但 score_value is None
                    # if is_accumulating:
                    #     print(f"[Indoor Sound Eval] 条文 {current_main_section} (累加型起点) 自身行无分数。累加器保持为0。")

            elif not main_section_match and current_main_section and is_accumulating and score_value is not None:
                accumulated_score_for_current_section += score_value
                print(f"[Indoor Sound Eval Row {r_idx}] 条文 {current_main_section} (累加子项) 增加 {score_value}。当前累加: {accumulated_score_for_current_section}")
            # else:
                # print(f"[Indoor Sound Eval Row {r_idx}] 非主条文行，或非累加模式，或无分数，不进行累加操作。main_match:{main_section_match is not None}, curr_main_sec:{current_main_section}, is_acc:{is_accumulating}, score_val:{score_value}")

        if current_main_section and is_accumulating: # 确保是累加型才保存
            scores[current_main_section] = accumulated_score_for_current_section
            print(f"[Indoor Sound Eval] 循环结束. 保存最后一个累加型条文 {current_main_section} 得分: {accumulated_score_for_current_section}")
        # elif current_main_section and current_main_section not in scores:
             # 对于最后一个条文是非累加型的，它的分数应该在其行被处理时已记录
             # print(f"[Indoor Sound Eval] 循环结束. 最后一个条文 {current_main_section} (非累加型) 已记录或无分数.")


    except Exception as e:
        print(f"Error opening or processing document for Indoor Sound Evaluation {docx_path}: {e}")
        import traceback
        print(traceback.format_exc())
        return {}
        
    print(f"[Indoor Sound Eval] 最终提取分数: {scores}")
    return scores

def parse_report_scores(docx_path: str) -> dict:
    """
    自动检测 Word 文档类型并调用相应的解析函数提取得分。

    Args:
        docx_path: Word 文档的路径。

    Returns:
        一个字典，包含提取到的 {条文号: 得分}，如果无法识别类型或出错则返回空字典。
    """
    try:
        document = Document(docx_path)
        paragraphs_to_check = document.paragraphs[:50] # 检查前50段
        text_for_type_check = "\n".join([p.text for p in paragraphs_to_check])
        tables_to_check = document.tables[:30]

        # --- 风环境报告检查 ---
        wind_keywords = ["室外风环境", "风速模拟", "冬季工况", "夏季工况"]
        if any(keyword in text_for_type_check for keyword in wind_keywords):
            return extract_wind_environment_scores(docx_path)

        # --- 室内声环境评价报告检查 ---
        is_indoor_sound_eval = False 
        fn_keyword_indoor_sound_eval = "室内声环境分析报告"
        indoor_sound_eval_keywords_text = ["声环境评价结果", "GB/T 50378-2024", "室内声环境设计"] 
        indoor_sound_eval_table_headers = ["检查项", "评价依据", "结论/得分"]

        if fn_keyword_indoor_sound_eval in docx_path.lower():
            print(f"[Parser] 检测到室内声环境评价报告 (基于文件名): {docx_path}")
            found_specific_table_for_filename_match = False
            for table in tables_to_check:
                if len(table.rows) > 0 and len(table.columns) >= len(indoor_sound_eval_table_headers):
                    headers = [cell.text.strip() for cell in table.rows[0].cells]
                    if indoor_sound_eval_table_headers[0] in headers[0] and \
                       indoor_sound_eval_table_headers[1] in headers[1] and \
                       indoor_sound_eval_table_headers[2] in headers[2]:
                        found_specific_table_for_filename_match = True
                        print(f"[Parser] ...文件名匹配，且特征表格已找到。")
                        break
            if not found_specific_table_for_filename_match:
                print(f"[Parser] ...文件名匹配，但特征表格未在检查范围内找到。仍按此类型处理（基于文件名优先）。")
            is_indoor_sound_eval = True # 优先相信文件名

        elif any(keyword in text_for_type_check for keyword in indoor_sound_eval_keywords_text):
            print(f"[Parser] 文本关键词匹配室内声环境评价报告，正在确认表格: {docx_path}")
            for table in tables_to_check:
                if len(table.rows) > 0 and len(table.columns) >= len(indoor_sound_eval_table_headers):
                    headers = [cell.text.strip() for cell in table.rows[0].cells]
                    if indoor_sound_eval_table_headers[0] in headers[0] and \
                       indoor_sound_eval_table_headers[1] in headers[1] and \
                       indoor_sound_eval_table_headers[2] in headers[2]:
                        is_indoor_sound_eval = True
                        print(f"[Parser] ...文本关键词和表格结构均匹配。")
                        break
            if not is_indoor_sound_eval:
                print(f"[Parser] ...文本关键词匹配，但特征表格未找到。不按此类型处理。")
        
        if is_indoor_sound_eval:
            return extract_indoor_sound_evaluation_scores(docx_path)
        # --- 室内声环境评价报告检查结束 ---

        # --- 构件隔声性能报告检查 ---
        insulation_keywords = ["构件隔声", "空气声", "撞击声"]
        insulation_table_keywords = ["得分统计表"] 
        is_insulation = False
        
        text_check_passed_ins = any(keyword in text_for_type_check for keyword in insulation_keywords)
        
        if text_check_passed_ins:
            print(f"[Parser] 文本关键词匹配构件隔声性能报告，正在确认表格: {docx_path}")
            found_insulation_table = False
            for table in tables_to_check: # Iterate through tables_to_check
                 table_text_lower = "\n".join(cell.text.strip().lower() for row in table.rows for cell in row.cells)
                 if any(keyword.lower() in table_text_lower for keyword in insulation_table_keywords):
                     found_insulation_table = True
                     break
            if found_insulation_table:
                is_insulation = True
                print(f"[Parser] ...文本关键词和特征表格均匹配构件隔声性能报告。")
            else:
                print(f"[Parser] ...文本关键词匹配构件隔声性能报告，但特征表格未找到。不按此类型处理。")
        
        if is_insulation:
            return extract_sound_insulation_scores(docx_path)
        # --- 构件隔声性能报告检查结束 ---
        
        # --- 绿色建材报告检查 ---
        green_material_keywords = ["绿色建材", "建材应用比例"]
        is_green_material = False
        if any(keyword in text_for_type_check for keyword in green_material_keywords):
            is_green_material = True
            # 这种报告格式简单，文本关键词足够判断，无需检查表格

        if is_green_material:
            print(f"Detected report type: Green Material for {docx_path}")
            return extract_green_material_score(docx_path) # 调用新函数

        # --- 噪声报告检查 ---
        # 使用更精确的关键词，移除 "声环境"
        noise_keywords = ["室外噪声", "环境噪声", "室外声环境"]
        noise_table_keywords = ["噪声最大值", "噪声限值", "得分情况"]
        is_noise = False
        text_check_passed = any(keyword in text_for_type_check for keyword in noise_keywords)
        # print(f"DEBUG [Noise]: Text check passed: {text_check_passed}") # Remove Debug print
        if text_check_passed:
            is_noise = True
        else:
            # print(f"DEBUG [Noise]: Checking tables...") # Remove Debug print
            for i, table in enumerate(tables_to_check):
                 if len(table.rows) > 0:
                    header_text = "\n".join(cell.text.strip() for cell in table.rows[0].cells) # strip() added
                    # print(f"DEBUG [Noise]: Table {i} header: {header_text}") # Remove Debug print
                    table_keywords_found = any(keyword in header_text for keyword in noise_table_keywords)
                    # print(f"DEBUG [Noise]: Table {i} - Table keywords found: {table_keywords_found}") # Remove Debug print
                    if table_keywords_found:
                         is_noise = True
                         # print(f"DEBUG [Noise]: Matched table {i}") # Remove Debug print
                         break
                 else:
                     # print(f"DEBUG [Noise]: Table {i} has no rows.") # Remove Debug print
                     pass # Keep pass or remove print
            # if not is_noise:
            #      print(f"DEBUG [Noise]: Table check failed.") # Remove Debug print

        if is_noise:
            # print(f"Detected report type: Noise Report for {docx_path}") # Keep useful print?
            return extract_noise_report_scores(docx_path)

        # --- 空气质量报告检查 ---
        air_quality_keywords = ["室内空气质量评价", "空气污染物浓度", "空气质量分析报告"] # 添加关键词
        air_quality_table_keywords = ["评价依据", "结论/得分"]
        is_air_quality = False
        text_check_passed = any(keyword in text_for_type_check for keyword in air_quality_keywords)
        # print(f"DEBUG [Air Quality]: Text check passed: {text_check_passed}") # Removed Debug print
        if text_check_passed:
            is_air_quality = True
        else:
            # print(f"DEBUG [Air Quality]: Checking tables...") # Removed Debug print
            for i, table in enumerate(tables_to_check):
                if len(table.rows) > 0:
                    header_text = "\n".join(cell.text.strip() for cell in table.rows[0].cells) # strip() added
                    # print(f"DEBUG [Air Quality]: Table {i} header: {header_text}") # Removed Debug print
                    table_keywords_found = any(keyword in header_text for keyword in air_quality_table_keywords)
                    specific_keyword_found = "室内空气质量" in header_text
                    # print(f"DEBUG [Air Quality]: Table {i} - Table keywords found: {table_keywords_found}, Specific keyword found: {specific_keyword_found}") # Removed Debug print
                    if table_keywords_found and specific_keyword_found:
                        is_air_quality = True
                        # print(f"DEBUG [Air Quality]: Matched table {i}") # Removed Debug print
                        break
                else:
                    # print(f"DEBUG [Air Quality]: Table {i} has no rows.") # Removed Debug print
                    pass # Keep pass or remove print
            # if not is_air_quality:
            #     print(f"DEBUG [Air Quality]: Table check failed.") # Removed Debug print

        if is_air_quality:
            # print(f"Detected report type: Air Quality for {docx_path}") # Keep this potentially useful info?
            return extract_air_quality_scores(docx_path)

        # --- 自然通风报告检查 ---
        ventilation_keywords = ["室内自然通风", "通风模拟", "换气次数"]
        ventilation_table_keywords = ["条文", "得分"]
        is_ventilation = False
        if any(keyword in text_for_type_check for keyword in ventilation_keywords):
            is_ventilation = True
        else:
            # 检查表格是否包含特定列名
            for table in tables_to_check:
                 if len(table.rows) > 0:
                    header_text = "\n".join(cell.text.strip() for cell in table.rows[0].cells)
                    if all(keyword in header_text for keyword in ventilation_table_keywords):
                         is_ventilation = True
                         break

        if is_ventilation:
             print(f"Detected report type: Natural Ventilation for {docx_path}")
             return extract_natural_ventilation_scores(docx_path) # 调用新函数

        # --- 室内污染物浓度报告检查 (复用自然通风的提取逻辑) ---
        pollutant_keywords = ["污染物浓度", "室内空气质量"]
        # 表格结构与自然通风一致
        is_pollutant_report = False
        # 优先检查文本关键词，同时结合表格结构确认
        if any(keyword in text_for_type_check for keyword in pollutant_keywords):
            # 进一步确认表格结构是否匹配
            for table in tables_to_check:
                 if len(table.rows) > 0:
                    header_text = "\n".join(cell.text.strip() for cell in table.rows[0].cells)
                    # 使用与自然通风相同的表格关键词
                    if all(keyword in header_text for keyword in ventilation_table_keywords):
                         is_pollutant_report = True
                         break

        if is_pollutant_report:
            print(f"Detected report type: Indoor Pollutant (using ventilation extractor) for {docx_path}")
            return extract_natural_ventilation_scores(docx_path) # 复用提取函数

        # --- 背景噪声报告检查 (复用自然通风的提取逻辑) ---
        # IMPORTANT: This generic background noise check should be AFTER the specific indoor_sound_evaluation_scores check
        # if there's a risk of "室内声环境" keyword overlap and different parsing logic.
        background_noise_keywords = ["背景噪声", "室内声环境"] # "室内声环境"关键词可能与上面冲突
        # ... (后续的背景噪声检查逻辑保持不变) ...

        # --- 未能识别 ---
        print(f"Warning: Could not determine report type for {docx_path}.")
        return {}

    except Exception as e:
        print(f"Error processing document {docx_path} in main parser: {e}")
        return {}

if __name__ == '__main__':
    # 创建一个测试函数，方便本地调试
    test_file_path = '../test/室内声环境分析报告.docx' # 请确保路径正确
    # 注意: 这里的相对路径是相对于 utils 目录的

    # 为了在工作空间根目录运行脚本时也能找到测试文件，使用绝对路径或更健壮的路径处理
    import os
    workspace_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__))) # 获取项目根目录
    test_file_path_abs = os.path.join(workspace_root, 'test', '室内声环境分析报告.docx')

    if os.path.exists(test_file_path_abs):
        extracted_scores = parse_report_scores(test_file_path_abs)
        print(f"从 {test_file_path_abs} 提取的得分:")
        print(extracted_scores)
    else:
        print(f"测试文件未找到: {test_file_path_abs}")

   