import os
import re
import shutil
import traceback
import base64
import io
import sys
import platform
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io
from flask import jsonify, send_file
from models import Project
import logging

logger = logging.getLogger('greenscore')

def replace_placeholders_with_format(paragraph, project_info):
    """
    在保留格式的情况下替换段落中的占位符（除项目概况外的其他占位符）
    
    参数:
        paragraph: 要处理的段落对象
        project_info: 包含替换值的字典
    """
    # 使用正则表达式找出所有占位符
    pattern = r'\{([^{}]+)\}'
    matches = re.finditer(pattern, paragraph.text)
    
    # 如果没有找到占位符，直接返回
    if not re.search(pattern, paragraph.text):
        return
        
    # 对于找到的每个占位符，保留段落的格式进行替换
    for match in list(re.finditer(pattern, paragraph.text)):
        placeholder = match.group(0)  # 完整的占位符，例如 {项目名称}
        key = match.group(1)  # 占位符中的键，例如 项目名称
        
        # 跳过项目概况占位符和示意图占位符，它们有专门的处理方法
        if key == "项目概况" or key == "示意图":
            continue
            
        if key in project_info and project_info[key]:
            # 获取占位符的起始位置和结束位置
            start = match.start()
            end = match.end()
            
            # 在指定位置替换文本
            runs = paragraph.runs
            
            # 记录当前的位置和处理的字符数
            current_pos = 0
            for run in runs:
                run_len = len(run.text)
                run_start = current_pos
                run_end = run_start + run_len
                
                # 检查这个run是否包含占位符的部分
                if run_end > start and run_start < end:
                    # 计算占位符在这个run中的相对位置
                    rel_start = max(0, start - run_start)
                    rel_end = min(run_len, end - run_start)
                    
                    # 替换这个run中的占位符部分
                    run.text = run.text[:rel_start] + str(project_info[key]) + run.text[rel_end:]
                
                current_pos += run_len
                
                # 替换完成后退出循环
                if current_pos > end:
                    break
        else:
            # 如果没有对应的值，则将占位符替换为空字符串
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, "")

def replace_placeholders_simple(paragraph, project_info):
    """
    使用简单的字符串替换方法替换段落中的项目概况占位符
    
    参数:
        paragraph: 要处理的段落对象
        project_info: 包含替换值的字典
    
    返回:
        是否成功替换了占位符
    """
    # 记录替换前的文本
    original_text = paragraph.text
    
    # 检查是否包含项目概况占位符
    placeholder = "{项目概况}"
    if placeholder in original_text and "项目概况" in project_info:
        logger.info(f"发现项目概况占位符: {original_text}")
        
        # 方法1: 尝试直接在每个run中替换文本
        replaced = False
        for run in paragraph.runs:
            if placeholder in run.text:
                # 替换run中的占位符
                run.text = run.text.replace(placeholder, project_info["项目概况"])
                replaced = True
        
        # 方法2: 如果方法1未成功（占位符跨runs），尝试清空所有runs并重设内容
        if not replaced and placeholder in paragraph.text:
            logger.info("占位符可能跨越多个run，尝试重构整个段落")
            
            # 保存第一个run的样式
            if len(paragraph.runs) > 0:
                first_run = paragraph.runs[0]
                # 清空所有runs
                for run in paragraph.runs:
                    run.text = ""
                
                # 将替换后的文本设置到第一个run
                first_run.text = original_text.replace(placeholder, project_info["项目概况"])
                replaced = True
        
        # 方法3: 如果上述方法都失败，则尝试直接替换段落文本
        if not replaced:
            logger.info("尝试使用更直接的方法重建段落")
            # 清空段落
            for run in paragraph.runs:
                run.text = ""
            
            # 创建新的run并设置替换后的文本
            new_text = original_text.replace(placeholder, project_info["项目概况"])
            new_run = paragraph.add_run(new_text)
        
        return True
    
    return False

def replace_image_placeholders(paragraph, doc, project_info):
    """
    替换段落中的示意图占位符为实际图片
    
    参数:
        paragraph: 要处理的段落对象
        doc: Word文档对象
        project_info: 包含替换值的字典（包括图片数据）
    
    返回:
        是否成功替换了占位符
    """
    # 先检查是否为编号式示意图占位符（如 {示意图1}、{示意图2} 等）
    numbered_placeholder_pattern = r'\{示意图(\d+)\}'
    numbered_match = re.search(numbered_placeholder_pattern, paragraph.text)
    
    if numbered_match:
        img_num = numbered_match.group(1)
        placeholder = f"{{示意图{img_num}}}"
        img_key = f"示意图{img_num}"
        
        if img_key in project_info:
            logger.info(f"发现编号示意图占位符: {placeholder}")
            return process_image_placeholder(paragraph, placeholder, project_info[img_key])
    
    # 检查是否包含常规示意图占位符
    placeholder = "{示意图}"
    if placeholder in paragraph.text:
        logger.info(f"发现示意图占位符: {paragraph.text}")
        
        # 获取所有图片数据
        image_data_list = []
        
        # 检查常规示意图键
        if "示意图" in project_info:
            image_data_list.append(project_info["示意图"])
        
        # 检查示意图数组键（只添加索引1及以后的图片，避免重复添加第一张图片）
        if "示意图数组" in project_info and isinstance(project_info["示意图数组"], list) and len(project_info["示意图数组"]) > 1:
            # 从索引1开始添加，避免重复添加第一张图片
            image_data_list.extend(project_info["示意图数组"][1:])
            logger.info(f"从示意图数组添加了{len(project_info['示意图数组'])-1}张图片")
        
        # 如果找到图片数据
        if image_data_list:
            # 获取父元素和当前段落的索引
            parent = paragraph._p.getparent()
            idx = parent.index(paragraph._p)
            
            # 先处理第一张图片，替换原占位符段落
            success = process_image_placeholder(paragraph, placeholder, image_data_list[0])
            
            # 如果有更多图片，就在原段落后创建新段落并添加图片
            for i in range(1, len(image_data_list)):
                # 创建新段落
                new_para = doc.add_paragraph()
                
                # 添加图片
                process_image_placeholder(new_para, placeholder, image_data_list[i])
                
                # 移动新段落到正确的位置（在原段落之后）
                new_p = new_para._p
                parent.remove(new_p)
                parent.insert(idx + i, new_p)
            
            return success
    
    return False

def process_image_placeholder(paragraph, placeholder, image_data):
    """
    处理示意图占位符，将其替换为图片
    
    参数:
        paragraph: 包含占位符的段落
        placeholder: 占位符文本
        image_data: 图片数据（Base64编码）
    
    返回:
        处理是否成功
    """
    try:
        # 检查图片数据是否为Data URL格式（例如data:image/jpeg;base64,/9j/...）
        if image_data.startswith('data:'):
            # 提取MIME类型和Base64编码的实际内容
            mime_type, base64_data = image_data.split(',', 1)
            # 不需要MIME类型，只需Base64内容
        else:
            base64_data = image_data
            
        # 解码Base64数据
        image_bytes = base64.b64decode(base64_data)
        image_stream = io.BytesIO(image_bytes)
        
        # 清空原有段落内容
        for run in paragraph.runs:
            run.text = ""
        
        # 添加图片到段落
        run = paragraph.add_run()
        run.add_picture(image_stream, width=Cm(8))  # 默认宽度为8厘米，可以根据需要调整
        
        # 设置段落对齐方式为居中
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        return True
    except Exception as e:
        logger.error(f"处理示意图时出错: {str(e)}")
        logger.error(traceback.format_exc())
        
        # 如果处理失败，保留原占位符或清空它
        for run in paragraph.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, "")
        
        return False

def process_document_placeholders(doc, project_info, app=None):
    """
    处理文档中的所有占位符
    
    参数:
        doc: Word文档对象
        project_info: 包含替换值的字典
        app: Flask应用实例（可选）
    """
    # 替换段落中的项目概况占位符
    overview_count = 0
    for para in doc.paragraphs:
        if "{项目概况}" in para.text:
            if replace_placeholders_simple(para, project_info):
                overview_count += 1
    
    # 替换表格中的项目概况占位符
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if "{项目概况}" in para.text:
                        if replace_placeholders_simple(para, project_info):
                            overview_count += 1
    
    # 替换段落中的示意图占位符
    image_count = 0
    
    # 查找所有段落中的示意图占位符（包括编号式占位符）
    pattern = r'\{示意图(\d*)\}'
    
    for para in doc.paragraphs:
        if re.search(pattern, para.text):
            if replace_image_placeholders(para, doc, project_info):
                image_count += 1
    
    # 替换表格中的示意图占位符
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if re.search(pattern, para.text):
                        if replace_image_placeholders(para, doc, project_info):
                            image_count += 1
    
    # 处理其他占位符
    for para in doc.paragraphs:
        replace_placeholders_with_format(para, project_info)
    
    # 处理表格中的其他占位符
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_placeholders_with_format(para, project_info)
    
    if app and overview_count > 0:
        app.logger.info(f"已处理 {overview_count} 处项目概况内容")
    elif overview_count > 0:
        logger.info(f"已处理 {overview_count} 处项目概况内容")
    
    if app and image_count > 0:
        app.logger.info(f"已处理 {image_count} 处示意图")
    elif image_count > 0:
        logger.info(f"已处理 {image_count} 处示意图")
    
    return overview_count

def update_toc(doc_path, app=None):
    """
    更新Word文档的目录
    
    参数:
        doc_path: Word文档路径
        app: Flask应用实例（可选），用于日志记录
    
    返回:
        更新是否成功
    """
    if platform.system() != 'Windows':
        if app:
            app.logger.warning("目录更新功能仅在Windows环境下可用")
        else:
            logger.warning("目录更新功能仅在Windows环境下可用")
        return False
    
    try:
        import win32com.client
        import pythoncom
        
        if app:
            app.logger.info(f"正在更新文档目录: {doc_path}")
        else:
            logger.info(f"正在更新文档目录: {doc_path}")
        
        # 初始化COM环境
        pythoncom.CoInitialize()
        
        try:
            # 使用Word COM对象打开文档并更新目录
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False  # 后台运行
            
            # 打开文档
            doc = word.Documents.Open(os.path.abspath(doc_path))
            
            # 更新所有域（包括页码、交叉引用等）
            word.ActiveDocument.Fields.Update()
            
            # 检查文档是否包含目录
            if doc.TablesOfContents.Count >= 1:
                # 更新第一个目录
                doc.TablesOfContents(1).Update()
                if app:
                    app.logger.info("目录已成功更新")
                else:
                    logger.info("目录已成功更新")
            else:
                if app:
                    app.logger.info("文档中未找到目录")
                else:
                    logger.info("文档中未找到目录")
            
            # 保存文档
            doc.Close(SaveChanges=True)
            word.Quit()
            
            return True
        finally:
            # 清理COM环境，确保在所有操作完成后调用
            pythoncom.CoUninitialize()
            
    except ImportError:
        if app:
            app.logger.warning("无法导入win32com库，请确保在Windows环境下安装了pywin32")
        else:
            logger.warning("无法导入win32com库，请确保在Windows环境下安装了pywin32")
        return False
    except Exception as e:
        error_msg = f"更新文档目录时出错: {str(e)}"
        if app:
            app.logger.error(error_msg)
            app.logger.error(traceback.format_exc())
        else:
            logger.error(error_msg)
            logger.error(traceback.format_exc())
        return False

def generate_decorative_cost_report_doc(data, app=None):
    """
    生成装饰性构件造价比例计算书
    
    参数:
        data: 包含项目信息和计算数据的字典
        app: Flask应用实例，用于日志记录
    
    返回:
        生成的文档路径或错误信息
    """
    try:
        if not data:
            return None, '没有接收到数据'
            
        template_file = data.get('templateFile', '装饰性构件造价比例计算书.docx')
        project_info = data.get('projectInfo', {})
        rows = data.get('rows', [])
        table_rows = data.get('tableRows', [])  # 获取新的表格数据格式
        project_id = data.get('projectId')
        
        if app:
            app.logger.info(f"生成装饰性构件造价比例计算书: 项目ID={project_id}")
        else:
            logger.info(f"生成装饰性构件造价比例计算书: 项目ID={project_id}")
        
        # 如果未提供项目信息但提供了项目ID，尝试获取项目信息
        if (not project_info or len(project_info) == 0) and project_id:
            project = Project.query.get(project_id)
            if project:
                project_info = {
                    '项目名称': project.name if project.name else '',
                    '项目地点': project.location if project.location else '',
                    '建设单位': project.construction_unit if project.construction_unit else '',
                    '设计单位': project.design_unit if project.design_unit else '',
                    '总建筑面积': project.total_building_area if project.total_building_area else '',
                }
                if app:
                    app.logger.info(f"已从数据库获取项目信息: {project_info}")
                else:
                    logger.info(f"已从数据库获取项目信息: {project_info}")
        
        # 构建模板文件路径
        template_path = os.path.join('static', 'templates', template_file)
        abs_template_path = os.path.abspath(template_path)
        
        if not os.path.exists(abs_template_path):
            # 尝试在不同位置查找模板文件
            alt_paths = [
                template_file,
                os.path.join('templates', template_file),
                os.path.join('static', template_file),
                os.path.join(os.path.dirname(os.path.dirname(__file__)), 'static', 'templates', template_file),
            ]
            
            for path in alt_paths:
                if os.path.exists(path):
                    template_path = path
                    abs_template_path = os.path.abspath(path)
                    if app:
                        app.logger.info(f"在替代位置找到模板文件: {abs_template_path}")
                    else:
                        logger.info(f"在替代位置找到模板文件: {abs_template_path}")
                    break
            else:
                error_msg = f"所有可能的模板路径均不存在"
                if app:
                    app.logger.error(error_msg)
                else:
                    logger.error(error_msg)
                return None, f'模板文件 {template_file} 不存在'

        # 确保临时目录存在
        os.makedirs('temp', exist_ok=True)
        
        # 复制模板文件到临时位置，以免修改原始模板
        output_filename = f"{project_info.get('项目名称', '项目')}-装饰性构件造价比例计算书.docx"
        temp_output_path = os.path.join('temp', output_filename)
        shutil.copy2(abs_template_path, temp_output_path)
        
        # 打开复制的文件进行编辑
        doc = Document(temp_output_path)
        
        # 生成项目概况文本
        overview_text = f"本项目规划建设用地面积为{project_info.get('总用地面积', '')}㎡，"
        overview_text += f"总建筑面积为{project_info.get('总建筑面积', '')}㎡，"
        overview_text += f"容积率为{project_info.get('容积率', '')}，"
        overview_text += f"绿地率为{project_info.get('绿地率', '')}%，"
        overview_text += f"建筑密度{project_info.get('建筑密度', '')}%。"
        
        # 将生成的文本添加到项目信息词典中
        project_info['项目概况'] = overview_text
        
        # 添加当前日期作为设计日期
        current_date = datetime.now().strftime('%Y年%m月%d日')
        project_info['设计日期'] = current_date
        
        # 处理行数据中的图片，添加到项目信息中提供给占位符替换函数
        if rows and len(rows) > 0:
            # 收集所有图片数据到数组
            all_images = []
            img_count = 0
            
            for row in rows:
                if 'imageData' in row and row['imageData']:
                    img_count += 1
                    all_images.append(row['imageData'])
                    # 保留编号的图片数据以兼容其他代码
                    project_info[f'示意图{img_count}'] = row['imageData']
                    
                    # 如果是第一张图片，也用常规键保存
                    if img_count == 1:
                        project_info['示意图'] = row['imageData']
                    
                    if app:
                        app.logger.info(f"从行数据中获取图片信息 #{img_count}")
                    else:
                        logger.info(f"从行数据中获取图片信息 #{img_count}")
            
            # 如果有多张图片，保存到示意图数组字段
            if len(all_images) > 1:
                project_info['示意图数组'] = all_images
                if app:
                    app.logger.info(f"共收集了 {len(all_images)} 张图片")
                else:
                    logger.info(f"共收集了 {len(all_images)} 张图片")
        
        # 处理文档中的所有占位符
        process_document_placeholders(doc, project_info, app)
        
        # 处理表格
        found_calc_table = False
        table_placeholder_patterns = ['{计算表}', '【计算表】', '{表格}', '【表格】']
                    
        # 找到 "综上对本项目装饰性构件所占比例汇总如下" 这句话所在的段落
        table_insert_para = None
        for i, para in enumerate(doc.paragraphs):
            if '综上对本项目装饰性构件所占比例汇总如下' in para.text:
                table_insert_para = para
                found_calc_table = True
                logger.info(f"找到表格插入位置: 段落{i+1}")
                break
        
        # 首先，尝试找到显式的表格占位符
        if not found_calc_table:
            for para in doc.paragraphs:
                for pattern in table_placeholder_patterns:
                    if pattern in para.text:
                        table_insert_para = para
                        found_calc_table = True
                        logger.info(f"找到表格占位符: {para.text}")
                        break
                if found_calc_table:
                    break
        
        # 如果找到了表格插入位置，插入表格
        if found_calc_table and table_insert_para and table_rows:
            # 保存段落的样式属性
            para_style = table_insert_para.style
            
            # 创建新的计算表
            new_table = doc.add_table(rows=1, cols=4)
            new_table.style = 'Table Grid'  # 使用表格网格样式
            
            # 设置表头
            header_cells = new_table.rows[0].cells
            header_texts = ["子项名称", "装饰性构件造价（万元）", "单栋建筑总造价（万元）", "装饰性构件造价占单栋建筑总造价的比例（%）"]
            
            # 应用表头样式和文本
            for cell, text in zip(header_cells, header_texts):
                cell_para = cell.paragraphs[0]
                cell_para.style = para_style  # 使用原占位符段落的样式
                cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = cell_para.add_run(text)
                run.bold = True
                
            # 添加数据行
            for row_data in table_rows:
                row_cells = new_table.add_row().cells
                
                # 子项名称
                para = row_cells[0].paragraphs[0]
                para.style = para_style
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.add_run(row_data.get('subItem', ''))
                
                # 装饰性构件造价
                para = row_cells[1].paragraphs[0]
                para.style = para_style
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.add_run(str(row_data.get('decorativeCost', '0.00')))
                
                # 单栋建筑总造价
                para = row_cells[2].paragraphs[0]
                para.style = para_style
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.add_run(str(row_data.get('totalCost', '0.00')))
                
                # 占比
                para = row_cells[3].paragraphs[0]
                para.style = para_style
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.add_run(str(row_data.get('percentage', '0.00%')))
            
            # 获取表格的XML并保存
            calculation_table_xml = new_table._element
            
            # 插入位置处理
            parent = table_insert_para._element.getparent()
            index = parent.index(table_insert_para._element)
            
            # 在段落后插入表格（不删除原段落）
            parent.insert(index + 1, calculation_table_xml)
        
        # 保存文档
        doc.save(temp_output_path)
        
        # 尝试更新文档目录
        update_toc(temp_output_path, app)
        
        return temp_output_path, None
    
    except Exception as e:
        error_msg = f"生成装饰性构件造价比例计算书失败: {str(e)}"
        if app:
            app.logger.error(error_msg)
            app.logger.error(traceback.format_exc())
        else:
            logger.error(error_msg)
            logger.error(traceback.format_exc())
        return None, str(e) 