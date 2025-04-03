import os
import re
from loguru import logger
from .extract_word_info import extract_project_info

def extract_doc_info(file_path):
    """
    从Word文档中提取项目信息的通用封装
    
    Args:
        file_path (str): Word文档的路径（.doc或.docx格式）
    
    Returns:
        dict: 包含提取出的项目信息
    """
    if not file_path:
        logger.error(f"文件路径为空")
        return None
        
    if not os.path.exists(file_path):
        logger.error(f"文件不存在: {file_path}")
        return None
    
    try:
        # 检查文件类型
        file_ext = os.path.splitext(file_path)[1].lower()
        
        # 处理没有扩展名的情况
        if not file_ext:
            logger.warning(f"文件没有扩展名: {file_path}，尝试作为.docx处理")
            # 尝试将其视为.docx文件处理
            file_ext = '.docx'
        
        if file_ext not in ['.doc', '.docx']:
            logger.error(f"不支持的文件格式: {file_ext}，仅支持.doc和.docx")
            return None
        
        # 调用详细提取函数
        info = extract_project_info(file_path)
        
        if info:
            # 标准化数据
            clean_info = {}
            for key, value in info.items():
                # 确保value是字符串
                if value is None:
                    value = ""
                    
                # 移除数字中的逗号
                if key in ["建筑面积", "建筑高度"]:
                    clean_info[key] = str(value).replace(',', '')
                else:
                    clean_info[key] = str(value)
            
            # 地下层数处理：如果没有地下层数，设为"0"
            if "地下层数" in clean_info and (not clean_info["地下层数"] or clean_info["地下层数"].strip() == ""):
                clean_info["地下层数"] = "0"
            
            # 检查表格中是否明确存在项目名称字段但值为空
            is_empty_project_name = False
            if "项目名称" in info and (info["项目名称"] is None or info["项目名称"].strip() == ""):
                # 如果提取的原始信息中项目名称字段存在但为空，标记为真正的空项目名称
                is_empty_project_name = True
                logger.warning("检测到文档中项目名称字段存在但内容为空")
            
            # 只有在非空项目名称的情况下才尝试补充
            if not is_empty_project_name and (not clean_info.get("项目名称") or len(clean_info.get("项目名称", "")) < 3):
                # 尝试从文件名中提取项目名称
                filename = os.path.basename(file_path)
                filename_without_ext = os.path.splitext(filename)[0]
                # 清理文件名中的常见词汇
                cleaned_filename = re.sub(r'(节能设计|规定性指标|计算报告书|审图答复|专篇|设计说明|建筑节能|\(.*?\))', '', filename_without_ext).strip()
                if cleaned_filename and len(cleaned_filename) > 3:
                    clean_info["项目名称"] = cleaned_filename
                # 如果有建设单位信息但没有项目名称，可以使用建设单位信息
                elif clean_info.get("建设单位") and len(clean_info.get("建设单位", "")) > 3:
                    clean_info["项目名称"] = clean_info["建设单位"] + "建设项目"
            
            # 确保所有项都有值，即使是空字符串
            for key in ["项目名称", "项目地点", "气候分区", "建筑分类", "结构形式", 
                        "建筑朝向", "建筑面积", "建筑层数", "地下层数", "建筑高度", 
                        "设计单位", "建设单位"]:
                if key not in clean_info:
                    clean_info[key] = ""
            
            # 如果检测到确实是空项目名称，则保持为空
            if is_empty_project_name:
                clean_info["项目名称"] = ""
            
            logger.info(f"成功从文件 {os.path.basename(file_path)} 提取信息: {clean_info}")
            return clean_info
        else:
            # 如果提取失败，尝试使用更通用的方法从文件中提取信息
            logger.warning(f"使用标准方法从 {os.path.basename(file_path)} 提取信息失败，尝试通用方法")
            generic_info = extract_generic_info(file_path)
            if generic_info and any(generic_info.values()):
                logger.info(f"通用方法成功从 {os.path.basename(file_path)} 提取信息: {generic_info}")
                return generic_info
            
            logger.error(f"无法从文件 {os.path.basename(file_path)} 提取任何信息")
            return None
    
    except Exception as e:
        logger.error(f"提取文档信息时出错: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return None

def extract_generic_info(file_path):
    """
    通用的提取方法，用于处理非标准格式的Word文档
    
    Args:
        file_path (str): Word文档的路径
    
    Returns:
        dict: 包含提取出的项目信息
    """
    generic_info = {
        "项目名称": "",
        "项目地点": "",
        "气候分区": "",
        "建筑分类": "",
        "结构形式": "",
        "建筑朝向": "",
        "建筑面积": "",
        "建筑层数": "",
        "地下层数": "0",  # 默认值
        "建筑高度": "",
        "设计单位": "",
        "建设单位": ""
    }
    
    try:
        file_ext = os.path.splitext(file_path)[1].lower()
        
        # 从文件名中提取项目名称
        filename = os.path.basename(file_path)
        filename_without_ext = os.path.splitext(filename)[0]
        # 清理常见词汇
        cleaned_filename = re.sub(r'(节能设计|规定性指标|计算报告书|审图答复|专篇|设计说明|建筑节能|\(.*?\))', '', filename_without_ext).strip()
        if cleaned_filename and len(cleaned_filename) > 3:
            generic_info["项目名称"] = cleaned_filename
        
        # 处理.docx文件
        if file_ext == '.docx':
            from docx import Document
            try:
                doc = Document(file_path)
                # 搜索文档中所有可能的关键字
                search_in_docx(doc, generic_info)
            except Exception as e:
                logger.error(f"处理.docx文件时出错: {str(e)}")
        
        # 处理.doc文件
        elif file_ext == '.doc' and os.name == 'nt':  # 仅在Windows下支持.doc
            try:
                import win32com.client
                import pythoncom
                
                # 初始化COM组件
                pythoncom.CoInitialize()
                
                try:
                    word = win32com.client.Dispatch("Word.Application")
                    word.Visible = False
                    doc = word.Documents.Open(os.path.abspath(file_path))
                    
                    # 提取文本
                    all_text = []
                    for i in range(1, doc.Paragraphs.Count + 1):
                        all_text.append(doc.Paragraphs(i).Range.Text.strip())
                    
                    # 搜索关键字
                    search_in_text(all_text, generic_info)
                    
                    # 关闭文档和Word应用
                    doc.Close(False)
                    word.Quit()
                finally:
                    # 确保释放COM资源
                    pythoncom.CoUninitialize()
            except Exception as e:
                logger.error(f"处理.doc文件时出错: {str(e)}")
        
        # 确保地下层数有值
        if not generic_info["地下层数"]:
            generic_info["地下层数"] = "0"
            
        return generic_info
        
    except Exception as e:
        logger.error(f"通用提取方法出错: {str(e)}")
        return generic_info

def search_in_docx(doc, info):
    """
    从docx文档中搜索关键信息
    
    Args:
        doc: Document对象
        info: 要填充的信息字典
    """
    # 从段落中提取
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    search_in_text(paragraphs, info)
    
    # 从表格中提取
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) >= 2:
                # 提取第一列单元格内容，去除冒号
                key = cells[0].strip().rstrip(':')
                value = cells[1].strip()
                
                # 匹配关键字
                if key == "项目名称" or key == "工程名称":
                    if value and len(value) > 3:
                        info["项目名称"] = value
                elif key == "项目地点" or key == "工程地点" or key == "建设地点":
                    info["项目地点"] = value
                elif key == "气候分区" or key == "气候区划":
                    info["气候分区"] = value
                elif key == "建筑类型" or key == "建筑分类":
                    info["建筑分类"] = value
                elif key == "结构形式" or key == "结构类型":
                    info["结构形式"] = value
                elif key == "建筑朝向" or key == "北向角度":
                    info["建筑朝向"] = value
                elif key == "建筑面积" or key == "总建筑面积":
                    match = re.search(r"([0-9,.]+)", value)
                    if match:
                        info["建筑面积"] = match.group(1)
                elif key == "建筑层数" or key == "层数":
                    info["建筑层数"] = value
                    # 尝试提取地下层数
                    underground_match = re.search(r"地下[：:\s]*(\d+)[层\s]", value)
                    if underground_match:
                        info["地下层数"] = underground_match.group(1)
                elif key == "地下层数":
                    match = re.search(r"(\d+)", value)
                    if match:
                        info["地下层数"] = match.group(1)
                elif key == "建筑高度":
                    match = re.search(r"([0-9,.]+)", value)
                    if match:
                        info["建筑高度"] = match.group(1)
                elif key == "设计单位":
                    info["设计单位"] = value
                elif key == "建设单位":
                    info["建设单位"] = value

def search_in_text(paragraphs, info):
    """
    从文本段落中搜索关键信息
    
    Args:
        paragraphs: 文本段落列表
        info: 要填充的信息字典
    """
    for text in paragraphs:
        text = text.strip()
        if not text:
            continue
        
        # 项目名称
        if ("项目名称" in text or "工程名称" in text) and not info["项目名称"]:
            match = re.search(r"(?:项目名称|工程名称)[：:]\s*(.+?)(?:\s*$|，|,|；|;|\(|（)", text)
            if match:
                info["项目名称"] = match.group(1).strip()
            else:
                match = re.search(r"(?:项目名称|工程名称)[：:]\s*(.+)", text)
                if match:
                    info["项目名称"] = match.group(1).strip()
        
        # 项目地点
        if ("项目地点" in text or "工程地点" in text or "建设地点" in text) and not info["项目地点"]:
            match = re.search(r"(?:项目地点|工程地点|建设地点)[：:]\s*(.+?)(?:\s|$|，|,)", text)
            if match:
                info["项目地点"] = match.group(1).strip()
        
        # 气候分区
        if ("气候分区" in text or "气候区划" in text) and not info["气候分区"]:
            match = re.search(r"(?:气候分区|气候区划)[：:]\s*(.+?)(?:\s|$|，|,)", text)
            if match:
                info["气候分区"] = match.group(1).strip()
        
        # 建筑分类
        if ("建筑类型" in text or "建筑分类" in text) and not info["建筑分类"]:
            match = re.search(r"(?:建筑类型|建筑分类)[：:]\s*(.+?)(?:\s|$|，|,)", text)
            if match:
                info["建筑分类"] = match.group(1).strip()
        
        # 结构形式
        if ("结构形式" in text or "结构类型" in text) and not info["结构形式"]:
            match = re.search(r"(?:结构形式|结构类型)[：:]\s*(.+?)(?:\s|$|，|,)", text)
            if match:
                info["结构形式"] = match.group(1).strip()
        
        # 建筑朝向
        if ("建筑朝向" in text or "北向角度" in text) and not info["建筑朝向"]:
            match = re.search(r"(?:建筑朝向|北向角度)[：:]\s*(.+?)(?:\s|$|，|,)", text)
            if match:
                info["建筑朝向"] = match.group(1).strip()
        
        # 建筑面积
        if (("建筑面积" in text and "总建筑面积" not in text) or "总建筑面积" in text) and not info["建筑面积"]:
            match = re.search(r"(?:建筑面积|总建筑面积)[：:]\s*([0-9,.]+)\s*(?:m²|平方米)?", text)
            if match:
                info["建筑面积"] = match.group(1).strip()
        
        # 建筑层数
        if (("层数" in text and "建筑层数" not in text and "地下层数" not in text) or "建筑层数" in text) and not info["建筑层数"]:
            match = re.search(r"(?:建筑层数|层数)[：:]\s*(.+?)(?:\s|$|，|,)", text)
            if match:
                info["建筑层数"] = match.group(1).strip()
                # 尝试提取地下层数
                underground_match = re.search(r"地下[：:\s]*(\d+)[层\s]", match.group(1))
                if underground_match:
                    info["地下层数"] = underground_match.group(1)
        
        # 地下层数
        if "地下层数" in text and not info["地下层数"]:
            match = re.search(r"地下层数[：:]\s*(\d+)", text)
            if match:
                info["地下层数"] = match.group(1).strip()
        
        # 建筑高度
        if "建筑高度" in text and not info["建筑高度"]:
            match = re.search(r"建筑高度[：:]\s*([0-9,.]+)\s*(?:m|米)?", text)
            if match:
                info["建筑高度"] = match.group(1).strip()
        
        # 设计单位
        if "设计单位" in text and not info["设计单位"]:
            match = re.search(r"设计单位[：:]\s*(.+?)(?:\s|$|，|,|。)", text)
            if match:
                info["设计单位"] = match.group(1).strip()
        
        # 建设单位
        if "建设单位" in text and not info["建设单位"]:
            match = re.search(r"建设单位[：:]\s*(.+?)(?:\s|$|，|,|。)", text)
            if match:
                info["建设单位"] = match.group(1).strip()

# 示例用法
if __name__ == "__main__":
    import sys
    
    if len(sys.argv) < 2:
        print("用法: python word_extractor.py <word文件路径>")
        sys.exit(1)
    
    file_path = sys.argv[1]
    result = extract_doc_info(file_path)
    
    if result:
        print("\n提取到的项目信息:")
        for key, value in result.items():
            print(f"{key}: {value}")
    else:
        print("提取信息失败") 