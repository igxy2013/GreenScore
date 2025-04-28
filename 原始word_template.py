from docx import Document
import json
from flask import jsonify
import re
import os
from flask import current_app
from docx.oxml import parse_xml
from docx.oxml.ns import qn, nsmap
from docx.shared import Pt

def replace_placeholders(template_path, data):
    try:
        print(f"\n=== 开始处理文档 ===")
        print(f"模板文件路径: {template_path}")
        print(f"\n接收到的数据:")
        print(json.dumps(data, ensure_ascii=False, indent=2))
        
        # 加载Word模板
        doc = Document(template_path)
        
        # 处理项目信息
        project_fields = [
            '项目名称', '设计单位', '建设单位', '总建筑面积', '星级目标', '建筑类型', 
            '建筑总分', '结构总分', '给排水总分', '电气总分', '暖通总分', '景观总分',
            '建筑创新总分', '结构创新总分', '给排水创新总分', '电气创新总分', '暖通创新总分', '景观创新总分',
            '项目总分', '创新总分', '项目地点', '建筑高度', '建筑层数',
            '安全耐久总分', '生活便利总分', '健康舒适总分', '资源节约总分', '环境宜居总分', '提高与创新总分',
            '设计日期', '环境健康与节能总分', '环境健康与节能创新总分',
            '总用地面积', '地上建筑面积', '地下建筑面积', '容积率', '建筑密度', '绿地率',
            '气候区划', '项目编号', '建设情况', '空调类型', '有无电梯', '有无地下车库',
            '有无垃圾用房', '有无景观水体', '是否全装修', '公建类型', '绿地向公众开放'
        ]
        
        print("\n可用字段列表:")
        for field in project_fields:
            print(f"- {field}")
        
        # 添加占位符映射关系
        placeholder_mapping = {
            '建创总分': '建筑创新总分',
            '结创总分': '结构创新总分',
            '暖创总分': '暖通创新总分',
            '电创总分': '电气创新总分',
            '水创总分': '给排水创新总分',
            '景创总分': '景观创新总分',
            '创新总分': '提高与创新总分',
            '节能总分': '环境健康与节能总分',
            '节创总分': '环境健康与节能创新总分'
        }
        
        # 获取文档中的所有书签
        print("\n处理文档书签:")
        bookmarks_dict = {}
        for bookmark_start in doc.element.xpath('//w:bookmarkStart'):
            bookmark_name = bookmark_start.get(qn('w:name'))
            if not bookmark_name:
                continue
            
            # 获取书签的内容范围
            bookmark_id = bookmark_start.get(qn('w:id'))
            bookmark_end = doc.element.xpath(f'//w:bookmarkEnd[@w:id="{bookmark_id}"]')[0]
            
            # 创建一个范围对象来表示书签的内容
            bookmark_range = bookmark_start.getnext()
            while bookmark_range is not None and bookmark_range is not bookmark_end:
                if bookmark_range.tag.endswith('}r') or bookmark_range.tag.endswith('}p'):
                    break
                bookmark_range = bookmark_range.getnext()
            
            if bookmark_range is not None:
                bookmarks_dict[bookmark_name] = bookmark_range
                print(f"找到书签: {bookmark_name}")
            else:
                print(f"跳过无效书签: {bookmark_name}")
        
        # 处理数字格式的书签（条文号对应的得分）
        for bookmark_name, bookmark_range in bookmarks_dict.items():
            if re.match(r'\d+(?:\.\d+)*?', bookmark_name) or bookmark_name.startswith('f'):
                # 在数据中查找对应的得分
                for item in data:
                    # 将条文号转换为新的书签格式（去掉点号并添加f前缀）
                    bookmark_format = 'f' + str(item.get('条文号')).replace('.', '')
                    # 同时检查原始条文号格式和带f前缀的格式
                    if bookmark_format == bookmark_name or str(item.get('条文号')) == bookmark_name:
                        try:
                            # 创建新的文本运行
                            new_text = str(item.get('得分', ''))
                            new_run = parse_xml(f'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:rPr><w:sz w:val="21"/></w:rPr><w:t>{new_text}</w:t></w:r>')
                            
                            # 替换书签内容
                            if bookmark_range is not None:
                                parent = bookmark_range.getparent()
                                if parent is not None:
                                    parent.replace(bookmark_range, new_run)
                                    print(f"替换书签 {bookmark_name} 的值为: {new_text}")
                        except Exception as e:
                            print(f"替换书签 {bookmark_name} 时出错: {str(e)}")
                        break

    
        # 处理设计日期书签（包括带数字后缀的）
        from datetime import datetime
        current_date = datetime.now().strftime("%Y年%m月%d日")
        # 使用正则表达式匹配设计日期及其带数字后缀的变体
        date_pattern = re.compile(r'^设计日期[0-9]*$')
        for bookmark_name in list(bookmarks_dict.keys()):
            if date_pattern.match(bookmark_name):
                # 创建新的文本运行
                new_run = parse_xml(f'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:rPr><w:sz w:val="21"/></w:rPr><w:t>{current_date}</w:t></w:r>')
                
                # 替换书签内容
                bookmark_range = bookmarks_dict[bookmark_name]
                if bookmark_range is not None:
                    parent = bookmark_range.getparent()
                    if parent is not None:
                        parent.replace(bookmark_range, new_run)
                        print(f"处理设计日期书签: {bookmark_name} -> {current_date}")

        # 处理标准字段书签（包括带数字后缀的书签）
        for field in project_fields:
            field_value = None # Flag to check if value is determined

            # --- 星级目标 ---
            if field == '星级目标':
                target = ''
                if data and isinstance(data[0], dict):
                    target = data[0].get('星级目标', '') # Use .get() for safety
                    field_value = f'基本级{"■" if target == "基本级" else "□"}一星级{"■" if target == "一星级" else "□"}二星级{"■" if target == "二星级" else "□"}三星级{"■" if target == "三星级" else "□"}'
                field_pattern = re.compile(r'^星级目标[0-9]*$')
                print(f"准备处理星级目标字段: {field_value}")

            # --- 绿色星级 ---
            elif field == '绿色星级':
                target = ''
                if data and isinstance(data[0], dict):
                    target = data[0].get('星级目标', '') # Based on 星级目标
                if target == '基本级': field_value = '基本级'
                elif target == '一星级': field_value = '一星级'
                elif target == '二星级': field_value = '二星级'
                elif target == '三星级': field_value = '三星级'
                else: field_value = target
                field_pattern = re.compile(r'^绿色星级[0-9]*$')
                print(f"准备处理绿色星级字段: {field_value}")

            # --- 建筑类型 ---
            elif field == '建筑类型':
                building_type = ''
                if data and isinstance(data[0], dict):
                    building_type = data[0].get('建筑类型', '')
                    field_value = f'居住建筑{"■" if building_type == "居住建筑" else "□"} 公共建筑{"■" if building_type == "公共建筑" else "□"} 居住+公建{"■" if building_type == "居住+公共建筑" else "□"}'
                field_pattern = re.compile(r'^建筑类型[0-9]*$')
                print(f"准备处理建筑类型字段: {field_value}")

            # --- General Fields Handling ---
            else:
                # For general fields, value is determined inside the loop below
                field_pattern = re.compile(f"^{re.escape(field)}[0-9]*$")

            # If field_value is determined (special cases) or it's a general field, look for matching bookmarks
            if field_value is not None or field not in ['星级目标', '绿色星级', '建筑类型']:
                processed_field = False
                for bookmark_name in list(bookmarks_dict.keys()):
                    # Check if the bookmark name matches the current field or its pattern
                    if bookmark_name == field or field_pattern.match(bookmark_name):
                        # Ensure the bookmark still exists in the dictionary before accessing
                        if bookmark_name not in bookmarks_dict:
                            print(f"警告: 尝试处理书签 {bookmark_name} 但它不在字典中 (可能已被处理?)")
                            continue

                        current_field_value = field_value # Use pre-calculated value for special fields

                        # If it's a general field, get the value now
                        if current_field_value is None:
                            if data and isinstance(data[0], dict):
                                # Use the base 'field' name to get data
                                current_field_value = str(data[0].get(field, ''))
                            else:
                                current_field_value = '' # Default to empty if no data
                            print(f"处理书签: {bookmark_name} (数据字段: {field}) -> {current_field_value}")
                        else:
                            # For special fields, value was already calculated
                             print(f"处理书签: {bookmark_name} -> {current_field_value}")

                        try:
                            # Create new run with the determined value
                            new_run = parse_xml(f'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:rPr><w:sz w:val="21"/></w:rPr><w:t>{current_field_value}</w:t></w:r>')
                            bookmark_range = bookmarks_dict[bookmark_name]
                            parent = bookmark_range.getparent()
                            if parent is not None:
                                parent.replace(bookmark_range, new_run)
                                processed_field = True
                        except Exception as e:
                            print(f"处理书签 {bookmark_name} 时出错: {e}")

                # Optional: Print message if a special field bookmark was expected but not found
                if field in ['星级目标', '绿色星级', '建筑类型'] and not processed_field:
                     print(f"未在文档中找到 '{field}' 或其变体书签。")
        
        # 处理成都项目总分书签
        if '成都项目总分' in bookmarks_dict:
            field_value = ''
            if data and isinstance(data[0], dict):
                project_score = float(data[0].get('项目总分', '0'))
                field_value = format((project_score + 400) / 10, '.1f')
            # 创建新的文本运行
            new_run = parse_xml(f'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:rPr><w:sz w:val="21"/></w:rPr><w:t>{field_value}</w:t></w:r>')
            
            # 替换书签内容
            bookmark_range = bookmarks_dict['成都项目总分']
            if bookmark_range is not None:
                parent = bookmark_range.getparent()
                if parent is not None:
                    parent.replace(bookmark_range, new_run)
        
        # 处理映射的简写书签（包括带数字后缀的）
        for short_name, full_name in placeholder_mapping.items():
            # 使用正则表达式匹配简写字段名及其带数字后缀的变体
            short_name_pattern = re.compile(f"{short_name}[0-9]*$")
            for bookmark_name in list(bookmarks_dict.keys()):
                if short_name_pattern.match(bookmark_name):
                    field_value = ''
                    if data and isinstance(data[0], dict):
                        field_value = str(data[0].get(full_name, ''))
                        print(f"处理简写书签: {bookmark_name} -> {field_value}")
                    # 创建新的文本运行
                    new_run = parse_xml(f'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:rPr><w:sz w:val="21"/></w:rPr><w:t>{field_value}</w:t></w:r>')
                    
                    # 替换书签内容
                    bookmark_range = bookmarks_dict[bookmark_name]
                    if bookmark_range is not None:
                        parent = bookmark_range.getparent()
                        if parent is not None:
                            parent.replace(bookmark_range, new_run)
                            print(f"替换简写书签: {bookmark_name} -> {field_value}")

        # 处理段落中的占位符
        print("\n=== 检查文档中的占位符 ===")
        print("检查段落中的占位符:")
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text
            if '{' in text:
                print(f"\n段落 {i}:")
                print(f"原始文本: {text}")
                
                # 使用简单的文本替换方法
                new_text = text
                
                # 查找条文号占位符
                for match in re.finditer(r'\{(\d+\.\d+\.\d+(?:措施)?)\}', text):
                    placeholder = match.group(1)
                    placeholder_with_braces = match.group(0)
                    
                    # 获取值
                    field_value = ''
                    is_measure = placeholder.endswith('措施')
                    clause_number = placeholder[:-2] if is_measure else placeholder
                    
                    if data:
                        for item in data:
                            if str(item.get('条文号', '')) == clause_number:
                                if is_measure:
                                    field_value = str(item.get('技术措施', ''))
                                    print(f"替换条文措施占位符 {{{placeholder}}} -> {field_value}")
                                else:
                                    field_value = str(item.get('得分', ''))
                                    print(f"替换条文得分占位符 {{{placeholder}}} -> {field_value}")
                                break
                    
                    # 替换占位符
                    new_text = new_text.replace(placeholder_with_braces, field_value)
                
                # 查找标准字段占位符
                for field in project_fields:
                    placeholder = '{' + field + '}'
                    if placeholder in text:
                        field_value = ''
                        if data and isinstance(data[0], dict):
                            field_value = str(data[0].get(field, ''))
                            print(f"替换占位符 {placeholder} -> {field_value}")
                        
                        # 替换占位符
                        new_text = new_text.replace(placeholder, field_value)
                
                # 查找映射的简写字段
                for short_name, full_name in placeholder_mapping.items():
                    placeholder = '{' + short_name + '}'
                    if placeholder in text:
                        field_value = ''
                        if data and isinstance(data[0], dict):
                            field_value = str(data[0].get(full_name, ''))
                            print(f"替换简写占位符 {placeholder} -> {field_value}")
                        
                        # 替换占位符
                        new_text = new_text.replace(placeholder, field_value)
                
                # 如果文本已更改，创建新段落内容
                if new_text != text:
                    # 清空段落
                    for run in paragraph.runs:
                        run.text = ""
                    
                    # 添加新文本，并应用宋体和指定字号
                    new_run = paragraph.add_run(new_text)
                    new_run.font.name = '宋体'
                    new_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    
                    # 检查是否包含条文号相关内容或特定的总分字段，如果是则使用小四号，否则使用四号
                    special_fields = ['安全耐久总分', '健康舒适总分', '生活便利总分', '资源节约总分', '环境宜居总分', '创新总分']
                    if re.search(r'\d+\.\d+\.\d+(?:措施)?', text) or any(f"{{{field}}}" in text for field in special_fields):
                        new_run.font.size = Pt(12)  # 小四字体大小为12磅
                    else:
                        new_run.font.size = Pt(14)  # 四号字体大小为14磅

        # 处理表格中的占位符
        print("\n检查表格中的占位符:")
        for i, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        if '{' in text:
                            print(f"\n表格 {i}, 行 {row_idx}, 列 {cell_idx}:")
                            print(f"原始文本: {text}")
                            
                            # 使用简单的文本替换方法
                            new_text = text
                            
                            # 查找条文号占位符
                            for match in re.finditer(r'\{(\d+\.\d+\.\d+(?:措施)?)\}', text):
                                placeholder = match.group(1)
                                placeholder_with_braces = match.group(0)
                                
                                # 获取值
                                field_value = ''
                                is_measure = placeholder.endswith('措施')
                                clause_number = placeholder[:-2] if is_measure else placeholder
                                
                                if data:
                                    for item in data:
                                        if str(item.get('条文号', '')) == clause_number:
                                            if is_measure:
                                                field_value = str(item.get('技术措施', ''))
                                                print(f"替换条文措施占位符 {{{placeholder}}} -> {field_value}")
                                            else:
                                                field_value = str(item.get('得分', ''))
                                                print(f"替换条文得分占位符 {{{placeholder}}} -> {field_value}")
                                            break
                                
                                # 替换占位符
                                new_text = new_text.replace(placeholder_with_braces, field_value)
                            
                            # 查找标准字段占位符
                            for field in project_fields:
                                placeholder = '{' + field + '}'
                                if placeholder in text:
                                    field_value = ''
                                    if data and isinstance(data[0], dict):
                                        field_value = str(data[0].get(field, ''))
                                        print(f"替换占位符 {placeholder} -> {field_value}")
                                    
                                    # 替换占位符
                                    new_text = new_text.replace(placeholder, field_value)
                            
                            # 查找映射的简写字段
                            for short_name, full_name in placeholder_mapping.items():
                                placeholder = '{' + short_name + '}'
                                if placeholder in text:
                                    field_value = ''
                                    if data and isinstance(data[0], dict):
                                        field_value = str(data[0].get(full_name, ''))
                                        print(f"替换简写占位符 {placeholder} -> {field_value}")
                                    
                                    # 替换占位符
                                    new_text = new_text.replace(placeholder, field_value)
                            
                            # 如果文本已更改，创建新段落内容
                            if new_text != text:
                                # 清空段落
                                for run in paragraph.runs:
                                    run.text = ""
                                
                                # 添加新文本，并应用宋体和指定字号
                                new_run = paragraph.add_run(new_text)
                                new_run.font.name = '宋体'
                                new_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                                
                                # 检查是否包含条文号相关内容或特定的总分字段，如果是则使用小四号，否则使用四号
                                special_fields = ['安全耐久总分', '健康舒适总分', '生活便利总分', '资源节约总分', '环境宜居总分', '创新总分']
                                if re.search(r'\d+\.\d+\.\d+(?:措施)?', text) or any(f"{{{field}}}" in text for field in special_fields):
                                    new_run.font.size = Pt(12)  # 小四字体大小为12磅
                                else:
                                    new_run.font.size = Pt(14)  # 四号字体大小为14磅

        # 确保temp目录存在
        os.makedirs('temp', exist_ok=True)
        
        # 生成带时间戳的输出文件名
        from datetime import datetime
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        output_path = os.path.join('temp', f'report_{timestamp}.docx')
        
        # 保存修改后的文档
        doc.save(output_path)
        print(f"\n文档已保存到: {output_path}")
        
        # 检查文件是否成功创建
        if not os.path.exists(output_path):
            raise Exception(f"文件保存失败: {output_path}")
            
        return output_path
    except Exception as e:
        print(f"处理文档时出错：{str(e)}")
        import traceback
        print(traceback.format_exc())
        return f'处理文档时出错：{str(e)}'
def modify_square_chars_font(doc):
    
    # 定义要查找的字符列表
    target_chars = ["■", "□"]
    
    # 处理文档正文中的段落
    for paragraph in doc.paragraphs:
        process_runs(paragraph.runs, target_chars)
    
    # 处理文档中的表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_runs(paragraph.runs, target_chars)
    
def process_runs(runs, target_chars):
    """处理run集合，修改目标字符的字体"""
    for run in runs:
        # 检查run中是否包含任何一个目标字符
        if any(char in run.text for char in target_chars):
            # 设置字体为宋体
            run.font.name = "宋体"
            # 对于中文字体，还需要设置对应的字体族
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')    
def process_template(data):
    try:
        from flask import current_app
        import os
        
        # 根据评价标准选择模板文件
        standard = data[0].get('评价标准', '成都市标')  # 默认使用成都市标
        
        if standard == '国标':
            return None  # 国标不进行任何操作
        
        # 获取星级目标
        star_rating_target = data[0].get('星级目标', '')
        
        # 根据评价标准和星级目标选择模板文件
        if standard == '四川省标' and star_rating_target == '基本级':
            template_file = 'sichuan_template-basic.docx'
        elif standard == '四川省标':
            template_file = 'sichuan_template.docx'
        else:
            template_file = 'chengdu_template.docx'
        
        # 获取模板文件的完整路径
        template_path = os.path.join(current_app.static_folder, 'templates', template_file)
        print(f"处理模板文件: {template_path}, 评价标准: {standard}, 星级目标: {star_rating_target}")
        
        # 检查模板文件是否存在
        if not os.path.exists(template_path):
            raise Exception(f"模板文件不存在: {template_path}")
                
        # 处理模板
        output_path = replace_placeholders(template_path, data)
        if not isinstance(output_path, str) or not output_path.endswith('.docx'):
            raise Exception(f"模板处理失败：{output_path}")
        
        # 返回生成的文档路径
        return output_path
        
    except Exception as e:
        print(f"处理模板失败: {str(e)}")
        import traceback
        print(traceback.format_exc())
        raise Exception(f'生成Word文档失败：{str(e)}')