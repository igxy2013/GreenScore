from docx import Document
import json
from flask import jsonify
import re
import os
from flask import current_app
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import qn, nsmap
from docx.shared import Pt, Inches
from copy import deepcopy
from datetime import datetime
import lxml.etree
import traceback
import platform # 用于检查操作系统
import sys # 用于检查平台
try:
    import pythoncom # 导入 pythoncom 用于 CoInitialize
except ImportError:
    # 在非 Windows 平台上，pythoncom 可能不存在，但这没关系，因为我们会在后面检查平台
    pythoncom = None 

def modify_square_chars_font(doc):
    
    # 定义要查找的字符列表
    target_chars = ["■", "□"]
    
    # 处理文档正文中的段落
    for paragraph in doc.paragraphs:
        process_runs(paragraph.runs, target_chars)
    
    # 处理文档中的表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_runs(paragraph.runs, target_chars)
    
def process_runs(runs, target_chars):
    """处理run集合，修改目标字符的字体"""
    for run in runs:
        # 检查run中是否包含任何一个目标字符
        if any(char in run.text for char in target_chars):
            # 设置字体为宋体
            run.font.name = "宋体"
            # 对于中文字体，还需要设置对应的字体族
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')    

def process_template(data):
    
    try:
        # --- 添加设计日期 --- 
        if data and isinstance(data[0], dict):
            # 检查 "设计日期" 是否已存在，如果不存在或为空，则添加/更新
            if not data[0].get("设计日期"):
                current_date = datetime.now().strftime("%Y年%m月%d日") # 修改日期格式为 年月日
                data[0]["设计日期"] = current_date
                print(f"[word_template.py] 已自动添加设计日期: {current_date}")
        # --- 设计日期添加结束 ---

        from flask import current_app
        import os
        
        # 根据评价标准选择模板文件
        standard = data[0].get('评价标准', '') 
        # 获取星级目标
        star_rating_target = data[0].get('星级目标', '')
        output_paths = [] # 用于存储生成的文档路径
        # --- 修改：存储生成的文档信息 (路径, 基础名) --- 
        output_info_list = [] 
        # --- 修改结束 ---

        if standard == '国标':
            if "安徽" in data[0].get('项目地点', '') :
                template_file = '安徽绿色建筑审查表.docx'
                template_path = os.path.join(current_app.static_folder, 'templates', template_file)
                print(f"处理模板文件: {template_path}, 评价标准: {standard}")
                if not os.path.exists(template_path):
                    raise Exception(f"模板文件不存在: {template_path}")
                # --- 修改：处理返回值 --- 
                result = replace_placeholders(template_path, data)
                if result:
                    output_path, base_name = result
                    if not isinstance(output_path, str) or not output_path.endswith('.docx'):
                         print(f"警告: 国标模板处理返回无效路径 '{output_path}'，跳过添加。")
                    else:
                         output_info_list.append((output_path, base_name))
                else:
                     print(f"警告: 国标模板处理失败，跳过添加。")
                # --- 修改结束 ---
            else:
                return None  # 国标（非安徽）不进行任何操作
        elif standard == '四川省标':
            if star_rating_target == '基本级':
                # 处理第一个模板
                template1_file = '四川省绿建审查表-基本级.docx'
            else:
                template1_file = '四川省绿建审查表-提高级.docx'
            template1_path = os.path.join(current_app.static_folder, 'templates', template1_file)
            print(f"处理第一个模板文件 (四川省标): {template1_path}")
            if not os.path.exists(template1_path):
                raise Exception(f"模板文件不存在: {template1_path}")
            # --- 修改：处理返回值 --- 
            result1 = replace_placeholders(template1_path, data)
            if result1:
                output1_path, base1_name = result1
                if not isinstance(output1_path, str) or not output1_path.endswith('.docx'):
                    print(f"警告: 处理第一个模板 {template1_file} 返回无效路径 '{output1_path}'，跳过添加。")
                else:
                    output_info_list.append((output1_path, base1_name))
            else:
                print(f"警告: 处理第一个模板 {template1_file} 失败，跳过添加。")
            # --- 修改结束 ---
            
            # 处理第二个模板
            template2_file = '水系统规划设计评审表.docx'
            template2_path = os.path.join(current_app.static_folder, 'templates', template2_file)
            print(f"处理第二个模板文件 (四川省标): {template2_path}")
            if not os.path.exists(template2_path):
                raise Exception(f"模板文件不存在: {template2_path}")
            # --- 修改：处理返回值 --- 
            result2 = replace_placeholders(template2_path, data)
            if result2:
                output2_path, base2_name = result2
                if not isinstance(output2_path, str) or not output2_path.endswith('.docx'):
                    print(f"警告: 处理第二个模板 {template2_file} 返回无效路径 '{output2_path}'，跳过添加。")
                else:
                    output_info_list.append((output2_path, base2_name))
            else:
                print(f"警告: 处理第二个模板 {template2_file} 失败，跳过添加。")
            # --- 修改结束 ---

        elif standard == '成都市标':
            # 定义成都市标需要处理的模板文件列表
            template_files = [
                '成都市绿色建筑审查表.docx',
                '水系统规划设计申报表.docx',
                '专项报告申报一览表.docx'
            ]

            for idx, template_file in enumerate(template_files):
                template_path = os.path.join(current_app.static_folder, 'templates', template_file)
                print(f"处理第 {idx+1} 个模板文件 (成都市标): {template_path}")
                
                if not os.path.exists(template_path):
                    # 如果任何一个模板文件不存在，可以选择是跳过还是报错
                    # 这里选择打印警告并跳过该文件
                    print(f"警告: 模板文件不存在: {template_path}, 将跳过此文件。")
                    continue # 跳过当前文件，继续处理下一个
                    # 或者可以选择报错: raise Exception(f"模板文件不存在: {template_path}")
                
                # 调用 replace_placeholders 处理模板
                # --- 修改：处理返回值 --- 
                result = replace_placeholders(template_path, data)
                
                # 检查模板是否处理成功
                if result:
                    output_path, base_name = result
                    if not isinstance(output_path, str) or not output_path.endswith('.docx'):
                         print(f"警告: 处理模板 {template_file} 返回无效路径 '{output_path}'，跳过添加。")
                    else:
                        # 将成功生成的文档信息添加到列表
                        output_info_list.append((output_path, base_name))
                        print(f"已成功处理并添加: {output_path}")
                else:
                    print(f"错误: 调用 replace_placeholders 处理模板 {template_file} 时返回 None，表示内部处理失败。请检查之前的日志查找具体错误。跳过添加。")
                # --- 修改结束 ---

        else:
            # 对于其他未知的标准，或者不需要处理的情况
             print(f"未知的评价标准或无需处理: {standard}")
             return None # 或返回空列表 []，取决于后续逻辑

        # --- 添加最终返回日志 --- 
        # --- 修改：返回信息列表 --- 
        final_return_value = output_info_list if output_info_list else None
        print(f"[word_template.py] process_template 即将返回: {final_return_value}")
        print(f"[word_template.py] process_template 即将返回类型: {type(final_return_value)}")
        # --- 结束日志 --- 
        # --- 修改结束 --- 
        # 返回生成的文档信息列表
        return final_return_value
        
    except Exception as e:
        print(f"处理模板整体过程失败: {str(e)}") # 修改日志信息
        print(traceback.format_exc())
        return None # 返回 None 表示处理失败

def replace_bookmarks_in_doc(doc, project_info, score_data, standard):
    """
    处理 Word 文档中的书签替换。
    使用原始word_template.py中的查找和替换逻辑。
    根据 standard 决定得分的显示方式。
    """
    print(f"--- 开始处理书签 (使用原始逻辑, 标准: {standard}) ---")

    # 1. 定义字段和映射 (来自原始文件)
    project_fields = [
        '项目名称', '设计单位', '建设单位', '总建筑面积', '星级目标', '建筑类型', # Note: 星级目标/建筑类型 handled specially below
        '建筑总分', '结构总分', '给排水总分', '电气总分', '暖通总分', '景观总分',
        '建筑创新总分', '结构创新总分', '给排水创新总分', '电气创新总分', '暖通创新总分', '景观创新总分',
        '项目总分', '创新总分', '项目地点', '建筑高度', '建筑层数',
        '安全耐久总分', '生活便利总分', '健康舒适总分', '资源节约总分', '环境宜居总分', '提高与创新总分',
        # '设计日期', # Handled separately
        '环境健康与节能总分', '环境健康与节能创新总分',
        '总用地面积', '地上建筑面积', '地下建筑面积', '容积率', '建筑密度', '绿地率',
        '气候区划', '项目编号', '建设情况', '空调类型', '有无电梯', '有无地下车库',
        '有无垃圾用房', '有无景观水体', '是否全装修', '公建类型', '绿地向公众开放',
        '绿色星级', # Handled specially below
        '标准项目总分' # Calculated in replace_placeholders and added to project_info
    ]

    placeholder_mapping = {
        '建创总分': '建筑创新总分',
        '结创总分': '结构创新总分',
        '暖创总分': '暖通创新总分',
        '电创总分': '电气创新总分',
        '水创总分': '给排水创新总分',
        '景创总分': '景观创新总分',
        '创新总分': '提高与创新总分',
        '节能总分': '环境健康与节能总分',
        '节创总分': '环境健康与节能创新总分'
    }

    # 2. 查找书签并存储范围 (来自原始文件, 稍作修改)
    print("  开始扫描文档查找书签标记...")
    bookmarks_found_count = 0
    bookmark_pairs = {} # Store pairs of {name: (start_element, end_element)} if possible
    bookmark_starts = {} # Fallback: Store only start elements

    try:
        all_start_elements = list(doc.element.iter(qn('w:bookmarkStart')))
        all_end_elements_by_id = {elem.get(qn('w:id')): elem for elem in doc.element.iter(qn('w:bookmarkEnd'))}

        print(f"  初步找到 {len(all_start_elements)} 个 w:bookmarkStart 和 {len(all_end_elements_by_id)} 个 w:bookmarkEnd 标记。")

        for start_element in all_start_elements:
            name = start_element.get(qn('w:name'))
            id = start_element.get(qn('w:id'))
            if not name or name.startswith('_') or not id: # 忽略无名/内部/无ID书签
                continue

            bookmarks_found_count += 1
            bookmark_starts[name] = start_element # Always store start element as fallback
            
            # Try to find the matching end element
            end_element = all_end_elements_by_id.get(id)
            if end_element is not None:
                # Store the pair if found
                bookmark_pairs[name] = (start_element, end_element)
            # else:
            #    print(f"  警告: 未找到书签 '{name}' (ID: {id}) 对应的结束标记。")

    except Exception as e:
        print(f"  错误: 查找书签标记时出错: {e}")
        traceback.print_exc()
        return # 无法继续处理

    print(f"  扫描完成，找到 {bookmarks_found_count} 个有效书签名称。找到 {len(bookmark_pairs)} 对完整的起止标记。开始处理替换...")
    processed_bookmarks = set()
    bookmarks_replaced_count = 0

    # --- 辅助函数：执行替换 (新逻辑) --- 
    def perform_replacement(bookmark_name, value_str):
        nonlocal bookmarks_replaced_count
        if bookmark_name in processed_bookmarks: # Skip if already processed
            return False
        
        start_element = None
        end_element = None
        
        if bookmark_name in bookmark_pairs:
            start_element, end_element = bookmark_pairs[bookmark_name]
        elif bookmark_name in bookmark_starts:
            start_element = bookmark_starts[bookmark_name]
            # print(f"  信息: 书签 '{bookmark_name}' 仅找到开始标记，将执行插入操作。")
        else:
            return False # Bookmark not found

        try:
            # 创建新的 run XML 元素
            new_run = parse_xml(f'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:rPr><w:sz w:val="21"/></w:rPr><w:t>{value_str}</w:t></w:r>')

            # --- 关键替换逻辑：尝试删除旧内容 (如果起止在同一父级) --- 
            can_delete = False
            if end_element is not None:
                parent_start = start_element.getparent()
                parent_end = end_element.getparent()
                if parent_start is not None and parent_start == parent_end:
                    # print(f"  书签 '{bookmark_name}' 起止标记在同一父级内，尝试删除旧 run。")
                    can_delete = True
                    elements_to_remove = []
                    current = start_element.getnext()
                    while current is not None and current != end_element:
                        if current.tag == qn('w:r'):
                            elements_to_remove.append(current)
                        # elif current.tag == qn('w:p'): # Stop if we hit another paragraph (shouldn't happen if parent is p)
                        #    print(f"  警告: 书签 '{bookmark_name}' 范围内遇到意外的段落，停止删除。")
                        #    can_delete = False
                        #    break
                        current = current.getnext()
                    
                    if can_delete:
                        for elem in elements_to_remove:
                            try:
                                parent_start.remove(elem)
                                # print(f"    Removed old run for bookmark '{bookmark_name}'")
                            except ValueError:
                                pass # Element already removed or not child of parent_start
                # else:
                    # print(f"  书签 '{bookmark_name}' 起止标记不在同一父级，跳过删除旧 run。")

            # 在 bookmarkStart 之后插入新的 run
            start_element.addnext(new_run)
            bookmarks_replaced_count += 1
            processed_bookmarks.add(bookmark_name)
            # print(f"    Replaced bookmark '{bookmark_name}' (Deletion attempted: {can_delete})")
            return True

        except Exception as e:
            print(f"  错误: 替换书签 '{bookmark_name}' 时出错: {e}")
            traceback.print_exc()
            processed_bookmarks.add(bookmark_name) # Mark as processed to avoid retrying
        return False

    # 3. 处理数字/条文号书签 (来自原始文件)
    print("  处理条文号相关书签...")
    for bookmark_name in list(bookmark_starts.keys()): # Iterate over keys found in doc
        is_clause_bookmark = re.match(r'\d+(?:\.\d+)*?', bookmark_name) or bookmark_name.startswith('f')
        if is_clause_bookmark and bookmark_name not in processed_bookmarks:
            for item in score_data:
                clause_num = item.get('条文号', '')
                if not clause_num: continue
                bookmark_format_f = 'f' + str(clause_num).replace('.', '')
                # 检查书签名称是否匹配 条文号 或 f+条文号
                if bookmark_name == str(clause_num) or bookmark_name == bookmark_format_f:
                    display_score = str(item.get('得分', ''))
                    if standard == '四川省标':
                        if display_score == '达标': display_score = '√'
                        elif display_score == '—' or display_score == '不达标': display_score = '×'
                        elif display_score == '0': display_score = '/'
                    if perform_replacement(bookmark_name, display_score):
                        break # 找到匹配项并处理后，跳出内层 score_data 循环

    # 4. 处理设计日期书签 (来自原始文件)
    print("  处理设计日期书签...")
    current_date = datetime.now().strftime("%Y年%m月%d日")
    date_pattern = re.compile(r'^设计日期[0-9]*$')
    for bookmark_name in list(bookmark_starts.keys()):
        if date_pattern.match(bookmark_name):
            perform_replacement(bookmark_name, current_date)

    # 5. 处理标准字段书签 (来自原始文件, 包含特殊和通用字段)
    print("  处理项目信息字段书签...")
    handled_special_fields = {'星级目标', '绿色星级', '建筑类型'} # 这些需要特殊格式化
    for field in project_fields:
        field_pattern = re.compile(f"^{re.escape(field)}[0-9]*$")
        field_value_str = None # 存储最终要插入的字符串

        # 特殊字段格式化
        if field == '星级目标':
            target = project_info.get('星级目标', '')
            field_value_str = f'基本级{"■" if target == "基本级" else "□"}一星级{"■" if target == "一星级" else "□"}二星级{"■" if target == "二星级" else "□"}三星级{"■" if target == "三星级" else "□"}'
        elif field == '绿色星级':
            target = project_info.get('星级目标', '') # 绿色星级基于星级目标
            if target == '基本级': field_value_str = '基本级'
            elif target == '一星级': field_value_str = '一星级'
            elif target == '二星级': field_value_str = '二星级'
            elif target == '三星级': field_value_str = '三星级'
            else: field_value_str = target # 其他情况直接显示
        elif field == '建筑类型':
            building_type = project_info.get('建筑类型', '')
            field_value_str = f'居住建筑{"■" if building_type == "居住建筑" else "□"} 公共建筑{"■" if building_type == "公共建筑" else "□"} 居住+公建{"■" if building_type == "居住+公共建筑" else "□"}'
        # 通用字段获取值
        elif field not in handled_special_fields:
             # 直接从 project_info 获取，因为 标准项目总分 已被添加
             field_value_str = str(project_info.get(field, ''))

        # 查找匹配的书签并替换
        if field_value_str is not None: # 确保有值可替换
            found_match_for_field = False
            for bookmark_name in list(bookmark_starts.keys()):
                if bookmark_name == field or field_pattern.match(bookmark_name):
                    if perform_replacement(bookmark_name, field_value_str):
                        found_match_for_field = True
                        # Don't break here, allow replacing multiple suffixed bookmarks (e.g., 项目名称1, 项目名称2)
            # if not found_match_for_field and field in project_fields:
            #      print(f"  信息: 未找到字段 '{field}' 对应的书签 (或其带后缀变体)。")

    # 6. 处理映射的简写书签 (来自原始文件)
    print("  处理简写映射书签...")
    for short_name, full_name in placeholder_mapping.items():
        short_name_pattern = re.compile(f"^{re.escape(short_name)}[0-9]*$")
        field_value = str(project_info.get(full_name, ''))
        for bookmark_name in list(bookmark_starts.keys()):
            if short_name_pattern.match(bookmark_name):
                perform_replacement(bookmark_name, field_value)

    print(f"--- 书签处理结束 (使用原始逻辑)。成功替换 {bookmarks_replaced_count} 个书签。 ---")
    # 报告未处理的书签 (可选)
    unprocessed_bookmarks = set(bookmark_starts.keys()) - processed_bookmarks
    if unprocessed_bookmarks:
         print(f"  信息: 以下 {len(unprocessed_bookmarks)} 个在文档中找到的书签未被替换 (可能无需替换或数据缺失): {list(unprocessed_bookmarks)[:20]}...")


def update_toc_with_com(doc_path):
    """
    使用 COM 接口更新 Word 文档的目录。
    仅在 Windows 且安装了 pywin32 和 Microsoft Word 时有效。
    """
    print("--- 尝试使用 COM 更新 Word 目录 ---")
    if platform.system() != "Windows":
        print("  非 Windows 系统，跳过 COM 目录更新。")
        return False
    
    # 检查 pythoncom 是否可用 (理论上 Windows + pywin32 应该有)
    if pythoncom is None:
         print("  错误：pythoncom 模块无法导入，无法初始化 COM。")
         return False

    try:
        import win32com.client as win32
        print("  成功导入 win32com.client。")
    except ImportError:
        print("  错误：未找到 pywin32 库，无法使用 COM 更新目录。请确保已安装 pywin32。")
        return False

    word = None
    doc = None
    updated = False
    coinitialized = False # 标记是否成功初始化 COM
    try:
        # --- 初始化 COM --- 
        try:
            pythoncom.CoInitialize()
            coinitialized = True
            print("  COM 已初始化。")
        except pythoncom.com_error as init_err:
            # 有时 COM 可能已被当前线程初始化 (例如嵌套调用)
            # 检查错误码，如果是 S_FALSE (0x1)，表示已初始化，可以继续
            # 其他错误则是有问题的
            if init_err.hresult == 0x1: # S_FALSE
                print("  COM 已被当前线程初始化，继续操作。")
                coinitialized = True # 视为已初始化，以便稍后反初始化
            else:
                print(f"  COM 初始化失败: {init_err}")
                return False # 初始化失败，无法继续

        print(f"  尝试启动 Word 并打开文档: {doc_path}")
        # 尝试连接到已运行的 Word 实例，如果失败则启动新实例
        try:
             word = win32.GetActiveObject("Word.Application")
             print("  已连接到活动的 Word 实例。")
        except pythoncom.com_error:
             # GetActiveObject 失败通常是因为没有活动实例
             print("  未找到活动的 Word 实例，尝试启动新的实例。")
             try:
                 word = win32.Dispatch("Word.Application")
                 print("  已启动新的 Word 实例。")
             except pythoncom.com_error as dispatch_err:
                 print(f"  启动新的 Word 实例失败: {dispatch_err}")
                 return False # 无法获取 Word 实例

        # 设置 Word 可见性 (False 为后台运行)
        word.Visible = False 
        word.DisplayAlerts = False # 禁止显示警告信息

        doc = word.Documents.Open(doc_path)
        print(f"  文档已打开。")

        if doc.TablesOfContents.Count >= 1:
            print(f"  找到 {doc.TablesOfContents.Count} 个目录。开始更新...")
            for i, toc in enumerate(doc.TablesOfContents):
                 try:
                     toc.Update()
                     print(f"    目录 {i+1} 更新成功。")
                     updated = True # 标记至少有一个目录成功更新
                 except Exception as toc_update_err:
                     print(f"    错误：更新目录 {i+1} 失败: {toc_update_err}")
            if updated:
                 print("  所有找到的目录已尝试更新。")
            else:
                 print("  未成功更新任何目录 (可能更新过程中出错)。")
        else:
            print("  文档中未找到目录 (TablesOfContents)。")

        if updated: # 只有在成功更新了至少一个目录后才保存
            print("  正在保存文档...")
            doc.Save()
            print("  文档保存成功。")
        else:
            print("  由于没有目录被成功更新，文档未重新保存。")

        return updated # 返回是否至少有一个目录被更新

    except Exception as e:
        print(f"  COM 操作过程中发生错误: {e}")
        print(traceback.format_exc())
        return False # 出错则返回 False
    finally:
        # --- 确保 Word 被关闭和释放 ---
        try:
            if doc is not None:
                doc.Close(SaveChanges=0) # 0 表示不保存更改（因为前面已经Save了或无需保存）
                print("  Word 文档已关闭。")
        except Exception as close_err:
            print(f"  关闭 Word 文档时出错: {close_err}")
        
        try:
            if word is not None:
                 # 尝试让 Word 退出，即使失败也要继续反初始化 COM
                 word.Quit()
                 print("  Word 应用已退出。")
        except Exception as quit_err:
            print(f"  退出 Word 应用时出错: {quit_err}")
        
        # 释放 COM 对象 (可选但推荐)
        word = None
        doc = None

        # --- 反初始化 COM --- 
        if coinitialized:
            try:
                pythoncom.CoUninitialize()
                print("  COM 已反初始化。")
            except Exception as uninit_err:
                 print(f"  COM 反初始化时出错: {uninit_err}")
        
        print("--- COM 目录更新尝试结束 ---")


def replace_placeholders(template_path, data):
    """
    替换 Word 文档中的占位符或书签。
    根据评价标准选择不同的处理方式。
    完成后尝试更新目录。
    返回: tuple (output_path, base_template_name) 或 None
    """
    output_path = None # 初始化为 None
    base_template_name = None # 初始化模板基础名称
    try:
        print(f"\n=== 开始处理文档模板 ===")
        print(f"模板文件路径: {template_path}")
        # --- 获取模板基础名称 --- 
        base_template_name = os.path.splitext(os.path.basename(template_path))[0]
        # --- 获取结束 ---
        
        if not data:
            print("错误：传入的数据为空")
            return None

        doc = Document(template_path)
        project_info = data[0] if isinstance(data[0], dict) else {}
        score_data = data[1:]
        standard = project_info.get('评价标准', '') # 获取评价标准

        # --- 图片替换逻辑 (保持不变) ---
        birdview_image_relative_path = project_info.get('鸟瞰图_path')
        image_placeholder = "{鸟瞰图}" # 图片占位符
        image_replaced = False # 标记是否成功替换了图片

        if birdview_image_relative_path:
            try:
                absolute_image_path = os.path.join(current_app.root_path, birdview_image_relative_path)
                image_exists = os.path.exists(absolute_image_path)
                print(f"[word_template.py] 检查图片绝对路径: {absolute_image_path}, 是否存在: {image_exists}")
            except Exception as path_err:
                print(f"[word_template.py] 构造或检查图片绝对路径时出错: {path_err}")
                image_exists = False
                absolute_image_path = None

            if image_exists and absolute_image_path:
                print(f"找到有效鸟瞰图路径，准备替换占位符/标记: {absolute_image_path}")
                # 假设图片仍使用文本占位符 "{鸟瞰图}"
                for paragraph in doc.paragraphs:
                    if image_placeholder in paragraph.text:
                        # (此处省略未改变的段落图片替换细节)
                        inline = paragraph.runs
                        full_text = "".join(run.text for run in inline)
                        if image_placeholder in full_text:
                            start_index = full_text.find(image_placeholder)
                            end_index = start_index + len(image_placeholder)
                            # ... (定位 run 的逻辑) ...
                            start_run_idx, start_offset, end_run_idx, end_offset = -1,-1,-1,-1 #简化表示
                            current_pos = 0
                            for i, run in enumerate(inline): # Find start/end run/offset
                                run_len = len(run.text)
                                if start_run_idx == -1 and start_index < current_pos + run_len: start_run_idx=i; start_offset=start_index-current_pos
                                if end_run_idx == -1 and end_index <= current_pos + run_len: end_run_idx=i; end_offset=end_index-current_pos; break
                                current_pos += run_len

                            if start_run_idx != -1 and end_run_idx != -1:
                                # (此处省略未改变的插入图片、清理run、处理剩余文本的细节)
                                try:
                                    # --- 修正和改进的清理与插入逻辑 ---
                                    # 1. 处理结束 Run (保留占位符之后的部分)
                                    if start_run_idx != end_run_idx:
                                        if end_run_idx < len(inline): # 检查索引是否有效
                                            end_run = inline[end_run_idx]
                                            end_run.text = end_run.text[end_offset:] # 保留占位符结束位置之后的部分

                                    # 2. 清理中间 Runs (完全清空)
                                    # 从后往前清，理论上更安全，虽然在这里影响不大
                                    for i in range(end_run_idx - 1, start_run_idx, -1):
                                        if i < len(inline): # 检查索引是否有效
                                            inline[i].text = ""

                                    # 3. 处理开始 Run (保留占位符之前的部分，并插入图片)
                                    start_run = inline[start_run_idx]
                                    start_run.text = start_run.text[:start_offset] # 保留占位符开始位置之前的部分

                                    # 4. 在开始 Run 的末尾插入图片
                                    start_run.add_picture(absolute_image_path, width=Inches(5.0))
                                    
                                    image_replaced = True
                                    print(f"  成功替换并插入图片于段落。")
                                    break # 成功替换后，跳出当前段落的处理
                                    # --- 修正逻辑结束 ---
                                except Exception as img_err: print(f"段落图片插入/清理出错: {img_err}")
                        if image_replaced: break
                if not image_replaced:
                    for table in doc.tables: # Check tables
                        # (此处省略未改变的表格图片替换逻辑)
                         for row in table.rows:
                             for cell in row.cells:
                                 for paragraph in cell.paragraphs:
                                      if image_placeholder in paragraph.text:
                                            inline = paragraph.runs
                                            full_text = "".join(run.text for run in inline)
                                            if image_placeholder in full_text:
                                                start_index = full_text.find(image_placeholder)
                                                end_index = start_index + len(image_placeholder)
                                                start_run_idx, start_offset, end_run_idx, end_offset = -1,-1,-1,-1 #简化表示
                                                current_pos = 0
                                                for i, run in enumerate(inline): # Find start/end run/offset
                                                    run_len = len(run.text)
                                                    if start_run_idx == -1 and start_index < current_pos + run_len: start_run_idx=i; start_offset=start_index-current_pos
                                                    if end_run_idx == -1 and end_index <= current_pos + run_len: end_run_idx=i; end_offset=end_index-current_pos; break
                                                    current_pos += run_len

                                                if start_run_idx != -1 and end_run_idx != -1:
                                                    try:
                                                        # --- 修正和改进的清理与插入逻辑 (表格内) ---
                                                        # 1. 处理结束 Run
                                                        if start_run_idx != end_run_idx:
                                                            if end_run_idx < len(inline):
                                                                end_run = inline[end_run_idx]
                                                                end_run.text = end_run.text[end_offset:]

                                                        # 2. 清理中间 Runs
                                                        for i in range(end_run_idx - 1, start_run_idx, -1):
                                                            if i < len(inline):
                                                                inline[i].text = ""

                                                        # 3. 处理开始 Run (清空并插入)
                                                        start_run = inline[start_run_idx]
                                                        start_run.text = start_run.text[:start_offset]

                                                        # 4. 插入图片
                                                        start_run.add_picture(absolute_image_path, width=Inches(5.0))

                                                        image_replaced = True
                                                        print(f"  成功替换并插入图片于表格单元格。")
                                                        break # 成功替换后，跳出当前单元格内段落的处理
                                                        # --- 修正逻辑结束 ---
                                                    except Exception as img_err: print(f"表格图片插入/清理出错: {img_err}")
                                      if image_replaced: break # 跳出单元格循环
                                 if image_replaced: break # 跳出表格行循环
                             if image_replaced: break # 跳出表格循环
                         if image_replaced: break # 跳出最外层表格循环 (以防万一)

            elif birdview_image_relative_path:
                 print(f"警告: 提供的鸟瞰图路径检查失败或文件不存在: {birdview_image_relative_path} (绝对路径: {absolute_image_path})")
        else:
             print("[word_template.py] 未找到 '鸟瞰图_path'，跳过图片替换。")
        # --- 图片替换逻辑结束 ---

        # --- 根据标准选择处理方式 ---
        if standard in ['四川省标', '成都市标']:
            # 计算并添加 '标准项目总分' 到 project_info 以便书签处理函数使用
            try:
                project_total_score = float(project_info.get('项目总分', '0'))
                standard_total_score = (project_total_score + 400) / 10
                project_info['标准项目总分'] = f"{standard_total_score:.1f}"
                print(f"  计算得到 '标准项目总分' (for bookmarks): {project_info['标准项目总分']}")
            except ValueError:
                print(f"  错误: '项目总分' '{project_info.get('项目总分')}' 无法计算 '标准项目总分'")
                project_info['标准项目总分'] = "" # 添加空字符串以防万一
            except Exception as calc_err: # 添加通用错误捕获
                 print(f"  计算 '标准项目总分' 时发生未知错误: {calc_err}")
                 project_info['标准项目总分'] = ""

            # 调用书签处理函数 (现在使用原始逻辑)
            replace_bookmarks_in_doc(doc, project_info, score_data, standard)
            print(f"已调用书签处理函数完成 {standard} 模板。")

        elif standard == '国标' and "安徽" in project_info.get('项目地点', ''):
            # --- 执行当前的 {占位符} 文本替换逻辑 (保持不变) ---
            # ... (此处省略未改变的占位符处理代码) ...
            print("--- 开始处理文本占位符 (国标安徽) ---")
            text_replacements = {}
            # ... (准备 text_replacements, 包括计算 标准项目总分) ...
            # ... (遍历段落和表格, 查找和替换占位符) ...
            # --- VVV 占位符替换逻辑 (保持不变) VVV --- 
            for key, value in project_info.items():
                if key.endswith('_path'): continue
                placeholder = f"{{{key}}}"
                text_replacements[placeholder] = str(value) if value is not None else ''
            try:
                # Calculate standard score for placeholder replacement
                project_total_score_placeholder = "{项目总分}"
                if project_total_score_placeholder in text_replacements:
                   project_total_score_str = text_replacements[project_total_score_placeholder]
                   project_total_score = float(project_total_score_str)
                   standard_total_score = (project_total_score + 400) / 10
                   text_replacements["{标准项目总分}"] = f"{standard_total_score:.1f}"
            except Exception: text_replacements["{标准项目总分}"] = ""
            for item in score_data:
                clause_num = item.get("条文号", "")
                if clause_num:
                   text_replacements[f"{{{clause_num}}}"] = str(item.get("得分", "0"))
                   text_replacements[f"{{{clause_num}措施}}"] = item.get("技术措施", "")
                   text_replacements[f"{{{clause_num}达标}}"] = item.get("是否达标", "")
                   text_replacements[f"{{{clause_num}分类}}"] = item.get("分类", "")
            # (Loop through paragraphs and tables, find runs, replace placeholders)
            elements_to_process = list(doc.paragraphs)
            for table in doc.tables: # Add table paragraphs
                for row in table.rows: 
                    for cell in row.cells: elements_to_process.extend(cell.paragraphs)
            for p in elements_to_process:
                 inline = p.runs
                 # (Detailed run-level placeholder replacement logic remains here)
                 runs_text = [(run, run.text) for run in inline] 
                 current_paragraph_text = "".join(rt[1] for rt in runs_text)
                 if '{' in current_paragraph_text:
                     for placeholder, replacement_text in text_replacements.items():
                         while placeholder in current_paragraph_text:
                             # ... (find start/end run/offset) ...
                             start_index=current_paragraph_text.find(placeholder); end_index=start_index+len(placeholder)
                             start_run_idx,start_offset,end_run_idx,end_offset=-1,-1,-1,-1; current_pos=0
                             for i, (run, text) in enumerate(runs_text):
                                 run_len=len(text)
                                 if start_run_idx==-1 and start_index<current_pos+run_len: start_run_idx=i; start_offset=start_index-current_pos
                                 if end_run_idx==-1 and end_index<=current_pos+run_len: end_run_idx=i; end_offset=end_index-current_pos; break
                                 current_pos+=run_len
                             if start_run_idx!=-1 and end_run_idx!=-1:
                                 # (Apply replacement across runs, clean middle, handle end run)
                                 try:
                                     target_run, target_text = runs_text[start_run_idx]
                                     target_run.text = target_text[:start_offset] + replacement_text
                                     runs_text[start_run_idx] = (target_run, target_run.text)
                                     for i in range(start_run_idx + 1, end_run_idx): # Clean middle runs
                                         if i < len(runs_text): runs_text[i][0].text = ""; runs_text[i] = (runs_text[i][0], "")
                                     if start_run_idx != end_run_idx: # Handle end run
                                         if end_run_idx < len(runs_text):
                                             end_run, end_text = runs_text[end_run_idx]
                                             end_run.text = end_text[end_offset:]
                                             runs_text[end_run_idx] = (end_run, end_run.text)
                                     else: # Ends in the same run
                                         target_run.text += target_text[end_offset:]
                                         runs_text[start_run_idx] = (target_run, target_run.text)
                                     current_paragraph_text = "".join(rt[1] for rt in runs_text) # Update text for while loop
                                 except Exception as replace_err: print(f"占位符替换错误: {replace_err}"); break
                             else: break # Stop if placeholder not located
            # --- ^^^ 占位符替换逻辑结束 ^^^ ---
            print("--- 文本占位符处理完成 (国标安徽) ---")

        else:
            print(f"未匹配到有效的处理逻辑 (标准: '{standard}', 地点: '{project_info.get('项目地点', '')}')，跳过内容替换。")


        # --- 通用后续处理 ---
        # 处理特殊字符字体 (保持不变)
        modify_square_chars_font(doc)
        print("方框符号字体处理完成。")

        # 保存处理后的文档 (保持不变)
        output_dir = current_app.config.get('EXPORT_FOLDER', 'static/exports')
        os.makedirs(output_dir, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S%f")
        project_name = project_info.get('项目名称', '未知项目')
        safe_project_name = re.sub(r"[\\/*?:'\"<>|]", "_", project_name)
        output_filename = f"{safe_project_name}_{base_template_name}_{timestamp}.docx"
        output_path = os.path.join(output_dir, output_filename)
        try:
            print(f"正在保存初次替换后的文档: {output_path}")
            doc.save(output_path)
            print(f"文档已初步保存: {output_path}")
            
            # --- 更新目录 --- 
            # 需要绝对路径来传递给 COM
            # absolute_output_path = os.path.abspath(output_path) 
            # update_toc_with_com(absolute_output_path) # <-- Temporarily comment out COM call
            print("[word_template.py] 跳过 COM 目录更新 (临时禁用)")
            # --- 目录更新结束 ---

            print(f"=== 完成处理文档模板 ===")
            return (output_path, base_template_name) # 返回元组

        except Exception as save_err:
            print(f"保存文档时出错: {save_err}")
            print(traceback.format_exc())
            print(f"=== 处理文档失败 (保存阶段) ===")
            output_path = None # 保存失败，重置为 None
            return None

    except Exception as e:
        print(f"处理文档模板时发生意外错误: {str(e)}")
        print(traceback.format_exc())
        print(f"=== 处理文档失败 ===")
        return None

def replace_simple_placeholders_in_runs(runs, data_dict):
    """在给定的 runs 列表中替换简单的 {key} 占位符"""
    full_text = "".join(run.text for run in runs)
    
    # 使用正则表达式查找所有 {key} 形式的占位符
    placeholders = re.findall(r"\{([^{}]+?)\}", full_text)
    
    if not placeholders:
        return # 没有找到占位符，直接返回
    
    modified = False
    for key in placeholders:
        # 检查 data_dict 中是否有对应的键
        if key in data_dict:
            value = str(data_dict[key]) if data_dict[key] is not None else "" # 转换为字符串，None 转为空字符串
            
            # 构建要替换的占位符字符串，例如 "{项目名称}"
            placeholder_to_replace = f"{{{key}}}"
            
            # --- 复杂的替换逻辑，处理跨 run 的占位符 --- 
            # 因为一个占位符可能跨越多个 run，例如 "{项" 在一个 run，"目名称}" 在另一个 run
            # 我们需要迭代 runs，构建当前文本，查找并替换
            current_run_index = 0
            while current_run_index < len(runs):
                # 从当前 run 开始构建文本片段，直到找到完整的占位符或结束
                temp_text = ""
                runs_involved_indices = []
                for i in range(current_run_index, len(runs)):
                    temp_text += runs[i].text
                    runs_involved_indices.append(i)
                    if placeholder_to_replace in temp_text:
                        break # 找到了完整的占位符
               
                # 如果找到了占位符
                if placeholder_to_replace in temp_text:
                    first_run_index = runs_involved_indices[0]
                    last_run_index = runs_involved_indices[-1]
                   
                    # 在第一个 run 中执行替换
                    first_run_text = runs[first_run_index].text
                    start_index = temp_text.find(placeholder_to_replace)
                    # 计算占位符在第一个 run 中的起始位置
                    placeholder_start_in_first_run = start_index - sum(len(runs[j].text) for j in range(current_run_index, first_run_index))
                   
                    # 替换逻辑：只在第一个run中进行替换，并清空后续相关run的文本
                    runs[first_run_index].text = first_run_text[:placeholder_start_in_first_run] + value + first_run_text[placeholder_start_in_first_run + len(placeholder_to_replace):]
                   
                    # 清空参与构成占位符的后续runs的文本
                    for i in range(first_run_index + 1, last_run_index + 1):
                        # 需要确保不超出 runs_involved_indices 的范围
                        if i in runs_involved_indices:
                             runs[i].text = ""
                   
                    modified = True
                    # 从替换后的下一个 run 开始继续检查（或从当前 run 重新开始以处理同一 run 内的多个占位符）
                    # 为简单起见，我们从第一个被修改的 run 之后开始，这可能错过同一 run 的后续占位符，但更安全
                    current_run_index = first_run_index + 1
                else:
                    # 如果从 current_run_index 开始没有找到完整的占位符，移动到下一个 run
                    current_run_index += 1
            # --- 复杂替换逻辑结束 --- 
           
    # 如果进行了修改，可以考虑合并相邻的、格式相同的空 run 或 run，但这比较复杂，暂时省略
    return modified

def replace_generic_placeholders(doc_path, data_dict):
    """
    打开 Word 文档，并替换所有段落和表格中的简单 {key} 占位符。
    
    参数:
        doc_path: Word 文档的路径。
        data_dict: 包含占位符键和替换值的字典。
    返回:
        bool: 如果进行了任何替换则返回 True，否则返回 False。
    """
    try:
        doc = Document(doc_path)
        print(f"[word_template.py] 开始通用占位符替换: {doc_path}")
        print(f"  替换数据: {list(data_dict.keys())}")
        
        overall_modified = False
        
        # 遍历文档中的所有段落
        for para in doc.paragraphs:
            if replace_simple_placeholders_in_runs(para.runs, data_dict):
                overall_modified = True
               
        # 遍历文档中的所有表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if replace_simple_placeholders_in_runs(para.runs, data_dict):
                            overall_modified = True
                           
        if overall_modified:
            doc.save(doc_path)
            print(f"[word_template.py] 通用占位符替换完成并保存: {doc_path}")
        else:
            print(f"[word_template.py] 未执行任何通用占位符替换: {doc_path}")
           
        return overall_modified

    except Exception as e:
        print(f"[word_template.py] 通用占位符替换失败: {doc_path}, 错误: {e}")
        print(traceback.format_exc())
        return False # 表示替换失败