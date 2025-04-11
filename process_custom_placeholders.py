import os
import base64
import io
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def process_custom_placeholders(docx_path, data):
    """
    处理Word文档中的自定义占位符，特别是图片和表格
    
    参数:
        docx_path: 已经处理过基本文本替换的Word文档路径
        data: 包含地图截图和公交站点列表的数据字典
    
    返回:
        处理后的文档路径
    """
    try:
        print("\n=== 开始处理自定义占位符 ===")
        
        # 加载文档
        doc = Document(docx_path)
        
        # 处理地图截图占位符
        map_image = data.get('地图截图', None)
        has_map_placeholder = False
        
        # 处理段落中的占位符
        for i, paragraph in enumerate(doc.paragraphs):
            if '{地图截图}' in paragraph.text:
                has_map_placeholder = True
                print(f"找到地图截图占位符，在段落 {i}")
                
                # 保存占位符段落位置
                map_placeholder_index = i
                
                # 清空段落内容
                paragraph.clear()
                
                # 如果有地图截图数据，添加图片
                if map_image and isinstance(map_image, str):
                    try:
                        # 准备图片数据
                        if "base64," in map_image:
                            map_image = map_image.split("base64,")[1]
                        
                        # 解码base64数据
                        image_data = base64.b64decode(map_image)
                        image_stream = io.BytesIO(image_data)
                        
                        # 直接在占位符段落添加图片
                        run = paragraph.add_run()
                        # 添加图片，设置合适的宽度
                        image = run.add_picture(image_stream, width=Cm(15))
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        print("已添加地图截图到占位符位置")
                    except Exception as e:
                        print(f"添加图片时出错: {str(e)}")
                        # 添加错误提示
                        run = paragraph.add_run("[图片数据处理失败]")
                        run.italic = True
                        run.font.color.rgb = (255, 0, 0)  # 红色
                        # 设置宋体小四
                        run.font.size = Pt(12)
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                else:
                    print("没有找到有效的地图截图数据")
                    run = paragraph.add_run("[无地图数据]")
                    # 设置宋体小四
                    run.font.size = Pt(12)
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                break
        
        # 处理公交站点列表占位符
        stations = data.get('stations', [])
        has_stations_placeholder = False
        
        # 再次遍历所有段落（因为插入图片可能会改变段落顺序）
        all_paragraphs = list(doc.paragraphs)  # 创建段落列表的副本
        for i, paragraph in enumerate(all_paragraphs):
            if '{公交站点列表}' in paragraph.text:
                has_stations_placeholder = True
                print(f"找到公交站点列表占位符，在段落 {i}")
                
                # 保存占位符的引用
                stations_paragraph = paragraph
                
                # 清空段落内容
                paragraph.clear()
                
                # 如果有站点数据，创建表格
                if stations and len(stations) > 0:
                    try:
                        # 添加标题
                        run = paragraph.add_run("周边公交站点列表")
                        run.bold = True
                        # 设置宋体小四(12磅)
                        run.font.size = Pt(12)
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # 在当前段落后插入表格
                        table = doc.add_table(rows=len(stations) + 1, cols=5)
                        table.style = 'Table Grid'
                        
                        # 获取表格元素并将其移动到占位符段落后面
                        tbl_element = table._element
                        paragraph._p.addnext(tbl_element)
                        
                        # 设置表头
                        header_cells = table.rows[0].cells
                        header_cells[0].text = '序号'
                        header_cells[1].text = '站点名称'
                        header_cells[2].text = '类型'
                        header_cells[3].text = '距离(米)'
                        header_cells[4].text = '详情'
                        
                        # 设置表头格式 - 使用宋体小四
                        for cell in header_cells:
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = cell.paragraphs[0].runs[0]
                            run.font.bold = True
                            run.font.size = Pt(12)  # 小四字号12磅
                            run.font.name = '宋体'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        
                        # 填充表格数据
                        for i, station in enumerate(stations):
                            row_cells = table.rows[i + 1].cells
                            row_cells[0].text = str(i + 1)
                            row_cells[1].text = station.get('name', '')
                            row_cells[2].text = station.get('type', '')
                            row_cells[3].text = str(station.get('distance', ''))
                            row_cells[4].text = station.get('detail', '')
                            
                            # 设置单元格对齐方式
                            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                            # 设置字体为宋体小四
                            for cell in row_cells:
                                for para in cell.paragraphs:
                                    for run in para.runs:
                                        run.font.size = Pt(12)  # 小四字号12磅
                                        run.font.name = '宋体'
                                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                                    
                                    # 如果段落没有run，需要创建一个
                                    if not para.runs:
                                        run = para.add_run(para.text)
                                        para.text = ""
                                        run.font.size = Pt(12)  # 小四字号12磅
                                        run.font.name = '宋体'
                                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        
                        print(f"已添加公交站点表格到占位符位置，共{len(stations)}行")
                    except Exception as e:
                        print(f"创建站点表格时出错: {str(e)}")
                        # 添加错误提示
                        run = paragraph.add_run("[站点数据处理失败]")
                        run.italic = True
                        run.font.color.rgb = (255, 0, 0)  # 红色
                        # 设置宋体小四
                        run.font.size = Pt(12)
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                else:
                    print("没有找到有效的站点数据")
                    run = paragraph.add_run("未找到周边800米范围内的公交站点")
                    # 设置宋体小四
                    run.font.size = Pt(12)
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                break
        
        # 如果没有找到占位符，输出警告
        if not has_map_placeholder:
            print("警告: 文档中未找到'{地图截图}'占位符")
        
        if not has_stations_placeholder:
            print("警告: 文档中未找到'{公交站点列表}'占位符")
            
        # 设置所有文本为宋体小四
        set_all_text_to_simsum_xiaosi(doc)
        
        # 生成新的输出文件名
        output_dir = os.path.dirname(docx_path)
        base_name = os.path.basename(docx_path)
        name_parts = os.path.splitext(base_name)
        # 使用相同的文件名
        output_path = os.path.join(output_dir, f"{name_parts[0]}{name_parts[1]}")
        
        # 保存文档
        doc.save(output_path)
        print(f"自定义占位符处理完成，文件已保存: {output_path}")
        
        return output_path
    except Exception as e:
        print(f"处理自定义占位符时出错: {str(e)}")
        import traceback
        print(traceback.format_exc())
        return docx_path  # 发生错误时返回原始文件路径

def set_all_text_to_simsum_xiaosi(doc):
    """
    将文档中的所有新添加的文本设置为宋体小四
    
    参数:
        doc: Document对象
    """
    # 根据段落内容判断是否是新添加的内容
    for paragraph in doc.paragraphs:
        # 如果段落内容包含我们可能添加的特定文本，则设置字体
        if any(marker in paragraph.text for marker in 
              ['周边公交站点列表', '[图片数据处理失败]', '[无地图数据]', 
               '未找到周边800米范围内的公交站点', '[站点数据处理失败]']):
            for run in paragraph.runs:
                run.font.name = '宋体'
                run.font.size = Pt(12)  # 小四字号12磅
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                
    # 检查表格中的文本
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    # 为表格内容设置字体
                    for run in paragraph.runs:
                        run.font.name = '宋体'
                        run.font.size = Pt(12)  # 小四字号12磅
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体') 